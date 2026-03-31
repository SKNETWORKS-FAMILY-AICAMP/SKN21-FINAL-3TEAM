"""
QA 테스트: 문서생성 LoRA 성능 테스트
=====================================
Phase 0: 환경 준비 (RunPod, .env, DB 시딩)
Phase 1: 기본 템플릿 (회의록/보고서/제안서) — LoRA 또는 API fallback
Phase 2: 커스텀 템플릿 업로드 + 생성

실행 방법:
  cd /c/SKN21-FINAL-3TEAM
  python tests/qa_doc_generate_lora.py

  # 백엔드 서버 띄운 상태에서 API E2E도 테스트:
  python tests/qa_doc_generate_lora.py --api
"""

import sys
import os
import json
import time
import asyncio
import argparse
from pathlib import Path
from datetime import datetime

# 프로젝트 루트 설정
PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))
sys.path.insert(0, str(PROJECT_ROOT / "backend"))

from dotenv import load_dotenv
load_dotenv(PROJECT_ROOT / ".env")

# ── 결과 저장용 ──
RESULTS: list[dict] = []
PASS_COUNT = 0
FAIL_COUNT = 0
WARN_COUNT = 0


def record(test_name: str, status: str, details: str = "", data: dict = None):
    """테스트 결과 기록"""
    global PASS_COUNT, FAIL_COUNT, WARN_COUNT
    icon = {"PASS": "✅", "FAIL": "❌", "WARN": "⚠️", "INFO": "ℹ️"}.get(status, "•")
    print(f"  {icon} [{status}] {test_name}: {details}")
    if status == "PASS":
        PASS_COUNT += 1
    elif status == "FAIL":
        FAIL_COUNT += 1
    elif status == "WARN":
        WARN_COUNT += 1
    RESULTS.append({"test": test_name, "status": status, "details": details, "data": data})


# ════════════════════════════════════════════════
# Phase 0: 환경 준비
# ════════════════════════════════════════════════

def phase0_check_env():
    """Phase 0: 환경 변수 및 RunPod 엔드포인트 확인"""
    print("\n" + "=" * 60)
    print("Phase 0: 환경 준비")
    print("=" * 60)

    # 0-1. RunPod 엔드포인트
    print("\n[0-1] RunPod 엔드포인트 확인...")
    vllm_url = os.getenv("VLLM_BASE_URL", "")
    vllm_key = os.getenv("VLLM_API_KEY", "")
    runpod_available = False

    if vllm_url and vllm_key:
        try:
            import httpx
            resp = httpx.get(
                f"{vllm_url}/models",
                headers={"Authorization": f"Bearer {vllm_key}"},
                timeout=15,
            )
            if resp.status_code == 200:
                models_data = resp.json()
                record("RunPod 엔드포인트", "PASS",
                       f"응답 OK, models={json.dumps(models_data, ensure_ascii=False)[:200]}")
                runpod_available = True
            else:
                record("RunPod 엔드포인트", "WARN",
                       f"HTTP {resp.status_code} — GPU 꺼짐 가능성, API fallback 사용")
        except Exception as e:
            record("RunPod 엔드포인트", "WARN", f"타임아웃/에러: {e} — API fallback 사용")
    else:
        record("RunPod 엔드포인트", "WARN", "VLLM_BASE_URL 또는 VLLM_API_KEY 미설정")

    # 0-2. .env 설정
    print("\n[0-2] .env 설정 확인...")
    env_checks = {
        "DOC_AGENT_MODE": ("sllm", os.getenv("DOC_AGENT_MODE", "")),
        "DOC_SLLM_TASKS": ("generate 포함", os.getenv("DOC_SLLM_TASKS", "")),
        "DOC_LORA_TASKS": ("generate 포함", os.getenv("DOC_LORA_TASKS", "")),
        "VLLM_USE_LORA": ("true", os.getenv("VLLM_USE_LORA", "")),
        "VLLM_MODEL": ("설정됨", os.getenv("VLLM_MODEL", "")),
        "OPENAI_API_KEY": ("설정됨 (fallback용)", os.getenv("OPENAI_API_KEY", "")[:10] + "..."),
    }
    for key, (expected, actual) in env_checks.items():
        if actual and actual != "...":
            record(f".env {key}", "PASS", f"{expected} → 실제: {actual[:30]}")
        else:
            record(f".env {key}", "FAIL", f"기대: {expected}, 실제: '{actual}'")

    return runpod_available


# ════════════════════════════════════════════════
# Phase 1: 기본 템플릿 테스트 (직접 Agent 함수 호출)
# ════════════════════════════════════════════════

TEST_INPUTS = {
    "meeting_minutes": {
        "user_input": (
            "제목: 2026년 3월 Sprint 회고 회의\n"
            "날짜: 2026-03-16\n"
            "참석자: 신지용, 윤경은, 진승언, 안혜빈, 문지영\n"
            "이번 스프린트에서 완료한 작업: 1) 문서 에이전트 LoRA 파인튜닝 완료 "
            "2) 프론트엔드 채팅 UI 개선 3) Google Calendar 연동 완료. "
            "논의 사항: vLLM 서빙 안정성 이슈로 RunPod Serverless로 전환 검토. "
            "API 응답 시간이 평균 5초에서 2초로 단축됨. "
            "다음 스프린트에서 E2E 테스트 진행 예정. "
            "경은이 RAG 파이프라인 성능 개선 필요하다고 제안. "
            "혜빈이 Google Sheets 연동 버그 수정 완료 보고."
        ),
        "checks": {
            "str_fields": ["title", "summary"],
            "list_fields": ["decisions", "action_items"],
            "action_items_struct": ["task", "assignee", "due_date"],
            "min_action_items": 2,
            "summary_not_copy": True,
        },
    },
    "report": {
        "user_input": (
            "제목: AI 모듈 개발 주간 업무보고서\n"
            "날짜: 2026-03-16\n"
            "보고 대상: 팀장 김철수. 부서: AI개발팀. 보고 유형: 주간. "
            "이번 주 주요 업무: 1) LoRA v2 파인튜닝 학습 완료 "
            "2) vLLM RunPod Serverless 배포 테스트 "
            "3) RAG 하이브리드 검색 정확도 85% 달성. "
            "이슈: RunPod GPU cold start 문제 (대기 15초). "
            "향후 계획: E2E 통합 테스트, 프론트엔드 연동."
        ),
        "checks": {
            "str_fields": ["title", "overview"],
            "extract_fields": {
                "department": "AI개발팀",
                "report_to": "김철수",
            },
            "list_fields": ["tasks"],
            "tasks_struct": ["item", "assignee", "progress", "start_date", "end_date"],
            "min_tasks": 3,
            "mention_fields": {"issues": "cold start", "next_plan": "E2E"},
        },
    },
    "proposal": {
        "user_input": (
            "제목: 사내 AI 업무 자동화 시스템 도입 제안서\n"
            "날짜: 2026-03-16\n"
            "제출처: 경영지원팀. 제안사: 듀듀 AI Labs. 담당자: 신지용. "
            "제안 배경: 반복 업무 월 320시간 소요. "
            "제안 목적: LangGraph 멀티에이전트로 문서/규정/일정 자동화. "
            "현황: 문서 작성 2시간, 규정 확인 30분 소요. "
            "제안 내용: 1) AI 문서 자동생성 2) RAG 규정 검색 3) 캘린더 연동. "
            "추진일정: 1단계-설계(3월), 2단계-개발(4-5월), 3단계-테스트(6월), 4단계-배포(7월). "
            "예산: 서버 50만/월, GPU 30만/월, 인건비 500만/월. "
            "기대효과: 문서작성 80% 단축, 규정판단 95% 정확도."
        ),
        "checks": {
            "str_fields": ["title", "purpose", "content", "expected_effect"],
            "extract_fields": {
                "submit_to": "경영지원팀",
                "company": "듀듀",
                "manager": "신지용",
            },
            "list_fields": ["schedule", "budget"],
            "schedule_struct": ["item", "phase1"],
            "budget_struct": ["item", "amount"],
            "min_schedule": 3,
            "min_budget": 3,
        },
    },
}


async def phase1_direct_agent_tests():
    """Phase 1: document_agent 직접 호출 테스트"""
    print("\n" + "=" * 60)
    print("Phase 1: 기본 템플릿 테스트 (Agent 직접 호출)")
    print("=" * 60)

    from ai.agents.document_agent import generate_document

    for template_type, test_spec in TEST_INPUTS.items():
        label = {"meeting_minutes": "회의록", "report": "보고서", "proposal": "제안서"}[template_type]
        print(f"\n{'─' * 40}")
        print(f"[테스트 1-{list(TEST_INPUTS).index(template_type)+1}] {label} 생성")
        print(f"{'─' * 40}")

        t0 = time.time()
        try:
            result = await generate_document(
                category=template_type,
                user_input=test_spec["user_input"],
            )
            elapsed = time.time() - t0
            record(f"{label} 생성 호출", "PASS", f"{elapsed:.2f}s")
        except Exception as e:
            record(f"{label} 생성 호출", "FAIL", f"에러: {e}")
            import traceback
            traceback.print_exc()
            continue

        # ── 결과 분석 ──
        data = result.get("data", {})
        checks = test_spec["checks"]
        model_name = result.get("model_name", "unknown")
        record(f"{label} 사용 모델", "INFO", model_name)

        # JSON 파싱 성공 여부
        if "content" in data and len(data) == 1:
            record(f"{label} JSON 파싱", "FAIL", "fallback {'content': raw} 반환됨")
        else:
            record(f"{label} JSON 파싱", "PASS", f"keys={list(data.keys())}")

        # 문자열 필드 확인
        for field in checks.get("str_fields", []):
            val = data.get(field, "")
            if val and len(str(val)) > 5:
                record(f"{label} {field}", "PASS", f"길이={len(str(val))}자")
            else:
                record(f"{label} {field}", "FAIL", f"비어있거나 너무 짧음: '{val}'")

        # 추출 필드 확인 (report_to, department 등)
        for field, expected_substr in checks.get("extract_fields", {}).items():
            val = str(data.get(field, ""))
            if expected_substr in val:
                record(f"{label} {field} 추출", "PASS", f"'{expected_substr}' 포함 → '{val}'")
            else:
                record(f"{label} {field} 추출", "WARN", f"기대 '{expected_substr}', 실제 '{val}'")

        # 배열 필드 확인
        for field in checks.get("list_fields", []):
            val = data.get(field, [])
            if isinstance(val, list) and len(val) > 0:
                record(f"{label} {field} 배열", "PASS", f"{len(val)}개 항목")
            else:
                record(f"{label} {field} 배열", "FAIL", f"타입={type(val).__name__}, 값={str(val)[:100]}")

        # action_items 구조 확인 (회의록)
        if "action_items_struct" in checks:
            ai_list = data.get("action_items", [])
            min_count = checks.get("min_action_items", 2)
            if len(ai_list) >= min_count:
                record(f"{label} action_items 수량", "PASS", f"{len(ai_list)}개 >= {min_count}")
            else:
                record(f"{label} action_items 수량", "FAIL", f"{len(ai_list)}개 < {min_count}")
            if ai_list and isinstance(ai_list[0], dict):
                for key in checks["action_items_struct"]:
                    if key in ai_list[0]:
                        record(f"{label} action_items.{key}", "PASS", f"값: {ai_list[0][key]}")
                    else:
                        record(f"{label} action_items.{key}", "FAIL", f"키 없음. keys={list(ai_list[0].keys())}")

        # tasks 구조 확인 (보고서)
        if "tasks_struct" in checks:
            tasks = data.get("tasks", [])
            min_count = checks.get("min_tasks", 3)
            if len(tasks) >= min_count:
                record(f"{label} tasks 수량", "PASS", f"{len(tasks)}개 >= {min_count}")
            else:
                record(f"{label} tasks 수량", "WARN", f"{len(tasks)}개 < {min_count}")
            if tasks and isinstance(tasks[0], dict):
                for key in checks["tasks_struct"]:
                    if key in tasks[0]:
                        record(f"{label} tasks.{key}", "PASS", f"값: {tasks[0][key]}")
                    else:
                        record(f"{label} tasks.{key}", "FAIL", f"키 없음. keys={list(tasks[0].keys())}")

        # schedule/budget 구조 확인 (제안서)
        if "schedule_struct" in checks:
            sch = data.get("schedule", [])
            min_count = checks.get("min_schedule", 3)
            if len(sch) >= min_count:
                record(f"{label} schedule 수량", "PASS", f"{len(sch)}개 >= {min_count}")
            else:
                record(f"{label} schedule 수량", "WARN", f"{len(sch)}개 < {min_count}")
            if sch and isinstance(sch[0], dict):
                for key in checks["schedule_struct"]:
                    if key in sch[0]:
                        record(f"{label} schedule.{key}", "PASS", f"값: {sch[0][key]}")
                    else:
                        record(f"{label} schedule.{key}", "FAIL", f"키 없음")

        if "budget_struct" in checks:
            bud = data.get("budget", [])
            min_count = checks.get("min_budget", 3)
            if len(bud) >= min_count:
                record(f"{label} budget 수량", "PASS", f"{len(bud)}개 >= {min_count}")
            else:
                record(f"{label} budget 수량", "WARN", f"{len(bud)}개 < {min_count}")
            if bud and isinstance(bud[0], dict):
                for key in checks["budget_struct"]:
                    if key in bud[0]:
                        record(f"{label} budget.{key}", "PASS", f"값: {bud[0][key]}")
                    else:
                        record(f"{label} budget.{key}", "FAIL", f"키 없음")

        # mention 필드 (보고서 issues/next_plan)
        for field, keyword in checks.get("mention_fields", {}).items():
            val = str(data.get(field, "")).lower()
            if keyword.lower() in val:
                record(f"{label} {field} 내용", "PASS", f"'{keyword}' 언급됨")
            else:
                record(f"{label} {field} 내용", "WARN", f"'{keyword}' 미언급. 실제: {val[:80]}")

        # summary가 입력 복붙이 아닌지 확인
        if checks.get("summary_not_copy"):
            summary = str(data.get("summary", ""))
            raw_input = test_spec["user_input"]
            if summary and summary != raw_input and len(summary) < len(raw_input) * 0.8:
                record(f"{label} summary 품질", "PASS", "입력 복붙 아님, 요약됨")
            elif not summary:
                record(f"{label} summary 품질", "FAIL", "summary 비어있음")
            else:
                record(f"{label} summary 품질", "WARN", "입력과 유사 — 요약 품질 낮음")

        # DOCX 파일 존재 확인
        docx_path = result.get("docx_path", "")
        if docx_path and Path(docx_path).exists():
            docx_size = Path(docx_path).stat().st_size
            record(f"{label} DOCX 생성", "PASS", f"크기={docx_size:,}B, 경로={docx_path}")

            # DOCX 테이블 수 확인
            try:
                from docx import Document as DocxDoc
                doc = DocxDoc(docx_path)
                table_count = len(doc.tables)
                expected_tables = {"meeting_minutes": 5, "report": 7, "proposal": 9}
                exp = expected_tables.get(template_type, 1)
                if table_count >= exp * 0.6:  # 60% 이상이면 PASS
                    record(f"{label} DOCX 테이블 수", "PASS", f"{table_count}개 (기대: ~{exp})")
                else:
                    record(f"{label} DOCX 테이블 수", "WARN", f"{table_count}개 (기대: ~{exp})")

                # 빈 셀 비율 확인
                total_cells = 0
                empty_cells = 0
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            total_cells += 1
                            if not cell.text.strip():
                                empty_cells += 1
                if total_cells > 0:
                    empty_ratio = empty_cells / total_cells
                    if empty_ratio < 0.3:
                        record(f"{label} DOCX 빈 셀 비율", "PASS", f"{empty_ratio:.1%} ({empty_cells}/{total_cells})")
                    else:
                        record(f"{label} DOCX 빈 셀 비율", "WARN", f"{empty_ratio:.1%} ({empty_cells}/{total_cells})")
            except Exception as e:
                record(f"{label} DOCX 분석", "WARN", f"python-docx 분석 실패: {e}")
        else:
            record(f"{label} DOCX 생성", "FAIL", f"파일 없음: {docx_path}")

        # data dump (디버깅용)
        print(f"\n  📋 LLM 응답 데이터 (요약):")
        for k, v in data.items():
            if isinstance(v, list):
                print(f"    {k}: [{len(v)}개 항목]")
                for i, item in enumerate(v[:3]):
                    print(f"      [{i}] {json.dumps(item, ensure_ascii=False)[:120]}")
            elif isinstance(v, str) and len(v) > 100:
                print(f"    {k}: {v[:100]}...")
            else:
                print(f"    {k}: {v}")


# ════════════════════════════════════════════════
# Phase 1-API: HTTP API 기반 E2E 테스트
# ════════════════════════════════════════════════

async def phase1_api_tests(base_url: str = "http://localhost:8000"):
    """Phase 1 API: 백엔드 HTTP 엔드포인트 테스트"""
    print("\n" + "=" * 60)
    print("Phase 1-API: HTTP API E2E 테스트")
    print("=" * 60)

    import httpx

    # Health check
    try:
        async with httpx.AsyncClient(timeout=5) as client:
            resp = await client.get(f"{base_url}/health")
            if resp.status_code != 200:
                record("백엔드 Health", "FAIL", f"HTTP {resp.status_code}")
                return
            record("백엔드 Health", "PASS", "서버 실행 중")
    except Exception as e:
        record("백엔드 Health", "FAIL", f"서버 미실행: {e}")
        print("  → 백엔드를 먼저 실행하세요: cd backend && uvicorn app.main:app --reload --port 8000")
        return

    async with httpx.AsyncClient(base_url=base_url, timeout=120) as client:
        # 회원가입 + 로그인
        test_email = f"qa_test_{int(time.time())}@test.com"
        try:
            reg_resp = await client.post("/api/v1/auth/register", json={
                "email": test_email, "password": "test1234", "name": "QA테스터", "team": "개발"
            })
            if reg_resp.status_code in (201, 409):
                record("회원가입", "PASS", f"HTTP {reg_resp.status_code}")
            else:
                record("회원가입", "FAIL", f"HTTP {reg_resp.status_code}: {reg_resp.text[:100]}")
        except Exception as e:
            record("회원가입", "FAIL", str(e))

        login_resp = await client.post("/api/v1/auth/login", json={
            "email": test_email, "password": "test1234"
        })
        if login_resp.status_code != 200:
            # 기존 계정 시도
            login_resp = await client.post("/api/v1/auth/login", json={
                "email": "admin@test.com", "password": "admin1234"
            })
        if login_resp.status_code != 200:
            record("로그인", "FAIL", f"HTTP {login_resp.status_code}")
            return

        token = login_resp.json()["access_token"]
        headers = {"Authorization": f"Bearer {token}"}
        record("로그인", "PASS", f"토큰 발급됨")

        # 시스템 템플릿 확인
        tpl_resp = await client.get("/api/v1/documents/templates/", headers=headers)
        if tpl_resp.status_code == 200:
            templates = tpl_resp.json()
            system_tpls = [t for t in templates if t.get("is_system")]
            record("시스템 템플릿 조회", "PASS", f"전체 {len(templates)}개, 시스템 {len(system_tpls)}개")
        else:
            record("시스템 템플릿 조회", "FAIL", f"HTTP {tpl_resp.status_code}")

        # API를 통한 문서 생성 테스트
        api_test_cases = [
            {
                "name": "회의록 (API)",
                "body": {
                    "template_type": "meeting_minutes",
                    "title": "2026년 3월 Sprint 회고 회의",
                    "date": "2026-03-16",
                    "attendees": ["신지용", "윤경은", "진승언", "안혜빈", "문지영"],
                    "content": (
                        "이번 스프린트에서 완료한 작업: 1) 문서 에이전트 LoRA 파인튜닝 완료 "
                        "2) 프론트엔드 채팅 UI 개선 3) Google Calendar 연동 완료. "
                        "논의 사항: vLLM 서빙 안정성 이슈로 RunPod Serverless로 전환 검토. "
                        "API 응답 시간이 평균 5초에서 2초로 단축됨. "
                        "다음 스프린트에서 E2E 테스트 진행 예정."
                    ),
                },
            },
            {
                "name": "보고서 (API)",
                "body": {
                    "template_type": "report",
                    "title": "AI 모듈 개발 주간 업무보고서",
                    "date": "2026-03-16",
                    "content": (
                        "보고 대상: 팀장 김철수. 부서: AI개발팀. 보고 유형: 주간. "
                        "이번 주 주요 업무: 1) LoRA v2 파인튜닝 학습 완료 "
                        "2) vLLM RunPod Serverless 배포 테스트 "
                        "3) RAG 하이브리드 검색 정확도 85% 달성."
                    ),
                },
            },
            {
                "name": "제안서 (API)",
                "body": {
                    "template_type": "proposal",
                    "title": "사내 AI 업무 자동화 시스템 도입 제안서",
                    "date": "2026-03-16",
                    "content": (
                        "제출처: 경영지원팀. 제안사: 듀듀 AI Labs. 담당자: 신지용. "
                        "제안 배경: 반복 업무 월 320시간 소요. "
                        "제안 내용: 1) AI 문서 자동생성 2) RAG 규정 검색 3) 캘린더 연동. "
                        "추진일정: 1단계-설계(3월), 2단계-개발(4-5월). "
                        "예산: 서버 50만/월, GPU 30만/월, 인건비 500만/월. "
                        "기대효과: 문서작성 80% 단축."
                    ),
                },
            },
        ]

        for tc in api_test_cases:
            print(f"\n  → {tc['name']}...")
            t0 = time.time()
            try:
                gen_resp = await client.post(
                    "/api/v1/documents/generate",
                    json=tc["body"],
                    headers=headers,
                )
                elapsed = time.time() - t0
                if gen_resp.status_code == 200:
                    result = gen_resp.json()
                    record(f"{tc['name']} 생성", "PASS",
                           f"{elapsed:.2f}s, doc_id={result.get('document_id', '?')}, "
                           f"model={result.get('model_name', '?')}")

                    # DOCX 다운로드 테스트
                    dl_url = result.get("download_url", "")
                    if dl_url:
                        dl_resp = await client.get(dl_url, headers=headers)
                        if dl_resp.status_code == 200:
                            record(f"{tc['name']} DOCX 다운로드", "PASS",
                                   f"크기={len(dl_resp.content):,}B")
                        else:
                            record(f"{tc['name']} DOCX 다운로드", "FAIL",
                                   f"HTTP {dl_resp.status_code}")

                    # 데이터 확인
                    data = result.get("data", {})
                    keys = list(data.keys())
                    record(f"{tc['name']} 데이터 키", "INFO", f"{keys}")
                else:
                    record(f"{tc['name']} 생성", "FAIL",
                           f"HTTP {gen_resp.status_code}: {gen_resp.text[:200]}")
            except Exception as e:
                record(f"{tc['name']} 생성", "FAIL", f"에러: {e}")


# ════════════════════════════════════════════════
# Phase 2: 커스텀 템플릿 테스트
# ════════════════════════════════════════════════

async def phase2_custom_template_test():
    """Phase 2: 커스텀 DOCX 업로드 + 생성 (직접 호출)"""
    print("\n" + "=" * 60)
    print("Phase 2: 커스텀 템플릿 테스트")
    print("=" * 60)

    # 커스텀 DOCX 양식 생성 (테스트용)
    print("\n[2-1] 테스트용 커스텀 DOCX 양식 생성...")
    try:
        from docx import Document as DocxDoc

        test_docx_path = str(PROJECT_ROOT / "tests" / "test_custom_template.docx")
        doc = DocxDoc()
        doc.add_heading("프로젝트 킥오프 회의록", level=1)
        doc.add_paragraph("프로젝트명: {{project_name}}")
        doc.add_paragraph("회의 일시: {{meeting_date}}")
        doc.add_paragraph("참석자: {{attendees}}")
        doc.add_paragraph("회의 목적: {{purpose}}")
        doc.add_paragraph("논의 사항: {{discussion}}")
        doc.add_paragraph("결정 사항: {{decisions}}")
        doc.add_paragraph("다음 단계: {{next_steps}}")
        doc.save(test_docx_path)
        record("커스텀 DOCX 양식 생성", "PASS", test_docx_path)
    except Exception as e:
        record("커스텀 DOCX 양식 생성", "FAIL", str(e))
        return

    # template_extractor 필드 추출 테스트
    print("\n[2-2] 필드 추출 테스트...")
    try:
        from ai.document_parser.template_extractor import extract_template_fields, fields_to_prompt
        fields = extract_template_fields(test_docx_path)
        if fields and len(fields) >= 3:
            record("extract_template_fields()", "PASS",
                   f"{len(fields)}개 필드: {[f.get('key','?') for f in fields]}")
        else:
            record("extract_template_fields()", "WARN",
                   f"추출 필드 수 부족: {len(fields) if fields else 0}개")
            # 수동 필드 정의 (fallback)
            fields = [
                {"key": "project_name", "description": "프로젝트 이름"},
                {"key": "meeting_date", "description": "회의 일시"},
                {"key": "attendees", "description": "참석자 목록"},
                {"key": "purpose", "description": "회의 목적"},
                {"key": "discussion", "description": "논의 사항"},
                {"key": "decisions", "description": "결정 사항 목록"},
                {"key": "next_steps", "description": "다음 단계"},
            ]

        prompt = fields_to_prompt(fields)
        if prompt and len(prompt) > 50:
            record("fields_to_prompt()", "PASS", f"프롬프트 길이={len(prompt)}자")
        else:
            record("fields_to_prompt()", "FAIL", f"프롬프트 비어있음: '{prompt[:50]}'")
    except Exception as e:
        record("필드 추출", "FAIL", str(e))
        import traceback
        traceback.print_exc()
        return

    # 커스텀 템플릿으로 직접 LLM 호출 테스트
    print("\n[2-3] 커스텀 필드 기반 문서 생성 (LLM 직접 호출)...")
    try:
        from ai.llm.prompts import DOC_GENERATE_SLLM_PROMPT
        from ai.agents.document_agent import _call_llm

        field_spec = fields_to_prompt(fields)
        user_prompt = (
            f"다음 내용을 바탕으로 문서를 JSON 형식으로 작성해주세요.\n\n"
            f"[문서 유형] 프로젝트 킥오프 회의록\n\n"
            f"[필드 명세]\n{field_spec}\n\n"
            f"[회의 내용]\n"
            f"프로젝트 킥오프 미팅. 참석자: 김영희, 이철수. "
            f"목적: Q2 프로젝트 범위 확정. "
            f"논의: MVP 기능으로 제한 합의. 일정 4~6월."
        )

        t0 = time.time()
        result_str = await _call_llm(DOC_GENERATE_SLLM_PROMPT, user_prompt, json_mode=True, task="generate")
        elapsed = time.time() - t0
        record("커스텀 LLM 호출", "PASS", f"{elapsed:.2f}s, 응답 길이={len(result_str)}자")

        # JSON 파싱
        try:
            data = json.loads(result_str)
            field_keys = [f.get("key", "") for f in fields]
            matched = [k for k in field_keys if k in data]
            match_ratio = len(matched) / len(field_keys) if field_keys else 0
            record("커스텀 JSON 파싱", "PASS",
                   f"매칭 {len(matched)}/{len(field_keys)} ({match_ratio:.0%}): {matched}")
            if match_ratio >= 0.6:
                record("커스텀 필드 매칭률", "PASS", f"{match_ratio:.0%} >= 60%")
            else:
                record("커스텀 필드 매칭률", "WARN", f"{match_ratio:.0%} < 60%")

            print(f"\n  📋 커스텀 LLM 응답:")
            for k, v in data.items():
                val_str = json.dumps(v, ensure_ascii=False) if not isinstance(v, str) else v
                print(f"    {k}: {val_str[:120]}")

        except json.JSONDecodeError:
            record("커스텀 JSON 파싱", "FAIL", f"파싱 실패. 원본: {result_str[:200]}")

    except Exception as e:
        record("커스텀 LLM 호출", "FAIL", str(e))
        import traceback
        traceback.print_exc()

    # 테스트 파일 정리
    try:
        os.unlink(test_docx_path)
    except Exception:
        pass


# ════════════════════════════════════════════════
# 결과 요약 + 파일 저장
# ════════════════════════════════════════════════

def print_summary():
    """테스트 결과 요약 출력"""
    print("\n" + "=" * 60)
    print("QA 테스트 결과 요약")
    print("=" * 60)
    print(f"  ✅ PASS: {PASS_COUNT}")
    print(f"  ❌ FAIL: {FAIL_COUNT}")
    print(f"  ⚠️  WARN: {WARN_COUNT}")
    print(f"  총 {len(RESULTS)}개 항목")

    if FAIL_COUNT > 0:
        print("\n  ❌ 실패 항목:")
        for r in RESULTS:
            if r["status"] == "FAIL":
                print(f"    - {r['test']}: {r['details']}")

    if WARN_COUNT > 0:
        print("\n  ⚠️  경고 항목:")
        for r in RESULTS:
            if r["status"] == "WARN":
                print(f"    - {r['test']}: {r['details']}")


def save_results():
    """결과를 JSON 파일로 저장"""
    output_dir = PROJECT_ROOT / "tests" / "qa_results"
    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = output_dir / f"qa_doc_generate_{timestamp}.json"

    summary = {
        "timestamp": timestamp,
        "pass": PASS_COUNT,
        "fail": FAIL_COUNT,
        "warn": WARN_COUNT,
        "total": len(RESULTS),
        "env": {
            "DOC_AGENT_MODE": os.getenv("DOC_AGENT_MODE", ""),
            "DOC_SLLM_TASKS": os.getenv("DOC_SLLM_TASKS", ""),
            "DOC_LORA_TASKS": os.getenv("DOC_LORA_TASKS", ""),
            "VLLM_USE_LORA": os.getenv("VLLM_USE_LORA", ""),
        },
        "results": RESULTS,
    }
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)
    print(f"\n  📄 결과 저장: {output_path}")
    return output_path


# ════════════════════════════════════════════════
# Main
# ════════════════════════════════════════════════

async def main():
    parser = argparse.ArgumentParser(description="QA 테스트: 문서생성 LoRA 성능")
    parser.add_argument("--api", action="store_true", help="HTTP API E2E 테스트도 실행")
    parser.add_argument("--api-only", action="store_true", help="API 테스트만 실행")
    parser.add_argument("--base-url", default="http://localhost:8000", help="백엔드 URL")
    args = parser.parse_args()

    print("╔══════════════════════════════════════════════════════════╗")
    print("║   QA 테스트: 문서생성 LoRA 성능 테스트                  ║")
    print("║   날짜: " + datetime.now().strftime("%Y-%m-%d %H:%M:%S") + "                        ║")
    print("╚══════════════════════════════════════════════════════════╝")

    # Phase 0
    runpod_available = phase0_check_env()

    if not args.api_only:
        # Phase 1: 직접 Agent 호출
        await phase1_direct_agent_tests()

        # Phase 2: 커스텀 템플릿
        await phase2_custom_template_test()

    if args.api or args.api_only:
        # Phase 1-API: HTTP E2E
        await phase1_api_tests(args.base_url)

    # 결과 요약
    print_summary()
    result_path = save_results()

    # 종료 코드
    if FAIL_COUNT > 0:
        print(f"\n  ⛔ {FAIL_COUNT}개 FAIL — 수정 필요")
        sys.exit(1)
    else:
        print(f"\n  🎉 모든 테스트 통과!")
        sys.exit(0)


if __name__ == "__main__":
    asyncio.run(main())
