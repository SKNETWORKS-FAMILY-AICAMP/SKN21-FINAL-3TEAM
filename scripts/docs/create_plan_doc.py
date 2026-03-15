# -*- coding: utf-8 -*-
"""기획서 v2 생성 스크립트 - 중간발표 기준 업데이트"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

FONT = "Arial Unicode MS"
doc = Document()

# ─── 페이지 설정 (기존 문서와 동일) ───
section = doc.sections[0]
section.page_width = Emu(7562215)
section.page_height = Emu(10689590)
section.left_margin = Emu(914400)
section.right_margin = Emu(914400)
section.top_margin = Emu(810260)
section.bottom_margin = Emu(346710)


# ─── 헬퍼 함수 ───
def add_heading_styled(text, level=1):
    sizes = {1: 21, 2: 15, 3: 11, 4: 12}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(sizes.get(level, 10))
    run.bold = True
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if level == 3:
        run.font.color.rgb = RGBColor(0x43, 0x43, 0x43)
    elif level == 4:
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.style = doc.styles[f"Heading {level}"]
    return p


def add_para(text, bold=False, size=10, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    return p


def add_bullet(text, size=10, indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if indent_level > 0:
        p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.name = FONT
        run.font.size = Pt(9)
        run.bold = True
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "D5E8F0"
        })
        shading.append(shading_elem)
    # Data
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = FONT
            run.font.size = Pt(9)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return table


# ══════════════════════════════════════
# 표지
# ══════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

add_heading_styled("SK네트웍스 Family AI 과정 21기", 1)
add_heading_styled("최종 프로젝트 기획서", 1)

for _ in range(6):
    doc.add_paragraph()

add_heading_styled("산출물 단계", 2)
add_para("기획")
add_heading_styled("평가 산출물", 2)
add_para("최종 프로젝트 기획서")
add_heading_styled("제출 일자", 2)
add_para("2026.02.06 (초판) / 2026.02.27 (중간발표 반영 개정)")
add_heading_styled("깃허브 경로", 2)
add_para("https://github.com/SKNETWORKS-FAMILY-AICAMP/SKN21-FINAL-3TEAM.git")
add_heading_styled("작성 팀원", 2)
add_para("문지영, 신지용, 안혜빈, 윤경은, 진승언")

doc.add_page_break()

# ══════════════════════════════════════
# 1. 프로젝트 주제
# ══════════════════════════════════════
add_para("1. 프로젝트 주제", bold=True, size=15)
add_heading_styled("WorkFlow Agent (듀듀): 사내 AI 업무 자동화 시스템", 3)
add_para("LangGraph 기반 멀티 Agent 오케스트레이션으로 규정 판단 · 문서 처리 · 일정 관리를 자동화하는 AI 플랫폼")

# ── 1.1. 프로젝트 배경 및 정의 ──
add_heading_styled("1.1. 프로젝트 배경 및 정의", 2)
add_heading_styled("1.1.1. 배경", 3)
add_para("기업 내 단순 반복 업무를 지능형 AI를 통해 자동화하여 실무자의 생산성을 극대화하려는 시장의 니즈가 급증하고 있다. 10~100인 규모의 IT 스타트업과 중소기업에서는 전담 HR·법무팀 없이 운영되는 경우가 대부분이며, 다음과 같은 문제에 직면한다.")

add_bullet("복잡해진 사내 규정과 정보의 파편화: 인사, 보안 등 사내 규정이 방대해져 기존 시스템으로는 정확한 규정검색, 판단, 예외 상황 대응이 어려움")
add_bullet("의사결정 병목: Slack에서 CTO나 팀 리더에게 직접 질문 → 핵심 인력의 업무 시간이 반복 질문 대응에 소모 (건당 평균 30분~1시간)")
add_bullet("수작업 중심의 행정 업무 병목: 회의록 작성, 문서 초안 기획, 복수의 앱(메일, 캘린더 등)을 오가는 수동적 일정 관리에 과도한 리소스 소모")

add_table(
    ["지표", "수치", "설명"],
    [
        ["충분한 집중 시간 부족", "68%", "핵심 업무에 몰입할 수 있는 Focus Time이 없다고 응답"],
        ["과도한 정보 탐색 소요", "62%", "필요한 정보나 데이터를 찾는 데 너무 많은 시간이 낭비됨"],
        ["커뮤니케이션 소모 비중", "57%", "업무 시간 중 회의·이메일·채팅에 쓰이는 시간 (실제 생산 활동 43%)"],
    ],
)
add_para("* 출처: Microsoft Work Trend Index (2023)", size=8)

add_heading_styled("1.1.2. 프로젝트 정의", 3)
add_para("본 프로젝트는 사내 규정 판단, 문서 분석, 일정 관리를 지능형 AI 에이전트가 처리하는 시스템을 구축한다. 단순 LLM 호출을 넘어, LangGraph를 통한 유연한 오케스트레이션과 sLLM 파인튜닝으로 도메인 특화 성능을 극대화한다.")
add_para("핵심 기능은 다음과 같다:")
add_bullet("판단 Agent: RAG 기반 다중 규정 교차 판단 + 근거 + 대안 제시")
add_bullet("문서 Agent: 문서 생성(템플릿 기반) / 요약 / 검색 / QA (4개 세부 Intent)")
add_bullet("일정 Agent: 자연어 → 일정 자동 등록/조회 + Google Calendar·Tasks·Gmail·Sheets 통합")
add_bullet("대화형 챗봇: Intent 분류(KoELECTRA) → LangGraph Orchestrator → 적절한 Agent 라우팅")
add_para("핵심 차별점:")
add_bullet("보안: 사내 데이터 외부 유출 차단을 위한 프라이빗 sLLM 운영")
add_bullet("정밀: RAG 기반의 근거 중심 규정 교차 판단 (3중 보조장치: 환각 탐지 + 조항 검증 + Confidence 보정)")
add_bullet("통합: Google 워크스페이스 4종(Calendar, Tasks, Gmail, Sheets) 연동")

# ── 1.2. 핵심 가치 ──
add_heading_styled("1.2. 핵심 가치", 2)
add_heading_styled("1.2.1. 의사결정 속도 혁신", 3)
add_bullet("As-Is: Slack에서 팀 리더/CTO에게 질문 → 답변 대기 → 정확성 불확실 → 30분~1시간")
add_bullet("To-Be: AI 어시스턴트에 질문 → 즉시 답변 + 해당 규정 원문 + 대안 제시")

add_heading_styled("1.2.2. 책임 소재의 명확화", 3)
add_bullet("모든 판단에 규정/가이드라인 출처를 제공하여 혼란 해소")
add_bullet("조직 성장 과정에서 발생하는 규정 관련 분쟁 및 노무 이슈에 대비한 기록 확보")
add_bullet("ISO 27001, ISMS 등 인증 시 내부 통제 체계 증빙으로 활용 가능")

add_heading_styled("1.2.3. 문서 업무의 자동화", 3)
add_bullet("회의록 → 자동 구조화(결정사항, Action Item, 기한)")
add_bullet("JD, 보고서, 제안서 등 반복 문서 → 템플릿 기반 자동 생성")
add_bullet("긴 문서 → 핵심 요약 자동 제공")

add_heading_styled("1.2.4. 신입사원 온보딩 가속화", 3)
add_bullet("신입사원이 규정을 몰라도, AI가 실시간으로 가이드")
add_bullet("온보딩 문서가 별도로 없는 스타트업 환경에서, AI가 기존 규정 + 회의 결정사항 기반으로 즉각 가이드")

doc.add_page_break()

# ══════════════════════════════════════
# 2. 문제 정의
# ══════════════════════════════════════
add_heading_styled("2. 문제 정의", 2)
add_heading_styled("2.1. 현황 분석 및 문제점", 3)

add_heading_styled("2.1.1. 규정 문서의 산재", 4)
add_bullet("규정 자체가 존재하지 않거나, 창업 초기 작성 후 방치된 상태")
add_bullet("노션, 구글 독스, Confluence, GitHub 등에 산발적으로 흩어져 있어 파악이 어려움")
add_bullet("취업규칙은 노무사가 작성했지만, 실무진이 읽어본 적 없음")

add_heading_styled("2.1.2. 의사결정의 병목 현상", 4)
add_bullet('개발자: "테스트 서버에 외부 API 키 올려도 돼?" → CTO에게 Slack DM → CTO 미팅 중 → 2시간 대기')
add_bullet('신규 입사자: "재택 며칠까지 가능하죠?" → 인사 담당에게 질문 → 불명확한 답변')
add_bullet('팀 리더: "인턴에게 AWS 콘솔 접근 줘도 되나?" → 정보보안 규정 어디있는지 모름 → 그냥 부여')
add_para("→ 결과: 규정 미준수 → 추후 보안사고 또는 노무 이슈 발생")

add_heading_styled("2.1.3. 회의록 관리의 형해화", 4)
add_bullet("회의록 자체를 작성하지 않는 경우가 대부분 (Slack 스레드가 회의록 대체)")
add_bullet("Action Item은 구두 합의 → 1주일 후 누락")
add_bullet("동일한 안건이 2~3번 재논의되는 것이 일상")

add_heading_styled("2.2. 해결방안", 3)
add_table(
    ["구분", "기존 방식 (As-Is)", "도입 후 변화 (To-Be)", "효과"],
    [
        ["규정 검색", "Slack에서 물어봄 / Notion 검색", "AI 기반 Hybrid Search (BM25 + Vector)", "답변 시간 30분→10초"],
        ["의사결정", "CTO/팀리더에게 DM", "AI 판단 + 규정 근거 + 대안 제시", "핵심 인력 업무 방해 제거"],
        ["회의록 관리", "안 씀 / Slack이 대체", "자동 구조화 + 검색 가능", "결정사항 추적 가능"],
        ["Action Item", "구두 합의 (누락 多)", "자동 추출 + Google Calendar 등록", "실행률 향상"],
        ["일정 관리", "수동 앱 전환", "자연어 → 자동 등록 + Google 4종 연동", "행정 업무 자동화"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════
# 3. 시장조사 및 BM
# ══════════════════════════════════════
add_heading_styled("3. 시장조사 및 BM", 2)
add_heading_styled("3.1. 시장 동향 및 전망", 3)
add_heading_styled("3.1.1. 기업 AI 어시스턴트 시장의 성장", 4)
add_bullet("국내 중소기업 AI 도입률은 약 5% 수준 (2024 중소벤처기업부) → 진입 기회가 큰 블루오션")
add_bullet("ChatGPT Enterprise($60/인/월) 등 기존 솔루션은 중소기업에 가격 부담")
add_bullet('Gap: 중소기업이 부담 없이 도입 가능하면서, "우리 회사 규정"을 이해하는 솔루션은 부재')

add_para("3.1.2. sLLM 및 온프레미스 AI의 부상", bold=True, size=11)
add_bullet("클라우드 SaaS (1차 타겟): GPU 서버 없이 즉시 사용 가능, 월 구독형")
add_bullet("온프레미스 / Private Cloud (2차 타겟): 핀테크·헬스케어 등 보안 민감 기업 대상, sLLM(8B)으로 자체 구축 가능")

add_heading_styled("3.2. 경쟁사 분석", 3)
add_table(
    ["구분", "Notion AI", "Otter.ai", "ChatGPT Enterprise", "WorkFlow Agent (본 프로젝트)"],
    [
        ["사내 규정 이해", "X (범용)", "X", "X", "O (규정 업로드 후 학습)"],
        ["회의록 → Action Item", "△ (수동)", "△ (요약만)", "△ (범용요약)", "O (자동 추출 + 일정 등록)"],
        ["규정 기반 판단", "X", "X", "X", "O (Yes/No + 근거 + 대안)"],
        ["Google 워크스페이스 연동", "X", "X", "△", "O (Calendar+Tasks+Gmail+Sheets)"],
        ["한국어 성능", "△", "△", "O", "O (한국어 특화 sLLM)"],
        ["도입 비용 (월/인)", "$10", "$8.75", "$60", "₩5,000 ~ 15,000 (예상)"],
    ],
)

add_heading_styled("3.3. 차별화 전략", 3)
add_para("타겟: 10~100인 규모 IT 스타트업 / 중소기업", bold=True)
add_bullet("전담 HR·법무 인력이 없는 조직")
add_bullet("IT 특화 규정(개발 가이드라인, 정보보안)이 있는 조직")
add_bullet("빠른 성장으로 규정 체계화가 시급한 조직")

add_para('"판단"을 제공하는 유일한 도구', bold=True)
add_bullet("타 솔루션: 문서 요약, 검색만 제공")
add_bullet("WorkFlow Agent: Yes/No + 사내 규정 근거 + 실행 가능한 대안 제시")
add_bullet("3중 보조장치(환각 탐지, 조항 존재 검증, Confidence 보정)로 신뢰성 확보")

add_heading_styled("3.4. 비즈니스 모델", 2)
add_table(
    ["플랜", "대상", "월 가격", "포함 기능"],
    [
        ["Starter", "1~10인", "무료", "규정 3개, 월 100회 질의"],
        ["Team", "11~50인", "₩50,000/월", "규정 무제한, 회의록 분석, Action Item"],
        ["Business", "51~100인", "₩150,000/월", "전체 기능 + 관리자 대시보드 + Google 연동"],
        ["Enterprise", "100인+", "별도 협의", "온프레미스 + 커스텀 + sLLM 전용 배포"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════
# 4. 시스템 구성 기획 (대폭 업데이트)
# ══════════════════════════════════════
add_heading_styled("4. 시스템 구성 기획", 2)
add_heading_styled("4.1. 전체 아키텍처", 3)
add_para("3-Layer 아키텍처 + 외부 서비스 연동 구조로 구성된다.", size=10)
add_para("")
add_para("전체 파이프라인:", bold=True)
add_para("사용자 입력 → Intent 분류(KoELECTRA) → LangGraph Orchestrator → Agent 처리 → SSE 스트리밍 응답")
add_para("")

add_table(
    ["Layer", "기술", "역할"],
    [
        ["AI Engine", "LangGraph, KoELECTRA, Qdrant, GPT/Claude API, BM25, ko-sbert-nli, Docling+PaddleOCR", "Intent 분류, Agent 오케스트레이션, RAG, 문서 파싱"],
        ["Backend", "FastAPI, PostgreSQL, SQLAlchemy+Alembic, JWT, Google OAuth, Redis, SSE", "API 서버, DB, 인증, 캐시, 실시간 스트리밍"],
        ["Frontend", "React 18(Vite), Zustand, TanStack Query, Tailwind CSS, FullCalendar", "UI/UX, 상태관리, 캘린더, SSE 수신"],
        ["Infra", "AWS(EC2+S3+RDS), Docker, GitHub Actions, RunPod A100", "배포, 컨테이너, CI/CD, GPU 학습"],
    ],
)

add_heading_styled("4.2. 핵심 Agent 설계", 3)

add_heading_styled("4.2.1. Intent 분류 체계 (8개)", 4)
add_para("사용자의 자연어 질문을 KoELECTRA 모델로 8개 카테고리로 분류하여 적절한 Agent로 라우팅한다.")
add_table(
    ["Intent", "호출 Agent", "설명"],
    [
        ["judgment", "Judgment Agent", "규정 판단/정보 질의"],
        ["doc_search", "Document Agent", "문서 검색"],
        ["doc_generate", "Document Agent", "문서 생성 (회의록 포함)"],
        ["doc_summary", "Document Agent", "문서 요약"],
        ["doc_qa", "Document Agent", "문서 QA (근거 기반 답변)"],
        ["schedule_add", "Schedule Agent", "일정 추가"],
        ["schedule_view", "Schedule Agent", "일정 조회"],
        ["general", "General Response", "일반 대화"],
    ],
)

add_heading_styled("4.2.2. Judgment Agent (규정 판단)", 4)
add_para("역할: RAG 기반 다중 규정 교차 판단 + 근거 제시 + 대안 제시")
add_para("처리 흐름:", bold=True)
add_bullet("사용자 질문 → RAG Hybrid Search (규정 문서, top_k=10)")
add_bullet("LLM 판단 및 구조화 JSON 생성")
add_bullet("3중 보조장치 검증: 환각 탐지 / 조항 존재 검증 / Confidence 보정")
add_bullet("최종 판단 결과 출력")
add_para("출력 예시:", bold=True)
add_para('질문: "인턴에게 AWS 콘솔 접근 권한을 줘도 되나요?"', indent=True)
add_para("→ 조건부 가능 / 근거: 정보보안 규정 3.2조, 개발 가이드 5.1조, 인사 2.3조", indent=True)
add_para("→ 대안: 테스트 환경 한정 부여 / Confidence: 0.85", indent=True)

add_heading_styled("4.2.3. Document Agent (문서 처리)", 4)
add_para("역할: 4개 세부 기능(생성/요약/검색/QA)을 처리하는 통합 문서 Agent")
add_table(
    ["기능", "입력", "처리", "출력"],
    [
        ["doc_generate", "템플릿 ID + 사용자 입력", "템플릿 로드 → LLM 초안 생성", "초안 + 추가 입력 항목"],
        ["doc_summary", "문서 ID", "문서 로드 → 회사 요약 포맷 적용", "구조화된 요약문"],
        ["doc_search", "검색 쿼리 + 필터", "RAG Hybrid Search (업로드 문서)", "관련 문서 목록 + 추천"],
        ["doc_qa", "질문", "RAG → 근거 추출 → LLM 답변", "답변 + 인용 출처"],
    ],
)
add_para("지원 템플릿: 회의록, 보고서, 채용공고(JD), 제안서")
add_para("문서 파싱: Docling (디지털 PDF) + PaddleOCR (스캔 문서) + DOCX 파서")

add_heading_styled("4.2.4. Schedule Agent (일정 관리)", 4)
add_para("역할: 자연어 기반 일정 등록/조회 + Google 워크스페이스 4종 통합")
add_table(
    ["Google 서비스", "기능"],
    [
        ["Google Calendar", "일정 등록/조회/수정 + Meet 자동 생성"],
        ["Google Tasks", "할 일(Action Item) 동기화"],
        ["Gmail", "일정 알림/초대 메일 발송"],
        ["Google Sheets", "일정 추적 시트 자동 업데이트"],
    ],
)
add_para("통합 OAuth 2.0으로 한 번 인증하면 4개 서비스 모두 사용 가능")

add_heading_styled("4.3. RAG 시스템 설계", 3)
add_heading_styled("4.3.1. Hybrid Search 전략", 4)
add_table(
    ["방식", "장점", "단점", "사용 시점"],
    [
        ["BM25 (rank_bm25)", "키워드 정확 매칭", "의미 이해 불가", '조항 번호 검색 ("3.2조")'],
        ["Vector Search (Qdrant)", "의미 유사도 검색", "키워드 누락 가능", '자연어 질문 ("신입 권한")'],
    ],
)
add_para("→ BM25 + Vector Search를 RRF(Reciprocal Rank Fusion)로 합산하여 Recall과 Precision 동시 확보")
add_para("→ bge-reranker-v2-m3 리랭커 구현 완료 (성능 최적화 후 적용 예정)")

add_heading_styled("4.3.2. Vector DB 구성", 4)
add_bullet("DB: Qdrant Cloud (코사인 유사도, 768차원)")
add_bullet("Embedding 모델: jhgan/ko-sbert-nli (한국어 특화)")
add_bullet("Chunk 전략: 규정 문서 → 조항 단위 분할 (장/조 메타데이터 보존) / 회의록·일반 문서 → 문단 단위 분할")
add_para("메타데이터 필터링:", bold=True)
add_bullet('source: "regulations" (규정) / "documents" (업로드 문서)')
add_bullet('doc_type: "HR", "IT", "governance", "general", "meeting_minutes"')

add_heading_styled("4.3.3. RAG 검색 대상", 4)
add_table(
    ["Agent/기능", "RAG", "검색 대상", "비고"],
    [
        ["Judgment", "O", "사내 규정 (source=regulations)", "규정만 검색"],
        ["doc_search / doc_qa", "O", "업로드 문서 (source=documents)", "업로드 문서만 검색"],
        ["doc_generate / doc_summary", "X", "-", "사용자 입력/명시적 문서 기반"],
        ["Schedule / General", "X", "-", "-"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════
# 5. 모델링 계획 (대폭 업데이트)
# ══════════════════════════════════════
add_heading_styled("5. 모델링 계획", 2)

add_heading_styled("5.1. 개발 전략: LLM API 먼저 → sLLM 교체", 3)
add_para("파인튜닝 먼저 하면 input/output이 바뀔 때마다 데이터를 다시 만들어야 하므로, LLM API로 기능을 완성한 뒤 sLLM으로 교체하는 전략을 채택한다.")
add_table(
    ["단계", "내용", "상태"],
    [
        ["1단계", "설계 · 환경 세팅", "✅ 완료"],
        ["2단계", "LLM API(GPT/Claude)로 전체 기능 구현", "✅ 완료"],
        ["3단계", "Agent 개발 — LLM API 기반 실제 동작 확인", "✅ 대부분 완료"],
        ["4단계", "Intent 파인튜닝 (단일질문 분류)", "✅ 완료 (v2 7단계 실험)"],
        ["", "Agent 파인튜닝 (판단/문서 특화)", "← 중간발표 후 진행"],
        ["5단계", "복합질문 강화 + sLLM(vLLM) 교체", "← 예정"],
        ["6단계", "통합 테스트 → 배포", "← 예정"],
    ],
)

add_heading_styled("5.2. sLLM 모델 선정", 3)
add_para("Base 모델: Kanana-1.5-8B (벤치마크 종합 0.652, 한국어 특화)", bold=True)
add_bullet("성능: KMMLU 한국어 벤치마크 기준 선정 완료")
add_bullet("속도: 8B 파라미터 → A100 40GB로 실시간 추론 가능")
add_bullet("서빙: vLLM (OpenAI 호환 API + LoRA 핫스왑 + 스트리밍)")
add_bullet("라이선스: 상업적 이용 가능")

add_heading_styled("5.2.1. Fine-tuning 전략 (LoRA)", 4)
add_table(
    ["대상", "방식", "상태"],
    [
        ["LoRA v1 — 판단 Agent 특화", "규정 판단 input/output 기반 학습", "LLM API 기능 확정 후 진행"],
        ["LoRA v2 — 문서 Agent 특화", "문서 생성/요약 input/output 기반 학습", "LLM API 기능 확정 후 진행"],
    ],
)
add_para("현재 단계: GPT-4o-mini / Claude Sonnet 4 API로 기능 구현 중 → input/output 형태 확정 후 파인튜닝 데이터 수집 예정")

add_heading_styled("5.3. Intent Classification (완료)", 3)
add_para("모델: koelectra-base-v3 (파인튜닝)", bold=True)
add_bullet("분류 카테고리: 8개 (judgment, doc_search, doc_generate, doc_summary, doc_qa, schedule_add, schedule_view, general)")
add_bullet("Test F1: 97.88% / Adversarial F1: 87.58%")
add_bullet("과신뢰 오분류 69% 감소 (Label Smoothing 0.1 적용)")
add_bullet("추론 속도: 7.9ms (GPU)")

add_para("7단계 체계적 실험:", bold=True)
add_para("Baseline → 32-point Grid Search → 최종 평가 → 오분류 분석 → 타겟 보강 → Label Smoothing → 시나리오 검증")

doc.add_page_break()

# ══════════════════════════════════════
# 6. 사용 데이터 (대폭 업데이트)
# ══════════════════════════════════════
add_heading_styled("6. 사용 데이터", 2)

add_heading_styled("6.1. 규정 문서 (Knowledge Base)", 3)
add_para("가상 회사(듀듀테크) 규정을 직접 작성하여 Knowledge Base로 활용:")
add_table(
    ["분류", "규정명", "주요 내용"],
    [
        ["HR", "급여규정", "급여체계, 직급별 기본급(S1~M2), 직책수당, 시간외수당, 상여금"],
        ["HR", "교육훈련규정", "신입사원 입문교육(2주), 법정 의무교육, 자격증 지원"],
        ["HR", "출장규정", "국내/해외 출장비, 숙박비 한도, 일비/식비"],
        ["HR", "복리후생규정", "건강검진, 의료비 지원, 경조금, 자녀학자금, 주거 지원"],
        ["HR", "징계규정", "징계 종류(견책~해고), 중징계 사유, 소명 절차"],
        ["IT", "개인정보처리규정", "수집 원칙, 암호화(SSL/TLS, 해시), CCTV 30일 보관, 침해사고 72h 신고"],
        ["IT", "듀듀테크 사내규정(종합)", "정보보호, 접근통제, 소프트웨어 개발보안, 오픈소스 정책 등"],
        ["거버넌스", "윤리강령", "이해충돌, 부정청탁, 금품수수 한도(건당 5만/연 30만)"],
    ],
)
add_para("총 232개 청크로 Qdrant에 인덱싱 완료 (조항 단위 분할, 장/조 메타데이터 보존)")

add_heading_styled("6.2. Intent 분류 데이터 (v2 실험 완료)", 3)
add_table(
    ["구분", "건수", "비고"],
    [
        ["기본 데이터 (GPT-4o + Claude)", "2,299개", "8개 intent × 2 LLM"],
        ["경계 쌍 데이터", "600개", "10쌍 × 30개 × 2 LLM"],
        ["타겟 보강 데이터", "98개", "오분류 패턴 기반"],
        ["학습 합계 (Train)", "2,425개", "Val 285 / Test 286"],
        ["Adversarial 테스트셋", "450개", "7단계 실험 완료"],
    ],
)

add_heading_styled("6.3. 샘플 문서 데이터", 3)
add_para("시스템 데모 및 테스트용 샘플 문서 30건을 Qdrant에 시딩:")
add_table(
    ["유형", "건수", "예시"],
    [
        ["회의록", "10건", "스프린트 킥오프, 보안 점검, 기술 스택 변경, OKR 수립 등"],
        ["보고서", "10건", "월간 보고서, 성능 분석, QA 테스트 결과, 기술 부채 등"],
        ["제안서", "10건", "AI 챗봇 고도화, 클라우드 마이그레이션, ISMS 준비 등"],
    ],
)

add_heading_styled("6.4. 회의록 샘플", 3)
add_para("문서 Agent 테스트용 가상 회의록 8건 (JSON 형식, data/proceedings/)")

doc.add_page_break()

# ══════════════════════════════════════
# 7. 역할 분담 (업데이트)
# ══════════════════════════════════════
add_heading_styled("7. 역할 분담 (R&R)", 2)
add_heading_styled("7.1. 팀 구성 (5인)", 3)
add_table(
    ["역할", "담당자", "주요 업무", "산출물"],
    [
        ["PM", "신지용",
         "전체 일정 관리 및 스프린트 운영\nIntent Classification 구현 (KoELECTRA)\nLangGraph Orchestrator 설계\nSSE 스트리밍 + API 스키마\nAWS 배포 및 CI/CD",
         "시스템 아키텍처 문서\nOrchestrator 코드\nIntent 모델 (F1 97.88%)"],
        ["AI Engineer", "윤경은",
         "Judgment Agent 개발\nHybrid RAG 파이프라인 구축 (BM25+Qdrant+RRF)\nLoRA v1 파인튜닝 (판단 특화)\nvLLM 서빙 구축",
         "Judgment Agent 코드\nRAG 파이프라인\n정확도 평가 리포트"],
        ["AI Engineer", "진승언",
         "Document Agent 개발 (4개 서브 기능)\n문서 템플릿 시스템 (회의록/보고서/JD/제안서)\n문서 파싱 모듈 (Docling+PaddleOCR)\nLoRA v2 파인튜닝 (문서 특화)",
         "Document Agent 코드\n템플릿 시스템\n문서 파싱 모듈"],
        ["Backend", "안혜빈",
         "PostgreSQL DB 설계 (12 테이블)\nJWT 인증 + Google OAuth 2.0\nSchedule Agent + Google Services 4종 통합\n관리자 API + 통계",
         "DB 스키마 (12 tables)\nSchedule Agent API\nGoogle 통합 서비스"],
        ["Frontend", "문지영",
         "React 18(Vite) 기반 UI/UX (11페이지)\n챗봇 카드 UI + SSE 수신\nFullCalendar 일정 관리\n반응형 디자인 + 다크모드",
         "React 컴포넌트\nFigma 디자인\n사용자 가이드"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════
# 8. 핵심 성과 요약
# ══════════════════════════════════════
add_heading_styled("8. 핵심 성과 요약 (중간발표 시점)", 2)
add_table(
    ["영역", "성과"],
    [
        ["Intent 분류", "KoELECTRA 파인튜닝 — Test F1 97.88%, Adversarial F1 87.58%, 추론 7.9ms"],
        ["RAG 파이프라인", "BM25 + Vector(Qdrant) + RRF 하이브리드 검색 구현"],
        ["Google 4종 연동", "Calendar + Tasks + Gmail + Sheets 통합 OAuth"],
        ["Backend", "12 테이블 DB + JWT 인증 + SSE 실시간 스트리밍"],
        ["Frontend", "11 페이지 + 챗봇 카드 UI + FullCalendar 일정 관리"],
        ["실험", "7단계 체계적 실험 (Grid Search → Label Smoothing → 시나리오 검증)"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════
# 변경 내역
# ══════════════════════════════════════
add_heading_styled("변경 내역", 2)
add_para("본 문서는 1주차 기획서(2026.02.06)를 중간발표(2026.02.27) 시점 기준으로 개정한 버전이다.")
add_table(
    ["섹션", "변경 사항", "사유"],
    [
        ["1.1.1 배경", "Microsoft Work Trend Index 통계(68%/62%/57%) 추가", "문제 정의 근거 보강"],
        ["1.1.2 정의", "핵심 기능 4종 → 현행 Agent 구조로 재작성, 차별점(보안/정밀/통합) 추가", "실제 구현 반영"],
        ["2.2 해결방안", "Google 4종 연동, Hybrid Search 반영", "기능 확장 반영"],
        ["3.2 경쟁사", "Google 워크스페이스 연동 행 추가", "신규 기능"],
        ["4.1 아키텍처", "3-Layer 아키텍처(AI/Backend/Frontend/Infra) 전면 재작성", "LangGraph 기반 구현 완료"],
        ["4.2 Agent 설계", "Intent 8개 체계, Agent별 상세 설계 전면 재작성", "실제 구현 반영"],
        ["4.2.4 일정Agent", "Google 4종 연동 상세 추가 (기존: Phase 2 예정)", "구현 완료"],
        ["4.3 RAG", "ChromaDB → Qdrant, RRF 합산 방식 추가, 리랭커 추가", "기술 변경"],
        ["5.1 개발전략", "'LLM API 먼저 → sLLM 교체' 전략 신규 추가", "전략 구체화"],
        ["5.2 sLLM", "'추후 선정' → Kanana-1.5-8B 확정, LoRA v1/v2 계획 추가", "모델 선정 완료"],
        ["5.3 Intent", "klue/bert-base(4개) → koelectra-base-v3(8개), F1 97.88% 달성", "파인튜닝 완료"],
        ["6.1 규정데이터", "30개 조항 → 7개 규정 파일(HR5+IT2+거버넌스1), 232 청크", "데이터 대폭 확장"],
        ["6.2 Intent데이터", "신규: v2 2,425개 학습 데이터 + 450 Adversarial", "학습 데이터 구축 완료"],
        ["6.3 샘플문서", "신규: 30건 시딩 (회의록10+보고서10+제안서10)", "테스트 데이터 추가"],
        ["7.1 역할분담", "AI리드/AI서브 → AI Engineer 통일, 담당 업무 현행화", "역할 재정의"],
        ["8. 핵심성과", "신규 섹션: 중간발표 시점 성과 요약", "진행 상황 공유"],
    ],
)

# ─── 저장 ───
output_path = os.path.join("docs", "중간발표", "SK네트웍스 Family AI 과정 21기 최종 프로젝트 기획서_듀듀_v2_중간발표.docx")
doc.save(output_path)
print(f"저장 완료: {output_path}")
