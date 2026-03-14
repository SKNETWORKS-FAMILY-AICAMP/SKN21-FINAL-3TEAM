"""
4주차 산출물: 데이터 전처리 / 인공지능 학습 결과서 재작성 스크립트
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def add_word_toc(doc):
    """Word 자동 목차 필드(TOC) 삽입 — Word에서 열면 F9로 업데이트 가능"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-2" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run3._r.append(fldChar2)

    # 임시 텍스트 (Word에서 열어 F9 누르면 실제 목차로 대체됨)
    run4 = paragraph.add_run("[Word에서 이 문서를 열고 Ctrl+A → F9 를 눌러 목차를 업데이트하세요]")
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run4.font.size = Pt(9)
    run4.italic = True

    run5 = paragraph.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run5._r.append(fldChar3)

    return paragraph

OUTPUT_PATH = os.path.join(
    os.path.dirname(__file__), "..",
    "docs", "산출물", "4주차",
    "2.데이터 전처리_인공지능 학습 결과서_3팀.docx"
)

doc = Document()

# ── 스타일 헬퍼 ──────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "맑은 고딕"
    h.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    if level == 1:
        h.font.size = Pt(16)
    elif level == 2:
        h.font.size = Pt(13)
    else:
        h.font.size = Pt(11)


def add_table(headers, rows, col_widths=None):
    """테이블 추가 + 서식"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 헤더
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # 데이터
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    return p


# ════════════════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SK networks  |  Family AI Camp")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("데이터 전처리\n인공지능 학습 결과서")
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SKN Family AI Camp 21기 : 최종 프로젝트 3팀\nWorkFlow Agent (듀듀)")
run.font.size = Pt(13)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("작성: 3팀 -- 신지용(PM), 윤경은(AI), 진승언(AI), 안혜빈(Backend), 문지영(Frontend)")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 목차 (Word 자동 TOC 필드 + 수동 텍스트 병행)
# ════════════════════════════════════════════════════════
doc.add_heading("목 차", level=1)

# Word 자동 목차 필드 삽입 (Word에서 F9 누르면 Heading 기반 자동 생성)
add_word_toc(doc)
doc.add_paragraph()  # 간격

# 수동 목차 (자동 TOC가 업데이트되기 전 참조용)
toc_items = [
    ("1. 사용할 LLM 모델에 대한 설명", True),
    ("    1-1. 서비스 LLM 모델 (개발 단계: API 기반)", False),
    ("    1-2. 서비스 LLM 모델 (배포 단계: sLLM + vLLM)", False),
    ("    1-3. LLM 공통 모듈 (Factory Pattern)", False),
    ("2. RAG를 위해 선택한 임베딩 모델 후보와 선정 이유", True),
    ("    2-1. 임베딩 모델 -- jhgan/ko-sbert-nli", False),
    ("    2-2. Reranker 모델 -- BAAI/bge-reranker-v2-m3", False),
    ("    2-3. RAG 파이프라인 구조", False),
    ("3. 파인튜닝 -- Intent 분류 모델", True),
    ("    3-1. 후보 모델 비교 및 선정된 모델 선정 이유", False),
    ("    3-2. 모델 아키텍처 상세", False),
    ("4. 데이터 전처리에 사용된 인공지능 모델", True),
    ("    4-1. 학습 데이터 생성 모델 (GPT-4o + Claude Sonnet 4)", False),
    ("    4-2. 문서 파싱 모델 (Docling + PaddleOCR)", False),
    ("5. 학습 데이터 구성", True),
    ("    5-1. 데이터 생성 전략 -- 멀티 LLM 혼합형", False),
    ("    5-2. 데이터 분할 및 품질 검증", False),
    ("6. 학습(튜닝) 과정", True),
    ("    6-1. 7-Stage 실험 파이프라인", False),
    ("    6-2. 하이퍼파라미터 및 학습 설정", False),
    ("    6-3. Stage별 상세 결과", False),
    ("    6-4. 과적합/과소적합 방지", False),
    ("7. 성능 평가", True),
    ("    7-1. 최종 모델 성능 요약", False),
    ("    7-2. 시나리오 테스트 (정성 평가)", False),
    ("    7-3. 한계점 및 향후 과제", False),
]
for text, is_major in toc_items:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    if is_major:
        p.paragraph_format.space_before = Pt(6)
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(11)
    else:
        for run in p.runs:
            run.font.size = Pt(10)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 1. 사용할 LLM 모델에 대한 설명
# ════════════════════════════════════════════════════════
doc.add_heading("1. 사용할 LLM 모델에 대한 설명", level=1)
doc.add_paragraph(
    "WorkFlow Agent(듀듀)는 사용자의 업무 질의를 이해하고 판단/문서/일정 등의 Agent로 "
    "라우팅하여 답변을 생성하는 멀티 에이전트 시스템입니다. "
    "개발 전략으로 'LLM API 먼저 → sLLM(파인튜닝) 교체' 접근법을 채택하여, "
    "기능을 LLM API로 먼저 완성한 뒤 동일 인터페이스로 sLLM을 교체하는 구조입니다."
)

doc.add_heading("1-1. 서비스 LLM 모델 (개발 단계: API 기반)", level=2)
doc.add_paragraph(
    "개발 단계에서는 상용 LLM API를 활용하여 전체 기능을 구현하고 입출력 형태를 확정합니다."
)
add_table(
    ["모델", "제공사", "용도", "특징"],
    [
        ["GPT-4o", "OpenAI", "기본 LLM (판단/문서/일정 Agent)", "멀티모달, 빠른 추론, 한국어 우수"],
        ["Claude Sonnet 4", "Anthropic", "대체 LLM (동일 인터페이스)", "긴 컨텍스트, 정밀한 지시 수행"],
    ],
)
doc.add_paragraph(
    "환경변수 LLM_PROVIDER를 통해 openai / anthropic을 전환할 수 있으며, "
    "Agent 코드 변경 없이 LLM만 교체 가능한 구조입니다."
)

doc.add_heading("1-2. 서비스 LLM 모델 (배포 단계: sLLM + vLLM)", level=2)
doc.add_paragraph(
    "배포 단계에서는 파인튜닝된 sLLM을 vLLM 서버로 서빙하여 비용 절감과 응답 속도를 개선합니다."
)
add_table(
    ["모델", "크기", "특징", "용도"],
    [
        ["kaist-ai/Kanana-1.5-8B\n(베이스 모델)", "8B", "한국어 특화, 카카오 개발\n멘토 직접 추천", "LoRA 어댑터의 베이스"],
        ["LoRA v1 (판단 특화)", "어댑터", "규정 검색 결과 기반\nYes/No 판단 + 근거 생성", "Judgment Agent"],
        ["LoRA v2 (문서 특화)", "어댑터", "회의록 구조화, 문서 요약\n리스크 감지", "Document Agent"],
    ],
)
doc.add_paragraph(
    "vLLM의 OpenAI 호환 API를 사용하므로, LLM_PROVIDER='vllm'으로 전환하면 "
    "기존 코드 변경 없이 sLLM으로 교체됩니다. "
    "LoRA 핫스왑(with_lora)을 통해 하나의 베이스 모델에서 용도별 어댑터를 실시간 전환합니다."
)

doc.add_heading("1-3. LLM 공통 모듈 (Factory Pattern)", level=2)
doc.add_paragraph(
    "ai/llm/factory.py에 구현된 Factory + Singleton 패턴으로 "
    "전체 시스템에서 동일한 인터페이스로 LLM을 호출합니다."
)
add_table(
    ["Provider", "기본 모델", "환경변수", "비고"],
    [
        ["openai", "gpt-4o", "OPENAI_API_KEY\nOPENAI_MODEL", "개발 기본값"],
        ["anthropic", "claude-sonnet-4-20250514", "ANTHROPIC_API_KEY\nANTHROPIC_MODEL", "대체 LLM"],
        ["vllm", "kaist-ai/Kanana-1.5-8B", "VLLM_BASE_URL\nVLLM_MODEL", "배포 시 사용"],
    ],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 2. RAG 임베딩 모델
# ════════════════════════════════════════════════════════
doc.add_heading("2. RAG를 위해 선택한 임베딩 모델 후보와 선정 이유", level=1)
doc.add_paragraph(
    "Judgment Agent가 사내 규정을 검색하여 판단 근거를 제공하기 위해 "
    "RAG(Retrieval-Augmented Generation) 파이프라인을 구축하였습니다. "
    "핵심 구성 요소는 임베딩 모델(벡터 검색)과 Reranker(정밀 재정렬)입니다."
)

doc.add_heading("2-1. 임베딩 모델 — jhgan/ko-sbert-nli", level=2)
doc.add_paragraph("임베딩 모델 후보를 아래와 같이 비교한 후 최종 선정하였습니다.")

add_table(
    ["모델", "벡터 차원", "한국어 성능", "특징", "선정 여부"],
    [
        ["jhgan/ko-sbert-nli", "768", "우수", "한국어 NLI 데이터로 학습\nSentence-BERT 기반\n문장 유사도 계산 최적화", "최종 선정"],
        ["sentence-transformers/\nparaphrase-multilingual-\nMiniLM-L12-v2", "384", "보통", "다국어 지원 (50+언어)\n경량 모델", "미채택"],
        ["BAAI/bge-m3", "1024", "우수", "다국어 Sparse+Dense\n고차원으로 메모리 부담", "미채택"],
    ],
)

p = doc.add_paragraph()
run = p.add_run("선정 이유: ")
run.bold = True
doc.add_paragraph(
    "jhgan/ko-sbert-nli는 한국어 NLI(Natural Language Inference) 데이터셋으로 학습되어 "
    "한국어 문장 간 의미 유사도 계산에 최적화되어 있습니다. "
    "768차원 벡터로 검색 품질과 메모리 효율의 균형이 우수하며, "
    "Qdrant Vector DB와의 호환성도 검증되었습니다. "
    "SentenceTransformer 라이브러리로 싱글턴 패턴 + Lazy Loading을 적용하여 "
    "서비스 시작 시 메모리를 효율적으로 관리합니다."
)

doc.add_heading("2-2. Reranker 모델 — BAAI/bge-reranker-v2-m3", level=2)
doc.add_paragraph(
    "Hybrid Search(BM25 + Vector) 결과에 노이즈가 섞일 수 있어, "
    "Reranker로 검색 결과를 한 번 더 정밀하게 재정렬합니다."
)
add_table(
    ["항목", "설명"],
    [
        ["모델명", "BAAI/bge-reranker-v2-m3"],
        ["아키텍처", "Cross-Encoder (질문-문서 쌍 관련도 점수 산출)"],
        ["한국어 지원", "다국어 지원, 한국어 성능 우수 (BAAI 벤치마크 상위)"],
        ["입출력", "입력: (질문, 문서) 쌍 → 출력: 관련도 점수 (0~1)"],
        ["선정 이유", "멘토 추천, 경량이면서 높은 재정렬 정확도\nTop 20 → Top 5 압축으로 LLM 입력 품질 향상"],
    ],
    col_widths=[4, 13],
)

doc.add_heading("2-3. RAG 파이프라인 구조", level=2)
doc.add_paragraph("전체 RAG 파이프라인은 다음과 같이 동작합니다:")
add_bullet("사용자 질문 입력")
add_bullet("BM25 키워드 검색 (Top 15) + Vector 시맨틱 검색 (Top 15) 동시 실행")
add_bullet("RRF(Reciprocal Rank Fusion, k=60)로 합산 정렬 (Top 20)")
add_bullet("소스 다양성 적용 (max_per_source = 3)")
add_bullet("Reranker(bge-reranker-v2-m3)로 관련도 재정렬 (Top 5)")
add_bullet("최종 Top 5 문서를 LLM에 전달하여 답변 생성")

doc.add_paragraph()
add_table(
    ["구성 요소", "기술", "역할"],
    [
        ["키워드 검색", "BM25 (rank_bm25.BM25Okapi)\n+ kiwipiepy 한국어 토크나이저", "형태소 기반 정확 매칭"],
        ["벡터 검색", "jhgan/ko-sbert-nli (768-dim)\n+ Qdrant Vector DB", "의미 기반 유사도 검색"],
        ["스코어 결합", "RRF (Reciprocal Rank Fusion)", "두 검색 결과 합산 정렬"],
        ["재정렬", "BAAI/bge-reranker-v2-m3\n(Cross-Encoder)", "최종 관련도 정밀 평가"],
    ],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 3. 파인튜닝 — Intent 분류 모델
# ════════════════════════════════════════════════════════
doc.add_heading("3. 파인튜닝 — Intent 분류 모델", level=1)
doc.add_paragraph(
    "사용자의 입력을 8개 intent로 분류하여 적절한 Agent로 라우팅하는 것이 "
    "오케스트레이터의 핵심 역할입니다. 잘못된 라우팅 = 잘못된 답변이므로, "
    "Intent 분류기의 정확도가 전체 시스템 품질을 결정합니다."
)
add_table(
    ["Intent", "라우팅 대상", "설명"],
    [
        ["judgment", "Judgment Agent", "사내 규정/정책 판단 질의"],
        ["doc_search", "Document Agent", "문서 검색/조회"],
        ["doc_generate", "Document Agent", "문서 생성 (보고서, 제안서 등)"],
        ["doc_summary", "Document Agent", "문서 요약"],
        ["doc_qa", "Document Agent", "문서 기반 Q&A"],
        ["schedule_add", "Schedule Agent", "일정 등록"],
        ["schedule_view", "Schedule Agent", "일정 조회"],
        ["general", "General Handler", "일상 대화/인사"],
    ],
)

doc.add_heading("3-1. 후보 모델 비교 및 선정된 모델 선정 이유", level=2)
doc.add_paragraph(
    '"같은 파이프라인, 3가지 철학 — 어떤 것이 한국어 직장 챗봇에 최적인가?" '
    "라는 질문에 답하기 위해 3개 모델을 동일 조건에서 비교하였습니다."
)
add_table(
    ["모델", "파라미터", "아키텍처", "역할", "선정 이유"],
    [
        ["koelectra-base-v3\n(monologg)", "112.9M", "ELECTRA (RTD)\n12-layer Transformer",
         "최종 선택", "한국어 토큰 교체 감지 학습\n짧은 문장 구분에 유리"],
        ["klue/bert-base", "110.6M", "BERT (MLM)\n12-layer Transformer",
         "Baseline", "한국어 NLU 표준 모델"],
        ["monologg/distilkobert", "28.4M", "DistilBERT (KD)\n6-layer Transformer",
         "경량 후보", "4배 작고 3배 빠름\n'111M이 과한가?' 검증"],
    ],
)

p = doc.add_paragraph()
run = p.add_run("Baseline 비교 결과 (Stage 2): ")
run.bold = True
doc.add_paragraph("고정 HP: epochs=5, lr=2e-5, batch=16, seed=42")
add_table(
    ["순위", "모델", "Val F1", "모델 크기", "추론 속도", "학습 시간"],
    [
        ["1", "koelectra-base-v3", "0.9825", "431MB", "14.2ms", "860s"],
        ["2", "klue/bert-base", "0.9780", "422MB", "14.4ms", "808s"],
        ["3", "distilkobert", "0.9498", "108MB", "3.9ms", "243s"],
    ],
)

p = doc.add_paragraph()
run = p.add_run("최종 선정: KoELECTRA-base-v3")
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

doc.add_paragraph(
    "선정 근거:\n"
    "1) Adversarial 강건성 최고: Adv F1 86.04% (bert 85.17%, distilkobert 79.26%)\n"
    "2) 추론 속도 균형: 7.9ms (bert 10.4ms 대비 24% 빠름)\n"
    "3) Seed 안정성: 0.9874 +/- 0.0033 (3-seed 평균)\n"
    "4) ELECTRA의 RTD(Replaced Token Detection) 방식: "
    "토큰 교체 감지 사전학습이 짧은 한국어 문장 구분에 유리"
)

doc.add_paragraph(
    "BERT가 아닌 이유: Test F1은 bert(0.9756) > koelectra(0.9726)이지만, "
    "실전(Adversarial) F1에서 koelectra가 0.87%p 우위이며 추론 속도도 빠름.\n"
    "DistilKoBERT가 아닌 이유: 4배 작고 3배 빠르지만, "
    "Adv F1 79.26%로 7%p 열세 — 8개 intent 분류에는 111M 규모가 필요."
)

doc.add_heading("3-2. 모델 아키텍처 상세", level=2)
doc.add_paragraph(
    "KoELECTRA-base-v3-discriminator는 3-layer 구조입니다. "
    "총 112.9M 파라미터, 입력 max_length=64 토큰."
)
add_table(
    ["레이어", "구성", "설명"],
    [
        ["Embedding", "Token + Position + Type\nembedding (768-dim)",
         "입력 텍스트를 768차원 벡터로 변환\nWordPiece 토크나이저 (vocab 35K)"],
        ["Encoder", "Transformer x 12층\nMulti-Head Attention (12 heads)",
         "12개 Transformer 블록\n각 블록은 셀프어텐션 + 피드포워드"],
        ["Classification\nHead", "Dropout (0.1)\n+ Dense (768 -> 8)\n+ Softmax",
         "[CLS] 토큰 출력을 8개 intent\n확률 분포로 변환"],
    ],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 4. 데이터 전처리에 사용된 인공지능 모델
# ════════════════════════════════════════════════════════
doc.add_heading("4. 데이터 전처리에 사용된 인공지능 모델", level=1)
doc.add_paragraph(
    "본 프로젝트에서는 데이터 전처리 단계에서 다양한 인공지능 모델을 활용하였습니다. "
    "크게 (1) 학습 데이터 생성을 위한 생성 모델과 (2) 문서 파싱을 위한 모델로 나뉩니다."
)

doc.add_heading("4-1. 학습 데이터 생성 모델 — 합성 데이터 생성 (멀티 LLM)", level=2)
doc.add_paragraph(
    "Intent 분류 모델의 학습 데이터를 생성하기 위해 2개의 LLM을 동시에 활용하여 "
    "합성 데이터(Synthetic Data)를 생성하였습니다."
)
add_table(
    ["모델", "역할", "생성량", "사용 이유"],
    [
        ["GPT-4o (OpenAI)", "기본 데이터 150개/intent\n경계 쌍 300개\n적대적 232개",
         "intent별 ~150개\n총 ~1,099개 기본\n+ 경계/적대적",
         "다양한 한국어 표현 생성\n안정적 JSON 출력"],
        ["Claude Sonnet 4\n(Anthropic)", "기본 데이터 150개/intent\n경계 쌍 300개\n적대적 240개",
         "intent별 150개\n총 1,200개 기본\n+ 경계/적대적",
         "GPT와 다른 문체/스타일\n교차 검증용 다양성 확보"],
        ["Gemini Pro 3.1\n(Google, 웹)", "시나리오 테스트 70개\n추가 생성", "70개 (수동 검수 후\n100개로 확장)",
         "정성 평가용 시나리오\n제3의 LLM으로 편향 방지"],
    ],
)
p = doc.add_paragraph()
run = p.add_run("멀티 LLM 전략의 효과:")
run.bold = True
add_bullet("다양성: 단일 LLM은 비슷한 문체/패턴 반복 → 2개 LLM이 각각 다른 스타일의 한국어 생성")
add_bullet("교차 검증: 한 LLM이 생성한 데이터를 다른 LLM이 라벨 검증 → 편향 제거")
add_bullet("Cross-LLM 중복 0개 확인 → 모델 독립적으로 올바른 라벨 보장")

doc.add_heading("4-2. 문서 파싱 모델 — Docling + PaddleOCR", level=2)
doc.add_paragraph(
    "RAG 파이프라인에 문서를 색인하기 위해 다양한 형식의 문서를 파싱합니다. "
    "멘토 추천에 따라 Docling + PaddleOCR 조합을 채택하였습니다."
)
add_table(
    ["모델/도구", "용도", "처리 대상", "핵심 기능"],
    [
        ["Docling (IBM)", "PDF 구조화 파싱", "디지털 PDF",
         "테이블/헤더/본문 구조 인식\n마크다운 변환\n조항 단위 청킹 (제N조 기준)"],
        ["PaddleOCR", "한국어 OCR", "스캔 PDF / 이미지",
         "한국어 텍스트 추출 (최상급)\nuse_angle_cls=True (회전 보정)\nconfidence >= 0.5 필터링"],
        ["python-docx", "DOCX 파싱", "Word 문서", "문단/표/스타일 추출"],
    ],
)
doc.add_paragraph(
    "문서 파싱 라우팅 전략:\n"
    "- 디지털 PDF → Docling으로 구조화 파싱\n"
    "- Docling 결과 50자 미만 → 스캔 문서로 판단 → PaddleOCR Fallback\n"
    "- DOCX → python-docx로 파싱\n"
    "- 이미지 (.png, .jpg 등) → PaddleOCR 직접 추출\n"
    "- 청킹: 규정 문서는 조항(제N조) 단위, 일반 문서는 마크다운 헤딩 기반 분할"
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 5. 학습 데이터 구성
# ════════════════════════════════════════════════════════
doc.add_heading("5. 학습 데이터 구성", level=1)

doc.add_heading("5-1. 데이터 생성 전략 — 멀티 LLM 혼합형 (방법 C)", level=2)
doc.add_paragraph(
    "기존 v1 데이터는 BERT 실패 패턴 보정용으로 누적되어 특정 모델에 편향되었습니다. "
    "클린 실험을 위해 GPT-4o + Claude Sonnet 4 두 개 LLM으로 통일된 기준으로 재생성하였습니다."
)
add_table(
    ["데이터 종류", "생성 방식", "건수", "용도"],
    [
        ["Seed 문장", "수동 작성", "intent별 10개", "생성 앵커 (다양성 기준점)"],
        ["기본 데이터", "Claude 150 + GPT 150\n(intent별)", "2,299개\n(중복 제거 후)", "학습 데이터 주축"],
        ["경계 쌍", "GPT 300 + Claude 300\n(혼동 쌍 10종 x 30개)", "600개", "혼동 가능 intent 쌍 강화"],
        ["적대적 테스트", "GPT 232 + Claude 240\n(8유형별)", "450개\n(중복 제거 후)", "강건성 평가 전용\n(학습에 미사용)"],
        ["보강 데이터\n(Stage 5)", "오분류 분석 기반\n타겟 생성", "98개", "약점 intent 보강"],
        ["시나리오 테스트", "Gemini 웹(Pro 3.1)\n+ 수동 검수", "100개", "실전 라우팅 시뮬레이션\n(정성 평가)"],
    ],
)

p = doc.add_paragraph()
run = p.add_run("기본 데이터 intent별 분포:")
run.bold = True
add_table(
    ["Intent", "Claude", "GPT", "합계"],
    [
        ["judgment", "150", "127", "277"],
        ["doc_search", "150", "136", "286"],
        ["doc_generate", "150", "122", "272"],
        ["doc_summary", "150", "144", "294"],
        ["schedule_add", "150", "148", "298"],
        ["schedule_view", "150", "137", "287"],
        ["general", "150", "151", "301"],
        ["doc_qa", "150", "134", "284"],
        ["합계", "1,200", "1,099", "2,299"],
    ],
)
doc.add_paragraph("클래스 균형: max/min = 1.11 (< 1.2 기준 통과). Cross-LLM 중복: 0개.")

doc.add_heading("5-2. 데이터 분할 및 품질 검증", level=2)
add_table(
    ["구분", "건수", "비율", "용도"],
    [
        ["Train", "2,327 + 98 (보강) = 2,425", "~80%", "모델 학습"],
        ["Validation", "285", "~10%", "매 epoch 평가, HP 선택"],
        ["Test (Hold-out)", "286", "~10%", "Stage 4에서 1회만 사용"],
        ["Adversarial v2", "450", "별도", "강건성 평가 (학습 미사용)"],
        ["시나리오 테스트", "100", "별도", "정성 평가 (4유형)"],
    ],
)

p = doc.add_paragraph()
run = p.add_run("품질 검증 체크리스트:")
run.bold = True
add_bullet("JSON 유효성: 파싱 에러 0건")
add_bullet("라벨 유효성: 8개 허용 라벨만 사용")
add_bullet("중복 제거: GPT 내부 중복 59건 제거")
add_bullet("클래스 균형: max/min = 1.11 (< 1.2 기준 통과)")
add_bullet("테스트 누출: test ∩ train = 공집합")
add_bullet("Cross-LLM 중복: 0개")

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 6. 학습(튜닝) 과정
# ════════════════════════════════════════════════════════
doc.add_heading("6. 학습(튜닝) 과정", level=1)

doc.add_heading("6-1. 7-Stage 실험 파이프라인", level=2)
doc.add_paragraph(
    "체계적 실험을 위해 7단계 파이프라인을 설계하고, "
    "각 단계별 Gate를 통과해야 다음 단계로 진행하는 방식을 채택하였습니다."
)
add_table(
    ["Stage", "내용", "핵심 결과"],
    [
        ["Stage 1", "데이터 생성 + QA", "2,299개 기본 + 600 경계 + 450 적대적"],
        ["Stage 2", "Baseline 3모델 비교\n(동일 HP)", "koelectra 0.9825 > bert 0.9780\n> distilkobert 0.9498"],
        ["Stage 3", "Grid Search 32-point\n+ 3-seed 안정성", "Best: ep10/lr3e-5/bs16\nVal F1 0.9897"],
        ["Stage 4", "최종 평가\n(Adv, 속도, 통계)", "koelectra Adv F1 86.04%\nbert 85.17%, distilkobert 79.26%"],
        ["Stage 5", "오분류 분석 +\n타겟 보강 98개", "Adv F1: 86.04% → 87.84%\n(+1.80%p)"],
        ["Stage 6", "Label Smoothing 0.1\n+ 과신뢰 해소", "과신뢰 42건 → 13건 (-69%)\nAdv F1 87.58%"],
        ["Stage 7", "doc 경계 라벨 리뷰\n+ 시나리오 100개", "오류의 63%가 라벨 문제\n시나리오 85/100 (85.0%)"],
    ],
)

doc.add_heading("6-2. 하이퍼파라미터 및 학습 설정", level=2)
doc.add_paragraph("학습 방식: Full Fine-tuning (인코더 모델 + 소규모 데이터 → LoRA 불필요)")

p = doc.add_paragraph()
run = p.add_run("최종 학습 설정 (Stage 6 채택):")
run.bold = True
add_table(
    ["하이퍼파라미터", "값"],
    [
        ["Base Model", "monologg/koelectra-base-v3-discriminator"],
        ["Task", "SequenceClassification (8 labels)"],
        ["Epochs", "10"],
        ["Batch Size", "16"],
        ["Learning Rate", "3e-5"],
        ["Optimizer", "AdamW"],
        ["Loss Function", "CrossEntropyLoss + Label Smoothing 0.1"],
        ["Scheduler", "Linear (warmup_ratio=0.0)"],
        ["Max Length", "64 tokens"],
        ["Weight Decay", "0.01"],
        ["Seed", "42 (random/numpy/torch/cuda 4중 고정)"],
        ["Train Data", "2,425개 (기본 2,327 + 보강 98)"],
        ["Validation Data", "285개"],
        ["GPU", "RunPod RTX 4090 (24GB)"],
        ["학습 시간", "~71초"],
        ["FP16", "True (혼합 정밀도)"],
    ],
    col_widths=[5, 12],
)

p = doc.add_paragraph()
run = p.add_run("Grid Search 결과 (Stage 3):")
run.bold = True
doc.add_paragraph("32-point grid: epochs {3,5,7,10} x lr {1e-5,2e-5,3e-5,5e-5} x batch {16,32}")
add_table(
    ["순위", "epochs", "lr", "batch", "Val F1"],
    [
        ["1", "10", "3e-5", "16", "0.9897"],
        ["2", "3", "2e-5", "16", "0.9864"],
        ["2", "7", "2e-5", "32", "0.9864"],
        ["2", "10", "2e-5", "32", "0.9864"],
        ["5", "5", "3e-5", "16", "0.9862"],
    ],
)
doc.add_paragraph(
    "3-seed 안정성 검증: 0.9874 +/- 0.0033 (seed={42, 123, 456})\n"
    "Baseline(0.9825) → Best(0.9897): +0.72%p → '데이터 품질 > 하이퍼파라미터' 재확인."
)

doc.add_heading("6-3. Stage별 상세 결과", level=2)

p = doc.add_paragraph()
run = p.add_run("3모델 비교 (Stage 4: 전체 Best Config 적용):")
run.bold = True
add_table(
    ["모델", "Test Acc", "Test F1", "Adv F1", "추론 속도", "Bootstrap 95% CI"],
    [
        ["koelectra-base-v3\n(채택)", "0.9720", "0.9726", "0.8604", "7.9ms", "[0.952, 0.990]"],
        ["klue/bert-base", "--", "0.9756", "0.8517", "10.4ms", "[0.956, 0.992]"],
        ["distilkobert", "--", "0.9645", "0.7926", "2.8ms", "[0.940, 0.984]"],
    ],
)
doc.add_paragraph(
    "McNemar 검정: 3쌍 모두 n.s. (p>0.05) → Test 286개로는 통계적 유의차 검출 불가.\n"
    "동급일 때는 실용적 기준(Adv F1 +0.87%p, 추론 -24%)으로 koelectra 채택."
)

p = doc.add_paragraph()
run = p.add_run("Stage별 성능 변화 (KoELECTRA):")
run.bold = True
add_table(
    ["Stage", "내용", "Test F1", "Adv F1"],
    [
        ["Stage 2", "Baseline (ep5, lr2e-5, bs16)", "-- (Val 기준)", "-- (Val 기준)"],
        ["Stage 3", "Grid Search 32-point", "-- (Val 기준)", "-- (Val 기준)"],
        ["Stage 4", "최종 평가 (best config)", "0.9726", "0.8604"],
        ["Stage 5", "타겟 보강 +98개 재학습", "0.9788 (+0.62%p)", "0.8784 (+1.80%p)"],
        ["Stage 6 (채택)", "Label Smoothing 0.1", "0.9788 (유지)", "0.8758 (-0.26%p)"],
        ["Stage 7 (미채택)", "추가 보강 +25개", "0.9756 (-0.32%p)", "0.8712 (-0.46%p)"],
    ],
)

p = doc.add_paragraph()
run = p.add_run("최종 모델 Per-class F1 (Stage 6 — Adversarial):")
run.bold = True
add_table(
    ["Intent", "Precision", "Recall", "F1 Score"],
    [
        ["judgment", "0.982", "0.898", "0.938"],
        ["doc_search", "0.818", "0.900", "0.857"],
        ["doc_generate", "0.902", "0.885", "0.893"],
        ["doc_summary", "0.893", "0.943", "0.917"],
        ["schedule_add", "0.964", "0.946", "0.955"],
        ["schedule_view", "0.785", "0.911", "0.843"],
        ["general", "0.836", "0.836", "0.836"],
        ["doc_qa", "0.854", "0.695", "0.766"],
    ],
)

doc.add_heading("6-4. 과적합/과소적합 방지", level=2)
add_table(
    ["기법", "적용 내용", "효과"],
    [
        ["Label Smoothing", "CrossEntropyLoss에\nsmoothing=0.1 적용 (Stage 6)",
         "과신뢰 오분류 42건 → 13건 (-69%)\nThreshold 0.85로 정답/오답 분리 가능"],
        ["Dropout", "Classification Head에\nDropout(p=0.1) 적용",
         "Transformer 출력 일부를 랜덤 비활성화\n과적합 방지"],
        ["Weight Decay", "AdamW optimizer\nweight_decay=0.01",
         "L2 정규화 효과로 가중치 발산 방지"],
        ["Seed 고정", "random/numpy/torch/cuda\n4중 고정 (seed=42)",
         "3-seed 검증: 0.9874 +/- 0.0033\n학습 재현성 보장"],
        ["타겟 보강", "Adversarial 오분류 분석 후\n98개 약점 데이터 추가 (Stage 5)",
         "Adv F1 +1.80%p, doc_qa +7.9%p\n데이터 기반 약점 보완"],
    ],
)

p = doc.add_paragraph()
run = p.add_run("Label Smoothing 효과 상세:")
run.bold = True
add_table(
    ["항목", "Stage 5 (적용 전)", "Stage 6 (적용 후)", "변화"],
    [
        ["오분류 중 과신뢰(>90%)", "42건 (66.7%)", "13건 (23.2%)", "-69%"],
        ["정답 confidence 중앙값", "0.9968", "0.9366", "부드러운 분포"],
        ["오답 confidence 중앙값", "~0.90", "~0.64", "정답/오답 분리 가능"],
    ],
)
doc.add_paragraph(
    "이를 통해 Confidence Threshold 0.85 기반 라우팅 전략:\n"
    "- confidence >= 0.85: 해당 intent의 Agent 라우팅\n"
    "- 0.4 <= confidence < 0.85: clarify (top-3 후보 제시하여 되묻기)\n"
    "- confidence < 0.4: general 강제 라우팅"
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 7. 성능 평가
# ════════════════════════════════════════════════════════
doc.add_heading("7. 성능 평가", level=1)

doc.add_heading("7-1. 최종 모델 성능 요약", level=2)
add_table(
    ["항목", "값"],
    [
        ["모델", "monologg/koelectra-base-v3-discriminator"],
        ["학습 방식", "Full Fine-tuning + Label Smoothing 0.1"],
        ["Best Config", "epochs=10, lr=3e-5, batch=16, warmup=0.0"],
        ["파라미터", "112.9M"],
        ["모델 크기", "431MB"],
        ["추론 속도", "7.9ms mean / 8.3ms p95 (RTX 4090)"],
    ],
    col_widths=[5, 12],
)

add_table(
    ["평가셋", "Accuracy", "Precision", "Recall", "F1 Score"],
    [
        ["Validation (285개)", "0.9895", "0.9892", "0.9899", "0.9894"],
        ["Test (286개)", "0.9790", "0.9787", "0.9792", "0.9788"],
        ["Adversarial (450개)", "0.8756", "0.8792", "0.8768", "0.8758"],
    ],
)

p = doc.add_paragraph()
run = p.add_run("서비스 설정:")
run.bold = True
add_table(
    ["설정", "값", "효과"],
    [
        ["INTENT_CONFIDENCE_THRESHOLD", "0.85", "이하 → clarify (top-3 후보 제시)"],
        ["INTENT_FALLBACK_THRESHOLD", "0.4", "이하 → general 강제 라우팅"],
        ["과신뢰 오류 (>90% conf)", "13건", "Stage 5 대비 -69% (42→13)"],
    ],
)

doc.add_heading("7-2. 시나리오 테스트 (정성 평가)", level=2)
doc.add_paragraph(
    "실제 업무 환경을 시뮬레이션하는 100문장 시나리오 테스트를 수행하였습니다. "
    "Gemini 웹(Pro 3.1)으로 70개를 추가 생성하고 수동 검수하여 4유형 균형 배치하였습니다."
)
add_table(
    ["유형", "개수", "정답", "정확도", "설명"],
    [
        ["normal (표준)", "12", "12", "100.0%", "표준적인 업무 질의"],
        ["boundary (경계)", "33", "26", "78.8%", "intent 간 혼동 가능한 질의"],
        ["short (초단문)", "30", "28", "93.3%", "1~3어절 짧은 입력"],
        ["informal (비공식)", "25", "19", "76.0%", "슬랭/축약어/이모티콘"],
        ["전체", "100", "85", "85.0%", "Macro F1: 0.8497"],
    ],
)

p = doc.add_paragraph()
run = p.add_run("오분류 15건 패턴 분석:")
run.bold = True
add_table(
    ["패턴", "건수", "주요 예시"],
    [
        ["doc 경계 혼동", "6", '"확인해줘" → doc_qa vs doc_search\n"정리" → generate vs summary'],
        ["informal/슬랭", "4", '"쌉가능?", "어케됨?" → general로 오분류'],
        ["schedule add↔view", "2", '"~캘린더" 패턴이 add로 편향'],
        ["짧은 입력 애매", "2", '"문서 질문", "이름이 뭐야"'],
        ["기타", "1", '"번역도 돼?" (general → judgment)'],
    ],
)
doc.add_paragraph(
    "핵심 인사이트: normal 100% (표준 입력 완벽 처리), "
    "doc 경계 혼동 6건 중 4건은 동일 Document Agent 라우팅 → 실서비스 무해. "
    "informal이 최약점(76.0%)이나, 실서비스에서는 threshold 기반 clarify로 안전하게 라우팅."
)

doc.add_heading("7-3. 한계점 및 향후 과제", level=2)

p = doc.add_paragraph()
run = p.add_run("남은 약점:")
run.bold = True
add_bullet("doc_qa Adv F1 76.6% (8개 intent 중 최저)")
add_bullet("  단, 라벨 리뷰 결과 오류의 63%가 모델 문제가 아님 (Adjusted F1 ~85%)", level=1)
add_bullet("  doc 4개 intent 모두 동일 Agent 라우팅 → 실서비스 영향 제한적", level=1)
add_bullet("informal 시나리오 76.0% — 극단적 슬랭/축약어 인식 한계")
add_bullet("소량 보강의 한계 — 25개 타겟 보강으로는 유의미한 개선 없음 (Stage 7에서 확인)")

p = doc.add_paragraph()
run = p.add_run("향후 개선 방향:")
run.bold = True
add_bullet("대규모 doc 경계 데이터 확보 (25개가 아닌 200개+ 수준)")
add_bullet("informal/슬랭 학습 데이터 추가 (실서비스 로그 기반 수집)")
add_bullet("실서비스 로그 기반 재학습 (운영 데이터 축적 후 Fine-tuning 반복)")
add_bullet("Multi-intent 분해 (복합 질문 처리, 현재 비활성)")

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("핵심 교훈: ")
run.bold = True
run.font.size = Pt(12)
run = p.add_run('"데이터 품질 > 하이퍼파라미터 > 모델 아키텍처"')
run.font.size = Pt(12)
run.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("최종 업데이트: 2026-02-26")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("-- End of Document --")
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── 저장 ──────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"[OK] 문서 저장 완료: {OUTPUT_PATH}")
