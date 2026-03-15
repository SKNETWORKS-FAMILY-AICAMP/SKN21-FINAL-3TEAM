"""
4주차 산출물: 인공지능 데이터 전처리 결과서 재작성
기존 내용 유지 + 목차 추가 + 요구사항 구조 정리
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(
    os.path.dirname(__file__), "..",
    "docs", "산출물", "4주차",
    "2.데이터 전처리_인공지능 데이터 전처리 결과서_3팀.docx",
)

doc = Document()

# ── 스타일 ──
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for lv in range(1, 4):
    h = doc.styles[f"Heading {lv}"]
    h.font.name = "맑은 고딕"
    h.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    h.font.size = [None, Pt(16), Pt(13), Pt(11)][lv]


# ── 헬퍼 ──
def tbl(headers, rows, cw=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i + 1].cells[j]
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if cw:
        for i, w in enumerate(cw):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    return p

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_toc_field(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin"); r._r.append(fc1)
    r2 = p.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = ' TOC \\o "1-2" \\h \\z \\u '
    r2._r.append(it)
    r3 = p.add_run()
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "separate"); r3._r.append(fc2)
    r4 = p.add_run("[Ctrl+A -> F9 로 목차 업데이트]")
    r4.font.color.rgb = RGBColor(0x99, 0x99, 0x99); r4.font.size = Pt(9); r4.italic = True
    r5 = p.add_run()
    fc3 = OxmlElement("w:fldChar"); fc3.set(qn("w:fldCharType"), "end"); r5._r.append(fc3)


# ══════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SK networks  |  Family AI Camp"); r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("데이터 전처리\n인공지능 데이터 전처리 결과서"); r.font.size = Pt(26); r.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SKN Family AI Camp 21기 : 최종 프로젝트 3팀\nWorkFlow Agent (듀듀)"); r.font.size = Pt(13)
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("작성: 3팀 -- 신지용(PM), 윤경은(AI), 진승언(AI), 안혜빈(Backend), 문지영(Frontend)")
r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_page_break()

# ══════════════════════════════════════════════
# 목차
# ══════════════════════════════════════════════
doc.add_heading("목 차", level=1)
add_toc_field(doc)
doc.add_paragraph()
toc = [
    ("1. 데이터셋 개요", True),
    ("    1-1. 대상 데이터 요약", False),
    ("    1-2. Intent 카테고리 (8개)", False),
    ("    1-3. 데이터 구성 분류", False),
    ("2. 원본 데이터 샘플 (예시)", True),
    ("    2-1. 기본 데이터 샘플", False),
    ("    2-2. 경계 쌍 데이터 샘플", False),
    ("    2-3. Adversarial 테스트 데이터 샘플", False),
    ("3. 전처리 흐름도", True),
    ("    3-1. 전체 파이프라인", False),
    ("    3-2. 실험 기반 보강 루프 (Stage 5~7)", False),
    ("4. 세부 전처리 내용", True),
    ("    4-1. 결측치 제거 및 라벨 유효성 검증", False),
    ("    4-2. 중복 제거 (Deduplication)", False),
    ("    4-3. 클래스 균형 검증 (정규화)", False),
    ("    4-4. 텍스트 전처리 (preprocessing.py)", False),
    ("    4-5. 전처리 Ablation 실험 결과", False),
    ("    4-6. 데이터 분할 (Stratified Split)", False),
    ("    4-7. 토크나이징 (Tokenizer)", False),
    ("    4-8. 라벨 인코딩 (Label Encoding)", False),
    ("5. 전처리 전후 데이터 비교", True),
    ("    5-1. 전처리 전후 비교 (샘플)", False),
    ("    5-2. 최종 데이터 파일 구조", False),
    ("    5-3. 사용 도구", False),
]
for text, major in toc:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(2)
    if major:
        p.paragraph_format.space_before = Pt(6)
        for r in p.runs:
            r.bold = True; r.font.size = Pt(11)
    else:
        for r in p.runs:
            r.font.size = Pt(10)
doc.add_page_break()

# ══════════════════════════════════════════════
# 1. 데이터셋 개요
# ══════════════════════════════════════════════
doc.add_heading("1. 데이터셋 개요", level=1)
doc.add_paragraph(
    "Intent Classifier v2 학습을 위해 수집/생성한 데이터셋의 전처리 결과를 기술합니다. "
    "본 문서는 v2 실험 기준만을 다루며, v1 실험 내용은 포함하지 않습니다."
)

doc.add_heading("1-1. 대상 데이터 요약", level=2)
tbl(
    ["항목", "설명", "예시"],
    [
        ["모델", "koelectra-base-v3-discriminator\n(112.9M params, ELECTRA RTD)", "--"],
        ["학습 목표", "사용자 입력을 8개 intent로 분류\n-> 적합한 Agent로 라우팅", '"연차 써도 돼?" -> judgment'],
        ["데이터 형식", "JSONL (text + label 2필드)", '{"text":"...", "label":"judgment"}'],
        ["생성 방식", "멀티 LLM 합성\n(GPT-4o + Claude Sonnet 4)", "각 LLM에서 동일 수량 생성"],
        ["총 데이터", "2,996건\n(기본 2,299 + 경계 600 + 보강 98)", "Train 2,425 / Val 285 / Test 286"],
        ["평가 데이터", "Adversarial 450건 + 시나리오 100건", "초단문/오타/간접표현 등"],
        ["최종 성능", "Test F1 97.88%\nAdversarial F1 87.58%", "추론 속도 7.9ms (GPU)"],
    ],
)

doc.add_heading("1-2. Intent 카테고리 (8개)", level=2)
tbl(
    ["Intent", "라우팅 Agent", "설명", "예시"],
    [
        ["judgment", "Judgment Agent", "사내 규정/정책 판단 질의", '"인턴에게 AWS 권한 줘도 돼?"'],
        ["doc_search", "Document Agent", "문서 검색/조회", '"연차 규정 문서 찾아줘"'],
        ["doc_generate", "Document Agent", "문서 생성 (보고서, 제안서 등)", '"보고서 만들어줘"'],
        ["doc_summary", "Document Agent", "문서 요약", '"이 문서 요약해줘"'],
        ["doc_qa", "Document Agent", "문서 기반 Q&A", '"보고서에 매출이 얼마야?"'],
        ["schedule_add", "Schedule Agent", "일정 등록", '"내일 3시에 팀미팅 잡아줘"'],
        ["schedule_view", "Schedule Agent", "일정 조회", '"이번 주 일정 알려줘"'],
        ["general", "General Handler", "일상 대화/인사", '"안녕하세요"'],
    ],
)

doc.add_heading("1-3. 데이터 구성 분류", level=2)
tbl(
    ["구분", "건수", "용도", "생성 방식"],
    [
        ["기본 데이터", "2,299건", "8 intent x ~288개\n(학습 데이터 주축)", "GPT-4o 150개 + Claude 150개\n/ intent"],
        ["경계 쌍 데이터", "600건", "혼동 가능 intent 쌍\n구분력 강화", "10개 쌍 x 30건 x 2 LLM"],
        ["타겟 보강 (Stage 5)", "98건", "오분류 패턴 기반\n취약 intent 보강", "Adversarial 오분류 분석 후\n타겟 생성"],
        ["Adversarial 테스트셋", "450건", "분류기 강건성 평가", "GPT 232 + Claude 240\n(중복 제거)"],
        ["시나리오 테스트셋", "100건", "실서비스 시나리오\n정성 평가", "Gemini Pro 3.1 생성\n+ 수동 검수"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 2. 원본 데이터 샘플
# ══════════════════════════════════════════════
doc.add_heading("2. 원본 데이터 샘플 (예시)", level=1)
doc.add_paragraph("각 데이터 유형별 원본 샘플을 아래에 제시합니다. 모든 데이터는 동일한 JSONL 형식(text + label)입니다.")

doc.add_heading("2-1. 기본 데이터 샘플 (8개 intent, 각 1건)", level=2)
tbl(
    ["#", "text (원문)", "label"],
    [
        ["1", "지각 3번이면 경고 받아?", "judgment"],
        ["2", "연차 규정 문서 찾아줘", "doc_search"],
        ["3", "보고서 만들어줘", "doc_generate"],
        ["4", "이 문서 요약해줘", "doc_summary"],
        ["5", "보고서에 나온 납기일이 언제야?", "doc_qa"],
        ["6", "내일 3시에 팀미팅 잡아줘", "schedule_add"],
        ["7", "이번 주 일정 알려줘", "schedule_view"],
        ["8", "안녕하세요", "general"],
    ],
)

doc.add_heading("2-2. 경계 쌍 데이터 샘플 (혼동 가능 intent 쌍)", level=2)
tbl(
    ["#", "text (원문)", "label", "경계 쌍"],
    [
        ["1", "출장비 관련 규정 문서 찾아줘", "doc_search", "doc_search <-> doc_qa"],
        ["2", "출장비 한도가 얼마인지 알려줘", "doc_qa", "doc_search <-> doc_qa"],
        ["3", "보안 규정 위반하면 어떻게 돼?", "judgment", "judgment <-> doc_qa"],
        ["4", "보안 규정 내용 요약해줘", "doc_summary", "doc_summary <-> doc_qa"],
    ],
)

doc.add_heading("2-3. Adversarial 테스트 데이터 샘플 (모델 취약 유형)", level=2)
tbl(
    ["#", "text (원문)", "label", "어려운 이유"],
    [
        ["1", "연차 되나?", "judgment", "초단문 (2어절)"],
        ["2", "반차 쓸 수 있어?", "judgment", "간접 표현"],
        ["3", "재택 가능?", "judgment", "맥락 의존 초단문"],
        ["4", "ㅂㄱㅅ 써줘", "doc_generate", "초성 입력"],
        ["5", "회이록 정리해조", "doc_generate", "오타 + 비표준"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 3. 전처리 흐름도
# ══════════════════════════════════════════════
doc.add_heading("3. 전처리 흐름도", level=1)
doc.add_paragraph(
    "Intent Classifier v2의 전처리는 크게 데이터 생성 -> 품질 검증 -> "
    "텍스트 전처리 -> 분할의 4단계로 구성됩니다."
)

doc.add_heading("3-1. 전체 파이프라인", level=2)
code_block(
    "+-------------------------------------------------------------+\n"
    "|  Step 1: 데이터 생성 (generate_data.py)                     |\n"
    "|          GPT-4o (150/intent) + Claude (150/intent)           |\n"
    "|          + 경계 쌍 600건 + Adversarial 450건                 |\n"
    "+-------------------------------------------------------------+\n"
    "                              |\n"
    "                              v\n"
    "+-------------------------------------------------------------+\n"
    "|  Step 2: 품질 검증 (QA)                                     |\n"
    "|   - 결측치 검증 (text/label 필드 존재 확인)                  |\n"
    "|   - 라벨 유효성 (8개 intent 외 라벨 탐지)                   |\n"
    "|   - 중복 제거 (exact match, GPT 내부 59건 제거)             |\n"
    "|   - 클래스 균형 (max/min ratio = 1.11)                      |\n"
    "|   - 데이터 누출 검증 (test ∩ train = 공집합)                |\n"
    "+-------------------------------------------------------------+\n"
    "                              |\n"
    "                              v\n"
    "+-------------------------------------------------------------+\n"
    "|  Step 3: 텍스트 전처리 (preprocessing.py)                   |\n"
    "|   P4. 공백/특수문자 정리 (반복문자 축소, 연속공백 제거)      |\n"
    "|   P1. 맞춤법 교정 (업무용어 사전 기반)                      |\n"
    "|   P2. 초성 복원 (업무 초성 -> 완성형, ㅋㅎㅠ 보존)          |\n"
    "|   P3. 슬랭 정규화 (인터넷 슬랭/축약어 -> 표준어)            |\n"
    "|                                                             |\n"
    "|   * Ablation 결과: 전처리 유무 성능 동일                    |\n"
    "|     -> P4(공백 정리)만 서비스 기본 적용                     |\n"
    "+-------------------------------------------------------------+\n"
    "                              |\n"
    "                              v\n"
    "+-------------------------------------------------------------+\n"
    "|  Step 4: 데이터 분할 + 인코딩                               |\n"
    "|   - Stratified Split 80/10/10 (seed=42)                     |\n"
    "|   - WordPiece 토크나이징 (max_length=64)                    |\n"
    "|   - 라벨 인코딩 (8개 intent -> 0~7 정수)                   |\n"
    "+-------------------------------------------------------------+\n"
    "                              |\n"
    "                              v\n"
    "          Train 2,425 / Val 285 / Test 286\n"
    "          + Adversarial 450 (평가 전용)\n"
    "          + Scenario 100 (정성 평가 전용)"
)

doc.add_heading("3-2. 실험 기반 보강 루프 (Stage 5~7)", level=2)
doc.add_paragraph(
    "학습 완료 후 Adversarial 평가 결과를 분석하여 약점 intent를 "
    "타겟 보강하는 반복 루프를 수행합니다."
)
code_block(
    "학습 완료 모델\n"
    "     |\n"
    "     v\n"
    "Adversarial 450건 평가 -> 오분류 63건 추출\n"
    "     |\n"
    "     v\n"
    "오분류 패턴 분석 (short_text 47건, overconfident 42건, boundary 30건)\n"
    "     |\n"
    "     v\n"
    "타겟 보강 데이터 생성 (98건, 8 intent 취약점 집중)\n"
    "     |\n"
    "     v\n"
    "보강 데이터 QA (적대적 누출 0건, 중복 0건 확인)\n"
    "     |\n"
    "     v\n"
    "Train에 추가 -> 재학습 -> 재평가\n"
    "     |\n"
    "     v\n"
    "Adv F1: 86.04% -> 87.84% (+1.80%p)\n"
    "doc_qa: 71.0% -> 78.9% (+7.9%p, 최대 개선)"
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 4. 세부 전처리 내용
# ══════════════════════════════════════════════
doc.add_heading("4. 세부 전처리 내용", level=1)

# 4-1
doc.add_heading("4-1. 결측치 제거 및 라벨 유효성 검증", level=2)
doc.add_paragraph(
    "모든 데이터는 LLM API로 생성된 합성 데이터이므로 결측치가 구조적으로 발생하기 어렵습니다. "
    "그럼에도 불구하고 다음 검증을 자동화하여 적용하였습니다."
)
tbl(
    ["검증 항목", "검증 방법", "결과"],
    [
        ["text 필드 결측", "JSONL 라인별 text 키 존재\n+ 빈 문자열 여부 체크", "결측 0건"],
        ["label 필드 결측", "JSONL 라인별 label 키 존재 체크", "결측 0건"],
        ["라벨 유효성", "8개 intent 목록 이외 라벨 탐지", "비유효 0건"],
        ["빈 문자열", "text.strip() == '' 체크", "빈 문자열 0건"],
        ["JSON 파싱 오류", "json.loads() 실패 건수", "파싱 오류 0건"],
    ],
)

# 4-2
doc.add_heading("4-2. 중복 제거 (Deduplication)", level=2)
doc.add_paragraph("텍스트 해시 기반 완전 중복(exact match)을 제거하였습니다.")
tbl(
    ["Intent", "중복 전", "중복 제거", "중복 후"],
    [
        ["judgment", "277", "0", "277"],
        ["doc_search", "286", "0", "286"],
        ["doc_generate", "297", "25", "272"],
        ["doc_summary", "302", "8", "294"],
        ["doc_qa", "284", "0", "284"],
        ["schedule_add", "299", "1", "298"],
        ["schedule_view", "296", "9", "287"],
        ["general", "317", "16", "301"],
        ["합계", "2,358", "59 (2.5%)", "2,299"],
    ],
)
doc.add_paragraph(
    "Adversarial 데이터에서도 Train과의 중복 제거 적용: "
    "463건 -> 450건 (13건 중복 제거)"
)

# 4-3
doc.add_heading("4-3. 클래스 균형 검증 (정규화)", level=2)
doc.add_paragraph(
    "클래스 불균형은 분류기 편향을 유발할 수 있어, max/min ratio를 1.2 이하로 관리합니다. "
    "중복 제거 후 기본 데이터 기준으로 ratio = 1.11로 양호합니다."
)
tbl(
    ["Intent", "기본 데이터", "+ 경계 쌍", "+ 보강 (Stage 5)", "최종 Train"],
    [
        ["judgment", "277", "+60", "+10", "~307"],
        ["doc_search", "286", "+60", "+10", "~316"],
        ["doc_generate", "272", "+60", "+15", "~307"],
        ["doc_summary", "294", "+60", "+10", "~324"],
        ["doc_qa", "284", "+60", "+20", "~324"],
        ["schedule_add", "298", "+60", "+11", "~329"],
        ["schedule_view", "287", "+60", "+11", "~318"],
        ["general", "301", "+60", "+11", "~332"],
        ["합계", "2,299", "+600", "+98", "2,425 (Train)"],
    ],
)
bullet("경계 쌍 600건은 10개 쌍 x 30건 x 2 LLM, 각 intent에 균등 배분")
bullet("최종 Train 기준 클래스 균형: max/min ratio = 1.28x (양호)")

# 4-4
doc.add_heading("4-4. 텍스트 전처리 (preprocessing.py)", level=2)
doc.add_paragraph(
    "4단계 규칙 기반 전처리 파이프라인을 구현하였습니다. 각 단계는 개별 on/off가 가능하며, "
    "Ablation 실험으로 효과를 검증하였습니다."
)
doc.add_paragraph("실행 순서: P4(공백 정리) -> P1(맞춤법) -> P2(초성 복원) -> P3(슬랭 정규화)")

p = doc.add_paragraph()
r = p.add_run("P4. 공백/특수문자 정리 (Text Cleaning)"); r.bold = True
tbl(
    ["처리", "규칙", "예시 (Before -> After)"],
    [
        ["반복 문자 축소", '3회 이상 -> 1회\nre.sub(r"(.)\\1{2,}", r"\\1")', "ㅋㅋㅋㅋ -> ㅋ"],
        ["반복 특수문자", "2회 이상 -> 1회", "!!!!! -> !\n???? -> ?"],
        ["연속 공백", "다중 공백 -> 단일 공백", '"연차   규정" -> "연차 규정"'],
        ["양쪽 공백", "strip()", '" 연차 규정 " -> "연차 규정"'],
    ],
)

p = doc.add_paragraph()
r = p.add_run("P1. 맞춤법 교정 (Spell Check)"); r.bold = True
tbl(
    ["처리", "규칙", "예시 (Before -> After)"],
    [
        ["업무 용어 교정", "규칙 기반 사전 매핑\n(SPELL_CORRECTIONS dict)", '"회이록" -> "회의록"\n"보거서" -> "보고서"'],
        ["붙여쓰기 교정", "복합어 규칙 적용", '"일졍" -> "일정"\n"스케쥴" -> "스케줄"'],
        ["구어체 교정", "비표준 종결어미 정규화", '"해조" -> "해줘"\n"해줘어" -> "해줘"'],
    ],
)

p = doc.add_paragraph()
r = p.add_run("P2. 초성 복원 (Chosung Restoration)"); r.bold = True
tbl(
    ["처리", "규칙", "예시 (Before -> After)"],
    [
        ["업무 초성 복원", "CHOSUNG_MAP 사전 매핑\n(업무 관련 초성만 대상)", '"ㅎㅇㄹ" -> "회의록"\n"ㅂㄱㅅ" -> "보고서"'],
        ["일반 초성 보존", "반응형 초성(ㅋ, ㅎ, ㅠ 등)은\n변환하지 않음", '"ㅋㅋ" -> "ㅋㅋ" (유지)'],
    ],
)

p = doc.add_paragraph()
r = p.add_run("P3. 슬랭 정규화 (Slang Normalization)"); r.bold = True
tbl(
    ["처리", "규칙", "예시 (Before -> After)"],
    [
        ["인터넷 슬랭", "SLANG_MAP 사전 매핑", '"걍" -> "그냥"\n"겜" -> "게임"'],
        ["축약어", "구어체 축약 정규화", '"넹" -> "네"\n"넵" -> "네"'],
        ["신조어", "비표준 표현 정규화", '"고고" -> "하자"\n"오키" -> "알겠어"'],
    ],
)

# 4-5
doc.add_heading("4-5. 전처리 Ablation 실험 결과", level=2)
doc.add_paragraph(
    "Stage 4 최종 평가에서 5가지 전처리 설정(Config A~E)에 대해 Ablation 실험을 수행한 결과, "
    "전처리 유무에 따른 성능 차이가 없었습니다. "
    "koelectra-base-v3의 WordPiece 토크나이저가 이미 비정형 입력을 효과적으로 처리하기 때문입니다."
)
tbl(
    ["Config", "적용 단계", "Test F1", "Adv F1"],
    [
        ["A (원본)", "없음", "97.26%", "86.04%"],
        ["B", "P4 (공백만)", "97.26%", "86.04%"],
        ["C", "P4 + P1", "97.26%", "86.04%"],
        ["D", "P4 + P1 + P2", "97.26%", "86.04%"],
        ["E (전체)", "P4 + P1 + P2 + P3", "97.26%", "86.04%"],
    ],
)
p = doc.add_paragraph()
r = p.add_run("-> 결론: "); r.bold = True
p.add_run(
    "서비스 안정성을 위해 P4(공백 정리)만 기본 적용. "
    "전처리보다 데이터 품질이 성능에 결정적."
)

# 4-6
doc.add_heading("4-6. 데이터 분할 (Stratified Split)", level=2)
doc.add_paragraph(
    "Stratified Split으로 각 intent의 비율을 유지하면서 80/10/10으로 분할합니다. "
    "seed=42를 고정하여 재현성을 보장합니다."
)
code_block(
    "random.seed(42)\n"
    "for label, items in by_label.items():\n"
    "    random.shuffle(items)\n"
    "    n = len(items)\n"
    "    n_test = max(1, int(n * 0.1))\n"
    "    n_val  = max(1, int(n * 0.1))\n"
    "    test  += items[:n_test]\n"
    "    val   += items[n_test:n_test+n_val]\n"
    "    train += items[n_test+n_val:]"
)
tbl(
    ["분할", "건수", "비율", "용도"],
    [
        ["Train", "2,425건\n(기본 2,327 + 보강 98)", "~80%", "모델 학습"],
        ["Validation", "285건", "~10%", "학습 중 성능 모니터링"],
        ["Test", "286건", "~10%", "최종 정량 평가"],
        ["Adversarial", "450건", "별도", "강건성 평가 (학습 미사용)"],
        ["Scenario", "100건", "별도", "정성 평가 (학습 미사용)"],
    ],
)
doc.add_paragraph(
    "데이터 누출 검증: Train<->Val 중복 0건, Train<->Test 중복 0건, Val<->Test 중복 0건"
)

# 4-7
doc.add_heading("4-7. 토크나이징 (Tokenizer)", level=2)
tbl(
    ["항목", "설정"],
    [
        ["토크나이저", "koelectra-base-v3 WordPiece"],
        ["max_length", "64 tokens"],
        ["padding", '"max_length" (고정 길이 패딩)'],
        ["truncation", "True (64 초과 시 절단)"],
        ["special tokens", "[CLS] + text + [SEP] + [PAD]..."],
    ],
)
code_block(
    'tokenizer = AutoTokenizer.from_pretrained(\n'
    '    "monologg/koelectra-base-v3-discriminator"\n'
    ')\n'
    'encoded = tokenizer(\n'
    '    texts,\n'
    '    padding="max_length",\n'
    '    max_length=64,\n'
    '    truncation=True,\n'
    '    return_tensors="pt"\n'
    ')'
)

# 4-8
doc.add_heading("4-8. 라벨 인코딩 (Label Encoding)", level=2)
tbl(
    ["label (문자열)", "label_id (정수)", "Agent"],
    [
        ["judgment", "0", "Judgment Agent"],
        ["doc_search", "1", "Document Agent"],
        ["doc_generate", "2", "Document Agent"],
        ["doc_summary", "3", "Document Agent"],
        ["doc_qa", "4", "Document Agent"],
        ["schedule_add", "5", "Schedule Agent"],
        ["schedule_view", "6", "Schedule Agent"],
        ["general", "7", "General Handler"],
    ],
)
code_block(
    'LABEL2ID = {\n'
    '    "judgment": 0, "doc_search": 1, "doc_generate": 2,\n'
    '    "doc_summary": 3, "doc_qa": 4, "schedule_add": 5,\n'
    '    "schedule_view": 6, "general": 7\n'
    '}\n'
    'labels = [LABEL2ID[item["label"]] for item in dataset]'
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 5. 전처리 전후 데이터 비교
# ══════════════════════════════════════════════
doc.add_heading("5. 전처리 전후 데이터 비교", level=1)

doc.add_heading("5-1. 전처리 전후 비교 (샘플)", level=2)
tbl(
    ["#", "원본 (Before)", "전처리 후 (After)", "적용 단계"],
    [
        ["1", "회이록   정리해조", "회의록 정리해줘", "P4 -> P1"],
        ["2", "ㅂㄱㅅ 써줘!!!!", "보고서 써줘!", "P4 -> P2"],
        ["3", "걍 ㅇㅈ 알려줘ㅋㅋㅋ", "그냥 일정 알려줘ㅋ", "P4 -> P2 -> P3"],
        ["4", "연차  되나???", "연차 되나?", "P4"],
        ["5", "넹 스케쥴 확인해조", "네 스케줄 확인해줘", "P1 -> P3"],
    ],
)

doc.add_heading("5-2. 최종 데이터 파일 구조", level=2)
code_block(
    "data/training/intent_v2/\n"
    "|-- raw/                        # LLM별 원본 (GPT/Claude x 8 intent)\n"
    "|-- {intent}.jsonl x 8          # intent별 정제 데이터 (중복 제거 후)\n"
    "|-- boundary_pairs.jsonl        # 경계 쌍 데이터 (600건)\n"
    "|-- adversarial_v2.json         # Adversarial 테스트셋 (450건)\n"
    "|-- scenario_test.json          # 시나리오 테스트 (100건)\n"
    "|-- augmentation_stage5.jsonl   # 타겟 보강 데이터 (98건)\n"
    "|-- splits/\n"
    "|   |-- train.jsonl             # 2,425건 (80%)\n"
    "|   |-- val.jsonl               # 285건 (10%)\n"
    "|   +-- test.jsonl              # 286건 (10%)\n"
    "+-- DATA_QA_REPORT.md           # 품질 검증 보고서"
)

doc.add_heading("5-3. 사용 도구", level=2)
tbl(
    ["도구", "버전", "용도"],
    [
        ["Python", "3.11+", "전처리 스크립트 실행"],
        ["transformers", "4.40+", "토크나이저 + 모델 학습"],
        ["datasets", "2.19+", "HuggingFace Dataset 로드"],
        ["scikit-learn", "1.4+", "Stratified Split, 평가 메트릭"],
        ["accelerate", "0.29+", "HuggingFace 학습 가속"],
        ["jsonschema", "4.x", "JSONL 스키마 자동 검증"],
    ],
)

# ── 풋터 ──
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("최종 수정일: 2026-02-26"); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("-- End of Document --"); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save(OUTPUT)
print(f"[OK] 저장 완료: {OUTPUT}")
