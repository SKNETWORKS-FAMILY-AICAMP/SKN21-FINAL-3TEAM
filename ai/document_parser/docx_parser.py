"""
DOCX 파서 (팀원 C 담당)

python-docx를 사용하여 DOCX 문서에서 텍스트를 추출.
테이블이 있으면 마크다운 테이블로 변환.
"""

import logging

logger = logging.getLogger(__name__)


class DocxParser:
    """python-docx 기반 DOCX 파싱"""

    def parse(self, file_path: str) -> str:
        """
        DOCX -> 마크다운 형식 텍스트 추출

        - 문단: 스타일(Heading 1~3)에 따라 마크다운 헤딩 적용
        - 테이블: 마크다운 테이블로 변환
        - 빈 문단은 건너뜀

        Returns:
            마크다운 형식 텍스트
        """
        from docx import Document

        doc = Document(file_path)
        parts = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1]  # namespace 제거

            if tag == "p":
                # 문단 처리
                para = self._find_paragraph(doc, element)
                if para is None:
                    continue
                text = para.text.strip()
                if not text:
                    continue

                # 스타일에 따라 마크다운 헤딩
                style_name = para.style.name if para.style else ""
                if "Heading 1" in style_name:
                    parts.append(f"# {text}")
                elif "Heading 2" in style_name:
                    parts.append(f"## {text}")
                elif "Heading 3" in style_name:
                    parts.append(f"### {text}")
                else:
                    parts.append(text)

            elif tag == "tbl":
                # 테이블 처리
                table = self._find_table(doc, element)
                if table is None:
                    continue
                md_table = self._table_to_markdown(table)
                if md_table:
                    parts.append(md_table)

        result = "\n\n".join(parts)
        logger.info("DOCX 파싱 완료: %s (%d자)", file_path, len(result))
        return result

    def _find_paragraph(self, doc, element):
        """element에 대응하는 Paragraph 객체 찾기"""
        for para in doc.paragraphs:
            if para._element is element:
                return para
        return None

    def _find_table(self, doc, element):
        """element에 대응하는 Table 객체 찾기"""
        for table in doc.tables:
            if table._element is element:
                return table
        return None

    def _table_to_markdown(self, table) -> str:
        """docx Table -> 마크다운 테이블 문자열"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        # 첫 행을 헤더로
        col_count = len(rows[0])
        lines = []
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join(["---"] * col_count) + " |")

        for row in rows[1:]:
            # 셀 수 맞추기
            padded = row + [""] * (col_count - len(row))
            lines.append("| " + " | ".join(padded[:col_count]) + " |")

        return "\n".join(lines)
