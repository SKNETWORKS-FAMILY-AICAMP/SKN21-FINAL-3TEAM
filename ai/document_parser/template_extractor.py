"""
DOCX 양식 필드 추출기 v3 (구조 기반 + 메타데이터 자동 부여)

어떤 DOCX 양식이든 구조를 분석하여 필드를 자동 추출한다.

추출 전략:
  1. 테이블 전체 열 스캔: [라벨 | 빈칸] 패턴 탐지 (홀수 열도)
  2. 섹션 테이블: 1열 테이블의 헤더 행
  3. 병합 헤더: 모든 셀이 같은 텍스트인 행
  4. 헤딩 스타일: Heading 1~3
  5. "라벨:" 본문 패턴

키 생성:
  - FIELD_MAPPING에 있으면 → 영어 키 (title, date 등)
  - 없으면 → 정규화된 한글 라벨을 키로 사용 (커스텀 필드)

메타데이터 (v3):
  - group: "meta" (기본 정보) / "body" (본문 서술)
  - type: "date" / "text" / "textarea" / "array"
  - fill: "extract" (텍스트에서 추출) / "generate" (AI 생성)
  - desc_source: "mapping" / "partial" / "fallback" (description 출처)

사용법:
    from ai.document_parser.template_extractor import extract_template_fields
    fields = extract_template_fields("양식.docx")
"""
import json
import logging
import re

logger = logging.getLogger(__name__)

# ── 한국어 → 영어 키 매핑 ──

FIELD_MAPPING = {
    # 공통
    "제목": {"key": "title", "desc": "문서 제목"},
    "문서제목": {"key": "title", "desc": "문서 제목"},
    "일시": {"key": "date", "desc": "날짜 (YYYY-MM-DD 형식, 없으면 오늘 날짜)"},
    "날짜": {"key": "date", "desc": "날짜 (YYYY-MM-DD 형식, 없으면 오늘 날짜)"},
    "일자": {"key": "date", "desc": "날짜 (YYYY-MM-DD 형식, 없으면 오늘 날짜)"},
    "작성일": {"key": "date", "desc": "작성 날짜 (YYYY-MM-DD 형식)"},
    "작성일자": {"key": "date", "desc": "작성 날짜 (YYYY-MM-DD 형식)"},
    "작성자": {"key": "author", "desc": "작성자 이름 (없으면 빈 문자열)"},
    "기록자": {"key": "author", "desc": "기록자 이름 (없으면 빈 문자열)"},

    # 회의록
    "참석자": {"key": "attendees", "desc": "참석자 이름 배열 (없으면 빈 배열)"},
    "참석인원": {"key": "attendees", "desc": "참석자 이름 배열 (없으면 빈 배열)"},
    "회의참석자": {"key": "attendees", "desc": "참석자 이름 배열 (없으면 빈 배열)"},
    "회의자": {"key": "attendees", "desc": "회의 참석자 이름 배열 (없으면 빈 배열)"},
    "장소": {"key": "location", "desc": "장소 (없으면 빈 문자열)"},
    "회의장소": {"key": "location", "desc": "회의 장소 (없으면 빈 문자열)"},
    "회의명": {"key": "title", "desc": "회의 주제를 반영한 구체적인 제목"},
    "회의제목": {"key": "title", "desc": "회의 주제를 반영한 구체적인 제목"},
    "회의주제": {"key": "title", "desc": "회의 주제를 반영한 구체적인 제목"},
    "회의유형": {"key": "meeting_type", "desc": "회의 유형 ('정기', '비정기', '긴급' 중 하나)"},
    "회의종류": {"key": "meeting_type", "desc": "회의 유형 ('정기', '비정기', '긴급' 중 하나)"},
    "시간": {"key": "time", "desc": "시간 (예: '14:00~15:30')"},
    "회의시간": {"key": "time", "desc": "회의 시간 (예: '14:00~15:30')"},
    "회의일시": {"key": "date", "desc": "회의 날짜 (YYYY-MM-DD 형식)"},
    "진행자": {"key": "moderator", "desc": "진행자/사회자 이름 (없으면 빈 문자열)"},
    "사회자": {"key": "moderator", "desc": "진행자/사회자 이름 (없으면 빈 문자열)"},
    "안건": {"key": "agenda", "desc": "회의 안건 목록 (배열)"},
    "회의안건": {"key": "agenda", "desc": "회의 안건 목록 (배열)"},
    "회의내용": {"key": "content", "desc": "회의 내용을 상세하게 기술"},
    "논의내용": {"key": "content", "desc": "논의 내용을 상세하게 기술"},
    "논의사항": {"key": "content", "desc": "논의 사항을 상세하게 기술"},
    "주요내용": {"key": "content", "desc": "주요 내용을 상세하게 기술"},
    "내용": {"key": "content", "desc": "내용을 상세하게 기술"},
    "요약": {"key": "summary", "desc": "주요 내용을 3~5문장으로 요약"},
    "회의요약": {"key": "summary", "desc": "회의에서 논의된 주요 내용을 3~5문장으로 요약"},
    "결정사항": {"key": "decisions", "desc": "결정된 사항 목록 (배열, 없으면 빈 배열)"},
    "결정된사항": {"key": "decisions", "desc": "결정된 사항 목록 (배열, 없으면 빈 배열)"},
    "의결사항": {"key": "decisions", "desc": "결정된 사항 목록 (배열, 없으면 빈 배열)"},
    "결정": {"key": "decisions", "desc": "결정 사항 목록 (배열)"},
    "후속조치": {"key": "action_items", "desc": '후속 조치 목록 배열. 각 항목은 {"task": "할 일", "assignee": "담당자", "due_date": "기한"} 형태'},
    "실행계획": {"key": "action_items", "desc": '후속 조치 목록 배열. 각 항목은 {"task": "할 일", "assignee": "담당자", "due_date": "기한"} 형태'},
    "조치사항": {"key": "action_items", "desc": '후속 조치 목록 배열. 각 항목은 {"task": "할 일", "assignee": "담당자", "due_date": "기한"} 형태'},
    "actionitem": {"key": "action_items", "desc": '후속 조치 목록 배열. 각 항목은 {"task": "할 일", "assignee": "담당자", "due_date": "기한"} 형태'},
    "다음회의": {"key": "next_meeting", "desc": "다음 회의 일정 (없으면 빈 문자열)"},
    "차기회의": {"key": "next_meeting", "desc": "다음 회의 일정 (없으면 빈 문자열)"},
    "비고": {"key": "notes", "desc": "비고 사항 (없으면 빈 문자열)"},
    "특이사항": {"key": "notes", "desc": "비고 사항 (없으면 빈 문자열)"},
    "회의목적": {"key": "meeting_purpose", "desc": "회의 목적 (1~2문장)"},
    "비고다음회의일정": {"key": "notes", "desc": "비고 및 다음 회의 일정 (없으면 빈 문자열)"},
    "회의기록부": None,  # 문서 타이틀이지 필드가 아님

    # 보고서
    "부서": {"key": "department", "desc": "부서명 (없으면 빈 문자열)"},
    "부서명": {"key": "department", "desc": "부서명 (없으면 빈 문자열)"},
    "소속": {"key": "department", "desc": "부서명 (없으면 빈 문자열)"},
    "직급": {"key": "position", "desc": "직급 (없으면 빈 문자열)"},
    "직위": {"key": "position", "desc": "직급 (없으면 빈 문자열)"},
    "보고대상": {"key": "report_to", "desc": "보고 대상 (없으면 빈 문자열)"},
    "수신": {"key": "report_to", "desc": "보고 대상 (없으면 빈 문자열)"},
    "보고유형": {"key": "report_type", "desc": "'일일', '주간', '월간', '수시' 중 하나"},
    "보고종류": {"key": "report_type", "desc": "'일일', '주간', '월간', '수시' 중 하나"},
    "보고기간": {"key": "period", "desc": "보고 기간 (예: '2026년 3월 1주차')"},
    "기간": {"key": "period", "desc": "기간 (예: '2026년 3월 ~ 6월')"},
    "개요": {"key": "overview", "desc": "업무 내용을 요약한 보고 개요 (3~5문장)"},
    "보고개요": {"key": "overview", "desc": "업무 내용을 요약한 보고 개요 (3~5문장)"},
    "업무내용": {"key": "main_content", "desc": "업무 세부 내용을 항목별로 구체적으로 작성"},
    "세부내용": {"key": "main_content", "desc": "세부 내용을 항목별로 구체적으로 작성"},
    "상세내용": {"key": "main_content", "desc": "세부 내용을 항목별로 구체적으로 작성"},
    "진행업무": {"key": "tasks", "desc": '진행 업무 목록 배열. 각 항목은 {"item": "업무명", "assignee": "담당자", "progress": "진행률"} 형태'},
    "업무목록": {"key": "tasks", "desc": '진행 업무 목록 배열. 각 항목은 {"item": "업무명", "assignee": "담당자", "progress": "진행률"} 형태'},
    "성과": {"key": "achievements", "desc": "주요 성과 목록 (배열)"},
    "주요성과": {"key": "achievements", "desc": "주요 성과 목록 (배열)"},
    "이슈": {"key": "issues", "desc": "이슈 및 건의사항 (없으면 빈 문자열)"},
    "이슈사항": {"key": "issues", "desc": "이슈 및 건의사항 (없으면 빈 문자열)"},
    "건의사항": {"key": "issues", "desc": "이슈 및 건의사항 (없으면 빈 문자열)"},
    "향후계획": {"key": "next_plan", "desc": "향후 계획 (구체적으로 작성)"},
    "차주계획": {"key": "next_plan", "desc": "향후 계획 (구체적으로 작성)"},
    "결론": {"key": "conclusion", "desc": "결론 및 종합 의견"},

    # 제안서
    "제출일": {"key": "submit_date", "desc": "제출 날짜 (YYYY-MM-DD, 없으면 오늘 날짜)"},
    "제출처": {"key": "submit_to", "desc": "제출처 (없으면 빈 문자열)"},
    "제안사": {"key": "company", "desc": "제안사 이름 (없으면 빈 문자열)"},
    "회사명": {"key": "company", "desc": "회사 이름 (없으면 빈 문자열)"},
    "담당자": {"key": "manager", "desc": "담당자 이름 (없으면 빈 문자열)"},
    "연락처": {"key": "contact", "desc": "연락처 (없으면 빈 문자열)"},
    "제안배경": {"key": "background", "desc": "제안 배경 (2~3문장)"},
    "배경": {"key": "background", "desc": "배경 (2~3문장)"},
    "제안목적": {"key": "purpose", "desc": "제안 목적 및 필요성 (3~5문장)"},
    "목적": {"key": "purpose", "desc": "목적 및 필요성 (3~5문장)"},
    "제안내용": {"key": "content", "desc": "제안 내용을 항목별로 구체적으로 작성"},
    "추진일정": {"key": "schedule", "desc": '추진 일정 배열. 각 항목은 {"item": "추진항목", "phase1": "1단계", "phase2": "2단계", "phase3": "3단계", "phase4": "4단계"} 형태'},
    "일정": {"key": "schedule", "desc": '추진 일정 배열'},
    "예산": {"key": "budget", "desc": '예산 배열. 각 항목은 {"item": "항목", "quantity": "수량", "unit_price": "단가", "amount": "금액"} 형태'},
    "소요예산": {"key": "budget", "desc": '예산 배열'},
    "기대효과": {"key": "expected_effect", "desc": "기대 효과 (3~5문장)"},
    "리스크": {"key": "risks", "desc": '리스크 목록 배열. 각 항목은 {"description": "설명", "level": "상/중/하", "mitigation": "대응방안"} 형태'},
    "위험요소": {"key": "risks", "desc": '리스크 목록 배열'},
}

# 부분 매칭용: 긴 키부터 정렬 (제안목적 > 목적)
_SORTED_MAPPING_KEYS = sorted(FIELD_MAPPING.keys(), key=len, reverse=True)

# 플레이스홀더 패턴: 이런 텍스트가 있으면 "비어있는 셀"로 간주
_PLACEHOLDER_RE = re.compile(
    r"^[\s☐□✓✔○●◎\d.·\-~:：()（）년월일시분/]*$"
)

# 라벨이 아닌 것: 순수 숫자, 열 헤더 흔한 단어
_NON_LABEL_WORDS = {"구분", "내용", "비고", "항목", "합계", "소계", "총계", "no", "no.", "합", "계"}


def _normalize_label(text: str) -> str:
    """필드 레이블 정규화: 공백/특수문자/번호 접두사 제거"""
    # "1. 제안 목적 및 필요성" → "제안목적및필요성"
    text = re.sub(r"^\d+[\s.·)\-]+", "", text)  # 번호 접두사 제거
    text = re.sub(r"[\s·:：\-_/()（）\[\]【】,，]", "", text)
    return text.strip()


def _is_empty_cell(text: str) -> bool:
    """셀이 비어있거나 플레이스홀더인지 판별"""
    if not text:
        return True
    return bool(_PLACEHOLDER_RE.match(text))


def _is_valid_label(text: str) -> bool:
    """라벨로 유효한 텍스트인지 판별"""
    normalized = _normalize_label(text)
    if not normalized:
        return False
    if len(normalized) > 20:
        return False
    # 순수 숫자
    if re.match(r"^\d+$", normalized):
        return False
    # 한글이나 영문이 최소 1자 이상
    if not re.search(r"[가-힣a-zA-Z]", normalized):
        return False
    # 흔한 열 헤더 (단독으로 쓰일 때만 제외)
    if normalized.lower() in _NON_LABEL_WORDS:
        return False
    return True


def _match_field(label: str, use_mapping: bool = False) -> dict | None:
    """한국어 레이블 → 매핑된 필드 반환. 없으면 커스텀 키 자동 생성.

    Args:
        use_mapping: False면 FIELD_MAPPING을 무시하고 한글 라벨을 키로 사용.
                     단, description은 FIELD_MAPPING에 있으면 가져옴.
    """
    normalized = _normalize_label(label)
    if not normalized:
        return None

    # 라벨 내부 공백 정규화
    # 한글 사이 공백은 제거 ("참  석  인  원" → "참석인원"), 한글-영문 사이는 유지
    clean_label = re.sub(r'(?<=[가-힣])\s+(?=[가-힣])', '', label.strip()).strip()
    clean_label = re.sub(r'\s+', ' ', clean_label).strip()

    # 명시적 제외 항목 체크 (use_mapping 여부와 무관)
    if normalized in FIELD_MAPPING and FIELD_MAPPING[normalized] is None:
        return None

    # FIELD_MAPPING에서 description 가져오기 (키 매핑과 별개)
    desc = f"{clean_label} 내용을 작성"
    desc_source = "fallback"
    if normalized in FIELD_MAPPING and FIELD_MAPPING[normalized]:
        desc = FIELD_MAPPING[normalized]["desc"]
        desc_source = "mapping"
    else:
        normalized_lower = normalized.lower()
        for map_key in _SORTED_MAPPING_KEYS:
            mapping = FIELD_MAPPING[map_key]
            if mapping is None:
                continue
            if map_key in normalized_lower or normalized_lower in map_key:
                desc = mapping["desc"]
                desc_source = "partial"
                break

    if use_mapping:
        # 기존 로직: FIELD_MAPPING에서 영어 키 반환
        if normalized in FIELD_MAPPING and FIELD_MAPPING[normalized]:
            return {"key": FIELD_MAPPING[normalized]["key"], "label": clean_label, "description": desc, "desc_source": desc_source}

        normalized_lower = normalized.lower()
        for map_key in _SORTED_MAPPING_KEYS:
            mapping = FIELD_MAPPING[map_key]
            if mapping is None:
                continue
            if map_key in normalized_lower or normalized_lower in map_key:
                return {"key": mapping["key"], "label": clean_label, "description": desc, "desc_source": desc_source}

    # 한글 라벨을 키로 사용
    custom_key = re.sub(r"[^가-힣a-zA-Z0-9]", "_", normalized)
    custom_key = re.sub(r"_+", "_", custom_key).strip("_")
    if not custom_key:
        return None
    return {
        "key": custom_key,
        "label": clean_label,
        "description": desc,
        "desc_source": desc_source,
    }


def extract_template_fields(file_path: str, use_mapping: bool = False) -> list[dict]:
    """
    DOCX 양식에서 필드 명세를 추출한다.

    전략: 구조 기반 — 테이블 패턴, 헤딩, 본문 패턴을 분석하여
    어떤 양식이든 필드를 자동 탐지.

    Args:
        file_path: DOCX 파일 경로
        use_mapping: True면 FIELD_MAPPING으로 영어 키 변환, False면 한글 키 유지

    Returns:
        필드 목록: [{"key": "date", "label": "일시", "description": "..."}, ...]
    """
    from docx import Document

    doc = Document(file_path)
    fields = []
    seen_keys = set()
    seen_semantic_keys = set()  # FIELD_MAPPING 영어 key 기준 의미 중복 방지

    def _add_field(label: str):
        """라벨 유효성 검사 → 매핑 → 중복 체크 (키 + 의미) → 추가"""
        if not _is_valid_label(label):
            return
        field = _match_field(label, use_mapping=use_mapping)
        if not field or field["key"] in seen_keys:
            return
        # 의미 중복 체크: FIELD_MAPPING에서 같은 영어 key로 매핑되는 필드 방지
        normalized = _normalize_label(label)
        semantic_key = None
        if normalized in FIELD_MAPPING and FIELD_MAPPING[normalized]:
            semantic_key = FIELD_MAPPING[normalized]["key"]
        else:
            for map_key in _SORTED_MAPPING_KEYS:
                mapping = FIELD_MAPPING[map_key]
                if mapping is None:
                    continue
                if map_key in normalized.lower() or normalized.lower() in map_key:
                    semantic_key = mapping["key"]
                    break
        if semantic_key and semantic_key in seen_semantic_keys:
            return  # 의미 중복 — 스킵
        if semantic_key:
            seen_semantic_keys.add(semantic_key)
        seen_keys.add(field["key"])
        fields.append(field)

    # ── 1. 테이블에서 필드 추출 ──
    for table in doc.tables:
        rows = table.rows
        if not rows:
            continue
        num_cols = len(table.columns) if table.columns else 0

        # 1-a. 병합 헤더 탐지: 첫 행의 모든 셀이 같은 텍스트
        if len(rows) >= 2 and num_cols > 1:
            first_cells = [c.text.strip() for c in rows[0].cells]
            unique = set(first_cells)
            if len(unique) == 1 and first_cells[0]:
                _add_field(first_cells[0])

        # 1-b. 열 헤더 행 탐지: 첫 행의 거의 모든 셀에 짧은 텍스트
        #      (예: ['구분', '내용', '비고'] → 컬럼 헤더를 필드로 추출)
        header_row_idx = -1
        if len(rows) >= 2 and num_cols >= 2:
            first_cells = [c.text.strip() for c in rows[0].cells]
            text_cells = [c for c in first_cells if c and len(c) <= 10]
            if len(text_cells) >= num_cols * 0.6:
                header_row_idx = 0
                # 컬럼 헤더 추출은 배열 테이블 감지 후에 처리 (1-b3에서)

        # 1-b2. 배열 테이블 감지
        #   패턴 A: 첫 열에 같은 텍스트 반복 (결정사항|내용|진행일정)
        #   패턴 B: 병합 헤더 + No. 번호열 + 컬럼 헤더 (4.추진일정|No.|항목|1단계|...)
        _array_table_handled = False
        if len(rows) >= 3 and num_cols >= 2:
            first_col_texts = [re.sub(r'\s+', '', rows[ri].cells[0].text.strip()) for ri in range(len(rows))]
            from collections import Counter
            col0_counts = Counter(t for t in first_col_texts if t)

            # 패턴 A: 첫 열에 같은 텍스트 3행 이상 반복
            array_label = None
            array_header_row = None
            if col0_counts:
                most_common_label, count = col0_counts.most_common(1)[0]
                if count >= 3 and most_common_label:
                    for ri in range(len(rows)):
                        t = rows[ri].cells[0].text.strip()
                        if re.sub(r'\s+', '', t) == most_common_label:
                            array_label = t
                            array_header_row = ri
                            break

            # 패턴 B: 병합 헤더(R0) + No./번호 열(R1) + 데이터 행(R2+)
            if not array_label and len(rows) >= 4:
                r0_cells = [c.text.strip() for c in rows[0].cells]
                r1_cells = [c.text.strip() for c in rows[1].cells]
                r0_unique = set(r0_cells)
                r1_c0 = re.sub(r'\s+', '', r1_cells[0]).lower() if r1_cells else ""
                # R0: 병합 헤더 (모든 셀 같은 텍스트) + R1C0: "No." 또는 번호
                if len(r0_unique) == 1 and r0_cells[0] and r1_c0 in ("no", "no.", "번호"):
                    array_label = r0_cells[0]
                    array_header_row = 1  # 컬럼 헤더는 R1

            if array_label:
                # 컬럼 헤더 추출
                sub_keys = []
                if array_header_row is not None:
                    for ci in range(1, num_cols):
                        col_header = re.sub(r'\s+', ' ', rows[array_header_row].cells[ci].text.strip()).strip()
                        if col_header and len(col_header) <= 20 and re.search(r'[가-힣a-zA-Z]', col_header):
                            sub_keys.append(col_header)

                if _is_valid_label(array_label):
                    field = _match_field(array_label, use_mapping=use_mapping)
                    if field:
                        if field["key"] in seen_keys:
                            # 이미 추출된 필드 → sub_keys만 업데이트
                            if sub_keys:
                                for existing in fields:
                                    if existing["key"] == field["key"]:
                                        existing["sub_keys"] = sub_keys
                                        existing["description"] = f"각 항목은 {', '.join(sub_keys)} 필드를 가진 객체 배열"
                                        break
                                _array_table_handled = True
                                logger.info("배열 테이블 감지 (기존 필드 업데이트): %s → sub_keys=%s", array_label, sub_keys)
                        else:
                            if sub_keys:
                                field["description"] = f"각 항목은 {', '.join(sub_keys)} 필드를 가진 객체 배열"
                                field["sub_keys"] = sub_keys
                            else:
                                field["description"] = f"{field.get('label', '')} 목록 (배열)"
                            seen_keys.add(field["key"])
                            fields.append(field)
                            _array_table_handled = True
                            logger.info("배열 테이블 감지: %s → sub_keys=%s", array_label, sub_keys)

        # 1-b3. 배열 테이블이 아닌 header_row의 컬럼 헤더를 필드로 추출
        if header_row_idx >= 0 and not _array_table_handled:
            first_cells = [c.text.strip() for c in rows[header_row_idx].cells]
            for cell_text in first_cells:
                if cell_text:
                    _add_field(cell_text)

        for ri, row in enumerate(rows):
            cells = [c.text.strip() for c in row.cells]

            # 1-c. 1열 섹션 테이블: [헤더] / [빈칸]
            if num_cols == 1 and ri == 0 and cells[0]:
                _add_field(cells[0])
                continue

            # 1-c2. 다열 행에서 모든 셀이 같은 텍스트 = 병합 라벨 (예: "특이사항/건의 사항")
            if num_cols >= 2 and ri > 0:
                unique_texts = set(c for c in cells if c)
                if len(unique_texts) == 1:
                    _add_field(unique_texts.pop())
                    continue

            # 열 헤더 행이면 스킵
            if ri == header_row_idx:
                continue

            # 배열 테이블로 이미 처리된 경우 해당 테이블의 반복 라벨 스킵
            if _array_table_handled:
                cell0_normalized = re.sub(r'\s+', '', cells[0]) if cells[0] else ''
                if cell0_normalized and col0_counts.get(cell0_normalized, 0) >= 3:
                    continue

            # 1-d. 다열 테이블: 모든 열 스캔
            #      [라벨 | 빈칸] 또는 [라벨 | 플레이스홀더] 패턴
            ci = 0
            while ci < len(cells):
                cell_text = cells[ci]
                # 다음 셀이 비어있거나 플레이스홀더면 → 현재 셀은 라벨
                next_empty = (ci + 1 < len(cells) and _is_empty_cell(cells[ci + 1]))
                # 마지막 열이 아닌 경우에만 라벨 후보
                if cell_text and next_empty:
                    _add_field(cell_text)
                    ci += 2  # 라벨 + 값 셀 건너뛰기
                elif cell_text and ci + 1 < len(cells):
                    # 다음 셀에 텍스트가 있어도, 현재 셀이 짧고 다음이 길면 라벨일 수 있음
                    next_text = cells[ci + 1]
                    if (len(_normalize_label(cell_text)) <= 10
                            and len(next_text) > len(cell_text) * 2
                            and _is_valid_label(cell_text)):
                        _add_field(cell_text)
                        ci += 2
                    else:
                        ci += 1
                else:
                    ci += 1

    # ── 2. 헤딩에서 필드 추출 ──
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if not text or len(text) > 30:
            continue

        if "Heading" in style_name or style_name.startswith("제목"):
            _add_field(text)

    # ── 3. 본문에서 "항목명:" 패턴 추출 ──
    for para in doc.paragraphs:
        text = para.text.strip()
        match = re.match(r"^([가-힣a-zA-Z][\w\s]{0,15})\s*[:：]\s*$", text)
        if match:
            _add_field(match.group(1))

    # ── 4. 메타데이터 자동 부여 (v3) ──
    fields = [_infer_field_meta(f) for f in fields]

    logger.info("양식 필드 추출 완료: %s → %d개 필드 (meta=%d, body=%d)",
                file_path, len(fields),
                sum(1 for f in fields if f["group"] == "meta"),
                sum(1 for f in fields if f["group"] == "body"))
    return fields


def _infer_field_meta(field: dict) -> dict:
    """필드에 group/type/fill 메타데이터를 자동 부여한다.

    group: "meta" (기본 정보, 짧은 텍스트) / "body" (본문, 서술형)
    type: "date" / "text" / "textarea" / "array"
    fill: "extract" (사용자 텍스트에서 추출) / "generate" (AI가 작성)
    """
    label = field.get("label", "")
    desc = field.get("description", "")
    normalized = _normalize_label(label)

    # sub_keys가 이미 있으면 배열 확정
    if field.get("sub_keys"):
        return {**field, "group": "body", "type": "array", "fill": "generate"}

    # ── type 추론 ──
    desc_lower = desc.lower()
    label_and_desc = f"{normalized} {desc_lower}"

    # date: description/라벨에 날짜 관련 키워드
    _date_kw = ("날짜", "일자", "일시", "yyyy", "yy-mm", "date")
    # array: 배열/목록/각 항목은/객체 배열
    _array_kw = ("배열", "목록", "각 항목은", "객체 배열")
    # textarea: 서술형 필드
    _textarea_kw = ("문장", "구체적으로", "상세하게", "항목별", "기술")

    if any(kw in label_and_desc for kw in _date_kw):
        ftype = "date"
    elif any(kw in desc_lower for kw in _array_kw):
        ftype = "array"
    elif any(kw in desc_lower for kw in _textarea_kw):
        ftype = "textarea"
    else:
        ftype = "text"

    # ── group 추론 ──
    group = "meta"  # 기본값

    # 1차: 번호 접두사 ("1. 보고 개요" 등) → 거의 확실히 body
    if re.match(r"^\d+[\s.]", label):
        group = "body"
    # 2차: type이 textarea/array → body
    elif ftype in ("textarea", "array"):
        group = "body"
    # 3차: description에 서술형 키워드
    elif any(kw in desc_lower for kw in ("문장", "구체적", "상세", "항목별", "배열", "목록")):
        group = "body"
    # 4차: 라벨 글자 수 6자 이상 + description이 "내용을 작성" 이 아닌 경우 → body 후보
    # (짧은 라벨은 meta: "부서", "직급", 긴 라벨은 body: "제안목적및필요성", "이슈및건의사항")
    elif len(normalized) >= 6 and "없으면" not in desc_lower:
        group = "body"
    # 5차: 라벨에 body성 키워드 포함 (fallback)
    else:
        _body_kw = ("개요", "내용", "분석", "계획", "효과", "배경", "결론", "목적", "필요성",
                     "사양", "방법", "절차", "결과", "평가", "소감", "의견", "요약", "현황")
        if any(kw in normalized for kw in _body_kw):
            group = "body"

    # ── fill 추론 ──
    fill = "generate" if group == "body" else "extract"

    return {**field, "group": group, "type": ftype, "fill": fill}


async def enrich_descriptions(fields: list[dict], doc_type: str = "문서") -> list[dict]:
    """FIELD_MAPPING에 없는 필드(desc_source=fallback)의 description을 LLM으로 보강한다.

    업로드 시 1회만 호출. 실패하면 기존 description 유지.
    보강 후 type/group/fill 메타데이터를 재추론한다.
    """
    fallback_fields = [f for f in fields if f.get("desc_source") == "fallback"]
    if not fallback_fields:
        return fields  # 보강할 필드 없음

    # 이미 파악된 필드를 컨텍스트로 제공
    known_context = ", ".join(
        f'{f["label"]}({f["description"]})' for f in fields if f.get("desc_source") != "fallback"
    )
    unknown_labels = "\n".join(f'- {f["label"]}' for f in fallback_fields)

    prompt = (
        f"이 문서는 \"{doc_type}\" 양식입니다.\n"
        f"이미 파악된 필드: {known_context}\n\n"
        f"아래 필드가 어떤 내용을 담는지 한 줄(20자 이내)로 설명해주세요.\n"
        f"구체적 방법론이나 프레임워크는 언급하지 마세요.\n"
        f"JSON 형식으로 출력: {{\"필드라벨\": \"설명\", ...}}\n\n"
        f"{unknown_labels}"
    )

    try:
        from ai.agents.document._common import _call_llm
        result_str = await _call_llm(
            "당신은 문서 양식 분석 전문가입니다. 필드 설명을 간결하게 작성하세요.",
            prompt,
            json_mode=True,
            task="extract",
            temperature=0,
        )
        enriched = json.loads(result_str)

        # fallback 필드의 description 교체
        label_to_desc = {k.strip(): v.strip() for k, v in enriched.items() if v}
        for f in fields:
            if f.get("desc_source") == "fallback" and f["label"] in label_to_desc:
                f["description"] = label_to_desc[f["label"]]
                f["desc_source"] = "llm"
        logger.info("description 보강 완료: %d개 중 %d개 enriched", len(fallback_fields), len(label_to_desc))
    except Exception as e:
        logger.warning("description 보강 실패 (fallback 유지): %s", e)

    # description 변경 후 메타데이터 재추론
    fields = [_infer_field_meta(f) for f in fields]
    return fields


def fields_to_parsed_structure(fields: list[dict]) -> str:
    """추출된 필드를 parsed_structure JSON 문자열로 변환 (DB 저장용)"""
    return json.dumps(fields, ensure_ascii=False)


def fields_to_prompt(fields: list[dict]) -> str:
    """추출된 필드를 [필드 명세] 프롬프트 문자열로 변환 (sLLM 호출용)"""
    lines = []
    for f in fields:
        desc = f.get('description') or f.get('label', f['key'])
        lines.append(f"- {f['key']}: {desc}")
    return "\n".join(lines)
