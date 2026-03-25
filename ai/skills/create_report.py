from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ai.skills._docx_styles import (
    BLUE_HEADER, BLUE_LIGHT, BLUE_ALT, NAVY_RGB, WHITE_RGB,
    set_shading, set_valign, set_row_height,
    style_section_header, style_label_cell, style_value_cell,
    inject_cell_text as _inject, add_title_line,
)

# 하위 호환 별칭
_BLUE_HEADER = BLUE_HEADER
_BLUE_LIGHT = BLUE_LIGHT
_BLUE_ALT = BLUE_ALT
_NAVY_RGB = NAVY_RGB
_WHITE_RGB = WHITE_RGB
_set_shading = set_shading
_set_valign = set_valign
_add_title_line = add_title_line


def create_report(output_path: str = "tests/업무보고서_생성.docx", data: dict = None):
    """
    업무보고서 DOCX 생성

    data 필드:
        title        : 보고서 제목
        author       : 작성자
        date         : 작성일 (예: 2026-02-23)
        department   : 부서
        position     : 직급
        report_to    : 보고 대상
        report_type  : 보고 유형 (일일/주간/월간/수시)
        overview     : 보고 개요 (str)
        main_content : 주요 내용 (str)
        tasks        : 진행 현황 list[{item, assignee, progress, start_date, end_date}]
        issues       : 이슈 및 건의 사항 (str)
        next_plan    : 향후 계획 (str)
        attachments  : 첨부 자료 (str)
        notes        : 비고 (str)
    """
    doc = Document()

    # ── 페이지 여백 ──
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 제목 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("업  무  보  고  서")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = _NAVY_RGB

    # 제목 아래 파란 구분선
    _add_title_line(doc)

    # ── 표0: 기본 정보 (4행 4열) ──
    t0 = doc.add_table(rows=4, cols=4)
    t0.style = "Table Grid"
    for row in t0.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(5.5)
        row.cells[2].width = Cm(2.5)
        row.cells[3].width = Cm(5.5)

    # Row0: 보고서 제목 (값 셀 3개 병합)
    t0.rows[0].cells[1].merge(t0.rows[0].cells[3])
    style_label_cell(t0.rows[0].cells[0], "보고서 제목")
    style_value_cell(t0.rows[0].cells[1])
    set_row_height(t0.rows[0], 1.0)

    # Row1: 작성자 / 작성일
    style_label_cell(t0.rows[1].cells[0], "작성자")
    style_value_cell(t0.rows[1].cells[1])
    style_label_cell(t0.rows[1].cells[2], "작성일")
    style_value_cell(t0.rows[1].cells[3], "2026년    월    일")

    # Row2: 부서 / 직급
    style_label_cell(t0.rows[2].cells[0], "부서")
    style_value_cell(t0.rows[2].cells[1])
    style_label_cell(t0.rows[2].cells[2], "직급")
    style_value_cell(t0.rows[2].cells[3])

    # Row3: 보고 대상 / 보고 유형
    style_label_cell(t0.rows[3].cells[0], "보고 대상")
    style_value_cell(t0.rows[3].cells[1])
    style_label_cell(t0.rows[3].cells[2], "보고 유형")
    style_value_cell(t0.rows[3].cells[3], "☐ 일일  ☐ 주간  ☐ 월간  ☐ 수시")

    doc.add_paragraph()

    # ── 표1: 보고 개요 ──
    t1 = doc.add_table(rows=2, cols=1)
    t1.style = "Table Grid"
    style_section_header(t1.rows[0].cells[0], "1. 보고 개요")
    style_value_cell(t1.rows[1].cells[0])
    set_row_height(t1.rows[1], 3.0)

    doc.add_paragraph()

    # ── 표2: 주요 내용 ──
    t2 = doc.add_table(rows=2, cols=1)
    t2.style = "Table Grid"
    style_section_header(t2.rows[0].cells[0], "2. 주요 내용")
    style_value_cell(t2.rows[1].cells[0])
    set_row_height(t2.rows[1], 5.0)

    doc.add_paragraph()

    # ── 표3: 진행 현황 (동적 행: 섹션헤더 + 컬럼헤더 + 데이터 N행) ──
    tasks_raw = data.get("tasks", []) if data else []
    if isinstance(tasks_raw, dict):
        _tasks_list = list(tasks_raw.values())
    elif isinstance(tasks_raw, list):
        _tasks_list = tasks_raw
    else:
        _tasks_list = []
    _task_data_rows = max(len(_tasks_list), 3)
    t3 = doc.add_table(rows=2 + 1, cols=6)  # 헤더2행 + 첫 데이터행
    t3.style = "Table Grid"

    for i in range(1, 6):
        t3.rows[0].cells[0].merge(t3.rows[0].cells[i])
    style_section_header(t3.rows[0].cells[0], "3. 진행 현황")

    for i, h in enumerate(["No.", "업무 항목", "담당자", "진행률", "시작일", "완료 예정일"]):
        style_label_cell(t3.rows[1].cells[i], h)

    # 필요한 만큼 데이터 행 추가
    for _ in range(_task_data_rows - 1):
        t3.add_row()

    for r in range(2, 2 + _task_data_rows):
        row_bg = _BLUE_ALT if r % 2 == 0 else "FFFFFF"
        for c in range(6):
            _set_shading(t3.rows[r].cells[c], row_bg)
        style_value_cell(t3.rows[r].cells[0], str(r - 1))
        for c in range(1, 6):
            style_value_cell(t3.rows[r].cells[c])
        set_row_height(t3.rows[r], 1.0)

    doc.add_page_break()

    # ── 표4: 이슈 및 건의 사항 ──
    t4 = doc.add_table(rows=2, cols=1)
    t4.style = "Table Grid"
    style_section_header(t4.rows[0].cells[0], "4. 이슈 및 건의 사항")
    style_value_cell(t4.rows[1].cells[0])
    set_row_height(t4.rows[1], 3.0)

    doc.add_paragraph()

    # ── 표5: 향후 계획 ──
    t5 = doc.add_table(rows=2, cols=1)
    t5.style = "Table Grid"
    style_section_header(t5.rows[0].cells[0], "5. 향후 계획")
    style_value_cell(t5.rows[1].cells[0])
    set_row_height(t5.rows[1], 3.0)

    doc.add_paragraph()

    # ── 표6: 첨부 자료 / 비고 ──
    t6 = doc.add_table(rows=2, cols=2)
    t6.style = "Table Grid"
    style_label_cell(t6.rows[0].cells[0], "첨부 자료")
    style_value_cell(t6.rows[0].cells[1])
    style_label_cell(t6.rows[1].cells[0], "비고")
    style_value_cell(t6.rows[1].cells[1])

    doc.add_paragraph()

    # ── 표7: 결재란 ──
    t7 = doc.add_table(rows=2, cols=3)
    t7.style = "Table Grid"
    for i, h in enumerate(["작성", "검토", "승인"]):
        style_label_cell(t7.rows[0].cells[i], h)
        style_value_cell(t7.rows[1].cells[i])
        set_row_height(t7.rows[1], 2.0)

    # ── data 주입 ──
    if data:
        _inject(t0.rows[0].cells[1], data.get("title", ""))
        _inject(t0.rows[1].cells[1], data.get("author", ""))
        _inject(t0.rows[1].cells[3], data.get("date", ""))
        _inject(t0.rows[2].cells[1], data.get("department", ""))
        _inject(t0.rows[2].cells[3], data.get("position", ""))
        _inject(t0.rows[3].cells[1], data.get("report_to", ""))

        type_map = {
            "일일": "☑ 일일  ☐ 주간  ☐ 월간  ☐ 수시",
            "주간": "☐ 일일  ☑ 주간  ☐ 월간  ☐ 수시",
            "월간": "☐ 일일  ☐ 주간  ☑ 월간  ☐ 수시",
            "수시": "☐ 일일  ☐ 주간  ☐ 월간  ☑ 수시",
        }
        _inject(t0.rows[3].cells[3], type_map.get(data.get("report_type", ""), "☐ 일일  ☐ 주간  ☐ 월간  ☐ 수시"))

        _inject(t1.rows[1].cells[0], data.get("overview", ""))
        _inject(t2.rows[1].cells[0], data.get("main_content", ""))

        tasks = _tasks_list
        for r in range(2, 2 + _task_data_rows):
            task = tasks[r - 2] if r - 2 < len(tasks) else {}
            if not isinstance(task, dict):
                task = {}
            _inject(t3.rows[r].cells[1], task.get("item", ""))
            _inject(t3.rows[r].cells[2], task.get("assignee", ""))
            _inject(t3.rows[r].cells[3], task.get("progress", ""))
            _inject(t3.rows[r].cells[4], task.get("start_date", ""))
            _inject(t3.rows[r].cells[5], task.get("end_date", ""))

        _inject(t4.rows[1].cells[0], data.get("issues", ""))
        _inject(t5.rows[1].cells[0], data.get("next_plan", ""))
        _inject(t6.rows[0].cells[1], data.get("attachments", ""))
        _inject(t6.rows[1].cells[1], data.get("notes", ""))

    doc.save(output_path)
    print(f"업무보고서 생성 완료: {output_path}")


if __name__ == "__main__":
    create_report()
