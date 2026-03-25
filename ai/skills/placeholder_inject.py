"""
플레이스홀더 자동 삽입 — 원본 DOCX 빈 셀에 {{key}} 마커 삽입

업로드 시 1회 실행. 이후 docxtpl로 단순 치환하여 문서 생성.
"""
import copy
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

from ai.document_parser.template_extractor import (
    _is_empty_cell,
    _is_valid_label,
    _normalize_label,
)


def inject_placeholders(source_path: str, output_path: str, fields: list[dict]) -> dict:
    """원본 DOCX의 빈 셀에 {{key}} 플레이스홀더를 삽입한다.

    Args:
        source_path: 원본 DOCX 경로
        output_path: 플레이스홀더 삽입된 DOCX 저장 경로
        fields: parsed_structure 필드 목록 [{"key", "label", "sub_keys"?}, ...]

    Returns:
        {"success": bool, "injected_count": int, "injected_fields": [...]}
    """
    doc = Document(source_path)
    injected = []

    # 라벨 → 필드 매핑 구축 (정규화된 라벨로 검색)
    label_to_field = {}
    for f in fields:
        label = f.get("label", "")
        if label:
            norm = _normalize_label(label)
            if norm:
                label_to_field[norm] = f

    # ── 테이블 스캔: 라벨-빈칸 패턴에 {{key}} 삽입 ──
    for table in doc.tables:
        rows = table.rows
        num_cols = len(table.columns) if table.columns else 0

        # 배열 테이블 감지 (병합 헤더 + No. 패턴)
        array_field = None
        array_header_ri = None
        array_data_ri = None

        if len(rows) >= 3 and num_cols >= 2:
            r0_cells = [c.text.strip() for c in rows[0].cells]
            r0_unique = set(r0_cells)
            r1_c0 = re.sub(r'\s+', '', rows[1].cells[0].text.strip()).lower() if len(rows) > 1 else ""

            # 패턴 A: 병합 헤더 + No. 행
            if len(r0_unique) == 1 and r0_cells[0] and r1_c0 in ("no", "no.", "번호"):
                header_label = _normalize_label(r0_cells[0])
                if header_label in label_to_field:
                    array_field = label_to_field[header_label]
                    array_header_ri = 1
                    array_data_ri = 2

            # 패턴 B: C0에 같은 라벨 반복 + 컬럼 헤더 행
            # 예: 결정사항|내용|진행일정 / 결정사항|(빈)|(빈) / ...
            if not array_field and len(rows) >= 3 and num_cols >= 2:
                from collections import Counter
                c0_texts = [_normalize_label(rows[ri].cells[0].text.strip()) for ri in range(len(rows))]
                c0_counts = Counter(t for t in c0_texts if t)
                if c0_counts:
                    most_common, count = c0_counts.most_common(1)[0]
                    if count >= 3 and most_common in label_to_field:
                        # 첫 번째 반복 행 = 헤더 (컬럼명이 있는 행)
                        for ri in range(len(rows)):
                            if c0_texts[ri] == most_common:
                                # 이 행의 다른 셀에 컬럼 헤더가 있으면 헤더 행
                                other_texts = [rows[ri].cells[ci].text.strip()
                                              for ci in range(1, num_cols)
                                              if _is_valid_label(rows[ri].cells[ci].text.strip())]
                                if other_texts:
                                    array_field = label_to_field[most_common]
                                    array_header_ri = ri
                                    array_data_ri = ri + 1
                                    break

        if array_field and array_header_ri is not None:
            # ── 배열 테이블: {%tr for %} 루프 삽입 ──
            key = array_field["key"]
            sub_keys = array_field.get("sub_keys", [])

            if array_data_ri < len(rows) and sub_keys:
                data_row = rows[array_data_ri]

                # sub_key를 안전한 변수명으로 변환
                safe_sks = []
                for sk in sub_keys:
                    safe_sk = re.sub(r'[^가-힣a-zA-Z0-9]', '', sk)
                    # Jinja2 변수명은 숫자로 시작 불가 → 접두사 추가
                    if safe_sk and safe_sk[0].isdigit():
                        safe_sk = 'f_' + safe_sk
                    safe_sks.append(safe_sk)

                # C0가 No./번호 열인지 확인
                header_c0 = ""
                if array_header_ri is not None:
                    header_c0 = _normalize_label(rows[array_header_ri].cells[0].text.strip()).lower()
                is_no_col = header_c0 in ("no", "no.", "번호", "#", "")

                # 데이터 행 셀에 {{item.sub_key}} 삽입
                if is_no_col:
                    # 패턴 A: No. 열 있음 → C0에 번호, C1~에 sub_key
                    data_row.cells[0].text = "{{loop.index}}"
                    for ci in range(1, num_cols):
                        if ci < len(data_row.cells) and ci - 1 < len(safe_sks):
                            data_row.cells[ci].text = "{{item." + safe_sks[ci - 1] + "}}"
                else:
                    # 패턴 B: C0가 반복 라벨(세로병합) → C0 건드리지 않음, C1~에 sub_key
                    data_row.cells[0].text = ""  # C0 비움 (세로병합 라벨)
                    for ci in range(1, num_cols):
                        if ci < len(data_row.cells) and ci - 1 < len(safe_sks):
                            data_row.cells[ci].text = "{{item." + safe_sks[ci - 1] + "}}"

                # XML 레벨에서 {%tr for/endfor%} 별도 행 삽입
                data_tr = data_row._tr
                tbl_elem = table._tbl

                # for 행 생성 (데이터 행 복제 → 텍스트를 for 태그로)
                for_tr = copy.deepcopy(data_tr)
                for tc in for_tr.findall(qn('w:tc')):
                    for p in tc.findall(qn('w:p')):
                        for r in p.findall(qn('w:r')):
                            for t in r.findall(qn('w:t')):
                                t.text = ''
                first_tc = for_tr.findall(qn('w:tc'))[0]
                first_p = first_tc.findall(qn('w:p'))[0]
                first_r = first_p.findall(qn('w:r'))
                if first_r and first_r[0].findall(qn('w:t')):
                    first_r[0].findall(qn('w:t'))[0].text = "{%tr for item in " + key + " %}"
                else:
                    # run이 없으면 새로 생성
                    from lxml import etree
                    new_r = etree.SubElement(first_p, qn('w:r'))
                    new_t = etree.SubElement(new_r, qn('w:t'))
                    new_t.text = "{%tr for item in " + key + " %}"

                # endfor 행 생성 (원본 데이터 행에서 복제, 텍스트 전부 비운 후 endfor만)
                from lxml import etree
                end_tr = copy.deepcopy(data_tr)
                for tc in end_tr.findall(qn('w:tc')):
                    for p in tc.findall(qn('w:p')):
                        for r in p.findall(qn('w:r')):
                            for t in r.findall(qn('w:t')):
                                t.text = ''
                # 첫 셀에 endfor 태그
                end_first_tc = end_tr.findall(qn('w:tc'))[0]
                end_first_p = end_first_tc.findall(qn('w:p'))[0]
                end_first_r = end_first_p.findall(qn('w:r'))
                if end_first_r and end_first_r[0].findall(qn('w:t')):
                    end_first_r[0].findall(qn('w:t'))[0].text = "{%tr endfor %}"
                else:
                    new_r = etree.SubElement(end_first_p, qn('w:r'))
                    new_t = etree.SubElement(new_r, qn('w:t'))
                    new_t.text = "{%tr endfor %}"

                # 나머지 빈 데이터 행 먼저 제거 (for/endfor 삽입 전에!)
                rows_to_remove = []
                for ri in range(array_data_ri + 1, len(rows)):
                    row_cells_text = [c.text.strip() for c in rows[ri].cells]
                    has_content = any(t and not _is_empty_cell(t) for t in row_cells_text)
                    if not has_content:
                        rows_to_remove.append(rows[ri]._tr)
                for tr_remove in rows_to_remove:
                    tr_remove.getparent().remove(tr_remove)

                # 삽입: for행 → 데이터행 → endfor행
                data_tr.addprevious(for_tr)
                data_tr.addnext(end_tr)

                injected.append(key)
                print(f"[placeholder] 배열 루프 삽입: {key} ({len(safe_sks)}개 sub_keys)")
            continue

        # ── 일반 필드: 라벨-빈칸 패턴 ──
        for ri, row in enumerate(rows):
            cells = row.cells
            ci = 0
            seen_tcs = set()
            while ci < len(cells):
                tc_id = id(cells[ci]._tc)
                if tc_id in seen_tcs:
                    ci += 1
                    continue
                seen_tcs.add(tc_id)

                cell_text = cells[ci].text.strip()
                # 라벨 + 옆 빈칸 패턴
                if ci + 1 < len(cells):
                    next_tc_id = id(cells[ci + 1]._tc)
                    next_text = cells[ci + 1].text.strip()
                    if (cell_text and _is_valid_label(cell_text)
                            and _is_empty_cell(next_text)
                            and next_tc_id not in seen_tcs):
                        norm = _normalize_label(cell_text)
                        if norm in label_to_field:
                            field = label_to_field[norm]
                            key = field["key"]
                            if key not in injected:
                                # 빈 셀에 {{key}} 삽입
                                target_cell = cells[ci + 1]
                                target_cell.text = ""
                                p = target_cell.paragraphs[0]
                                p.clear()
                                p.add_run("{{" + key + "}}")
                                injected.append(key)
                                print(f"[placeholder] {{{{{{key}}}}}} 삽입: R{ri}C{ci+1} (라벨: {cell_text})")
                        seen_tcs.add(next_tc_id)
                        ci += 2
                        continue

                # 1열 섹션 테이블 (헤더 + 내용)
                if num_cols == 1 and ri == 0 and cell_text:
                    norm = _normalize_label(cell_text)
                    if norm in label_to_field and len(rows) >= 2:
                        field = label_to_field[norm]
                        key = field["key"]
                        if key not in injected:
                            value_cell = rows[1].cells[0]
                            if _is_empty_cell(value_cell.text.strip()):
                                value_cell.text = ""
                                p = value_cell.paragraphs[0]
                                p.clear()
                                p.add_run("{{" + key + "}}")
                                injected.append(key)
                                print(f"[placeholder] {{{{{{key}}}}}} 삽입: 섹션 T?R1C0 (헤더: {cell_text})")

                ci += 1

        # ── 패턴 2: 라벨 위 + 값 아래 (세로 패턴) ──
        # 예: R0=[작성 | 검토 | 승인], R1=[(빈)|(빈)|(빈)]
        #     R0=[업무 내용 | 비고], R1=[(빈)|(빈)]
        if len(rows) >= 2:
            for ri in range(len(rows) - 1):
                row_cells_text = [c.text.strip() for c in rows[ri].cells]
                next_row_cells = rows[ri + 1].cells

                # 이 행의 모든 셀이 라벨이고, 아래 행이 비어있으면
                seen_tcs_v = set()
                for ci_v in range(len(rows[ri].cells)):
                    tc_id_v = id(rows[ri].cells[ci_v]._tc)
                    if tc_id_v in seen_tcs_v:
                        continue
                    seen_tcs_v.add(tc_id_v)

                    label_text = rows[ri].cells[ci_v].text.strip()
                    if not label_text or not _is_valid_label(label_text):
                        continue

                    # 아래 셀이 비어있는지 확인
                    if ci_v < len(next_row_cells):
                        below_text = next_row_cells[ci_v].text.strip()
                        if _is_empty_cell(below_text):
                            norm = _normalize_label(label_text)
                            if norm in label_to_field:
                                field = label_to_field[norm]
                                key = field["key"]
                                if key not in injected:
                                    target_cell = next_row_cells[ci_v]
                                    target_cell.text = ""
                                    p = target_cell.paragraphs[0]
                                    p.clear()
                                    p.add_run("{{" + key + "}}")
                                    injected.append(key)
                                    print(f"[placeholder] {{{{{{key}}}}}} 삽입: T?R{ri+1}C{ci_v} (위 라벨: {label_text})")

    doc.save(output_path)
    print(f"[placeholder] 완료: {len(injected)}개 필드 삽입 → {output_path}")

    return {
        "success": len(injected) > 0,
        "injected_count": len(injected),
        "injected_fields": injected,
    }
