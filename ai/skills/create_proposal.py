from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 모던 프리미엄 다크 슬레이트 테마 ──
_BLUE_HEADER = "1E293B"   # 섹션 헤더 배경 (다크 네이비/슬레이트)
_BLUE_LIGHT  = "F1F5F9"   # 라벨 셀 배경 (밝고 연한 그레이 블루)
_BLUE_ALT    = "F8FAFC"   # 테이블 짝수 행 배경 (백색에 가까운 블루)
_NAVY_RGB    = RGBColor(0x1E, 0x29, 0x3B)
_WHITE_RGB   = RGBColor(0xFF, 0xFF, 0xFF)


def _set_shading(cell, fill_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)


def _set_valign(cell, align: str = "center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def set_row_height(row, height_cm: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_cm * 567)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def style_section_header(cell, text: str):
    """섹션 제목 셀: 파란 배경 + 흰 굵은 글씨"""
    _set_shading(cell, _BLUE_HEADER)
    _set_valign(cell)
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = _WHITE_RGB
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def style_label_cell(cell, text: str):
    """라벨 셀: 연한 파란 배경 + 굵은 글씨 + 가운데 정렬"""
    _set_shading(cell, _BLUE_LIGHT)
    _set_valign(cell)
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def style_value_cell(cell, text: str = ""):
    """값 셀: 흰 배경 + 기본 글씨"""
    _set_valign(cell)
    para = cell.paragraphs[0]
    para.clear()
    if text:
        run = para.add_run(text)
        run.font.size = Pt(10)


def _inject(cell, text: str):
    """값 셀에 데이터 주입"""
    cell.text = str(text) if text else ""
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].font.size = Pt(10)


def _add_title_line(doc):
    """제목 아래 파란 구분선 문단 추가"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _BLUE_HEADER)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def create_proposal(output_path: str = "tests/제안서_생성.docx", data: dict = None):
    """
    제안서 DOCX 생성

    data 필드:
        title           : 제안서 제목 (표지 부제목)
        submit_date     : 제출일 (예: 2026-02-23)
        submit_to       : 제출처
        company         : 제안사
        manager         : 담당자
        contact         : 연락처
        proposal_name   : 제안명
        background      : 제안 배경 (str)
        proposal_date   : 제안 일자
        period          : 제안 기간 (예: 2026년 3월 ~ 6월)
        proposer        : 제안사 (본문 기본정보)
        manager_contact : 담당자 / 연락처 (본문 기본정보)
        purpose         : 제안 목적 및 필요성 (str)
        analysis        : 현황 분석 (str)
        content         : 제안 내용 (str)
        schedule        : 추진 일정 list[{item, phase1, phase2, phase3, phase4}]
        budget          : 소요 예산 list[{item, quantity, unit_price, amount}]
        budget_total    : 예산 합계 (str)
        expected_effect : 기대 효과 (str)
        attachments     : 첨부 자료 (str)
        notes           : 비고 (str)
    """
    doc = Document()

    # ── 페이지 여백 ──
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 표지 ──
    for _ in range(6):
        doc.add_paragraph()

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ct_run = cover_title.add_run("제  안  서")
    ct_run.font.size = Pt(36)
    ct_run.font.bold = True
    ct_run.font.color.rgb = _NAVY_RGB

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st_run = subtitle.add_run(data.get("title", "[제안서 제목을 입력하세요]") if data else "[제안서 제목을 입력하세요]")
    st_run.font.size = Pt(16)
    st_run.font.color.rgb = _NAVY_RGB

    for _ in range(8):
        doc.add_paragraph()

    # ── 표0: 표지 정보 (5행 2열) ──
    t0 = doc.add_table(rows=5, cols=2)
    t0.style = "Table Grid"
    for row in t0.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(12.5)

    for ri, (label, placeholder) in enumerate([
        ("제출일", "2026년    월    일"),
        ("제출처", ""),
        ("제안사", ""),
        ("담당자", ""),
        ("연락처", ""),
    ]):
        style_label_cell(t0.rows[ri].cells[0], label)
        style_value_cell(t0.rows[ri].cells[1], placeholder)
        set_row_height(t0.rows[ri], 0.9)

    # ── 페이지 구분 ──
    doc.add_page_break()

    # ── 본문 제목 ──
    body_title = doc.add_paragraph()
    body_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bt_run = body_title.add_run("제  안  서")
    bt_run.font.size = Pt(22)
    bt_run.font.bold = True
    bt_run.font.color.rgb = _NAVY_RGB

    # 본문 제목 아래 파란 구분선
    _add_title_line(doc)

    # ── 표1: 기본 정보 (4행 4열) ──
    t1 = doc.add_table(rows=4, cols=4)
    t1.style = "Table Grid"
    for row in t1.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(5.5)
        row.cells[2].width = Cm(2.5)
        row.cells[3].width = Cm(5.5)

    # Row0: 제안명 (값 셀 3개 병합)
    t1.rows[0].cells[1].merge(t1.rows[0].cells[3])
    style_label_cell(t1.rows[0].cells[0], "제안명")
    style_value_cell(t1.rows[0].cells[1])
    set_row_height(t1.rows[0], 1.0)

    # Row1: 제안 배경 (값 셀 3개 병합)
    t1.rows[1].cells[1].merge(t1.rows[1].cells[3])
    style_label_cell(t1.rows[1].cells[0], "제안 배경")
    style_value_cell(t1.rows[1].cells[1])
    set_row_height(t1.rows[1], 1.2)

    # Row2: 제안 일자 / 제안 기간
    style_label_cell(t1.rows[2].cells[0], "제안 일자")
    style_value_cell(t1.rows[2].cells[1], "2026년    월    일")
    style_label_cell(t1.rows[2].cells[2], "제안 기간")
    style_value_cell(t1.rows[2].cells[3], "2026년  월  일 ~ 2026년  월  일")

    # Row3: 제안사 / 담당자·연락처
    style_label_cell(t1.rows[3].cells[0], "제안사")
    style_value_cell(t1.rows[3].cells[1])
    style_label_cell(t1.rows[3].cells[2], "담당자 / 연락처")
    style_value_cell(t1.rows[3].cells[3])

    doc.add_paragraph()

    # ── 표2: 1. 제안 목적 및 필요성 ──
    t2 = doc.add_table(rows=2, cols=1)
    t2.style = "Table Grid"
    style_section_header(t2.rows[0].cells[0], "1. 제안 목적 및 필요성")
    style_value_cell(t2.rows[1].cells[0])
    set_row_height(t2.rows[1], 4.0)

    doc.add_paragraph()

    # ── 표3: 2. 현황 분석 ──
    t3 = doc.add_table(rows=2, cols=1)
    t3.style = "Table Grid"
    style_section_header(t3.rows[0].cells[0], "2. 현황 분석")
    style_value_cell(t3.rows[1].cells[0])
    set_row_height(t3.rows[1], 4.0)

    doc.add_paragraph()

    # ── 표4: 3. 제안 내용 ──
    t4 = doc.add_table(rows=2, cols=1)
    t4.style = "Table Grid"
    style_section_header(t4.rows[0].cells[0], "3. 제안 내용")
    style_value_cell(t4.rows[1].cells[0])
    set_row_height(t4.rows[1], 5.0)

    doc.add_paragraph()

    # ── 표5: 4. 추진 일정 (동적 행: 섹션헤더 + 컬럼헤더 + 데이터 N행) ──
    schedule_raw = data.get("schedule", []) if data else []
    if isinstance(schedule_raw, dict):
        _sched_list = list(schedule_raw.values())
    elif isinstance(schedule_raw, list):
        _sched_list = schedule_raw
    else:
        _sched_list = []
    _sched_data_rows = max(len(_sched_list), 3)
    t5 = doc.add_table(rows=2 + 1, cols=6)  # 헤더2행 + 첫 데이터행
    t5.style = "Table Grid"

    for i in range(1, 6):
        t5.rows[0].cells[0].merge(t5.rows[0].cells[i])
    style_section_header(t5.rows[0].cells[0], "4. 추진 일정")

    for i, h in enumerate(["No.", "추진 항목", "1단계", "2단계", "3단계", "4단계"]):
        style_label_cell(t5.rows[1].cells[i], h)

    # 필요한 만큼 데이터 행 추가
    for _ in range(_sched_data_rows - 1):
        t5.add_row()

    for r in range(2, 2 + _sched_data_rows):
        row_bg = _BLUE_ALT if r % 2 == 0 else "FFFFFF"
        for c in range(6):
            _set_shading(t5.rows[r].cells[c], row_bg)
        style_value_cell(t5.rows[r].cells[0], str(r - 1))
        for c in range(1, 6):
            style_value_cell(t5.rows[r].cells[c])
        set_row_height(t5.rows[r], 1.0)

    doc.add_paragraph()

    # ── 표6: 5. 소요 예산 (동적 행: 섹션헤더 + 컬럼헤더 + 데이터 N행 + 합계행) ──
    budget_raw = data.get("budget", []) if data else []
    if isinstance(budget_raw, dict):
        _budget_list = list(budget_raw.values())
    elif isinstance(budget_raw, list):
        _budget_list = budget_raw
    else:
        _budget_list = []
    _budget_data_rows = max(len(_budget_list), 3)
    t6 = doc.add_table(rows=2 + 1, cols=5)  # 헤더2행 + 첫 데이터행
    t6.style = "Table Grid"

    for i in range(1, 5):
        t6.rows[0].cells[0].merge(t6.rows[0].cells[i])
    style_section_header(t6.rows[0].cells[0], "5. 소요 예산")

    for i, h in enumerate(["No.", "항목", "수량", "단가", "금액"]):
        style_label_cell(t6.rows[1].cells[i], h)

    # 필요한 만큼 데이터 행 추가
    for _ in range(_budget_data_rows - 1):
        t6.add_row()

    for r in range(2, 2 + _budget_data_rows):
        row_bg = _BLUE_ALT if r % 2 == 0 else "FFFFFF"
        for c in range(5):
            _set_shading(t6.rows[r].cells[c], row_bg)
        style_value_cell(t6.rows[r].cells[0], str(r - 1))
        for c in range(1, 5):
            style_value_cell(t6.rows[r].cells[c])
        set_row_height(t6.rows[r], 1.0)

    # 합계 행 추가
    t6.add_row()
    _total_idx = 2 + _budget_data_rows
    t6.rows[_total_idx].cells[0].merge(t6.rows[_total_idx].cells[3])
    style_label_cell(t6.rows[_total_idx].cells[0], "합계")
    style_value_cell(t6.rows[_total_idx].cells[4])

    doc.add_paragraph()

    # ── 표7: 6. 기대 효과 ──
    t7 = doc.add_table(rows=2, cols=1)
    t7.style = "Table Grid"
    style_section_header(t7.rows[0].cells[0], "6. 기대 효과")
    style_value_cell(t7.rows[1].cells[0])
    set_row_height(t7.rows[1], 4.0)

    doc.add_paragraph()

    # ── 표8: 첨부 자료 / 비고 ──
    t8 = doc.add_table(rows=2, cols=2)
    t8.style = "Table Grid"
    style_label_cell(t8.rows[0].cells[0], "첨부 자료")
    style_value_cell(t8.rows[0].cells[1])
    style_label_cell(t8.rows[1].cells[0], "비고")
    style_value_cell(t8.rows[1].cells[1])

    doc.add_paragraph()

    # ── 표9: 결재란 ──
    t9 = doc.add_table(rows=2, cols=3)
    t9.style = "Table Grid"
    for i, h in enumerate(["작성", "검토", "승인"]):
        style_label_cell(t9.rows[0].cells[i], h)
        style_value_cell(t9.rows[1].cells[i])
        set_row_height(t9.rows[1], 2.0)

    # ── data 주입 ──
    if data:
        # 표지
        _inject(t0.rows[0].cells[1], data.get("submit_date", ""))
        _inject(t0.rows[1].cells[1], data.get("submit_to", ""))
        _inject(t0.rows[2].cells[1], data.get("company", ""))
        _inject(t0.rows[3].cells[1], data.get("manager", ""))
        _inject(t0.rows[4].cells[1], data.get("contact", ""))

        # 기본 정보 (LoRA 출력 키 → 빌더 키 매핑)
        _inject(t1.rows[0].cells[1], data.get("proposal_name", "") or data.get("title", ""))
        _inject(t1.rows[1].cells[1], data.get("background", ""))
        _inject(t1.rows[2].cells[1], data.get("proposal_date", "") or data.get("submit_date", ""))
        _inject(t1.rows[2].cells[3], data.get("period", ""))
        _inject(t1.rows[3].cells[1], data.get("proposer", "") or data.get("company", ""))
        manager_contact = data.get("manager_contact", "")
        if not manager_contact:
            m = data.get("manager", "")
            c = data.get("contact", "")
            manager_contact = f"{m} / {c}" if m and c else m or c
        _inject(t1.rows[3].cells[3], manager_contact)

        # 본문 섹션
        _inject(t2.rows[1].cells[0], data.get("purpose", ""))
        _inject(t3.rows[1].cells[0], data.get("current_situation", "") or data.get("analysis", ""))
        _inject(t4.rows[1].cells[0], data.get("content", ""))

        # 추진 일정
        schedule = _sched_list
        for r in range(2, 2 + _sched_data_rows):
            item = schedule[r - 2] if r - 2 < len(schedule) else {}
            if not isinstance(item, dict):
                item = {}
            _inject(t5.rows[r].cells[1], item.get("item", ""))
            _inject(t5.rows[r].cells[2], item.get("phase1", ""))
            _inject(t5.rows[r].cells[3], item.get("phase2", ""))
            _inject(t5.rows[r].cells[4], item.get("phase3", ""))
            _inject(t5.rows[r].cells[5], item.get("phase4", ""))

        # 소요 예산
        budget = _budget_list
        for r in range(2, 2 + _budget_data_rows):
            item = budget[r - 2] if r - 2 < len(budget) else {}
            if not isinstance(item, dict):
                item = {}
            _inject(t6.rows[r].cells[1], item.get("item", ""))
            _inject(t6.rows[r].cells[2], item.get("quantity", ""))
            _inject(t6.rows[r].cells[3], item.get("unit_price", ""))
            _inject(t6.rows[r].cells[4], item.get("amount", ""))
        _inject(t6.rows[_total_idx].cells[4], data.get("budget_total", ""))

        # 기대 효과 / 첨부 / 비고
        _inject(t7.rows[1].cells[0], data.get("expected_effect", ""))
        _inject(t8.rows[0].cells[1], data.get("attachments", ""))
        _inject(t8.rows[1].cells[1], data.get("notes", ""))

    doc.save(output_path)
    print(f"제안서 생성 완료: {output_path}")


if __name__ == "__main__":
    create_proposal()
