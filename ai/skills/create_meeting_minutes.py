from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 모던 프리미엄 다크 슬레이트 테마 ──
_BLUE_HEADER = "1E293B"   # 섹션 헤더 배경 (다크 네이비/슬레이트)
_BLUE_LIGHT  = "F1F5F9"   # 라벨 셀 배경 (밝고 연한 그레이 블루)
_BLUE_ALT    = "F8FAFC"   # 테이블 짝수 행 배경 (백색에 가까운 블루)
_NAVY_RGB    = RGBColor(0x1E, 0x29, 0x3B)
_WHITE_RGB   = RGBColor(0xFF, 0xFF, 0xFF)


def _set_shading(cell, fill_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)


def _set_valign(cell, align: str = "center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def set_row_height(row, height_cm: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_cm * 567)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def style_section_header(cell, text: str):
    """섹션 제목 셀: 파란 배경 + 흰 굵은 글씨"""
    _set_shading(cell, _BLUE_HEADER)
    _set_valign(cell)
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = _WHITE_RGB
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def style_label_cell(cell, text: str):
    """라벨 셀: 연한 파란 배경 + 굵은 글씨 + 가운데 정렬"""
    _set_shading(cell, _BLUE_LIGHT)
    _set_valign(cell)
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def style_value_cell(cell, text: str = ""):
    """값 셀: 흰 배경 + 기본 글씨"""
    _set_valign(cell)
    para = cell.paragraphs[0]
    para.clear()
    if text:
        run = para.add_run(text)
        run.font.size = Pt(10)


def _inject_cell_text(cell, text: str):
    """셀에 데이터 주입"""
    cell.text = str(text) if text else ""
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].font.size = Pt(10)


def _add_title_line(doc):
    """제목 아래 파란 구분선 문단 추가"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _BLUE_HEADER)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def create_meeting_minutes(output_path: str = "회의록_test.docx", data: dict = None):
    doc = Document()

    # ── 페이지 여백 설정 ──
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 제목 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("회  의  록")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = _NAVY_RGB

    # 제목 아래 파란 구분선
    _add_title_line(doc)

    # ── 표1: 기본 정보 (5행 4열) ──
    t0 = doc.add_table(rows=5, cols=4)
    t0.style = "Table Grid"

    for row in t0.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(6.0)
        row.cells[2].width = Cm(2.5)
        row.cells[3].width = Cm(5.0)

    # Row 0: 회의 제목 (값 셀 3개 병합)
    t0.rows[0].cells[1].merge(t0.rows[0].cells[3])
    style_label_cell(t0.rows[0].cells[0], "회의 제목")
    style_value_cell(t0.rows[0].cells[1])
    set_row_height(t0.rows[0], 1.0)

    # Row 1: 회의 날짜 / 회의 시간
    style_label_cell(t0.rows[1].cells[0], "회의 날짜")
    style_value_cell(t0.rows[1].cells[1], "2026년    월    일 (   )")
    style_label_cell(t0.rows[1].cells[2], "회의 시간")
    style_value_cell(t0.rows[1].cells[3], "   :   ~   :  ")

    # Row 2: 회의 장소 / 회의 유형
    style_label_cell(t0.rows[2].cells[0], "회의 장소")
    style_value_cell(t0.rows[2].cells[1])
    style_label_cell(t0.rows[2].cells[2], "회의 유형")
    style_value_cell(t0.rows[2].cells[3])

    # Row 3: 참석자 (값 셀 3개 병합)
    t0.rows[3].cells[1].merge(t0.rows[3].cells[3])
    style_label_cell(t0.rows[3].cells[0], "참석자")
    style_value_cell(t0.rows[3].cells[1], "(쉼표로 구분하여 작성)")
    set_row_height(t0.rows[3], 1.2)

    # Row 4: 작성자 (값 셀 3개 병합)
    t0.rows[4].cells[1].merge(t0.rows[4].cells[3])
    style_label_cell(t0.rows[4].cells[0], "작성자")
    style_value_cell(t0.rows[4].cells[1])

    doc.add_paragraph()

    # ── 표2: 회의 내용 ──
    t1 = doc.add_table(rows=2, cols=1)
    t1.style = "Table Grid"
    style_section_header(t1.rows[0].cells[0], "회의 내용")
    style_value_cell(t1.rows[1].cells[0])
    set_row_height(t1.rows[1], 5.0)

    doc.add_paragraph()

    # ── 표3: 결정 사항 ──
    t2 = doc.add_table(rows=2, cols=1)
    t2.style = "Table Grid"
    style_section_header(t2.rows[0].cells[0], "결정 사항")
    style_value_cell(t2.rows[1].cells[0])
    set_row_height(t2.rows[1], 3.0)

    doc.add_paragraph()

    # ── 표4: Action Item (5행 5열: 섹션헤더 + 컬럼헤더 + 데이터3행) ──
    t3 = doc.add_table(rows=5, cols=5)
    t3.style = "Table Grid"

    # 첫 행: 섹션 헤더 (5열 병합)
    for i in range(1, 5):
        t3.rows[0].cells[0].merge(t3.rows[0].cells[i])
    style_section_header(t3.rows[0].cells[0], "Action Item")

    # 두 번째 행: 컬럼 헤더
    for i, h in enumerate(["No.", "Action Item", "담당자", "기한", "상태"]):
        style_label_cell(t3.rows[1].cells[i], h)

    # 데이터 행 3행 (교대 색상)
    for r in range(2, 5):
        row_bg = _BLUE_ALT if r % 2 == 0 else "FFFFFF"
        for c in range(5):
            _set_shading(t3.rows[r].cells[c], row_bg)
        style_value_cell(t3.rows[r].cells[0], str(r - 1))
        style_value_cell(t3.rows[r].cells[1])
        style_value_cell(t3.rows[r].cells[2])
        style_value_cell(t3.rows[r].cells[3])
        style_value_cell(t3.rows[r].cells[4])
        set_row_height(t3.rows[r], 1.0)

    doc.add_paragraph()

    # ── 표5: 비고 / 다음 회의 일정 ──
    t4 = doc.add_table(rows=2, cols=1)
    t4.style = "Table Grid"
    style_section_header(t4.rows[0].cells[0], "비고 / 다음 회의 일정")
    style_value_cell(t4.rows[1].cells[0])
    set_row_height(t4.rows[1], 2.0)

    # ── data가 있으면 셀에 실제 데이터 주입 ──
    if data:
        _inject_cell_text(t0.rows[0].cells[1], data.get("title", ""))
        _inject_cell_text(t0.rows[1].cells[1], data.get("date", ""))
        _inject_cell_text(t0.rows[1].cells[3], data.get("time", ""))
        _inject_cell_text(t0.rows[2].cells[1], data.get("location", ""))

        meeting_type = data.get("meeting_type", "")
        _inject_cell_text(t0.rows[2].cells[3], meeting_type)

        attendees = data.get("attendees", [])
        attendees_text = ", ".join(attendees) if isinstance(attendees, list) else str(attendees)
        _inject_cell_text(t0.rows[3].cells[1], attendees_text)
        _inject_cell_text(t0.rows[4].cells[1], data.get("author", ""))

        _inject_cell_text(t1.rows[1].cells[0], data.get("content", ""))

        decisions = data.get("decisions", [])
        if isinstance(decisions, list):
            parts = []
            for d in decisions:
                if isinstance(d, dict):
                    parts.append(d.get("decision", d.get("description", d.get("content", str(d)))))
                else:
                    parts.append(str(d))
            decisions_text = "\n".join(parts)
        else:
            decisions_text = str(decisions)
        _inject_cell_text(t2.rows[1].cells[0], decisions_text)

        action_items = data.get("action_items", [])
        for r in range(2, 5):
            ai_item = action_items[r - 2] if r - 2 < len(action_items) else {}
            _inject_cell_text(t3.rows[r].cells[1], ai_item.get("task", "") or ai_item.get("content", ""))
            _inject_cell_text(t3.rows[r].cells[2], ai_item.get("assignee", ""))
            _inject_cell_text(t3.rows[r].cells[3], ai_item.get("due_date", ""))
            status = ai_item.get("status", "") if ai_item else ""
            _inject_cell_text(t3.rows[r].cells[4], status)

        _inject_cell_text(t4.rows[1].cells[0], data.get("notes", ""))

    doc.save(output_path)
    print(f"회의록 생성 완료: {output_path}")


if __name__ == "__main__":
    create_meeting_minutes()
