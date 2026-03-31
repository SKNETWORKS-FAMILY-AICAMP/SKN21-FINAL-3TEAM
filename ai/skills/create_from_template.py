"""
범용 커스텀 양식 DOCX 빌더

업로드된 DOCX 양식의 필드 스펙을 기반으로, LLM이 생성한 데이터를
시스템 빌더와 동일한 스타일의 범용 레이아웃 DOCX로 생성한다.

필드 분류:
  1. field.group="meta" → 메타 테이블 (4열, 2쌍씩)
  2. field.group="body" → 본문 섹션 (번호 매김, 섹션 헤더 + 내용)
  3. 실제 데이터가 list → 배열 테이블 (컬럼 헤더 + 데이터 행)
"""
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 스타일 상수 ──
_HEADER_BG = "1E293B"
_NAVY = RGBColor(0x1E, 0x29, 0x3B)

# title로 인식할 key 목록
_TITLE_KEYS = {"title", "제목", "제안명", "보고서제목", "회의제목", "제안서제목",
               "회의제목", "문서제목", "보고서_제목", "제안서_제목"}


def _clean_label(label: str) -> str:
    """라벨 텍스트 정리: 줄바꿈 → 공백, 앞뒤 공백 제거"""
    return re.sub(r'\s*\n\s*', ' ', label).strip()


def create_generic_document(output_path: str, data: dict, fields: list[dict], doc_title: str = "문서") -> str:
    """
    범용 DOCX 생성 — 시스템 빌더(회의록/보고서/제안서)와 동일한 스타일

    field.group 메타데이터를 활용하여 메타/본문/배열을 자동 분류하고,
    원본 양식의 필드 순서를 유지하며, 빈 필드도 양식으로 표시한다.
    """
    from ai.skills._docx_styles import (
        style_section_header, style_label_cell, style_value_cell,
        set_row_height, set_shading as _set_shading, BLUE_ALT as _BLUE_ALT,
    )

    # ── group 메타데이터 보정 (기존 양식은 group이 없을 수 있음) ──
    from ai.document_parser.template_extractor import _infer_field_meta
    fields = [_infer_field_meta(f) if not f.get("group") else f for f in fields]

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 제목 추출 ──
    title = doc_title
    title_key_found = None
    for f in fields:
        if f["key"] in _TITLE_KEYS and data.get(f["key"]):
            title = str(data[f["key"]])
            title_key_found = f["key"]
            break

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = _NAVY

    # 구분선
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _HEADER_BG)
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── 필드 분류 (원본 순서 유지, 빈 필드 포함) ──
    meta_fields = []
    body_fields = []
    array_fields = []

    for f in fields:
        key = f["key"]
        if key == title_key_found:
            continue
        val = data.get(key)
        label = _clean_label(f.get("label", key))
        group = f.get("group", "")

        # 빈 값은 항상 meta (작은 셀로 표시, 큰 body 섹션 방지)
        is_empty = val is None or val == "" or val == []
        if is_empty:
            meta_fields.append((label, "", f))
            continue

        # 실제 데이터가 list인 경우
        if isinstance(val, list):
            # 짧은 문자열 배열 (참석자 등)은 메타 테이블에 쉼표로 표시
            if all(isinstance(v, str) for v in val) and sum(len(v) for v in val) < 80:
                meta_fields.append((label, ", ".join(val), f))
            else:
                array_fields.append((label, val, f))
        elif group == "meta":
            meta_fields.append((label, str(val), f))
        elif group == "body":
            body_fields.append((label, str(val), f))
        else:
            # fallback: 짧으면 meta, 길면 body
            str_val = str(val)
            if len(str_val) <= 50:
                meta_fields.append((label, str_val, f))
            else:
                body_fields.append((label, str_val, f))

    # ── 1. 메타 테이블 (4열: 라벨|값|라벨|값) ──
    if meta_fields:
        # title 키가 meta에 있으면 첫 행 전체 병합
        title_in_meta = None
        other_meta = []
        for item in meta_fields:
            if item[2]["key"] in _TITLE_KEYS and title_in_meta is None:
                title_in_meta = item
            else:
                other_meta.append(item)

        if title_in_meta:
            row_count = 1 + (len(other_meta) + 1) // 2
        else:
            row_count = (len(meta_fields) + 1) // 2

        if row_count == 0:
            row_count = 1

        t = doc.add_table(rows=row_count, cols=4)
        t.style = "Table Grid"

        for row in t.rows:
            row.cells[0].width = Cm(2.5)
            row.cells[1].width = Cm(6.0)
            row.cells[2].width = Cm(2.5)
            row.cells[3].width = Cm(5.0)
            set_row_height(row, 0.8)

        ri = 0
        if title_in_meta:
            # 첫 행: 전체 병합
            t.rows[0].cells[1].merge(t.rows[0].cells[3])
            style_label_cell(t.rows[0].cells[0], title_in_meta[0])
            style_value_cell(t.rows[0].cells[1], title_in_meta[1])
            set_row_height(t.rows[0], 1.0)
            ri = 1
            items_to_place = other_meta
        else:
            items_to_place = meta_fields

        # 2쌍씩 배치
        for i, (label, val, f) in enumerate(items_to_place):
            row_idx = ri + i // 2
            col_idx = (i % 2) * 2
            if row_idx < len(t.rows):
                style_label_cell(t.rows[row_idx].cells[col_idx], label)
                style_value_cell(t.rows[row_idx].cells[col_idx + 1], val)

        # 홀수면 마지막 행 오른쪽 빈칸
        if len(items_to_place) % 2 == 1:
            last_ri = ri + len(items_to_place) // 2
            if last_ri < len(t.rows):
                style_label_cell(t.rows[last_ri].cells[2], "")
                style_value_cell(t.rows[last_ri].cells[3], "")

        # 간격
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(6)
        spacer.paragraph_format.space_after = Pt(0)

    # ── 2. 본문 섹션 (번호 매김 + 빈 필드 유지) ──
    for idx, (label, val, f) in enumerate(body_fields, 1):
        t = doc.add_table(rows=2, cols=1)
        t.style = "Table Grid"

        # 라벨에 이미 번호가 있으면 그대로, 없으면 번호 추가
        if re.match(r"^\d+[\s.]", label):
            header_text = label
        else:
            header_text = f"{idx}. {label}"
        style_section_header(t.rows[0].cells[0], header_text)

        if val:
            style_value_cell(t.rows[1].cells[0], val)
            for para in t.rows[1].cells[0].paragraphs:
                para.paragraph_format.line_spacing = Pt(16)
            line_count = max(val.count('\n') + 1, len(val) // 60 + 1)
            height = max(2.0, min(line_count * 0.8, 8.0))
        else:
            style_value_cell(t.rows[1].cells[0], "")
            height = 1.5  # 빈 셀은 최소 높이만
        set_row_height(t.rows[1], height)

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(6)
        spacer.paragraph_format.space_after = Pt(0)

    # ── 3. 배열 테이블 ──
    for label, items, f in array_fields:
        if not items:
            continue

        if isinstance(items[0], dict):
            cols = list(items[0].keys())
            col_count = len(cols) + 1
            row_count = len(items) + 2

            t = doc.add_table(rows=row_count, cols=col_count)
            t.style = "Table Grid"

            # No. 열 좁게
            t.columns[0].width = Cm(1.0)

            # 섹션 헤더 (병합)
            for i in range(1, col_count):
                t.rows[0].cells[0].merge(t.rows[0].cells[i])
            style_section_header(t.rows[0].cells[0], label)

            # 컬럼 헤더
            style_label_cell(t.rows[1].cells[0], "No.")
            for ci, col in enumerate(cols):
                style_label_cell(t.rows[1].cells[ci + 1], col)

            # 데이터 행 (교대 색상)
            for ri, item in enumerate(items):
                row_idx = ri + 2
                bg = _BLUE_ALT if ri % 2 == 0 else "FFFFFF"
                for c in range(col_count):
                    _set_shading(t.rows[row_idx].cells[c], bg)
                style_value_cell(t.rows[row_idx].cells[0], str(ri + 1))
                for ci, col in enumerate(cols):
                    style_value_cell(t.rows[row_idx].cells[ci + 1], str(item.get(col, "")))
                set_row_height(t.rows[row_idx], 0.8)
        else:
            # 문자열 배열
            t = doc.add_table(rows=2, cols=1)
            t.style = "Table Grid"
            style_section_header(t.rows[0].cells[0], label)
            val = "\n".join(f"- {item}" for item in items)
            style_value_cell(t.rows[1].cells[0], val)
            set_row_height(t.rows[1], max(2.0, len(items) * 0.6))

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(6)
        spacer.paragraph_format.space_after = Pt(0)

    # ── 4. 결재란 ──
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(12)
    spacer.paragraph_format.space_after = Pt(0)

    t_sign = doc.add_table(rows=2, cols=3)
    t_sign.style = "Table Grid"
    for i, h in enumerate(["작성", "검토", "승인"]):
        style_label_cell(t_sign.rows[0].cells[i], h)
        style_value_cell(t_sign.rows[1].cells[i], "")
        set_row_height(t_sign.rows[1], 2.0)

    doc.save(output_path)
    print(f"[create_from_template] 범용 DOCX 생성 완료: {output_path}")
    return output_path
