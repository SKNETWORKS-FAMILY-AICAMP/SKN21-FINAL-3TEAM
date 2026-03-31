"""DOCX 공통 스타일링 유틸리티

회의록, 보고서, 제안서 빌더에서 공유하는 셀 스타일링 함수 및 테마 상수.
"""
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 모던 프리미엄 다크 슬레이트 테마 ──
BLUE_HEADER = "1E293B"   # 섹션 헤더 배경 (다크 네이비/슬레이트)
BLUE_LIGHT  = "F1F5F9"   # 라벨 셀 배경 (밝고 연한 그레이 블루)
BLUE_ALT    = "F8FAFC"   # 테이블 짝수 행 배경 (백색에 가까운 블루)
NAVY_RGB    = RGBColor(0x1E, 0x29, 0x3B)
WHITE_RGB   = RGBColor(0xFF, 0xFF, 0xFF)


def set_shading(cell, fill_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)


def set_valign(cell, align: str = "center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def set_row_height(row, height_cm: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_cm * 567)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def style_section_header(cell, text: str):
    """섹션 제목 셀: 파란 배경 + 흰 굵은 글씨"""
    set_shading(cell, BLUE_HEADER)
    set_valign(cell)
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = WHITE_RGB
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def style_label_cell(cell, text: str):
    """라벨 셀: 연한 파란 배경 + 굵은 글씨 + 가운데 정렬"""
    set_shading(cell, BLUE_LIGHT)
    set_valign(cell)
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def style_value_cell(cell, text: str = ""):
    """값 셀: 흰 배경 + 기본 글씨"""
    set_valign(cell)
    para = cell.paragraphs[0]
    para.clear()
    if text:
        run = para.add_run(text)
        run.font.size = Pt(10)


def inject_cell_text(cell, text: str):
    """셀에 데이터 주입"""
    cell.text = str(text) if text else ""
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].font.size = Pt(10)


def add_title_line(doc):
    """제목 아래 파란 구분선 문단 추가"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE_HEADER)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p
