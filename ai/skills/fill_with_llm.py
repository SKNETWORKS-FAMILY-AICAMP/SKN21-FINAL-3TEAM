"""
sLLM 매핑 보조 DOCX 채우기

원본 DOCX 양식을 열어서 셀 구조를 추출하고,
sLLM에게 "어느 셀에 어떤 data key를 넣을지" 매핑만 판단시킨 후,
값은 원본 data dict에서 직접 주입. 원본 레이아웃 100% 보존.
"""
import json
import re
from docx import Document
from docx.shared import Pt


def _is_placeholder(text: str) -> bool:
    """플레이스홀더 텍스트인지 판별 (빈칸 취급할 텍스트)
    - 숫자, 날짜 관련 패턴 (20 년 월 일, 14:00~15:30)
    - 체크박스 (☐ ☑)
    - 한글 2자 이하의 단독 날짜 단위는 라벨이므로 제외 ("일시", "장소" 등)
    """
    if not text:
        return True
    # 한글 글자가 있으면 라벨일 가능성 → 플레이스홀더가 아님
    # 단, "년", "월", "일", "시", "분"만 있으면 날짜 플레이스홀더
    korean_chars = re.findall(r'[가-힣]', text)
    if korean_chars:
        date_units = {'년', '월', '일', '시', '분'}
        non_date_korean = [c for c in korean_chars if c not in date_units]
        if non_date_korean:
            return False  # 날짜 단위 외 한글 있음 → 라벨
        # 날짜 단위만 있는 경우: 숫자가 함께 있어야 플레이스홀더 ("20년 3월" OK, "일 시" NO)
        has_digits = bool(re.search(r'\d', text))
        if not has_digits:
            return False  # 숫자 없이 "일 시"만 → 라벨
    # 날짜 단위 + 숫자/특수문자만 → 플레이스홀더
    return bool(re.match(r'^[\s☐□✓✔○●◎\d.·\-~:：()（）/년월일시분]*$', text))


def _extract_cell_structure(doc: Document) -> list[dict]:
    """원본 DOCX의 전체 셀 구조를 추출 (병합셀 감지, 공백 정규화)"""
    cells = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            seen_tcs = set()
            for ci, cell in enumerate(row.cells):
                tc_id = id(cell._tc)
                is_merged = tc_id in seen_tcs
                seen_tcs.add(tc_id)
                raw_text = cell.text.strip()
                text = re.sub(r'\s+', ' ', raw_text).strip() if raw_text else ""
                cells.append({
                    "pos": f"T{ti}R{ri}C{ci}",
                    "table": ti, "row": ri, "col": ci,
                    "text": text,
                    "is_empty": not text or _is_placeholder(text),
                    "is_merged_dup": is_merged,
                })
    return cells


# ── 영문 sub-key → 한글 컬럼 헤더 매칭 사전 ──
_SUB_KEY_ALIASES = {
    # 공통
    "item": ["항목", "추진항목", "업무", "내용", "구분"],
    "task": ["할일", "업무", "내용", "항목", "ActionItem", "Action Item"],
    "ActionItem": ["ActionItem", "Action Item", "할일", "업무", "내용"],
    "action_item": ["ActionItem", "Action Item", "할일", "업무"],
    "name": ["이름", "명칭", "항목"],
    # 일정
    "phase1": ["1단계", "1차", "phase1"],
    "phase2": ["2단계", "2차", "phase2"],
    "phase3": ["3단계", "3차", "phase3"],
    "phase4": ["4단계", "4차", "phase4"],
    "start_date": ["시작일", "시작"],
    "end_date": ["종료일", "완료일", "완료예정일", "마감"],
    # 예산
    "quantity": ["수량", "건수"],
    "unit_price": ["단가", "단위가격"],
    "amount": ["금액", "합계", "비용", "총액"],
    # 인원
    "assignee": ["담당자", "담당", "책임자"],
    "due_date": ["기한", "마감일", "완료일"],
    "progress": ["진행률", "진행상황", "진행"],
}


def _match_sub_keys_to_columns(
    data_keys: list[str],
    header_cols: dict[int, str],
    key_path: str,
    field_mapping: list[dict] | None,
    eng_to_kor: dict,
    fallback_start_col: int,
) -> dict[str, int]:
    """data dict의 sub-key를 테이블 컬럼 인덱스에 매핑 (이름 기반).

    매칭 우선순위:
      1) data key가 한글이면 → 헤더 텍스트와 직접 매칭
      2) data key가 영문이면 → _SUB_KEY_ALIASES로 한글 후보 → 헤더 매칭
      3) field_mapping.sub_keys 위치 매핑 (sub_keys[i] ↔ header → data_keys[i])
      4) 최종 fallback: No./번호 스킵 후 순차 배치
    """
    sk_to_col: dict[str, int] = {}
    used_cols: set[int] = set()

    def _try_match_header(candidate: str) -> int | None:
        """candidate 텍스트와 매칭되는 컬럼 인덱스 반환"""
        cand = re.sub(r'\s+', '', candidate)
        for col_i, h_text in header_cols.items():
            if col_i in used_cols:
                continue
            if cand == h_text or cand in h_text or h_text in cand:
                return col_i
        return None

    # ── 1) 직접 이름 매칭 (한글 key) + alias 매칭 (영문 key) ──
    for dk in data_keys:
        dk_norm = re.sub(r'\s+', '', dk)
        # 한글 직접 매칭
        col = _try_match_header(dk_norm)
        if col is not None:
            sk_to_col[dk] = col
            used_cols.add(col)
            continue
        # 영문 alias 매칭
        aliases = _SUB_KEY_ALIASES.get(dk.lower(), [])
        for alias in aliases:
            col = _try_match_header(alias)
            if col is not None:
                sk_to_col[dk] = col
                used_cols.add(col)
                break

    # ── 2) field_mapping.sub_keys 위치 매핑 (아직 매핑 안 된 key만) ──
    if field_mapping and len(sk_to_col) < len(data_keys):
        fm_entry = None
        for f in (field_mapping or []):
            if f["key"] == key_path:
                fm_entry = f
                break
        if not fm_entry:
            kor_key = eng_to_kor.get(key_path)
            if kor_key:
                fm_entry = next((f for f in field_mapping if f["key"] == kor_key), None)

        if fm_entry and fm_entry.get("sub_keys"):
            tpl_sub_keys = fm_entry["sub_keys"]
            # sub_keys[i] → 헤더 매칭으로 컬럼 확보, 그 컬럼을 data_keys[i]에 할당
            for i, tsk in enumerate(tpl_sub_keys):
                if i >= len(data_keys) or data_keys[i] in sk_to_col:
                    continue
                col = _try_match_header(re.sub(r'\s+', '', tsk))
                if col is not None:
                    sk_to_col[data_keys[i]] = col
                    used_cols.add(col)

    # ── 3) 최종 fallback: No./번호 스킵 후 순차 배치 ──
    if not sk_to_col:
        skip_labels = {'no', 'no.', '번호', '#'}
        start_col = fallback_start_col
        for col_i in sorted(header_cols.keys()):
            if header_cols[col_i].lower().strip('.') in skip_labels:
                continue
            start_col = col_i
            break
        for ski, dk in enumerate(data_keys):
            if dk not in sk_to_col:
                sk_to_col[dk] = start_col + ski

    return sk_to_col


def _format_value(val) -> str:
    """data 값을 문자열로 변환 (셀 주입용 — 쉼표 구분)"""
    if val is None:
        return ""
    if isinstance(val, list):
        parts = []
        for item in val:
            if isinstance(item, dict):
                parts.append(", ".join(f"{v}" for v in item.values() if v))
            else:
                parts.append(str(item))
        # 문자열 배열은 쉼표로, dict 배열은 줄바꿈으로
        if parts and isinstance(val[0], dict):
            return "\n".join(parts)
        return ", ".join(parts) if parts else ""
    if isinstance(val, dict):
        return ", ".join(f"{v}" for v in val.values() if v)
    return str(val)


def _build_mapping_prompt(cells: list[dict], data: dict, field_mapping: list[dict] = None) -> tuple[str, str]:
    """sLLM에게 보낼 매핑 프롬프트

    field_mapping: parsed_structure의 필드 목록 [{key, label, description}, ...]
    이걸 활용해서 "양식 라벨 ↔ data key" 연결 힌트를 제공.
    """
    # 양식 구조
    structure_lines = []
    for c in cells:
        if c.get("is_merged_dup"):
            continue
        if c["text"] and not c["is_empty"]:
            structure_lines.append(f'{c["pos"]}=[라벨]"{c["text"][:25]}"')
        else:
            structure_lines.append(f'{c["pos"]}=(빈칸)')

    # 라벨-빈칸 쌍 힌트
    label_cells = [c for c in cells if c["text"] and not c["is_empty"] and not c.get("is_merged_dup")]
    empty_cells = [c for c in cells if c["is_empty"] and not c.get("is_merged_dup")]

    pair_hints = []
    for c in label_cells:
        next_empty = next(
            (e for e in empty_cells
             if e["table"] == c["table"] and e["row"] == c["row"] and e["col"] == c["col"] + 1),
            None
        )
        if not next_empty:
            next_empty = next(
                (e for e in empty_cells
                 if e["table"] == c["table"] and e["row"] == c["row"] + 1 and e["col"] == c["col"]),
                None
            )
        if next_empty:
            pair_hints.append(f'  "{c["text"][:15]}" → {next_empty["pos"]}')

    # 핵심: 양식 라벨 ↔ data key 매핑 정보
    key_label_hints = []
    if field_mapping:
        for f in field_mapping:
            key = f.get("key", "")
            label = f.get("label", "")
            val = data.get(key)
            if val is None or val == "" or val == []:
                continue
            if isinstance(val, list):
                if val and isinstance(val[0], dict):
                    sub_keys = list(val[0].keys())
                    key_label_hints.append(f'  양식라벨="{label}" → data key="{key}" (dict배열 {len(val)}개, 하위키: {", ".join(sub_keys)}, 매핑형식: {key}[0].{sub_keys[0]})')
                else:
                    key_label_hints.append(f'  양식라벨="{label}" → data key="{key}" (문자열배열 {len(val)}개, 매핑형식: {key}[0], {key}[1]... 하위키 없음!)')
            else:
                key_label_hints.append(f'  양식라벨="{label}" → data key="{key}"')
    else:
        # field_mapping 없으면 data key만 나열
        for k, v in data.items():
            if v is None or v == "" or v == []:
                continue
            if isinstance(v, list):
                key_label_hints.append(f'  {k}: (배열, {len(v)}개 항목)')
            else:
                key_label_hints.append(f'  {k}: "{str(v)[:30]}"')

    # 배열 테이블 컬럼 헤더 감지 (같은 테이블에서 컬럼 헤더행 → 하위키 매핑 힌트)
    array_col_hints = []
    for f in (field_mapping or []):
        key = f.get("key", "")
        val = data.get(key)
        if isinstance(val, list) and val and isinstance(val[0], dict):
            sub_keys = list(val[0].keys())
            # 양식에서 이 배열의 컬럼 헤더를 찾기
            label_text = re.sub(r'\s+', ' ', f.get("label", "")).strip()
            for c in label_cells:
                if label_text and label_text in c["text"]:
                    # 같은 행의 다른 라벨 셀 = 컬럼 헤더
                    col_headers = [
                        lc["text"] for lc in label_cells
                        if lc["table"] == c["table"] and lc["row"] == c["row"] and lc["col"] != c["col"]
                    ]
                    if col_headers:
                        for ci, ch in enumerate(col_headers):
                            if ci < len(sub_keys):
                                array_col_hints.append(f'  "{ch}" 컬럼 → {key}[N].{sub_keys[ci]}')
                    break

    array_hint_text = "\n".join(array_col_hints) if array_col_hints else ""

    sys_prompt = (
        "당신은 문서 양식 전문가입니다. DOCX 양식의 셀 구조와 데이터가 주어지면, "
        "각 데이터 key를 어느 빈칸에 넣을지 매핑합니다.\n\n"
        "## 절대 규칙\n"
        "1. [라벨] 셀에는 절대 매핑하지 마세요\n"
        "2. (빈칸) 셀에만 매핑하세요\n"
        "3. '양식 라벨↔data key 매핑'의 key를 **그대로** 사용하세요. 번역하거나 변환하지 마세요!\n"
        "   예: data key가 '회의일시'면 반드시 '회의일시'로 쓰세요. 'meeting_date'로 바꾸지 마세요.\n"
        "4. 배열은 key[0], key[1] 인덱스로, dict배열은 key[0].subkey\n"
        "5. 반드시 JSON 배열로만 응답하세요\n\n"
        f"## 양식 라벨 ↔ data key 매핑\n" + "\n".join(key_label_hints) + "\n\n"
        + (f"## 배열 컬럼 ↔ 하위키 매핑\n{array_hint_text}\n\n" if array_hint_text else "")
        + f"## 라벨→빈칸 위치 힌트\n" + "\n".join(pair_hints) + "\n\n"
        '응답 형식: [{"pos":"T0R0C1","key":"date"}, {"pos":"T3R1C1","key":"decisions[0].content"}]'
    )

    user_prompt = (
        f"[양식 구조]\n" + "\n".join(structure_lines) + "\n\n"
        f"위 양식의 빈칸에 데이터 key를 매핑해서 JSON 배열로 반환하세요."
    )

    return sys_prompt, user_prompt


def _resolve_key(data: dict, key_path: str) -> str:
    """key 경로에서 값을 추출. 예: "decisions[0].content" """
    if key_path in data:
        return _format_value(data[key_path])

    match = re.match(r'^(\w+)\[(\d+)\](?:\.(\w+))?$', key_path)
    if match:
        base_key, idx, sub_key = match.group(1), int(match.group(2)), match.group(3)
        arr = data.get(base_key, [])
        if isinstance(arr, list) and idx < len(arr):
            item = arr[idx]
            if sub_key and isinstance(item, dict):
                return str(item.get(sub_key, ""))
            return _format_value(item)

    return ""


def _inject_to_cell(doc: Document, table_idx: int, row_idx: int, col_idx: int, value, preserve_label: bool = False):
    """원본 DOCX의 특정 셀에 값 주입

    preserve_label=True: 기존 텍스트(라벨)를 보존하고 아래에 줄바꿈 후 값 추가
    """
    # dict/list → 읽기 좋은 문자열로 변환
    if isinstance(value, dict):
        value = ", ".join(f"{v}" for v in value.values() if v)
    elif isinstance(value, list):
        parts = []
        for item in value:
            if isinstance(item, dict):
                parts.append(", ".join(f"{v}" for v in item.values() if v))
            else:
                parts.append(str(item))
        value = "\n".join(parts)
    value = str(value)
    try:
        table = doc.tables[table_idx]
        cell = table.rows[row_idx].cells[col_idx]
        if preserve_label:
            # 라벨 보존: 기존 텍스트 유지 + 줄바꿈 + 값 추가
            existing = cell.text.strip()
            if existing:
                para = cell.add_paragraph()
                run = para.add_run(str(value))
                run.font.size = Pt(10)
            else:
                cell.text = str(value)
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
        else:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.clear()
            cell.text = str(value)
            para = cell.paragraphs[0]
            if para.runs:
                para.runs[0].font.size = Pt(10)
    except (IndexError, AttributeError) as e:
        print(f"[fill_with_llm] 셀 주입 실패 T{table_idx}R{row_idx}C{col_idx}: {e}")


async def fill_docx_with_llm(template_path: str, output_path: str, data: dict, field_mapping: list[dict] = None) -> dict:
    """
    sLLM 매핑 보조로 원본 DOCX 양식을 채운다.

    Args:
        template_path: 원본 DOCX 양식 경로
        output_path: 출력 DOCX 경로
        data: LoRA가 생성한 데이터 dict (값은 여기서 직접 주입)
        field_mapping: parsed_structure의 필드 목록 (양식 라벨↔key 연결용)
    """
    from ai.agents.document._common import _call_llm

    doc = Document(template_path)
    cells = _extract_cell_structure(doc)
    print(f"[fill_with_llm] 양식 구조 추출: {len(cells)}개 셀")

    sys_prompt, user_prompt = _build_mapping_prompt(cells, data, field_mapping)

    print(f"[fill_with_llm] sLLM 매핑 요청 (key만)...")
    try:
        response = await _call_llm(sys_prompt, user_prompt, json_mode=True, task="generate")
        print(f"[fill_with_llm] sLLM 응답 ({len(response)}자): {response[:300]}")

        # JSON 파싱 (잘려도 복구)
        json_match = re.search(r'\[.*\]', response, re.DOTALL)
        if json_match:
            mappings = json.loads(json_match.group())
        else:
            cleaned = response.strip()
            if cleaned.startswith('[') and not cleaned.endswith(']'):
                last_brace = cleaned.rfind('}')
                if last_brace > 0:
                    cleaned = cleaned[:last_brace + 1] + ']'
            mappings = json.loads(cleaned)
    except Exception as e:
        print(f"[fill_with_llm] sLLM 매핑 실패: {e}")
        return {"success": False, "filled_count": 0, "total_cells": len(cells), "error": str(e)}

    if not isinstance(mappings, list):
        mappings = []

    # 영어→한글 key 역매핑 테이블 (field_mapping 기반으로 동적 생성)
    _ENGLISH_TO_KOREAN = {}
    if field_mapping:
        from ai.document_parser.template_extractor import FIELD_MAPPING, _normalize_label
        for f in field_mapping:
            data_key = f["key"]
            label = f.get("label", "")
            label_normalized = re.sub(r'\s+', '', label)
            # FIELD_MAPPING에서 이 라벨의 영어 key 찾기
            norm = _normalize_label(label)
            if norm in FIELD_MAPPING and FIELD_MAPPING[norm]:
                eng_key = FIELD_MAPPING[norm]["key"]
                _ENGLISH_TO_KOREAN[eng_key] = data_key
            # 라벨 자체도 역매핑
            _ENGLISH_TO_KOREAN[label_normalized] = data_key

    def _resolve_with_fallback(key_path: str) -> str:
        """영어 key → 한글 key fallback 포함 resolve"""
        val = _resolve_key(data, key_path)
        if val:
            return val
        # 영어→한글 역매핑 시도
        base_match = re.match(r'^(\w+?)(?:\[|$)', key_path)
        if base_match:
            eng_base = base_match.group(1)
            kor_base = _ENGLISH_TO_KOREAN.get(eng_base)
            if kor_base:
                kor_path = key_path.replace(eng_base, kor_base, 1)
                val = _resolve_key(data, kor_path)
                if val:
                    return val
        return ""

    # 매핑 결과로 원본 DOCX에 data 값 주입
    # 같은 key가 여러 셀에 매핑된 경우 배열 자동 인덱싱
    key_occurrence = {}  # key → 등장 횟수 (배열 자동 분배용)
    filled = 0
    for m in mappings:
        pos = m.get("pos", "")
        key_path = m.get("key", "")
        if not pos or not key_path:
            continue

        match = re.match(r"T(\d+)R(\d+)C(\d+)", pos)
        if not match:
            continue

        # 배열 자동 분배: 같은 key가 반복되면 인덱스 자동 부여
        value = _resolve_with_fallback(key_path)

        if not value:
            continue

        # value가 리스트 전체인지 확인 (단일 key로 배열이 resolve된 경우)
        raw_val = data.get(key_path)
        if raw_val is None:
            # 영어→한글 역매핑
            base_match2 = re.match(r'^(\w+?)$', key_path)
            if base_match2:
                kor_key = _ENGLISH_TO_KOREAN.get(key_path, key_path)
                raw_val = data.get(kor_key)

        ti, ri, ci = int(match.group(1)), int(match.group(2)), int(match.group(3))

        if isinstance(raw_val, list) and raw_val:
            # 같은 key가 전체 매핑에서 몇 번 나오는지 확인
            total_same_key = sum(1 for m2 in mappings if m2.get("key") == key_path)

            if total_same_key <= 1 and all(isinstance(v, dict) for v in raw_val):
                # dict 배열 (schedule, budget 등) → 테이블 헤더 기반 행 분배
                data_sub_keys = list(raw_val[0].keys())

                # ── 1) 테이블 컬럼 헤더행 탐색 (위/아래 모두) ──
                header_cols = {}  # col_index → header_text
                header_ri = None
                table_total_rows = len(doc.tables[ti].rows)

                def _find_column_header(search_rows):
                    """라벨이 2개 이상인 행을 컬럼 헤더로 판별"""
                    for check_ri in search_rows:
                        if check_ri < 0 or check_ri >= table_total_rows:
                            continue
                        row_cells_h = [c for c in cells
                                       if c["table"] == ti and c["row"] == check_ri
                                       and not c.get("is_merged_dup")]
                        lbl_count = sum(1 for c in row_cells_h if c["text"] and not c["is_empty"])
                        if lbl_count >= 3:  # No. + 항목 + 담당자 등 최소 3개
                            return check_ri, row_cells_h
                        if lbl_count >= 2 and len(row_cells_h) >= 3:
                            return check_ri, row_cells_h
                    return None, []

                # ri 자체 → ri 위 → ri 아래 순으로 탐색
                for search in [
                    [ri],
                    list(range(ri - 1, max(ri - 3, -1), -1)),
                    list(range(ri + 1, min(ri + 4, table_total_rows))),
                ]:
                    found_ri, found_cells = _find_column_header(search)
                    if found_ri is not None:
                        header_ri = found_ri
                        for c in found_cells:
                            if c["text"] and not c["is_empty"]:
                                header_cols[c["col"]] = re.sub(r'\s+', '', c["text"])
                        break

                # 데이터 시작 행: 헤더 바로 다음 행
                data_start_ri = (header_ri + 1) if header_ri is not None else ri
                # ri가 데이터 행보다 뒤에 있으면 ri 사용 (이미 데이터 행에 매핑된 경우)
                if data_start_ri <= ri and header_ri != ri:
                    data_start_ri = ri

                # ── 2) data key → 컬럼 인덱스 매핑 (이름 기반) ──
                sk_to_col = _match_sub_keys_to_columns(
                    data_sub_keys, header_cols, key_path,
                    field_mapping, _ENGLISH_TO_KOREAN, ci,
                )

                # ── 3) 행 분배: 병합/합계 행 스킵 ──
                table_row_count = len(doc.tables[ti].rows)
                for idx, item in enumerate(raw_val):
                    row_idx = data_start_ri + idx
                    if row_idx >= table_row_count:
                        break
                    # 병합 행 감지 → 스킵 (합계 등)
                    row_cells_check = [c for c in cells
                                       if c["table"] == ti and c["row"] == row_idx]
                    merged_count = sum(1 for c in row_cells_check if c.get("is_merged_dup"))
                    if merged_count > len(row_cells_check) // 2:
                        print(f"[fill_with_llm] T{ti}R{row_idx} 병합행 스킵")
                        break

                    for sk, sv_raw in item.items():
                        sv = str(sv_raw) if sv_raw else ""
                        if not sv:
                            continue
                        col_idx = sk_to_col.get(sk)
                        if col_idx is None:
                            continue
                        cell_pos = f"T{ti}R{row_idx}C{col_idx}"
                        try:
                            _inject_to_cell(doc, ti, row_idx, col_idx, sv)
                            filled += 1
                            print(f"[fill_with_llm] {cell_pos} ← {key_path}[{idx}].{sk} = {sv[:30]}")
                        except (IndexError, AttributeError):
                            pass
                continue  # 다음 매핑으로 (단일 셀 주입 스킵)
            elif total_same_key <= 1:
                # 문자열 배열 (참석자 등) → 한 셀에 쉼표로
                value = ", ".join(str(v) for v in raw_val)
            else:
                # 여러 행에 매핑 → 배열 자동 인덱싱 (결정사항 등)
                idx = key_occurrence.get(key_path, 0)
                key_occurrence[key_path] = idx + 1
                if idx < len(raw_val):
                    item = raw_val[idx]
                    value = _format_value(item) if not isinstance(item, str) else item
                else:
                    continue

        # 라벨 셀인지 체크 → 라벨이면 옆/아래 빈칸으로 리다이렉트
        target_cell = next((c for c in cells if c["pos"] == pos), None)
        is_label_cell = target_cell and target_cell["text"] and not target_cell["is_empty"]

        if is_label_cell:
            redirect = None
            # 병합 행인지 체크 (섹션 헤더: "3. 진행 현황" 등)
            ri_merged = sum(1 for c in cells if c["table"] == ti and c["row"] == ri and c.get("is_merged_dup"))
            is_section_header = ri_merged >= 2  # 대부분 셀이 병합됨

            if is_section_header:
                # 섹션 헤더 → 아래 행에서 빈 셀 탐색 (컬럼 헤더 스킵)
                for search_ri in range(ri + 1, ri + 5):
                    row_cells_s = [c for c in cells if c["table"] == ti and c["row"] == search_ri and not c.get("is_merged_dup")]
                    empty_s = [c for c in row_cells_s if c["is_empty"]]
                    label_s = [c for c in row_cells_s if c["text"] and not c["is_empty"]]
                    # 빈 셀이 절반 이상이고 라벨이 적으면 → 데이터 행
                    if empty_s and len(empty_s) >= len(row_cells_s) // 2 and len(label_s) <= 2:
                        redirect = empty_s[0]
                        break
            else:
                # 일반 라벨 → 옆 빈칸 (같은 행, col+1)
                redirect = next(
                    (c for c in cells if c["table"] == ti and c["row"] == ri
                     and c["col"] == ci + 1 and c["is_empty"] and not c.get("is_merged_dup")), None)
                # 아래 빈칸 (다음 행, 같은 col)
                if not redirect:
                    redirect = next(
                        (c for c in cells if c["table"] == ti and c["row"] == ri + 1
                         and c["col"] == ci and c["is_empty"] and not c.get("is_merged_dup")), None)

            if redirect:
                _inject_to_cell(doc, ti, redirect["row"], redirect["col"], value)
                filled += 1
                print(f"[fill_with_llm] {pos}→{redirect['pos']} (리다이렉트) ← {key_path} = {value[:40]}")
            else:
                # 빈칸 못 찾으면 스킵 — 보충 패스에서 올바른 위치에 배치
                print(f"[fill_with_llm] {pos} 라벨셀 스킵 (보충 대기) ← {key_path} = {value[:40]}")
        else:
            _inject_to_cell(doc, ti, ri, ci, value)
            filled += 1
            print(f"[fill_with_llm] {pos} ← {key_path} = {value[:40]}")

    # ── 규칙 기반 fallback: sLLM이 매핑 못 한 필드를 라벨 매칭으로 보충 ──
    # 실제 주입된 셀 위치 + 주입된 key(영어→한글 역매핑 포함) 추적
    injected_positions = set()
    injected_keys = set()
    for m in mappings:
        pos = m.get("pos", "")
        kp = m.get("key", "")
        if pos and kp and _resolve_with_fallback(kp):
            injected_positions.add(pos)
            # 영어→한글 역매핑으로 원본 key도 추적
            base = re.match(r'^(\w+)', kp).group(1) if re.match(r'^(\w+)', kp) else kp
            injected_keys.add(base)
            injected_keys.add(_ENGLISH_TO_KOREAN.get(base, base))

    if field_mapping:
        # 라벨 → 셀 매핑 (공백 완전 제거 정규화, 첫 번째만 유지)
        label_cells_map = {}
        for c in cells:
            if c["text"] and not c["is_empty"] and not c.get("is_merged_dup"):
                norm = re.sub(r'\s+', '', c["text"])
                if norm and norm not in label_cells_map:
                    label_cells_map[norm] = c

        for f in field_mapping:
            key = f.get("key", "")
            label_normalized = re.sub(r'\s+', '', f.get("label", ""))
            val = data.get(key)
            if not val or not label_normalized:
                continue

            # 라벨 셀 찾기
            matched_cell = label_cells_map.get(label_normalized)
            if not matched_cell:
                continue

            # 이 key가 sLLM에 의해 이미 주입됐는지 확인
            if key in injected_keys:
                continue
            label_pos = matched_cell["pos"]
            adjacent_pos = f"T{matched_cell['table']}R{matched_cell['row']}C{matched_cell['col']+1}"
            if adjacent_pos in injected_positions or label_pos in injected_positions:
                continue

            ti_m = matched_cell["table"]

            # dict 배열: 라벨이 반복되는 테이블의 데이터 행에 분배
            if isinstance(val, list) and val and isinstance(val[0], dict):
                # 같은 테이블에서 데이터를 넣을 첫 번째 빈 행 찾기
                # (병합 헤더 행, 컬럼 헤더 행을 스킵하고 빈 셀이 있는 첫 행)
                data_start_row = None
                for check_ri in range(len(doc.tables[ti_m].rows)):
                    row_cells = [c for c in cells if c["table"] == ti_m and c["row"] == check_ri and not c.get("is_merged_dup")]
                    empty_count = sum(1 for c in row_cells if c["is_empty"])
                    text_count = sum(1 for c in row_cells if c["text"] and not c["is_empty"])
                    # 빈 셀이 절반 이상이고 텍스트 셀이 2개 이하 → 데이터 행
                    if empty_count >= len(row_cells) * 0.5 and text_count <= 2:
                        data_start_row = check_ri
                        break
                if data_start_row is None:
                    data_start_row = matched_cell["row"] + 2  # fallback: 헤더+컬럼헤더 스킵

                # 헤더행에서 컬럼 매핑 구축
                fb_header_cols = {}
                for check_ri in range(data_start_row - 1, max(data_start_row - 3, -1), -1):
                    row_cells_h = [c for c in cells
                                   if c["table"] == ti_m and c["row"] == check_ri
                                   and not c.get("is_merged_dup")]
                    label_count = sum(1 for c in row_cells_h if c["text"] and not c["is_empty"])
                    if label_count >= 2:
                        for c in row_cells_h:
                            if c["text"] and not c["is_empty"]:
                                fb_header_cols[c["col"]] = re.sub(r'\s+', '', c["text"])
                        break

                data_sub_keys_fb = list(val[0].keys())
                fb_sk_to_col = _match_sub_keys_to_columns(
                    data_sub_keys_fb, fb_header_cols, key,
                    field_mapping, _ENGLISH_TO_KOREAN, matched_cell["col"] + 1,
                )

                table_row_count = len(doc.tables[ti_m].rows)
                for idx, item in enumerate(val):
                    row_idx = data_start_row + idx
                    if row_idx >= table_row_count:
                        break
                    # 병합 행 스킵 (합계 등)
                    row_cells_chk = [c for c in cells if c["table"] == ti_m and c["row"] == row_idx]
                    merged_cnt = sum(1 for c in row_cells_chk if c.get("is_merged_dup"))
                    if merged_cnt > len(row_cells_chk) // 2:
                        print(f"[fill_with_llm] (fallback) T{ti_m}R{row_idx} 병합행 스킵")
                        break
                    for sk, sv_raw in item.items():
                        sv = str(sv_raw) if sv_raw else ""
                        if not sv:
                            continue
                        col_idx = fb_sk_to_col.get(sk)
                        if col_idx is None:
                            continue
                        cell_pos = f"T{ti_m}R{row_idx}C{col_idx}"
                        if cell_pos not in injected_positions:
                            _inject_to_cell(doc, ti_m, row_idx, col_idx, sv)
                            filled += 1
                            print(f"[fill_with_llm] (fallback) {cell_pos} ← {key}[{idx}].{sk} = {sv[:30]}")
                continue

            # 단순 값: 라벨 옆 또는 아래 빈칸
            target = None
            for ec in cells:
                if ec.get("is_merged_dup"):
                    continue
                if ec["table"] == ti_m and ec["row"] == matched_cell["row"] and ec["col"] == matched_cell["col"] + 1 and ec["is_empty"]:
                    target = ec
                    break
            if not target:
                for ec in cells:
                    if ec.get("is_merged_dup"):
                        continue
                    if ec["table"] == ti_m and ec["row"] == matched_cell["row"] + 1 and ec["col"] == matched_cell["col"] and ec["is_empty"]:
                        target = ec
                        break

            value_str = _format_value(val)
            if target and target["pos"] not in injected_positions:
                _inject_to_cell(doc, target["table"], target["row"], target["col"], value_str)
                filled += 1
                print(f"[fill_with_llm] (fallback) {target['pos']} ← {key} = {value_str[:40]}")
            elif not target:
                # 빈칸 없음 → 라벨 셀에 내용 추가 (라벨 보존)
                _inject_to_cell(doc, matched_cell["table"], matched_cell["row"], matched_cell["col"], value_str, preserve_label=True)
                filled += 1
                print(f"[fill_with_llm] (fallback+preserve) {matched_cell['pos']} ← {key} = {value_str[:40]}")

    # ── 최종 보충: 라벨 옆 셀이 비어있으면 값 채우기 ──
    # sLLM이 잘못된 셀에 매핑한 경우를 보완 (참석자→작성자 셀 등)
    if field_mapping:
        label_cells_final = {}
        for c in cells:
            if c["text"] and not c["is_empty"] and not c.get("is_merged_dup"):
                norm = re.sub(r'\s+', '', c["text"])
                if norm and norm not in label_cells_final:
                    label_cells_final[norm] = c

        for f in field_mapping:
            key = f.get("key", "")
            val = data.get(key)
            if not val:
                continue
            label_norm = re.sub(r'\s+', '', f.get("label", ""))
            if not label_norm:
                continue

            # dict 배열이 아닌 단순 값: 라벨 옆 빈 셀이면 보충
            if not (isinstance(val, list) and val and isinstance(val[0], dict)):
                mc = label_cells_final.get(label_norm)
                if not mc:
                    continue
                # 옆 셀 (col+1) 확인
                adj = next((c for c in cells if c["table"] == mc["table"] and c["row"] == mc["row"]
                            and c["col"] == mc["col"] + 1 and not c.get("is_merged_dup")), None)
                if not adj:
                    # 아래 셀
                    adj = next((c for c in cells if c["table"] == mc["table"] and c["row"] == mc["row"] + 1
                                and c["col"] == mc["col"] and not c.get("is_merged_dup")), None)
                if adj:
                    cell_text = doc.tables[adj["table"]].rows[adj["row"]].cells[adj["col"]].text.strip()
                    value_str = _format_value(val)
                    # 비어있거나, 현재 값이 기대 값과 다르면 올바른 값으로 덮어쓰기
                    if not cell_text or cell_text != value_str.strip():
                        _inject_to_cell(doc, adj["table"], adj["row"], adj["col"], value_str)
                        filled += 1
                        print(f"[fill_with_llm] (보충) {adj['pos']} ← {key} = {value_str[:30]}")
                continue

            data_sub_keys_final = list(val[0].keys())
            label_norm = re.sub(r'\s+', '', f.get("label", ""))

            # 이 배열이 속한 테이블 찾기 (라벨 매칭)
            matched_ti = None
            for c in cells:
                cn = re.sub(r'\s+', '', c.get("text", ""))
                if cn and (label_norm in cn or cn in label_norm) and not c.get("is_merged_dup"):
                    matched_ti = c["table"]
                    break

            if matched_ti is None:
                continue

            # 해당 테이블에서 컬럼 헤더 찾기
            final_header_cols = {}
            final_header_ri = None
            t_rows = len(doc.tables[matched_ti].rows)
            for check_ri in range(t_rows):
                row_cells_h = [c for c in cells
                               if c["table"] == matched_ti and c["row"] == check_ri
                               and not c.get("is_merged_dup")]
                lbl_cnt = sum(1 for c in row_cells_h if c["text"] and not c["is_empty"])
                if lbl_cnt >= 3:
                    final_header_ri = check_ri
                    for c in row_cells_h:
                        if c["text"] and not c["is_empty"]:
                            final_header_cols[c["col"]] = re.sub(r'\s+', '', c["text"])
                    break

            if not final_header_cols or final_header_ri is None:
                continue

            final_sk_to_col = _match_sub_keys_to_columns(
                data_sub_keys_final, final_header_cols, key,
                field_mapping, _ENGLISH_TO_KOREAN, 1,
            )
            data_start = final_header_ri + 1

            for idx, item in enumerate(val):
                row_idx = data_start + idx
                if row_idx >= t_rows:
                    break
                for sk, sv_raw in item.items():
                    sv = str(sv_raw) if sv_raw else ""
                    if not sv:
                        continue
                    col_idx = final_sk_to_col.get(sk)
                    if col_idx is None:
                        continue
                    # 셀이 아직 비어있을 때만 보충
                    cell_text = doc.tables[matched_ti].rows[row_idx].cells[col_idx].text.strip()
                    if not cell_text:
                        _inject_to_cell(doc, matched_ti, row_idx, col_idx, sv)
                        filled += 1
                        print(f"[fill_with_llm] (보충) T{matched_ti}R{row_idx}C{col_idx} ← {key}[{idx}].{sk} = {sv[:30]}")

    doc.save(output_path)
    print(f"[fill_with_llm] 완료: {filled}개 셀 채움 → {output_path}")
    return {"success": True, "filled_count": filled, "total_cells": len(cells)}
