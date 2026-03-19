"""
문서-규정 검증 파이프라인 (Document + Judgment Agent 결합)

문서 agent가 생성한 문서를 판단 agent(regulation_checker)가 검증하는 구조.
비스트리밍으로 내부 JSON만 반환.

사용법:
    from ai.agents.regulation_validator import validate_document_regulations, check_content_regulations

    # 1. 문서 생성 후 규정 검증
    reg_result = await validate_document_regulations(data, "report", user_id=1)

    # 2. 문서 검색/요약 후 규정 연결
    reg_result = await check_content_regulations("USB 사용에 대한 내용...", user_id=1)
"""

import asyncio
import json
import logging
import time
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


# ── 문서 유형별 규정 체크 대상 필드 ──
_TEMPLATE_CHECK_FIELDS = {
    "meeting_minutes": {
        "fields": [
            ("decisions", "회의 결정사항"),
            ("action_items", "실행 항목"),
            ("content", "회의 내용"),
        ],
        "focus": "회의 결정사항과 후속 조치가 내부 규정에 부합하는지",
    },
    "report": {
        "fields": [
            ("main_content", "주요 내용"),
            ("tasks", "업무 항목"),
            ("issues", "이슈 사항"),
            ("next_plan", "향후 계획"),
        ],
        "focus": "보고서에 언급된 업무 내용과 계획이 내부 규정에 부합하는지",
    },
    "proposal": {
        "fields": [
            ("budget", "예산"),
            ("purpose", "제안 목적"),
            ("expected_effect", "기대효과"),
            ("content", "제안 내용"),
            ("schedule", "추진 일정"),
        ],
        "focus": "제안서의 목적, 내용, 예산, 일정이 내부 규정에 부합하는지",
    },
    "jd": {
        "fields": [
            ("benefits", "복리후생"),
            ("working_conditions", "근무조건"),
            ("requirements", "자격요건"),
            ("responsibilities", "직무 내용"),
        ],
        "focus": "채용공고의 복리후생, 근무조건, 자격요건이 내부 규정에 부합하는지",
    },
}


async def validate_document_regulations(
    data: Dict[str, Any],
    template_type: str,
    user_id: int | None = None,
) -> Dict[str, Any]:
    """문서 생성 후 규정 준수 검증 (비스트리밍)

    문서 agent가 생성한 data를 받아, 판단 agent(regulation_checker)로
    내부 규정 위반 여부를 체크하고 결과를 반환.

    Args:
        data: 생성된 문서 데이터 (LLM 결과)
        template_type: meeting_minutes, report, proposal, jd
        user_id: RAG 필터용

    Returns:
        {
            "checked": True,
            "notes": [{topic, query, result, reason, regulation, confidence}],
            "has_violations": bool,
            "has_conditions": bool,
            "summary": "규정 검증 결과 마크다운 텍스트",
        }
    """
    _t = time.time()
    logger.info("[RegValidator] 문서 규정 검증 시작 | type=%s", template_type)

    # 1. 규정 체크 대상 쿼리 추출 (LLM + 규칙 기반 하이브리드)
    queries = await _extract_regulation_queries(data, template_type)

    if not queries:
        logger.info("[RegValidator] 규정 체크 대상 없음")
        return _empty_result()

    # 2. 병렬 규정 체크 (regulation_checker 사용)
    from ai.agents.regulation_checker import regulation_check

    tasks = [regulation_check(q["query"], user_id=user_id) for q in queries]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    # 3. 결과 취합
    notes = _aggregate_results(queries, results)

    has_violations = any(n["result"] == "no" for n in notes)
    has_conditions = any(n["result"] == "conditional" for n in notes)
    summary = _format_regulation_summary(notes, has_violations, has_conditions)

    logger.info(
        "[RegValidator] 검증 완료 (%.2fs) | %d건 체크, %d건 관련, violations=%s, conditions=%s",
        time.time() - _t, len(queries), len(notes), has_violations, has_conditions,
    )

    return {
        "checked": True,
        "notes": notes,
        "has_violations": has_violations,
        "has_conditions": has_conditions,
        "summary": summary,
    }


async def check_content_regulations(
    content: str,
    user_id: int | None = None,
) -> Dict[str, Any]:
    """문서 검색/요약 결과에 대한 규정 연결 (비스트리밍)

    문서 내용에서 규정과 관련된 부분을 찾아 규정 정보를 연결합니다.

    Args:
        content: 문서 내용 또는 요약/QA 답변 텍스트
        user_id: RAG 필터용

    Returns:
        {
            "checked": True,
            "notes": [{topic, query, result, reason, regulation, confidence}],
            "summary": "규정 연결 결과 텍스트",
        }
    """
    _t = time.time()
    logger.info("[RegValidator] 컨텐츠 규정 연결 시작 | content_len=%d", len(content))

    # LLM으로 규정 관련 포인트 추출
    queries = await _extract_content_regulation_points(content)

    if not queries:
        return {"checked": True, "notes": [], "summary": ""}

    from ai.agents.regulation_checker import regulation_check

    tasks = [regulation_check(q["query"], user_id=user_id) for q in queries]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    notes = _aggregate_results(queries, results)
    summary = _format_content_regulation_notes(notes) if notes else ""

    logger.info(
        "[RegValidator] 컨텐츠 규정 연결 완료 (%.2fs) | %d건", time.time() - _t, len(notes),
    )

    return {"checked": True, "notes": notes, "summary": summary}


def append_regulation_section_to_docx(docx_path: str, notes: List[Dict]) -> None:
    """생성된 DOCX 파일에 규정 검증 결과 섹션 추가

    Args:
        docx_path: DOCX 파일 경로
        notes: regulation validation notes 리스트
    """
    if not notes:
        return

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor

        doc = DocxDocument(docx_path)

        # 구분선 역할 빈 줄
        doc.add_paragraph()

        # 제목
        heading = doc.add_heading("규정 검증 결과", level=2)

        # 위반 → 조건부 → 부합 순으로 정렬
        _ORDER = {"no": 0, "conditional": 1, "yes": 2}
        sorted_notes = sorted(notes, key=lambda n: _ORDER.get(n["result"], 3))

        for note in sorted_notes:
            result_label = {
                "no": "[위반]",
                "conditional": "[조건부]",
                "yes": "[부합]",
            }.get(note["result"], "[확인]")

            result_color = {
                "no": RGBColor(0xCC, 0x00, 0x00),        # 빨강
                "conditional": RGBColor(0xCC, 0x88, 0x00),  # 주황
                "yes": RGBColor(0x00, 0x88, 0x00),        # 초록
            }.get(note["result"], RGBColor(0x33, 0x33, 0x33))

            p = doc.add_paragraph()

            # 결과 라벨 (색상)
            run_label = p.add_run(f"{result_label} ")
            run_label.bold = True
            run_label.font.color.rgb = result_color
            run_label.font.size = Pt(10)

            # 주제
            run_topic = p.add_run(note["topic"])
            run_topic.bold = True
            run_topic.font.size = Pt(10)

            # 사유
            run_reason = p.add_run(f"\n{note['reason']}")
            run_reason.font.size = Pt(9)

            # 근거 규정
            if note.get("regulation"):
                run_reg = p.add_run(f"\n근거: {note['regulation']}")
                run_reg.font.size = Pt(9)
                run_reg.italic = True

        doc.save(docx_path)
        logger.info("[RegValidator] DOCX 규정 섹션 추가 완료: %s", docx_path)

    except Exception as e:
        logger.warning("[RegValidator] DOCX 규정 섹션 추가 실패: %s", e)


# ── 내부 함수 ──


def _empty_result() -> Dict[str, Any]:
    return {
        "checked": True,
        "notes": [],
        "has_violations": False,
        "has_conditions": False,
        "summary": "",
    }


def _aggregate_results(
    queries: List[Dict], results: list
) -> List[Dict[str, Any]]:
    """regulation_check 결과를 취합하여 notes 리스트 생성"""
    notes = []
    for query_info, result in zip(queries, results):
        if isinstance(result, Exception):
            logger.warning("[RegValidator] 규정 체크 실패: %s", result)
            continue
        if not result.get("checked"):
            continue
        # no_regulation → 관련 규정 없음 → 스킵
        if result.get("result") == "no_regulation":
            continue

        notes.append({
            "topic": query_info["topic"],
            "query": query_info["query"],
            "result": result["result"],       # yes | no | conditional
            "reason": result.get("reason", ""),
            "regulation": result.get("regulation"),
            "confidence": result.get("confidence", 0.0),
        })
    return notes


async def _extract_regulation_queries(
    data: Dict[str, Any],
    template_type: str,
) -> List[Dict[str, str]]:
    """문서 데이터에서 규정 체크 대상 쿼리 추출 (LLM + 규칙 기반 하이브리드)

    1차: LLM으로 문서 내용 분석하여 규정 관련 포인트 추출
    2차 (fallback): 템플릿 유형별 필드 기반 규칙 추출
    """
    # 문서 내용 수집
    check_config = _TEMPLATE_CHECK_FIELDS.get(template_type, {})
    target_fields = check_config.get("fields", [])
    focus = check_config.get("focus", "내부 규정에 부합하는지")

    content_parts = []
    for field_key, field_label in target_fields:
        val = data.get(field_key)
        if not val:
            continue
        text = _value_to_text(val)
        if text.strip():
            content_parts.append(f"[{field_label}]\n{text}")

    full_content = "\n\n".join(content_parts)

    # target field가 비어있으면 content/summary에서 추출
    if not full_content.strip():
        for fallback_key in ("content", "summary", "overview"):
            if data.get(fallback_key):
                full_content = str(data[fallback_key])
                break

    if not full_content.strip():
        return []

    # LLM으로 규정 체크 포인트 추출 시도
    queries = await _llm_extract_queries(full_content, focus)

    # LLM 실패 시 규칙 기반 fallback
    if not queries:
        queries = _rule_based_extract_queries(data, template_type)

    return queries[:5]  # 최대 5개


async def _llm_extract_queries(
    content: str,
    focus: str,
) -> List[Dict[str, str]]:
    """LLM으로 문서 내용에서 규정 체크 대상 추출"""
    sys_prompt = f"""당신은 문서 규정 검증 전문가입니다.
주어진 문서 내용에서 내부 규정 검증이 필요한 항목을 추출하세요.

## 검증 포커스
{focus}

## 추출 기준
- 근무조건 (근무시간, 재택근무, 유연근무 등)
- 비용/예산 (출장비, 교육비, 복리후생 등)
- 인사/채용 (채용조건, 수습기간, 해고 등)
- 보안/정보보호 (개인정보, USB 사용, 문서 보안 등)
- 결재/승인 (금액 기준, 결재 권한 등)
- 휴가/근태 (연차, 반차, 병가, 경조 등)

## 출력 형식 (JSON 배열만 출력)
[{{"topic": "항목명", "query": "규정 확인 질문"}}]

규칙:
- 최대 5개까지만 추출
- 규정과 무관한 내용은 스킵
- 각 query는 규정 판단이 가능한 구체적 질문으로 작성
- 규정 관련 항목이 없으면 빈 배열 [] 반환
- JSON 외 텍스트 금지"""

    user_prompt = f"다음 문서 내용에서 규정 검증이 필요한 항목을 추출하세요.\n\n{content[:5000]}"

    try:
        from ai.llm import get_llm
        llm = get_llm()
        response = await llm.generate(
            prompt=user_prompt,
            system_prompt=sys_prompt,
            json_mode=True,
            temperature=0.1,
            max_tokens=800,
        )

        raw = _strip_code_block(response.content)
        items = json.loads(raw)
        if isinstance(items, list):
            return [
                {"topic": item.get("topic", ""), "query": item.get("query", "")}
                for item in items
                if item.get("query")
            ][:5]
    except Exception as e:
        logger.warning("[RegValidator] LLM 쿼리 추출 실패: %s", e)

    return []


def _rule_based_extract_queries(
    data: Dict[str, Any],
    template_type: str,
) -> List[Dict[str, str]]:
    """규칙 기반 규정 체크 쿼리 추출 (LLM fallback)"""
    queries = []

    if template_type == "meeting_minutes":
        decisions = data.get("decisions", [])
        if decisions:
            text = _value_to_text(decisions)
            queries.append({
                "topic": "회의 결정사항",
                "query": f"다음 회의 결정사항이 내부 규정에 부합하는지 확인: {text[:500]}",
            })
        action_items = data.get("action_items", [])
        if action_items:
            text = _value_to_text(action_items)
            queries.append({
                "topic": "실행 항목",
                "query": f"다음 실행 항목이 내부 규정에 부합하는지 확인: {text[:500]}",
            })

    elif template_type == "report":
        for key, topic in [("main_content", "주요 내용"), ("tasks", "업무 항목")]:
            val = data.get(key)
            if val:
                text = _value_to_text(val)
                queries.append({
                    "topic": topic,
                    "query": f"다음 보고서 {topic}이 내부 규정에 부합하는지 확인: {text[:500]}",
                })

    elif template_type == "proposal":
        for key, topic in [("budget", "예산"), ("purpose", "제안 목적"), ("content", "제안 내용")]:
            val = data.get(key)
            if val:
                text = _value_to_text(val)
                queries.append({
                    "topic": topic,
                    "query": f"다음 제안서 {topic}이 내부 규정에 부합하는지 확인: {text[:500]}",
                })

    elif template_type == "jd":
        for key, topic in [("benefits", "복리후생"), ("working_conditions", "근무조건")]:
            val = data.get(key)
            if val:
                text = _value_to_text(val)
                queries.append({
                    "topic": topic,
                    "query": f"다음 채용공고 {topic}이 내부 규정에 부합하는지 확인: {text[:500]}",
                })

    # Fallback: content 필드
    if not queries:
        content = data.get("content", "") or data.get("summary", "")
        if content:
            queries.append({
                "topic": "문서 내용",
                "query": f"다음 문서 내용이 내부 규정에 부합하는지 확인: {str(content)[:500]}",
            })

    return queries


async def _extract_content_regulation_points(content: str) -> List[Dict[str, str]]:
    """문서 컨텐츠에서 규정 관련 포인트 추출 (검색/요약용)"""
    sys_prompt = """당신은 문서 분석 전문가입니다.
주어진 문서 내용에서 내부 규정과 관련될 수 있는 포인트를 추출하세요.

## 추출 대상
- 근무 관련 언급 (재택, 출장, 야근 등)
- 비용/금액 관련 언급
- 보안/정보보호 관련 언급
- 인사/채용 관련 언급
- 복리후생 관련 언급

## 출력 형식 (JSON 배열만 출력)
[{"topic": "관련 주제", "query": "해당 내용이 내부 규정에 부합하는가?"}]

규칙:
- 최대 3개까지만 추출
- 규정과 명확히 무관한 내용은 스킵
- 규정 관련 포인트가 없으면 빈 배열 [] 반환
- JSON 외 텍스트 금지"""

    user_prompt = f"다음 문서 내용에서 규정 관련 포인트를 추출하세요.\n\n{content[:3000]}"

    try:
        from ai.llm import get_llm
        llm = get_llm()
        response = await llm.generate(
            prompt=user_prompt,
            system_prompt=sys_prompt,
            json_mode=True,
            temperature=0.1,
            max_tokens=500,
        )

        raw = _strip_code_block(response.content)
        items = json.loads(raw)
        if isinstance(items, list):
            return [
                {"topic": item.get("topic", ""), "query": item.get("query", "")}
                for item in items
                if item.get("query")
            ][:3]
    except Exception as e:
        logger.warning("[RegValidator] 컨텐츠 규정 포인트 추출 실패: %s", e)

    return []


# ── 포맷팅 ──


def _format_regulation_summary(
    notes: List[Dict],
    has_violations: bool,
    has_conditions: bool,
) -> str:
    """규정 검증 결과를 마크다운 텍스트로 포맷 (문서 생성용)"""
    if not notes:
        return ""

    lines = ["\n---\n\n## 규정 검증 결과\n"]

    if has_violations:
        lines.append("**[주의] 규정 위반 사항이 발견되었습니다.**\n")
    elif has_conditions:
        lines.append("**[참고] 조건부 허용 사항이 있습니다.**\n")
    else:
        lines.append("**검토된 항목이 규정에 부합합니다.**\n")

    violations = [n for n in notes if n["result"] == "no"]
    conditionals = [n for n in notes if n["result"] == "conditional"]
    compliant = [n for n in notes if n["result"] == "yes"]

    if violations:
        lines.append("### [위반]")
        for n in violations:
            reg = f" ({n['regulation']})" if n.get("regulation") else ""
            lines.append(f"- **{n['topic']}**: {n['reason']}{reg}")
        lines.append("")

    if conditionals:
        lines.append("### [조건부]")
        for n in conditionals:
            reg = f" ({n['regulation']})" if n.get("regulation") else ""
            lines.append(f"- **{n['topic']}**: {n['reason']}{reg}")
        lines.append("")

    if compliant:
        lines.append("### [부합]")
        for n in compliant:
            reg = f" ({n['regulation']})" if n.get("regulation") else ""
            lines.append(f"- **{n['topic']}**: {n['reason']}{reg}")
        lines.append("")

    return "\n".join(lines)


def _format_content_regulation_notes(notes: List[Dict]) -> str:
    """문서 검색/요약 결과에 대한 규정 연결 노트 (doc_retrieve용)"""
    if not notes:
        return ""

    lines = ["\n\n---\n**관련 규정 안내**\n"]

    for n in notes:
        label = {"no": "[위반]", "conditional": "[조건부]", "yes": "[부합]"}.get(
            n["result"], ""
        )
        reg = f" - {n['regulation']}" if n.get("regulation") else ""
        lines.append(f"- {label} **{n['topic']}**: {n['reason']}{reg}")

    return "\n".join(lines)


# ── 유틸 ──


def _value_to_text(val) -> str:
    """딕셔너리/리스트/문자열을 검색 가능한 텍스트로 변환"""
    if isinstance(val, str):
        return val
    if isinstance(val, list):
        parts = []
        for item in val:
            if isinstance(item, dict):
                parts.append(json.dumps(item, ensure_ascii=False))
            else:
                parts.append(str(item))
        return "\n".join(parts)
    if isinstance(val, dict):
        return json.dumps(val, ensure_ascii=False)
    return str(val) if val else ""


def _strip_code_block(text: str) -> str:
    """LLM 응답에서 코드블록 마커 제거"""
    raw = text.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1] if "\n" in raw else raw
        if raw.endswith("```"):
            raw = raw[:-3]
        raw = raw.strip()
    return raw
