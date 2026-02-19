"""
데이터 수집 보고서 docx 생성 스크립트
SK네트웍스 Family AI과정 21기 최종 프로젝트 3조 — WorkFlow Agent (듀듀)

사용법:
    python generate_report.py

출력:
    1. 데이터 수집 및 저장_수집 데이터_3조.docx
"""

import json
import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 경로 설정 ──────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parents[2]  # docs/산출물/데이터수집 → project root
OUTPUT_FILE = SCRIPT_DIR / "1. 데이터 수집 및 저장_수집 데이터_3조.docx"

# 샘플 데이터 경로
INTENT_FILE = PROJECT_ROOT / "data" / "training" / "intent" / "judgment.jsonl"
MEETING_FILE = PROJECT_ROOT / "data" / "proceedings" / "hr_meeting_minutes_100.json"
BENCHMARK_FILE = PROJECT_ROOT / "data" / "evaluation" / "benchmark_testset.jsonl"

# ── 스타일 상수 ────────────────────────────────────────────────
FONT_NAME = "맑은 고딕"
FONT_NAME_EN = "Malgun Gothic"

COLOR_DARK_BLUE = RGBColor(0x1F, 0x3A, 0x5F)   # 진한 남색 (헤더)
COLOR_LIGHT_BLUE = RGBColor(0x44, 0x72, 0xC4)   # 파란색 (표 헤더)
COLOR_HEADER_BG = "2E75B6"                        # 표 헤더 배경
COLOR_ALT_ROW = "F2F7FB"                          # 표 교대행 배경
COLOR_LIGHT_GRAY = "F5F5F5"                       # 코드블록 배경
COLOR_WHITE = "FFFFFF"
COLOR_BLACK = "000000"
COLOR_BORDER = "B4C6E7"                           # 표 테두리

TITLE_SIZE = Pt(28)
SUBTITLE_SIZE = Pt(14)
HEADING1_SIZE = Pt(18)
HEADING2_SIZE = Pt(14)
BODY_SIZE = Pt(10.5)
SMALL_SIZE = Pt(9)
TABLE_SIZE = Pt(9.5)
CODE_SIZE = Pt(8.5)


# ── 유틸리티 함수 ──────────────────────────────────────────────

def set_font(run, name=FONT_NAME, size=BODY_SIZE, bold=False, color=None, italic=False):
    """폰트 속성 설정"""
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # 한글 폰트 설정
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), name)


def set_cell_shading(cell, color_hex):
    """셀 배경색 설정"""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """셀 테두리 설정"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)

    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            border = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" w:sz="4" '
                f'w:space="0" w:color="{val}"/>'
            )
            existing = tcBorders.find(qn(f"w:{side}"))
            if existing is not None:
                tcBorders.remove(existing)
            tcBorders.append(border)


def set_table_borders(table, color=COLOR_BORDER):
    """테이블 전체 테두리 설정"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def add_styled_paragraph(doc, text, size=BODY_SIZE, bold=False, color=None,
                          alignment=None, space_before=Pt(0), space_after=Pt(6),
                          italic=False, line_spacing=1.3):
    """스타일이 적용된 문단 추가"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    return p


def add_heading_styled(doc, text, level=1, color=COLOR_DARK_BLUE):
    """커스텀 스타일 헤딩"""
    sizes = {1: HEADING1_SIZE, 2: HEADING2_SIZE}
    size = sizes.get(level, BODY_SIZE)
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    pf = p.paragraph_format
    pf.space_before = Pt(18) if level == 1 else Pt(12)
    pf.space_after = Pt(8)
    # 하단 선 추가 (level 1만)
    if level == 1:
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="2E75B6"/>'
            f'</w:pBdr>'
        )
        p._element.get_or_add_pPr().append(pBdr)
    return p


def create_styled_table(doc, headers, rows, col_widths=None, header_color=COLOR_HEADER_BG):
    """스타일이 적용된 표 생성"""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # 헤더 행
    hdr_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        set_font(run, size=TABLE_SIZE, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
        # 수직 중앙 정렬
        cell.vertical_alignment = 1  # CENTER

    # 데이터 행
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_font(run, size=TABLE_SIZE)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = 1
            # 교대행 배경
            if row_idx % 2 == 1:
                set_cell_shading(cell, COLOR_ALT_ROW)

    # 열 너비 설정
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = width

    return table


def add_code_block(doc, text, label=""):
    """코드 블록 스타일 텍스트"""
    if label:
        add_styled_paragraph(doc, label, size=SMALL_SIZE, bold=True,
                             color=COLOR_LIGHT_BLUE, space_before=Pt(8))

    p = doc.add_paragraph()
    # 배경색 (문단 shading)
    pPr = p._element.get_or_add_pPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{COLOR_LIGHT_GRAY}" w:val="clear"/>'
    )
    pPr.append(shading)
    # 좌우 들여쓰기
    indent = parse_xml(
        f'<w:ind {nsdecls("w")} w:left="284" w:right="284"/>'
    )
    pPr.append(indent)

    run = p.add_run(text)
    set_font(run, name="Consolas", size=CODE_SIZE)
    # Consolas에도 한글 폰트 fallback
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is not None:
        rFonts.set(qn("w:eastAsia"), FONT_NAME)

    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15
    return p


def add_bullet_point(doc, text, level=0, bold_prefix=""):
    """불릿 포인트 (수동 구현)"""
    p = doc.add_paragraph()
    indent_val = 360 + (level * 360)
    hanging = 200

    pPr = p._element.get_or_add_pPr()
    ind = parse_xml(
        f'<w:ind {nsdecls("w")} w:left="{indent_val}" w:hanging="{hanging}"/>'
    )
    pPr.append(ind)

    bullet_char = "•" if level == 0 else "–"
    if bold_prefix:
        run = p.add_run(f"{bullet_char} {bold_prefix}")
        set_font(run, size=BODY_SIZE, bold=True)
        run2 = p.add_run(text)
        set_font(run2, size=BODY_SIZE)
    else:
        run = p.add_run(f"{bullet_char} {text}")
        set_font(run, size=BODY_SIZE)

    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.3
    return p


def add_numbered_item(doc, number, text, indent=360):
    """번호 목록 아이템"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    ind = parse_xml(
        f'<w:ind {nsdecls("w")} w:left="{indent}" w:hanging="200"/>'
    )
    pPr.append(ind)
    run = p.add_run(f"{number}. {text}")
    set_font(run, size=BODY_SIZE)
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.3
    return p


# ── 페이지 번호 ────────────────────────────────────────────────

def add_page_numbers(doc):
    """푸터에 페이지 번호 추가 (- N - 형태)"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 페이지 번호 필드 삽입: "- {PAGE} -"
    run_before = p.add_run("- ")
    set_font(run_before, size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))
    # PAGE 필드
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_field = p.add_run()
    run_field._element.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_field2 = p.add_run()
    run_field2._element.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_field3 = p.add_run()
    run_field3._element.append(fldChar2)
    set_font(run_field3, size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))
    run_after = p.add_run(" -")
    set_font(run_after, size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))


# ── Intent 건수 집계 ───────────────────────────────────────────

def count_intent_data():
    """Intent 파일에서 카테고리별 원본/증강 건수 집계"""
    intent_dir = PROJECT_ROOT / "data" / "training" / "intent"
    categories = [
        "judgment", "doc_search", "doc_generate",
        "meeting_generate", "schedule_add", "schedule_view", "general",
    ]
    # 카테고리 매칭용 키워드 (파일명 부분매칭)
    cat_keywords = {
        "judgment": ["judgment"],
        "doc_search": ["doc_search"],
        "doc_generate": ["doc_generate"],
        "meeting_generate": ["meeting_generate", "meeting"],
        "schedule_add": ["schedule_add"],
        "schedule_view": ["schedule_view"],
        "general": ["general"],
    }
    result = {}
    orig_total = 0
    matched_aug_files = set()

    for cat in categories:
        orig_file = intent_dir / f"{cat}.jsonl"
        orig_count = 0
        if orig_file.exists():
            with open(orig_file, "r", encoding="utf-8") as f:
                orig_count = sum(1 for line in f if line.strip())
        # 증강 건수 — 키워드 매칭
        aug_count = 0
        for kw in cat_keywords[cat]:
            for aug_file in intent_dir.glob(f"augment_*{kw}*.jsonl"):
                if aug_file not in matched_aug_files:
                    with open(aug_file, "r", encoding="utf-8") as f:
                        aug_count += sum(1 for line in f if line.strip())
                    matched_aug_files.add(aug_file)
        result[cat] = (orig_count, aug_count)
        orig_total += orig_count

    # cross-category: 위에서 매칭 안 된 augment 파일들
    cross_count = 0
    for fp in sorted(intent_dir.glob("augment_*.jsonl")):
        if fp not in matched_aug_files:
            with open(fp, "r", encoding="utf-8") as f:
                cross_count += sum(1 for line in f if line.strip())
    result["_cross"] = cross_count
    return result


# ── 데이터 로딩 ────────────────────────────────────────────────

def load_intent_samples(n=3):
    """Intent JSONL에서 샘플 로드"""
    samples = []
    files_and_labels = [
        (PROJECT_ROOT / "data/training/intent/judgment.jsonl", "judgment"),
        (PROJECT_ROOT / "data/training/intent/doc_search.jsonl", "doc_search"),
        (PROJECT_ROOT / "data/training/intent/schedule_add.jsonl", "schedule_add"),
    ]
    for fpath, _ in files_and_labels:
        if fpath.exists():
            with open(fpath, "r", encoding="utf-8") as f:
                line = f.readline().strip()
                if line:
                    samples.append(line)
    return samples[:n]


def load_meeting_sample():
    """회의록 JSON에서 첫 번째 샘플 로드"""
    if MEETING_FILE.exists():
        with open(MEETING_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            if data:
                return data[0]
    return None


# ── 문서 생성 ──────────────────────────────────────────────────

def build_document():
    doc = Document()

    # ── 페이지 설정 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 기본 스타일 수정
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    # [개선 2] 페이지 번호 추가
    add_page_numbers(doc)

    # ================================================================
    # 섹션 1: 표지 / 헤더
    # ================================================================
    # SK networks 로고 텍스트 대체
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_logo.add_run("SK networks")
    set_font(run, size=Pt(11), bold=True, color=RGBColor(0xE5, 0x00, 0x2B))
    run2 = p_logo.add_run("  |  Family AI Camp")
    set_font(run2, size=Pt(11), color=RGBColor(0x66, 0x66, 0x66))

    # 구분선
    p_line = doc.add_paragraph()
    pPr = p_line._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="E5002B"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # 빈 줄
    doc.add_paragraph()

    # 제목
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("데이터 수집")
    set_font(run, size=TITLE_SIZE, bold=True, color=COLOR_DARK_BLUE)
    p_title.paragraph_format.space_after = Pt(8)

    # 부제
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("SKN Family AI Camp 21기 : 최종 프로젝트 3조")
    set_font(run, size=SUBTITLE_SIZE, color=RGBColor(0x66, 0x66, 0x66))
    p_sub.paragraph_format.space_after = Pt(4)

    # 프로젝트명
    p_proj = doc.add_paragraph()
    p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_proj.add_run("WorkFlow Agent (듀듀)")
    set_font(run, size=Pt(12), color=RGBColor(0x44, 0x72, 0xC4), italic=True)
    p_proj.paragraph_format.space_after = Pt(16)

    # [개선 1] 팀원 정보 + 작성일자
    doc.add_paragraph()  # 간격
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_date.add_run("작성일: 2025-02-19")
    set_font(run, size=Pt(10), color=RGBColor(0x55, 0x55, 0x55))
    p_date.paragraph_format.space_after = Pt(4)

    p_team = doc.add_paragraph()
    p_team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_team.add_run("작성: 3조 — 신지용(PM), 윤경은(AI), 진승언(AI), 안혜빈(Backend), 문지영(Frontend)")
    set_font(run, size=Pt(9.5), color=RGBColor(0x55, 0x55, 0x55))
    p_team.paragraph_format.space_after = Pt(24)

    # ================================================================
    # 섹션 2: 프로젝트 주제 + 데이터 출처 요약표
    # ================================================================
    add_heading_styled(doc, "1. 프로젝트 주제", level=1)
    add_styled_paragraph(
        doc,
        "LangGraph 기반 사내 업무 자동화 멀티에이전트 시스템으로, "
        "규정 판단, 문서 생성/검색/요약, 회의록 구조화, 일정 관리를 "
        "AI 에이전트가 자동으로 처리합니다. "
        "klue/bert-base 기반 Intent Classifier가 사용자 의도를 7개 카테고리로 분류하고, "
        "LangGraph 오케스트레이터가 적합한 전문 에이전트(규정 판단 / 문서 / 일정)로 라우팅합니다. "
        "각 에이전트는 LLM API(현재) 또는 LoRA 파인튜닝 sLLM(vLLM 서빙, 추후)으로 동작하며, "
        "사내규정 RAG(ChromaDB + BM25 하이브리드 검색)와 Google Workspace 연동을 통해 "
        "실서비스 수준의 업무 자동화를 제공합니다.",
        size=BODY_SIZE, space_after=Pt(12),
    )

    add_heading_styled(doc, "2. 데이터 출처 요약", level=1)
    add_styled_paragraph(
        doc,
        "아래 표는 본 프로젝트에서 수집 및 생성한 전체 데이터의 요약입니다. "
        "총 4,600건 중 약 3,700건이 수집 완료되었으며, 나머지 약 900건은 추가 수집 예정입니다.",
        size=BODY_SIZE, space_after=Pt(8),
    )

    summary_headers = ["데이터", "출처", "데이터 형태", "정보 및 분량"]
    summary_rows = [
        ["Intent 분류", "자체 제작\n(GPT-4/Claude 생성 + 수동 검수)", "JSONL", "1,916건 (7개 의도)"],
        ["규정 판단", "자체 제작\n(가상 사내규정 기반)", "XLSX → JSONL", "1,000건"],
        ["규정 해석 Q&A", "자체 제작\n(가상 사내규정 기반)", "XLSX → JSONL", "500건"],
        ["회의록 분석", "자체 제작\n(GPT-4 생성)", "JSON", "800건"],
        ["규정 원본 (RAG)", "자체 제작\n(가상 사내규정 PDF)", "PDF", "1파일 (44청크)"],
        ["벤치마크 테스트셋", "자체 제작", "JSONL", "87건"],
        ["문서 검색/요약/생성/리스크", "수집 예정\n(GPT/Claude 생성)", "JSONL", "900건 예정"],
    ]
    summary_widths = [Cm(3.5), Cm(5), Cm(2.5), Cm(3.5)]
    create_styled_table(doc, summary_headers, summary_rows, summary_widths)

    # [개선 4] 데이터 → 모델 매핑 요약표
    add_styled_paragraph(doc, "", space_after=Pt(4))
    add_styled_paragraph(doc, "데이터-모델 매핑", size=BODY_SIZE, bold=True,
                         color=COLOR_LIGHT_BLUE, space_after=Pt(4))
    add_styled_paragraph(
        doc,
        "수집한 데이터가 어떤 모델의 학습에 사용되는지를 아래 표에 정리하였습니다.",
        space_after=Pt(6),
    )
    mapping_headers = ["모델", "베이스 모델", "학습 데이터", "건수"]
    mapping_rows = [
        ["Intent Classifier", "klue/bert-base", "Intent 분류 데이터", "1,916건"],
        ["LoRA v1 (규정 판단)", "Qwen2.5-7B 등", "규정 판단 + 규정 해석 Q&A", "1,500건"],
        ["LoRA v2 (문서 에이전트)", "Qwen2.5-7B 등",
         "회의록 + 문서 검색/요약/생성/리스크", "1,700건"],
    ]
    mapping_widths = [Cm(3.5), Cm(3), Cm(5), Cm(2.5)]
    t = create_styled_table(doc, mapping_headers, mapping_rows, mapping_widths)
    # 합계행 추가
    total_row = t.add_row()
    labels = ["합계", "", "", "~4,600건*"]
    for i, txt in enumerate(labels):
        cell = total_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        set_font(run, size=TABLE_SIZE, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = 1
        set_cell_shading(cell, "E2EFDA")  # 연한 녹색 합계행
    add_styled_paragraph(doc,
        "* RAG 원본(44청크), 벤치마크 테스트셋(87건) 별도",
        size=Pt(8.5), italic=True, color=RGBColor(0x88, 0x88, 0x88),
        space_after=Pt(8))

    # ================================================================
    # 섹션 3: 데이터 수집 계기
    # ================================================================
    add_heading_styled(doc, "3. 데이터 수집 계기", level=1)

    add_styled_paragraph(
        doc,
        "본 프로젝트는 3개의 모델을 파인튜닝하여 실서비스에 적용하는 것을 목표로 합니다. "
        "Intent Classifier(klue/bert-base), LoRA v1(규정 판단), LoRA v2(문서 에이전트) 모델 각각에 "
        "맞는 학습 데이터가 필요했습니다.",
        space_after=Pt(6),
    )
    add_styled_paragraph(
        doc,
        "사내 규정 판단(Yes/No/조건부), 회의록 구조화(자유 텍스트 → JSON), "
        "동적 필드 기반 문서 생성 등은 기존 공개 데이터셋이 존재하지 않는 도메인입니다. "
        "따라서 GPT-4 및 Claude API를 활용한 합성 데이터 생성 전략을 채택하였습니다.",
        space_after=Pt(6),
    )
    add_styled_paragraph(
        doc,
        "품질 확보를 위해 다음과 같은 3단계 프로세스를 적용했습니다.",
        space_after=Pt(4),
    )
    add_numbered_item(doc, 1,
        "카테고리별 예시 20건을 수작업으로 작성하여 품질 기준선(seed data)을 확립")
    add_numbered_item(doc, 2,
        "GPT-4 / Claude API에 시드 데이터와 상세 프롬프트를 제공하여 대량 생성")
    add_numbered_item(doc, 3,
        "2차 수동 검수를 통해 라벨 오류, 중복, 품질 미달 데이터 제거")

    add_styled_paragraph(
        doc,
        "\n또한, '듀듀테크놀로지'라는 가상 회사의 사내규정(30개 조항)을 직접 작성하여 "
        "일관된 도메인 컨텍스트를 기반으로 모든 데이터를 생성했습니다. "
        "이를 통해 규정 판단, 회의록, 문서 생성 데이터가 동일한 기업 맥락을 공유하도록 설계했습니다.",
        space_after=Pt(12),
    )

    # ================================================================
    # 섹션 4: 주요 데이터 상세 (3종)
    # ================================================================
    add_heading_styled(doc, "4. 주요 데이터 상세", level=1)

    # ── 4-1. Intent 분류 ──
    add_heading_styled(doc, "4-1. Intent 분류 학습 데이터 (1,916건) — 수집 완료", level=2)

    add_bullet_point(doc, "klue/bert-base 파인튜닝 → 7개 의도 분류 모델 학습", bold_prefix="용도: ")
    add_bullet_point(doc,
        "judgment, doc_search, doc_generate, meeting_generate, "
        "schedule_add, schedule_view, general",
        bold_prefix="의도 카테고리: ")
    add_bullet_point(doc,
        "data/training/intent/*.jsonl (원본 1,453건 + 증강 463건 = 1,916건)",
        bold_prefix="파일 위치: ")
    add_bullet_point(doc,
        "Eval F1 98.23%, Adversarial F1 90.07%",
        bold_prefix="학습 결과: ")

    # [개선 3] Intent 카테고리별 분포 표
    add_styled_paragraph(doc, "", space_after=Pt(2))
    add_styled_paragraph(doc, "카테고리별 분포", size=BODY_SIZE, bold=True,
                         color=COLOR_LIGHT_BLUE, space_after=Pt(4))
    intent_data = count_intent_data()
    intent_headers = ["의도 (Intent)", "원본", "증강", "합계"]
    intent_rows = []
    cat_labels = {
        "judgment": "judgment (규정 판단)",
        "doc_search": "doc_search (문서 검색)",
        "doc_generate": "doc_generate (문서 생성)",
        "meeting_generate": "meeting_generate (회의록)",
        "schedule_add": "schedule_add (일정 등록)",
        "schedule_view": "schedule_view (일정 조회)",
        "general": "general (일반 대화)",
    }
    sum_orig, sum_aug = 0, 0
    for cat in ["judgment", "doc_search", "doc_generate", "meeting_generate",
                "schedule_add", "schedule_view", "general"]:
        orig, aug = intent_data.get(cat, (0, 0))
        sum_orig += orig
        sum_aug += aug
        intent_rows.append([cat_labels[cat], str(orig), str(aug), str(orig + aug)])
    cross = intent_data.get("_cross", 0)
    if cross:
        intent_rows.append(["(공통 증강: 비정형/초성/경계)", "—", str(cross), str(cross)])
        sum_aug += cross
    intent_widths = [Cm(5.5), Cm(2), Cm(2), Cm(2)]
    it = create_styled_table(doc, intent_headers, intent_rows, intent_widths)
    # 합계행
    total_row = it.add_row()
    totals = ["합계", str(sum_orig), str(sum_aug), str(sum_orig + sum_aug)]
    for i, txt in enumerate(totals):
        cell = total_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        set_font(run, size=TABLE_SIZE, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = 1
        set_cell_shading(cell, "E2EFDA")

    add_styled_paragraph(doc, "", space_after=Pt(2))
    add_styled_paragraph(doc, "수집 방법", size=BODY_SIZE, bold=True,
                         color=COLOR_LIGHT_BLUE, space_after=Pt(4))
    add_numbered_item(doc, 1,
        "GPT-4에 카테고리별 프롬프트로 200건씩 초기 생성")
    add_numbered_item(doc, 2,
        "반말/존댓말, 짧은/긴 문장, 경계 케이스(20~30건/카테고리) 혼합")
    add_numbered_item(doc, 3,
        "수동 검수 후 라벨 수정 → 원본 v1.0 (1,453건)")
    add_numbered_item(doc, 4,
        "v1.2 증강: 인터넷 슬랭/초성/축약어 등 비정형 데이터 300건 추가")
    add_numbered_item(doc, 5,
        "v1.3 증강: Adversarial 테스트 기반 경계 케이스 타겟 163건 추가")

    add_styled_paragraph(doc, "", space_after=Pt(2))
    add_styled_paragraph(doc, "전처리 유의점", size=BODY_SIZE, bold=True,
                         color=COLOR_LIGHT_BLUE, space_after=Pt(4))
    add_bullet_point(doc, "한국어 정규화 (초성/숫자 치환, 특수문자 제거)")
    add_bullet_point(doc, "중복 제거 및 라벨 불균형 검사 (카테고리별 최소 150건 이상)")
    add_bullet_point(doc, "15%를 평가 데이터로 분리 (stratified split)")

    # ── 4-2. 규정 판단 ──
    add_heading_styled(doc, "4-2. 규정 판단 데이터 (1,000건 + Q&A 500건) — 원본 확보", level=2)

    add_bullet_point(doc, "LoRA v1 파인튜닝 → 사내 규정 기반 Yes/No/조건부 판단 + 설명형 답변", bold_prefix="용도: ")
    add_bullet_point(doc,
        "judgment_raw.xlsx (1,000건), regulation_qa_raw.xlsx (500건)",
        bold_prefix="파일: ")
    add_bullet_point(doc,
        "instruction / input(규정 원문 + 질문) / output(판단 결과 + 근거 + 대안)",
        bold_prefix="포맷: ")

    add_styled_paragraph(doc, "", space_after=Pt(2))
    add_styled_paragraph(doc, "수집 방법", size=BODY_SIZE, bold=True,
                         color=COLOR_LIGHT_BLUE, space_after=Pt(4))
    add_numbered_item(doc, 1,
        "듀듀테크놀로지 가상 사내규정 PDF 30개 조항을 기반으로 시나리오 설계")
    add_numbered_item(doc, 2,
        "각 조항별 판단 시나리오 30~50건씩 GPT-4로 생성 (총 1,000건)")
    add_numbered_item(doc, 3,
        "Q&A는 각 조항별 '이 규정의 의미는?', '이 경우 어떻게 해야 하나요?' 형태 15~20건씩 (총 500건)")
    add_numbered_item(doc, 4,
        "input에 규정 원문 포함 — 실서비스 RAG 검색 결과와 동일한 형태로 학습 가능")

    add_styled_paragraph(doc, "", space_after=Pt(2))
    add_styled_paragraph(doc, "전처리 유의점", size=BODY_SIZE, bold=True,
                         color=COLOR_LIGHT_BLUE, space_after=Pt(4))
    add_bullet_point(doc, "Excel(XLSX) → JSONL 변환 필요")
    add_bullet_point(doc, "input 필드에 규정 원문 텍스트 반드시 삽입 (RAG 실서비스와 동일 형태)")
    add_bullet_point(doc, "output에 판단(Yes/No/조건부) + 근거 조항 + 대안 구조 유지")

    # ── 4-3. 회의록 분석 ──
    add_heading_styled(doc, "4-3. 회의록 분석 데이터 (800건) — 원본 확보", level=2)

    add_bullet_point(doc, "LoRA v2 파인튜닝 → 회의 원문 → 구조화 JSON 추출", bold_prefix="용도: ")
    add_bullet_point(doc,
        "data/proceedings/*.json (7개 파일, 총 800건)",
        bold_prefix="파일: ")
    add_bullet_point(doc,
        "input(회의 원문 텍스트) → output(구조화 JSON)",
        bold_prefix="포맷: ")

    add_styled_paragraph(doc, "", space_after=Pt(2))
    add_styled_paragraph(doc, "수집 방법", size=BODY_SIZE, bold=True,
                         color=COLOR_LIGHT_BLUE, space_after=Pt(4))
    add_numbered_item(doc, 1,
        "GPT-4로 다양한 유형의 회의 시나리오 생성 (인사, 보안감사, 기획, 장애대응, 약식메모 등)")
    add_numbered_item(doc, 2,
        "유형별 100건씩 생성 후 수동 검수")
    add_numbered_item(doc, 3,
        "output 스키마 확정: 회의정보(제목/유형/핵심주제/일시), 결정사항, action_items(담당자/내용/기한/우선순위), 참석자(이름/직책)")

    add_styled_paragraph(doc, "", space_after=Pt(2))
    add_styled_paragraph(doc, "전처리 유의점", size=BODY_SIZE, bold=True,
                         color=COLOR_LIGHT_BLUE, space_after=Pt(4))
    add_bullet_point(doc, "JSON → JSONL 변환 (배열 → 개별 라인)")
    add_bullet_point(doc, "output JSON 스키마 일관성 검증 (필드명/타입 불일치 시 학습 불가)")
    add_bullet_point(doc, "'일시' 필드 매핑 확인 (빈 문자열 허용)")

    # ================================================================
    # 섹션 5: 나머지 데이터 요약 + 수집 예정
    # ================================================================
    add_heading_styled(doc, "5. 기타 데이터 및 수집 예정분", level=1)

    add_styled_paragraph(
        doc,
        "아래 표는 주요 3종 외의 기수집 데이터와 향후 수집 예정 데이터를 정리한 것입니다. "
        "수집 예정 데이터(900건)는 LoRA v2 문서 에이전트 파인튜닝에 사용됩니다.",
        space_after=Pt(8),
    )

    misc_headers = ["데이터", "건수", "모델", "상태", "수집 방법"]
    misc_rows = [
        ["규정 원본 PDF (RAG)", "44청크", "—", "V 완료", "가상 규정 PDF → ChromaDB 적재"],
        ["벤치마크 테스트셋", "87건", "—", "V 완료", "6개 카테고리별 테스트 문항"],
        ["문서 검색 답변", "200건", "LoRA v2", "[ ] 예정", "규정 조항 조합 → 검색결과+답변"],
        ["문서 요약 (동적 필드)", "300건", "LoRA v2", "[ ] 예정", "5종 필드 조합 × 60건씩"],
        ["문서 생성 (동적 필드)", "200건", "LoRA v2", "[ ] 예정", "4종 템플릿 × 50건씩"],
        ["리스크 감지", "200건", "LoRA v2", "[ ] 예정", "위반 100건 + 정상 100건"],
    ]
    misc_widths = [Cm(3.5), Cm(1.5), Cm(1.8), Cm(1.5), Cm(6)]
    mt = create_styled_table(doc, misc_headers, misc_rows, misc_widths)

    # [개선 5] 상태 셀 색상 (완료=녹색, 예정=회색) + 합계행
    for row_idx in range(len(misc_rows)):
        row = mt.rows[row_idx + 1]
        status_cell = row.cells[3]
        status_text = misc_rows[row_idx][3]
        if "완료" in status_text:
            set_cell_shading(status_cell, "E2EFDA")  # 연한 녹색
        elif "예정" in status_text:
            set_cell_shading(status_cell, "FFF2CC")  # 연한 노란색

    # 합계행
    total_row = mt.add_row()
    total_labels = [
        "합계 (주요 3종 제외)",
        "~987건",
        "",
        "완료 2 / 예정 4",
        "",
    ]
    for i, txt in enumerate(total_labels):
        cell = total_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        set_font(run, size=TABLE_SIZE, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = 1
        set_cell_shading(cell, "D6E4F0")  # 연한 파란색 합계행

    # ================================================================
    # 섹션 6: 원본 데이터 예시
    # ================================================================
    add_heading_styled(doc, "6. 원본 데이터 예시", level=1)
    add_styled_paragraph(
        doc,
        "각 주요 데이터의 실제 샘플을 아래에 제시합니다.",
        space_after=Pt(8),
    )

    # ── Intent 샘플 ──
    add_heading_styled(doc, "6-1. Intent 분류 데이터 (JSONL)", level=2)
    intent_samples = load_intent_samples()
    for sample in intent_samples:
        add_code_block(doc, sample)

    # ── 규정 판단 샘플 ──
    add_heading_styled(doc, "6-2. 규정 판단 데이터 (instruction / input / output)", level=2)
    # 벤치마크에서 규정 판단 샘플을 가져옴
    bench_sample = None
    if BENCHMARK_FILE.exists():
        with open(BENCHMARK_FILE, "r", encoding="utf-8") as f:
            line = f.readline().strip()
            if line:
                bench_sample = json.loads(line)

    if bench_sample:
        instruction_text = bench_sample.get("instruction", "")
        input_text = bench_sample.get("input", "")
        output_text = bench_sample.get("reference_output", "")

        add_code_block(doc, f"instruction: {instruction_text[:120]}...", label="[instruction]")
        # input은 첫 3줄만
        input_lines = input_text.split("\n")
        input_preview = "\n".join(input_lines[:6])
        if len(input_lines) > 6:
            input_preview += "\n..."
        add_code_block(doc, f"input:\n{input_preview}", label="[input]")
        add_code_block(doc, f"output: {output_text}", label="[output]")

    # ── 회의록 샘플 ──
    add_heading_styled(doc, "6-3. 회의록 분석 데이터 (input → output JSON)", level=2)
    meeting = load_meeting_sample()
    if meeting:
        # input 일부
        input_text = meeting.get("input", "")
        input_lines = input_text.split("\n")
        input_preview = "\n".join(input_lines[:8])
        if len(input_lines) > 8:
            input_preview += "\n..."
        add_code_block(doc, input_preview, label="[input — 회의 원문]")

        # output JSON
        output_data = meeting.get("output", {})
        output_json = json.dumps(output_data, ensure_ascii=False, indent=2)
        # 너무 길면 잘라냄
        output_lines = output_json.split("\n")
        if len(output_lines) > 25:
            output_preview = "\n".join(output_lines[:25]) + "\n  ..."
        else:
            output_preview = output_json
        add_code_block(doc, output_preview, label="[output — 구조화 JSON]")

    # ================================================================
    # 섹션 7: 데이터 전처리 시 유의점
    # ================================================================
    add_heading_styled(doc, "7. 데이터 전처리 시 유의해야 할 점", level=1)

    precautions = [
        ("규정 관련 데이터: ", "input에 규정 원문 필수 포함 (RAG 실서비스와 동일 형태로 학습해야 추론 시 성능 보장)"),
        ("동적 필드 방식: ", "input에 '필드:' 줄 포함 필수 — sLLM이 어떤 필드를 채워야 하는지 명시적으로 인식"),
        ("회의록 데이터: ", "output JSON 스키마 정확히 준수 (필드명 불일치 시 학습 불가)"),
        ("리스크 감지: ", "위반/정상 균등 분포 (각 100건씩, 1:1 비율)"),
        ("평가 데이터: ", "전체 데이터의 15%를 평가용으로 분리 (stratified split)"),
        ("JSON 포맷: ", "output 필드의 JSON 구조 엄격 준수 — 키 이름, 중첩 구조, 배열 형태 일관성 유지"),
    ]

    for bold_part, rest in precautions:
        add_bullet_point(doc, rest, bold_prefix=bold_part)

    # ================================================================
    # 섹션 8: 수집 자동화 절차
    # ================================================================
    add_heading_styled(doc, "8. 수집 자동화 절차", level=1)

    add_styled_paragraph(
        doc,
        "본 프로젝트의 데이터 수집은 LLM API 기반 자동 생성과 수동 검수를 조합한 "
        "반자동화(semi-automated) 파이프라인으로 수행됩니다. "
        "아래는 각 단계별 자동화 방법입니다.",
        space_after=Pt(8),
    )

    add_heading_styled(doc, "8-1. LLM API 배치 생성", level=2)
    add_bullet_point(doc,
        "OpenAI GPT-4 / Anthropic Claude API에 카테고리별 프롬프트를 전달하여 대량 생성",
        bold_prefix="방법: ")
    add_bullet_point(doc,
        "Python 스크립트에서 API를 반복 호출하고 결과를 JSONL로 직접 저장",
        bold_prefix="스크립트: ")
    add_bullet_point(doc,
        "카테고리당 200건씩 배치 생성 → 프롬프트에 시드 데이터 20건을 few-shot 예시로 포함",
        bold_prefix="배치 단위: ")
    add_bullet_point(doc,
        "API rate limit 대응을 위한 지수 백오프(exponential backoff) 재시도 로직 적용",
        bold_prefix="에러 처리: ")

    add_heading_styled(doc, "8-2. 데이터 변환 파이프라인", level=2)
    add_bullet_point(doc,
        "규정 판단 데이터: XLSX → JSONL 자동 변환 (pandas read_excel → to_json lines=True)",
        bold_prefix="XLSX→JSONL: ")
    add_bullet_point(doc,
        "회의록 데이터: JSON 배열 → 개별 JSONL 라인 변환 (json.load → 건별 json.dumps)",
        bold_prefix="JSON→JSONL: ")
    add_bullet_point(doc,
        "벤치마크 데이터: 6개 카테고리별 수동 작성 → 단일 JSONL로 병합",
        bold_prefix="병합: ")

    add_heading_styled(doc, "8-3. 품질 검증 자동화", level=2)
    add_bullet_point(doc,
        "JSONL 필드 존재 여부 및 타입 검증 (instruction/input/output 필드 필수 체크)",
        bold_prefix="스키마 검증: ")
    add_bullet_point(doc,
        "텍스트 해시 기반 완전 중복 제거 + 코사인 유사도 기반 유사 중복 탐지",
        bold_prefix="중복 검출: ")
    add_bullet_point(doc,
        "카테고리별 건수 자동 집계 및 불균형 비율 경고 (최소 150건 미만 시 알림)",
        bold_prefix="분포 검사: ")
    add_bullet_point(doc,
        "회의록 output JSON의 키/값 타입 일관성 자동 검증 (jsonschema validate)",
        bold_prefix="JSON 스키마: ")

    add_heading_styled(doc, "8-4. 증강 데이터 자동 생성", level=2)
    add_bullet_point(doc,
        "v1.2 증강: 인터넷 슬랭/초성/축약어 변환 규칙 기반 자동 생성 (300건)",
        bold_prefix="비정형 증강: ")
    add_bullet_point(doc,
        "v1.3 증강: Adversarial 테스트 결과에서 오분류된 패턴을 추출하여 타겟 증강 (163건)",
        bold_prefix="경계 케이스: ")
    add_bullet_point(doc,
        "증강 후 자동으로 train/eval 재분리 (stratified split 15%)",
        bold_prefix="자동 분리: ")

    # ================================================================
    # 섹션 9: 법적 요소 검토
    # ================================================================
    add_heading_styled(doc, "9. 법적 요소 검토", level=1)

    add_styled_paragraph(
        doc,
        "본 프로젝트에서 수집 및 생성한 데이터의 법적 적합성을 아래와 같이 검토하였습니다.",
        space_after=Pt(8),
    )

    add_heading_styled(doc, "9-1. 개인정보보호", level=2)
    add_bullet_point(doc,
        "모든 학습 데이터는 GPT-4 / Claude API로 생성한 합성(synthetic) 데이터로, "
        "실제 개인의 이름, 연락처, 주민등록번호 등 개인정보를 포함하지 않습니다.",
        bold_prefix="합성 데이터: ")
    add_bullet_point(doc,
        "도메인 컨텍스트로 사용된 '듀듀테크놀로지'는 가상의 회사이며, "
        "사내규정 30개 조항 또한 전부 가상으로 작성되어 실제 기업 정보와 무관합니다.",
        bold_prefix="가상 도메인: ")
    add_bullet_point(doc,
        "따라서 「개인정보보호법」 제2조의 개인정보에 해당하지 않으며, "
        "정보주체 동의 절차가 불필요합니다.",
        bold_prefix="법적 판단: ")

    add_heading_styled(doc, "9-2. 저작권 및 API 이용약관", level=2)
    add_bullet_point(doc,
        "OpenAI 이용약관(Terms of Use)에 따르면, API 출력물의 권리는 사용자에게 귀속되며 "
        "상업적 활용이 허용됩니다 (2024.01 기준 Section 3(a)).",
        bold_prefix="OpenAI: ")
    add_bullet_point(doc,
        "Anthropic 이용약관에 따르면, API를 통해 생성된 출력물에 대해 "
        "사용자가 권리를 보유하며 학습 데이터 목적의 활용이 허용됩니다.",
        bold_prefix="Anthropic: ")
    add_bullet_point(doc,
        "외부 공개 데이터셋(HuggingFace, Kaggle 등)은 사용하지 않았으므로, "
        "제3자 저작권 침해 위험이 없습니다.",
        bold_prefix="외부 데이터: ")

    add_heading_styled(doc, "9-3. 데이터 윤리", level=2)
    add_bullet_point(doc,
        "합성 데이터 생성 시 차별적 표현, 혐오 발언, 편향된 판단 기준이 포함되지 않도록 "
        "프롬프트에 명시적 가이드라인을 설정하였습니다.",
        bold_prefix="편향 방지: ")
    add_bullet_point(doc,
        "규정 판단 데이터의 Yes/No/조건부 분포를 균등하게 유지하여 "
        "특정 판단으로의 편향을 방지하였습니다.",
        bold_prefix="라벨 균형: ")
    add_bullet_point(doc,
        "생성된 데이터는 학습 목적으로만 사용되며, 실제 법률 자문이나 "
        "규정 해석의 근거로 활용되지 않습니다.",
        bold_prefix="사용 범위: ")

    # 법적 요소 요약 표
    add_styled_paragraph(doc, "", space_after=Pt(4))
    legal_headers = ["검토 항목", "해당 여부", "근거"]
    legal_rows = [
        ["개인정보 포함", "해당 없음", "전량 합성 데이터, 가상 도메인 사용"],
        ["제3자 저작권 침해", "해당 없음", "외부 데이터셋 미사용, 전량 자체 생성"],
        ["API 이용약관 준수", "준수", "OpenAI/Anthropic TOS 상 출력물 권리 사용자 귀속"],
        ["데이터 편향/윤리", "관리됨", "프롬프트 가이드라인 + 라벨 균형 유지"],
    ]
    create_styled_table(doc, legal_headers, legal_rows, [Cm(3.5), Cm(2.5), Cm(8)])

    # ── 마무리 ──
    doc.add_paragraph()  # 빈 줄
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_footer.add_run("— End of Document —")
    set_font(run, size=SMALL_SIZE, color=RGBColor(0x99, 0x99, 0x99), italic=True)

    return doc


# ── 메인 ───────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("데이터 수집 보고서 생성 시작")
    print("=" * 60)

    print(f"\n프로젝트 루트: {PROJECT_ROOT}")
    print(f"출력 파일: {OUTPUT_FILE}")

    # 데이터 파일 존재 확인
    for label, path in [
        ("Intent JSONL", INTENT_FILE),
        ("회의록 JSON", MEETING_FILE),
        ("벤치마크 JSONL", BENCHMARK_FILE),
    ]:
        status = "OK" if path.exists() else "NOT FOUND"
        print(f"  [{status}] {label}: {path}")

    print("\n문서 빌드 중...")
    doc = build_document()

    print(f"저장 중: {OUTPUT_FILE}")
    doc.save(str(OUTPUT_FILE))

    file_size = OUTPUT_FILE.stat().st_size / 1024
    print(f"\n완료! 파일 크기: {file_size:.1f} KB")
    print("=" * 60)


if __name__ == "__main__":
    main()
