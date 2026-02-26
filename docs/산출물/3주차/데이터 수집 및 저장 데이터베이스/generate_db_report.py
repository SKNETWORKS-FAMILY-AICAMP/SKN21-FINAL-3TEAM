"""
데이터베이스 설계문서 docx 생성 스크립트
SK네트웍스 Family AI과정 21기 최종 프로젝트 3조 — WorkFlow Agent (듀듀)

사용법:
    python generate_db_report.py

출력:
    1.데이터 수집 및 저장_데이터베이스 설계문서_3조.docx
"""

from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 경로 설정 ──────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).resolve().parent
OUTPUT_FILE = SCRIPT_DIR / "1.데이터 수집 및 저장_데이터베이스 설계문서_3조.docx"

# ── 스타일 상수 ────────────────────────────────────────────────
FONT_NAME = "맑은 고딕"

COLOR_DARK_BLUE = RGBColor(0x1F, 0x3A, 0x5F)
COLOR_LIGHT_BLUE = RGBColor(0x44, 0x72, 0xC4)
COLOR_HEADER_BG = "2E75B6"
COLOR_ALT_ROW = "F2F7FB"
COLOR_WHITE = "FFFFFF"
COLOR_BORDER = "B4C6E7"

TITLE_SIZE = Pt(28)
SUBTITLE_SIZE = Pt(14)
HEADING1_SIZE = Pt(18)
HEADING2_SIZE = Pt(14)
BODY_SIZE = Pt(10.5)
SMALL_SIZE = Pt(9)
TABLE_SIZE = Pt(9.5)


# ── 유틸리티 함수 (데이터수집 보고서에서 재활용) ─────────────────

def set_font(run, name=FONT_NAME, size=BODY_SIZE, bold=False, color=None, italic=False):
    """폰트 속성 설정"""
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), name)


def set_cell_shading(cell, color_hex):
    """셀 배경색 설정"""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading)


def set_table_borders(table, color=COLOR_BORDER):
    """테이블 전체 테두리 설정"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def add_styled_paragraph(doc, text, size=BODY_SIZE, bold=False, color=None,
                          alignment=None, space_before=Pt(0), space_after=Pt(6),
                          italic=False, line_spacing=1.3):
    """스타일이 적용된 문단 추가"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    return p


def add_heading_styled(doc, text, level=1, color=COLOR_DARK_BLUE):
    """커스텀 스타일 헤딩"""
    sizes = {1: HEADING1_SIZE, 2: HEADING2_SIZE}
    size = sizes.get(level, BODY_SIZE)
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    pf = p.paragraph_format
    pf.space_before = Pt(18) if level == 1 else Pt(12)
    pf.space_after = Pt(8)
    if level == 1:
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="2E75B6"/>'
            f'</w:pBdr>'
        )
        p._element.get_or_add_pPr().append(pBdr)
    return p


def create_styled_table(doc, headers, rows, col_widths=None, header_color=COLOR_HEADER_BG,
                         left_align_cols=None):
    """스타일이 적용된 표 생성. left_align_cols: 왼쪽 정렬할 열 인덱스 set"""
    if left_align_cols is None:
        left_align_cols = set()
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # 헤더 행
    hdr_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        set_font(run, size=TABLE_SIZE, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
        cell.vertical_alignment = 1

    # 데이터 행
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_font(run, size=TABLE_SIZE)
            if col_idx in left_align_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = 1
            if row_idx % 2 == 1:
                set_cell_shading(cell, COLOR_ALT_ROW)

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = width

    return table


def add_bullet_point(doc, text, level=0, bold_prefix=""):
    """불릿 포인트"""
    p = doc.add_paragraph()
    indent_val = 360 + (level * 360)
    hanging = 200
    pPr = p._element.get_or_add_pPr()
    ind = parse_xml(
        f'<w:ind {nsdecls("w")} w:left="{indent_val}" w:hanging="{hanging}"/>'
    )
    pPr.append(ind)
    bullet_char = "\u2022" if level == 0 else "\u2013"
    if bold_prefix:
        run = p.add_run(f"{bullet_char} {bold_prefix}")
        set_font(run, size=BODY_SIZE, bold=True)
        run2 = p.add_run(text)
        set_font(run2, size=BODY_SIZE)
    else:
        run = p.add_run(f"{bullet_char} {text}")
        set_font(run, size=BODY_SIZE)
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.3
    return p


def add_page_numbers(doc):
    """푸터에 페이지 번호 추가 (- N - 형태)"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_before = p.add_run("- ")
    set_font(run_before, size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_field = p.add_run()
    run_field._element.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_field2 = p.add_run()
    run_field2._element.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_field3 = p.add_run()
    run_field3._element.append(fldChar2)
    set_font(run_field3, size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))
    run_after = p.add_run(" -")
    set_font(run_after, size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))


# ── ERD 이미지 생성 ──────────────────────────────────────────

def generate_erd_image(output_path):
    """matplotlib으로 ERD 다이어그램 생성 — 도메인별 색상, 깔끔한 레이아웃"""

    # 도메인별 색상
    DOMAIN_COLORS = {
        "auth":     ("#2E75B6", "#D6E4F0"),  # 파랑 — 인증
        "doc":      ("#548235", "#E2EFDA"),  # 초록 — 문서
        "meeting":  ("#BF8F00", "#FFF2CC"),  # 황금 — 회의
        "ai":       ("#7030A0", "#E8D5F5"),  # 보라 — AI 판단
        "schedule": ("#C55A11", "#FCE4D6"),  # 주황 — 일정/연동
        "log":      ("#404040", "#E0E0E0"),  # 회색 — 로그
    }

    # 테이블 정보: (이름, 도메인, PK 표시, 설명 — 영문으로 폰트 호환)
    tables_info = {
        "users":                  ("auth",     "PK: id",  "User accounts & auth"),
        "oauth_tokens":           ("auth",     "PK: id",  "Google OAuth tokens (1:1)"),
        "documents":              ("doc",      "PK: id",  "Uploaded documents"),
        "document_templates":     ("doc",      "PK: id",  "Document templates"),
        "meetings":               ("meeting",  "PK: id",  "Meeting notes + AI summary"),
        "action_items":           ("meeting",  "PK: id",  "Action items + Google sync"),
        "regulations":            ("ai",       "PK: id",  "Company regulations (RAG)"),
        "judgments":              ("ai",       "PK: id",  "Regulation judgments"),
        "schedules":              ("schedule", "PK: id",  "Calendar events"),
        "google_sheet_trackers":  ("schedule", "PK: id",  "Google Sheets tracker"),
        "chat_logs":              ("log",      "PK: id",  "Chat conversation logs"),
    }

    # 배치: 4행 — users 중심, 도메인별 그룹핑, 넓은 간격
    # Row 0: users(중앙), oauth_tokens(우측)
    # Row 1: documents(좌), meetings(중), regulations(우)
    # Row 2: document_templates(좌), action_items(중), judgments(우)
    # Row 3: schedules(좌), google_sheet_trackers(중), chat_logs(우)
    positions = {
        "users":                  (7.5,  14),
        "oauth_tokens":           (14,   14),
        "documents":              (1.5,  10.5),
        "meetings":               (7.5,  10.5),
        "regulations":            (13.5, 10.5),
        "document_templates":     (1.5,  7),
        "action_items":           (7.5,  7),
        "judgments":              (13.5, 7),
        "schedules":              (1.5,  3.5),
        "google_sheet_trackers":  (7.5,  3.5),
        "chat_logs":              (13.5, 3.5),
    }

    # FK 관계: (from, to, fk_col, rel_type)
    relations = [
        ("documents",             "users",        "uploaded_by",    "N:1"),
        ("document_templates",    "users",        "uploaded_by",    "N:1"),
        ("meetings",              "users",        "created_by",     "N:1"),
        ("action_items",          "meetings",     "meeting_id",     "N:1"),
        ("action_items",          "users",        "assignee_id",    "N:1"),
        ("schedules",             "users",        "user_id",        "N:1"),
        ("schedules",             "action_items", "action_item_id", "N:1"),
        ("judgments",             "users",        "user_id",        "N:1"),
        ("chat_logs",             "users",        "user_id",        "N:1"),
        ("oauth_tokens",          "users",        "user_id",        "1:1"),
        ("google_sheet_trackers", "users",        "user_id",        "N:1"),
        ("google_sheet_trackers", "meetings",     "meeting_id",     "N:1"),
    ]

    fig, ax = plt.subplots(1, 1, figsize=(16, 14.5))
    ax.set_xlim(-0.5, 16)
    ax.set_ylim(1, 16.5)
    ax.set_aspect("equal")
    ax.axis("off")
    fig.patch.set_facecolor("white")

    BOX_W = 4.0
    BOX_H = 1.8
    CORNER_R = 0.15

    box_rects = {}

    # 박스 그리기
    for tbl_name, (cx, cy) in positions.items():
        domain, pk_text, desc = tables_info[tbl_name]
        header_color, body_color = DOMAIN_COLORS[domain]
        left = cx - BOX_W / 2
        bottom = cy - BOX_H / 2

        # 본체
        body = mpatches.FancyBboxPatch(
            (left, bottom), BOX_W, BOX_H,
            boxstyle=f"round,pad={CORNER_R}",
            facecolor=body_color, edgecolor=header_color, linewidth=1.5
        )
        ax.add_patch(body)

        # 헤더 영역 (상단 절반)
        header_h = BOX_H * 0.45
        header = mpatches.FancyBboxPatch(
            (left, cy + BOX_H / 2 - header_h), BOX_W, header_h,
            boxstyle=f"round,pad={CORNER_R}",
            facecolor=header_color, edgecolor=header_color, linewidth=1.5
        )
        ax.add_patch(header)

        # 테이블명 (크게, 볼드)
        ax.text(cx, cy + BOX_H / 2 - header_h / 2, tbl_name,
                ha="center", va="center", fontsize=10, fontweight="bold",
                color="white", fontfamily="monospace")

        # 설명 (하단)
        ax.text(cx, cy - BOX_H * 0.12, desc,
                ha="center", va="center", fontsize=8,
                color="#333333", fontfamily="sans-serif")

        box_rects[tbl_name] = (left, bottom, left + BOX_W, bottom + BOX_H)

    # 화살표 접점 계산
    def edge_point(rect, tx, ty):
        l, b, r, t = rect
        cx, cy = (l + r) / 2, (b + t) / 2
        dx, dy = tx - cx, ty - cy
        if abs(dx) < 0.001 and abs(dy) < 0.001:
            return cx, t
        hw, hh = (r - l) / 2, (t - b) / 2
        if abs(dx) * hh > abs(dy) * hw:
            s = hw / abs(dx)
        else:
            s = hh / abs(dy)
        return cx + dx * s, cy + dy * s

    # 관계 화살표 그리기
    drawn_offsets = {}
    for from_tbl, to_tbl, fk_col, rel_type in relations:
        r1 = box_rects[from_tbl]
        r2 = box_rects[to_tbl]
        c1 = ((r1[0] + r1[2]) / 2, (r1[1] + r1[3]) / 2)
        c2 = ((r2[0] + r2[2]) / 2, (r2[1] + r2[3]) / 2)

        # 같은 target으로 가는 화살표 오프셋 (겹침 방지)
        key = to_tbl
        offset_idx = drawn_offsets.get(key, 0)
        drawn_offsets[key] = offset_idx + 1
        rad = 0.05 + offset_idx * 0.04
        if offset_idx % 2 == 1:
            rad = -rad

        x1, y1 = edge_point(r1, c2[0], c2[1])
        x2, y2 = edge_point(r2, c1[0], c1[1])

        _, body_color = DOMAIN_COLORS[tables_info[from_tbl][0]]
        header_color = DOMAIN_COLORS[tables_info[from_tbl][0]][0]

        ax.annotate("",
            xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(
                arrowstyle="-|>", color=header_color,
                lw=1.2, shrinkA=2, shrinkB=2,
                connectionstyle=f"arc3,rad={rad}",
                alpha=0.7,
            ))

        # 관계 레이블 (중간점)
        mx = (x1 + x2) / 2 + rad * 2
        my = (y1 + y2) / 2 + rad * 2
        ax.text(mx, my, rel_type, fontsize=6.5, fontweight="bold",
                color=header_color, ha="center", va="center", alpha=0.8,
                bbox=dict(boxstyle="round,pad=0.15", facecolor="white",
                          edgecolor="none", alpha=0.85))

    # 제목
    ax.text(0, 16.2, "ERD: WorkFlow Agent Database (PostgreSQL 16)",
            fontsize=13, fontweight="bold", color="#1F3A5F", fontfamily="sans-serif")

    # 범례 (도메인별)
    legend_x = 0
    legend_y = 15.6
    legend_items = [
        ("auth",     "Auth"),
        ("doc",      "Document"),
        ("meeting",  "Meeting"),
        ("ai",       "AI/RAG"),
        ("schedule", "Schedule"),
        ("log",      "Log"),
    ]
    for i, (domain, label) in enumerate(legend_items):
        hc, bc = DOMAIN_COLORS[domain]
        x = legend_x + i * 2.6
        rect = mpatches.FancyBboxPatch(
            (x, legend_y - 0.15), 0.4, 0.3,
            boxstyle="round,pad=0.03", facecolor=hc, edgecolor=hc
        )
        ax.add_patch(rect)
        ax.text(x + 0.55, legend_y, label, fontsize=8, va="center",
                color="#333333", fontfamily="sans-serif")

    ax.text(0, 15.1, "11 tables  |  12 FK relations  |  Arrows: FK direction (child \u2192 parent)",
            fontsize=8, color="#888888", fontfamily="sans-serif")

    plt.tight_layout(pad=0.3)
    plt.savefig(str(output_path), dpi=200, bbox_inches="tight",
                facecolor="white", edgecolor="none")
    plt.close(fig)
    return output_path


# ── 테이블 정의 데이터 (ORM 모델 기반) ────────────────────────

# 각 테이블: (테이블명, 설명, [(컬럼명, 데이터타입, PK, FK, NotNull, 설명), ...])
# PK/FK/NotNull: "O" 또는 ""
# 공통 컬럼 created_at, updated_at은 모든 테이블에 추가

COMMON_TIMESTAMP_COLS = [
    ("created_at", "TIMESTAMP", "", "", "O", "생성 시각 (DEFAULT now())"),
    ("updated_at", "TIMESTAMP", "", "", "O", "수정 시각 (DEFAULT now())"),
]

TABLES = [
    ("users", "사용자 계정 및 인증 정보", [
        ("id", "INTEGER", "O", "", "O", "PK, AUTO INCREMENT"),
        ("email", "VARCHAR(255)", "", "", "O", "이메일 (UNIQUE, INDEX)"),
        ("hashed_password", "VARCHAR(255)", "", "", "O", "bcrypt 해시 비밀번호"),
        ("name", "VARCHAR(100)", "", "", "O", "사용자 이름"),
        ("is_admin", "BOOLEAN", "", "", "O", "관리자 여부 (DEFAULT false)"),
        ("is_active", "BOOLEAN", "", "", "O", "활성 상태 (DEFAULT true)"),
    ]),
    ("documents", "업로드 문서 메타데이터 및 상태 관리", [
        ("id", "INTEGER", "O", "", "O", "PK, AUTO INCREMENT"),
        ("title", "VARCHAR(500)", "", "", "O", "문서 제목"),
        ("file_path", "VARCHAR(1000)", "", "", "O", "저장 경로"),
        ("file_type", "VARCHAR(20)", "", "", "O", "파일 유형 (pdf/docx/txt)"),
        ("content", "TEXT", "", "", "", "파싱된 텍스트 내용 (NULLABLE)"),
        ("scope", "VARCHAR(10)", "", "", "O", "공개 범위 (DEFAULT 'company')"),
        ("uploaded_by", "INTEGER", "", "users.id", "O", "업로드한 사용자 FK"),
        ("status", "VARCHAR(20)", "", "", "O", "처리 상태 (DEFAULT 'processing')"),
    ]),
    ("document_templates", "문서 양식 템플릿 (기본 4종 + 사용자 커스텀)", [
        ("id", "INTEGER", "O", "", "O", "PK, AUTO INCREMENT"),
        ("name", "VARCHAR(500)", "", "", "O", "템플릿 이름"),
        ("description", "TEXT", "", "", "", "설명 (NULLABLE)"),
        ("file_path", "VARCHAR(1000)", "", "", "", "양식 파일 경로 (NULLABLE)"),
        ("file_type", "VARCHAR(20)", "", "", "", "파일 유형 (NULLABLE)"),
        ("parsed_structure", "TEXT", "", "", "", "AI 추출 양식 구조 JSON (NULLABLE)"),
        ("category", "VARCHAR(50)", "", "", "O", "카테고리 (DEFAULT 'custom')"),
        ("is_system", "BOOLEAN", "", "", "O", "기본 제공 여부 (DEFAULT false)"),
        ("scope", "VARCHAR(10)", "", "", "O", "공개 범위 (DEFAULT 'company')"),
        ("uploaded_by", "INTEGER", "", "users.id", "", "업로드 사용자 FK (NULLABLE)"),
        ("status", "VARCHAR(20)", "", "", "O", "처리 상태 (DEFAULT 'ready')"),
    ]),
    ("regulations", "사내 규정 원본 (RAG 소스)", [
        ("id", "INTEGER", "O", "", "O", "PK, AUTO INCREMENT"),
        ("title", "VARCHAR(500)", "", "", "O", "규정 제목"),
        ("category", "VARCHAR(100)", "", "", "O", "분류 (정보보안/인사/개발 등)"),
        ("article_number", "VARCHAR(50)", "", "", "O", "조항 번호"),
        ("content", "TEXT", "", "", "O", "규정 원문 전체"),
        ("version", "VARCHAR(20)", "", "", "O", "버전 (DEFAULT '1.0')"),
    ]),
    ("meetings", "회의록 원본 및 AI 분석 결과", [
        ("id", "INTEGER", "O", "", "O", "PK, AUTO INCREMENT"),
        ("title", "VARCHAR(500)", "", "", "O", "회의 제목"),
        ("raw_content", "TEXT", "", "", "O", "원본 회의록 텍스트"),
        ("summary", "TEXT", "", "", "", "AI 요약 (NULLABLE)"),
        ("decisions", "JSONB", "", "", "", "결정사항 JSON 배열 (NULLABLE)"),
        ("risk_level", "VARCHAR(20)", "", "", "", "리스크 수준 (NULLABLE)"),
        ("meeting_date", "TIMESTAMP", "", "", "", "회의 일시 (NULLABLE)"),
        ("created_by", "INTEGER", "", "users.id", "O", "작성자 FK"),
    ]),
    ("action_items", "회의 결정사항 실행 항목 + Google 서비스 연동", [
        ("id", "INTEGER", "O", "", "O", "PK, AUTO INCREMENT"),
        ("meeting_id", "INTEGER", "", "meetings.id", "", "회의 FK (NULLABLE)"),
        ("content", "VARCHAR(1000)", "", "", "O", "실행 항목 내용"),
        ("assignee", "VARCHAR(100)", "", "", "", "담당자 표시명 (NULLABLE)"),
        ("assignee_id", "INTEGER", "", "users.id", "", "내부 사용자 FK (NULLABLE)"),
        ("due_date", "TIMESTAMP", "", "", "", "마감일 (NULLABLE)"),
        ("priority", "VARCHAR(20)", "", "", "O", "우선순위 (DEFAULT 'medium')"),
        ("status", "VARCHAR(20)", "", "", "O", "상태 (DEFAULT 'pending')"),
        ("google_task_id", "VARCHAR(255)", "", "", "", "Google Tasks ID (NULLABLE)"),
        ("sheet_row_id", "INTEGER", "", "", "", "Google Sheets 행 번호 (NULLABLE)"),
        ("email_sent_at", "TIMESTAMP", "", "", "", "Gmail 발송 시각 (NULLABLE)"),
    ]),
    ("schedules", "캘린더 일정 + Google Calendar/Meet 연동", [
        ("id", "INTEGER", "O", "", "O", "PK, AUTO INCREMENT"),
        ("title", "VARCHAR(500)", "", "", "O", "일정 제목"),
        ("description", "VARCHAR(2000)", "", "", "", "상세 설명 (NULLABLE)"),
        ("start_time", "TIMESTAMP", "", "", "O", "시작 시간"),
        ("end_time", "TIMESTAMP", "", "", "", "종료 시간 (NULLABLE)"),
        ("schedule_type", "VARCHAR(50)", "", "", "O", "유형 (meeting/task/deadline)"),
        ("priority", "VARCHAR(20)", "", "", "O", "우선순위 (DEFAULT 'medium')"),
        ("google_event_id", "VARCHAR(255)", "", "", "", "Google Calendar ID (NULLABLE)"),
        ("google_meet_link", "VARCHAR(500)", "", "", "", "Google Meet 링크 (NULLABLE)"),
        ("action_item_id", "INTEGER", "", "action_items.id", "", "실행항목 FK (NULLABLE)"),
        ("user_id", "INTEGER", "", "users.id", "O", "소유자 FK"),
    ]),
    ("judgments", "규정 판단 Agent 결과 저장", [
        ("id", "INTEGER", "O", "", "O", "PK, AUTO INCREMENT"),
        ("question", "TEXT", "", "", "O", "사용자 질문 원문"),
        ("result", "VARCHAR(30)", "", "", "O", "판단 결과 (yes/no/conditional/no_regulation)"),
        ("confidence", "FLOAT", "", "", "O", "신뢰도 (0~1)"),
        ("reasoning", "TEXT", "", "", "O", "판단 근거"),
        ("conditions", "TEXT", "", "", "", "조건부 판단 시 조건 (NULLABLE)"),
        ("alternatives", "TEXT", "", "", "", "대안 제시 (NULLABLE)"),
        ("regulations_cited", "JSONB", "", "", "O", "참조 규정 JSON 배열 (GIN INDEX 고려)"),
        ("user_id", "INTEGER", "", "users.id", "O", "질문자 FK"),
    ]),
    ("chat_logs", "챗봇 대화 이력 (세션 기반)", [
        ("id", "INTEGER", "O", "", "O", "PK, AUTO INCREMENT"),
        ("session_id", "VARCHAR(50)", "", "", "O", "대화 세션 UUID (INDEX)"),
        ("user_id", "INTEGER", "", "users.id", "O", "사용자 FK"),
        ("user_message", "TEXT", "", "", "O", "사용자 입력 메시지"),
        ("intent", "VARCHAR(50)", "", "", "O", "분류된 의도"),
        ("intent_confidence", "FLOAT", "", "", "O", "Intent 신뢰도"),
        ("agent_type", "VARCHAR(50)", "", "", "O", "처리 Agent (judgment/document/schedule)"),
        ("agent_response", "TEXT", "", "", "O", "Agent 응답"),
        ("response_time_ms", "INTEGER", "", "", "", "응답 시간 ms (NULLABLE)"),
    ]),
    ("oauth_tokens", "Google OAuth 2.0 토큰 (AES-256 암호화 저장)", [
        ("id", "INTEGER", "O", "", "O", "PK, AUTO INCREMENT"),
        ("user_id", "INTEGER", "", "users.id", "O", "사용자 FK (UNIQUE)"),
        ("provider", "VARCHAR(50)", "", "", "O", "OAuth 제공자 (google)"),
        ("access_token", "TEXT", "", "", "O", "액세스 토큰 (AES-256 암호화)"),
        ("refresh_token", "TEXT", "", "", "", "리프레시 토큰 (AES-256 암호화, NULLABLE)"),
        ("expires_at", "TIMESTAMP", "", "", "", "토큰 만료 시각 (NULLABLE)"),
        ("scopes", "TEXT", "", "", "", "인가 범위 (콤마 구분, NULLABLE)"),
    ]),
    ("google_sheet_trackers", "Google Sheets 스프레드시트 연동 추적", [
        ("id", "INTEGER", "O", "", "O", "PK, AUTO INCREMENT"),
        ("user_id", "INTEGER", "", "users.id", "O", "사용자 FK"),
        ("spreadsheet_id", "VARCHAR(255)", "", "", "O", "스프레드시트 ID"),
        ("spreadsheet_url", "VARCHAR(500)", "", "", "O", "스프레드시트 URL"),
        ("sheet_name", "VARCHAR(255)", "", "", "O", "시트 이름 (DEFAULT 'Action Items')"),
        ("meeting_id", "INTEGER", "", "meetings.id", "", "회의 FK (NULLABLE)"),
    ]),
]

# FK 관계 정의서 데이터
RELATIONSHIPS = [
    ("fk_documents_users", "documents", "users", "N:1", "uploaded_by → users.id"),
    ("fk_document_templates_users", "document_templates", "users", "N:1", "uploaded_by → users.id (NULLABLE)"),
    ("fk_meetings_users", "meetings", "users", "N:1", "created_by → users.id"),
    ("fk_action_items_meetings", "action_items", "meetings", "N:1", "meeting_id → meetings.id (NULLABLE)"),
    ("fk_action_items_users", "action_items", "users", "N:1", "assignee_id → users.id (NULLABLE)"),
    ("fk_schedules_users", "schedules", "users", "N:1", "user_id → users.id"),
    ("fk_schedules_action_items", "schedules", "action_items", "N:1", "action_item_id → action_items.id (NULLABLE)"),
    ("fk_judgments_users", "judgments", "users", "N:1", "user_id → users.id"),
    ("fk_chat_logs_users", "chat_logs", "users", "N:1", "user_id → users.id"),
    ("fk_oauth_tokens_users", "oauth_tokens", "users", "1:1", "user_id → users.id (UNIQUE)"),
    ("fk_google_sheet_trackers_users", "google_sheet_trackers", "users", "N:1", "user_id → users.id"),
    ("fk_google_sheet_trackers_meetings", "google_sheet_trackers", "meetings", "N:1", "meeting_id → meetings.id (NULLABLE)"),
]


# ── 문서 생성 ──────────────────────────────────────────────────

def build_document():
    doc = Document()

    # 페이지 설정
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 기본 스타일
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    add_page_numbers(doc)

    # ================================================================
    # 섹션 1: 헤더 / 표지
    # ================================================================
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_logo.add_run("SK networks")
    set_font(run, size=Pt(11), bold=True, color=RGBColor(0xE5, 0x00, 0x2B))
    run2 = p_logo.add_run("  |  Family AI Camp")
    set_font(run2, size=Pt(11), color=RGBColor(0x66, 0x66, 0x66))

    # 구분선
    p_line = doc.add_paragraph()
    pPr = p_line._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="E5002B"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    doc.add_paragraph()

    # 제목
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("데이터베이스 설계문서")
    set_font(run, size=TITLE_SIZE, bold=True, color=COLOR_DARK_BLUE)
    p_title.paragraph_format.space_after = Pt(8)

    # 부제
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("SKN Family AI Camp 21기 : 최종 프로젝트 3조")
    set_font(run, size=SUBTITLE_SIZE, color=RGBColor(0x66, 0x66, 0x66))
    p_sub.paragraph_format.space_after = Pt(4)

    # 프로젝트명
    p_proj = doc.add_paragraph()
    p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_proj.add_run("WorkFlow Agent (듀듀)")
    set_font(run, size=Pt(12), color=RGBColor(0x44, 0x72, 0xC4), italic=True)
    p_proj.paragraph_format.space_after = Pt(16)

    doc.add_paragraph()

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_date.add_run("작성일: 2026-02-19")
    set_font(run, size=Pt(10), color=RGBColor(0x55, 0x55, 0x55))
    p_date.paragraph_format.space_after = Pt(4)

    p_team = doc.add_paragraph()
    p_team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_team.add_run("작성: 3조 \u2014 신지용(PM), 윤경은(AI), 진승언(AI), 안혜빈(Backend), 문지영(Frontend)")
    set_font(run, size=Pt(9.5), color=RGBColor(0x55, 0x55, 0x55))
    p_team.paragraph_format.space_after = Pt(24)

    # ================================================================
    # 섹션 2: 프로젝트 개요
    # ================================================================
    add_heading_styled(doc, "1. 프로젝트 개요", level=1)

    add_styled_paragraph(
        doc,
        "WorkFlow Agent(듀듀)는 LangGraph 기반 사내 업무 자동화 멀티에이전트 시스템입니다. "
        "Intent Classifier가 사용자 의도를 7개 카테고리(규정 판단, 문서 검색, 문서 생성, 회의록 생성, "
        "일정 등록, 일정 조회, 일반 대화)로 분류하고, 오케스트레이터가 전문 에이전트(Judgment / Document / Schedule)로 "
        "라우팅합니다. 데이터베이스는 PostgreSQL 16 기반 11개 테이블로 구성되며, "
        "비동기 ORM(SQLAlchemy 2.0 + asyncpg)을 통해 고성능 동시 처리를 지원합니다.",
        size=BODY_SIZE, space_after=Pt(12),
    )

    # 기술 스택 요약표
    add_heading_styled(doc, "1-1. 기술 스택 요약", level=2)
    tech_headers = ["항목", "기술"]
    tech_rows = [
        ["DBMS", "PostgreSQL 16"],
        ["ORM", "SQLAlchemy 2.0+ (async)"],
        ["Driver", "asyncpg"],
        ["Migration", "Alembic"],
        ["벡터 DB", "ChromaDB (RAG용, 별도 관리)"],
    ]
    create_styled_table(doc, tech_headers, tech_rows, [Cm(4), Cm(10)])

    add_styled_paragraph(doc, "", space_after=Pt(8))

    # 에이전트-테이블 매핑표
    add_heading_styled(doc, "1-2. 에이전트-테이블 매핑", level=2)
    agent_headers = ["컴포넌트", "역할", "READ", "WRITE"]
    agent_rows = [
        ["Orchestrator", "Intent 라우팅", "\u2014", "chat_logs"],
        ["Judgment Agent", "규정 판단", "regulations", "judgments"],
        ["Document Agent", "문서/회의록", "documents,\ndocument_templates", "meetings"],
        ["Schedule Agent", "일정/할일", "oauth_tokens,\naction_items", "schedules, action_items,\ngoogle_sheet_trackers"],
        ["Auth Service", "인증", "users, oauth_tokens", "users, oauth_tokens"],
    ]
    create_styled_table(doc, agent_headers, agent_rows,
                        [Cm(3), Cm(2.5), Cm(4.5), Cm(5.5)],
                        left_align_cols={2, 3})

    # ================================================================
    # 섹션 3: ERD — 다이어그램 + 엔터티 목록표
    # ================================================================
    add_heading_styled(doc, "2. ERD \u2014 엔터티 관계 다이어그램", level=1)

    # ERD 이미지 생성 및 삽입
    erd_path = SCRIPT_DIR / "erd_diagram.png"
    print("  ERD 다이어그램 생성 중...")
    generate_erd_image(erd_path)
    p_erd = doc.add_paragraph()
    p_erd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_erd.add_run().add_picture(str(erd_path), width=Cm(15))
    p_erd.paragraph_format.space_after = Pt(8)

    add_styled_paragraph(
        doc,
        "아래 표는 시스템을 구성하는 11개 엔터티(테이블)의 목록과 주요 역할입니다. "
        "상세한 FK 관계는 섹션 4(관계 정의서)에서 다룹니다.",
        space_after=Pt(8),
    )

    entity_headers = ["#", "테이블명", "PK", "주요 역할"]
    entity_rows = [
        ["1", "users", "id", "사용자 계정, 인증 정보, 관리자 구분"],
        ["2", "documents", "id", "업로드 문서 메타데이터 및 처리 상태"],
        ["3", "document_templates", "id", "문서 양식 템플릿 (기본 4종 + 커스텀)"],
        ["4", "regulations", "id", "사내 규정 원본 (RAG 검색 소스)"],
        ["5", "meetings", "id", "회의록 원본 + AI 요약/결정사항"],
        ["6", "action_items", "id", "회의 실행항목 + Google 서비스 연동"],
        ["7", "schedules", "id", "캘린더 일정 + Google Calendar/Meet"],
        ["8", "judgments", "id", "규정 판단 Agent 결과 (Yes/No/조건부)"],
        ["9", "chat_logs", "id", "챗봇 대화 이력 (세션 기반)"],
        ["10", "oauth_tokens", "id", "Google OAuth 2.0 토큰 (암호화)"],
        ["11", "google_sheet_trackers", "id", "Google Sheets 스프레드시트 추적"],
    ]
    create_styled_table(doc, entity_headers, entity_rows,
                        [Cm(1), Cm(4), Cm(1.5), Cm(9)],
                        left_align_cols={1, 3})

    # ================================================================
    # 섹션 4: 테이블 정의서 — 가로(Landscape) 섹션
    # ================================================================
    # 가로 섹션 전환
    landscape_sec = doc.add_section(WD_ORIENT.LANDSCAPE)
    landscape_sec.orientation = WD_ORIENT.LANDSCAPE
    landscape_sec.page_width = Cm(29.7)
    landscape_sec.page_height = Cm(21)
    landscape_sec.top_margin = Cm(2)
    landscape_sec.bottom_margin = Cm(2)
    landscape_sec.left_margin = Cm(2)
    landscape_sec.right_margin = Cm(2)

    add_heading_styled(doc, "3. 테이블 정의서", level=1)

    add_styled_paragraph(
        doc,
        "11개 테이블의 상세 컬럼 정의입니다. 모든 테이블은 created_at / updated_at "
        "타임스탬프 컬럼(DEFAULT now())을 공통으로 포함하며, 아래 표에서는 생략합니다. "
        "설명 컬럼에 UNIQUE, INDEX, DEFAULT 등 제약조건을 함께 기재합니다.",
        space_after=Pt(4),
    )

    # 통합 표 데이터 구성 (타임스탬프 컬럼 생략)
    COMPACT_SIZE = Pt(8)
    all_rows = []
    merge_ranges = []  # (start_row, end_row, tbl_name)
    current_row = 0
    for tbl_name, _tbl_desc, columns in TABLES:
        start = current_row
        for col_name, col_type, pk, fk, not_null, desc in columns:
            all_rows.append([col_name, col_type, pk, fk, not_null, desc])
            current_row += 1
        merge_ranges.append((start, current_row - 1, tbl_name))

    # 통합 표 생성
    table_def_headers = ["\ud14c\uc774\ube14\uba85", "\ucee8\ub7fc\uba85", "\ub370\uc774\ud130 \ud0c0\uc785", "PK", "FK", "NN", "\uc124\uba85"]
    ncols = len(table_def_headers)
    tbl = doc.add_table(rows=1 + len(all_rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)

    # 셀 여백 최소화 (상하 0.5pt, 좌우 2pt)
    tblPr = tbl._tbl.tblPr
    cell_mar = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'  <w:top w:w="10" w:type="dxa"/>'
        f'  <w:left w:w="40" w:type="dxa"/>'
        f'  <w:bottom w:w="10" w:type="dxa"/>'
        f'  <w:right w:w="40" w:type="dxa"/>'
        f'</w:tblCellMar>'
    )
    existing_mar = tblPr.find(qn("w:tblCellMar"))
    if existing_mar is not None:
        tblPr.remove(existing_mar)
    tblPr.append(cell_mar)

    # 헤더 행
    for i, h in enumerate(table_def_headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=COMPACT_SIZE, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        set_cell_shading(cell, COLOR_HEADER_BG)
        cell.vertical_alignment = 1

    # 데이터 행 (열 0은 병합 후 설정)
    for row_idx, row_data in enumerate(all_rows):
        row = tbl.rows[row_idx + 1]
        for col_idx in range(1, ncols):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(row_data[col_idx - 1]))
            set_font(run, size=COMPACT_SIZE)
            if col_idx in {1, 2, 4, 6}:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            cell.vertical_alignment = 1

    # 테이블명 셀 병합 + 교대 그룹 배경색
    group_bg = ["FFFFFF", "F2F7FB"]
    for g_idx, (start, end, tbl_name) in enumerate(merge_ranges):
        merged = tbl.cell(start + 1, 0)
        if start != end:
            merged = merged.merge(tbl.cell(end + 1, 0))
        merged.text = ""
        p = merged.paragraphs[0]
        run = p.add_run(tbl_name)
        set_font(run, size=COMPACT_SIZE, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        merged.vertical_alignment = 1
        bg = group_bg[g_idx % 2]
        if bg != "FFFFFF":
            set_cell_shading(merged, bg)
            for r in range(start + 1, end + 2):
                for c in range(1, ncols):
                    set_cell_shading(tbl.cell(r, c), bg)

    # 열 너비 (가로 모드: 25.7cm 사용 가능)
    col_widths = [Cm(3), Cm(3), Cm(2.5), Cm(0.8), Cm(2.8), Cm(0.8), Cm(12.8)]
    for row in tbl.rows:
        for i, w in enumerate(col_widths):
            if i < len(row.cells):
                row.cells[i].width = w

    # 세로(Portrait) 섹션 복원
    portrait_sec = doc.add_section(WD_ORIENT.PORTRAIT)
    portrait_sec.orientation = WD_ORIENT.PORTRAIT
    portrait_sec.page_width = Cm(21)
    portrait_sec.page_height = Cm(29.7)
    portrait_sec.top_margin = Cm(2.5)
    portrait_sec.bottom_margin = Cm(2.5)
    portrait_sec.left_margin = Cm(2.5)
    portrait_sec.right_margin = Cm(2.5)

    # ================================================================
    # 섹션 5: 관계 정의서
    # ================================================================
    add_heading_styled(doc, "4. 관계 정의서", level=1)

    add_styled_paragraph(
        doc,
        f"아래 표는 시스템 내 {len(RELATIONSHIPS)}개 FK 관계를 정리한 것입니다. "
        "users 테이블이 중심 엔터티로서 대부분의 테이블과 N:1 관계를 맺고 있으며, "
        "oauth_tokens만 1:1(UNIQUE) 관계입니다.",
        space_after=Pt(8),
    )

    rel_headers = ["관계명", "테이블 1", "테이블 2", "관계 유형", "비고"]
    rel_rows = []
    for rel_name, tbl1, tbl2, rel_type, note in RELATIONSHIPS:
        rel_rows.append([rel_name, tbl1, tbl2, rel_type, note])

    create_styled_table(doc, rel_headers, rel_rows,
                        [Cm(4.5), Cm(3), Cm(2.5), Cm(1.5), Cm(4.5)],
                        left_align_cols={0, 4})

    # ================================================================
    # 섹션 6: 샘플 데이터
    # ================================================================
    add_heading_styled(doc, "5. 샘플 데이터", level=1)

    add_styled_paragraph(
        doc,
        "핵심 3개 테이블(users, meetings, action_items)의 가상 샘플 데이터입니다. "
        "듀듀테크놀로지 도메인을 기반으로 작성되었습니다.",
        space_after=Pt(8),
    )

    # users 샘플
    add_heading_styled(doc, "5-1. users", level=2)
    users_headers = ["id", "email", "name", "is_admin", "is_active"]
    users_rows = [
        ["1", "admin@dudu-tech.com", "김철수", "true", "true"],
        ["2", "yhlee@dudu-tech.com", "이영희", "false", "true"],
    ]
    create_styled_table(doc, users_headers, users_rows,
                        [Cm(1), Cm(4.5), Cm(2), Cm(2), Cm(2)])

    add_styled_paragraph(doc, "", space_after=Pt(6))

    # meetings 샘플
    add_heading_styled(doc, "5-2. meetings", level=2)
    meetings_headers = ["id", "title", "summary", "risk_level", "created_by"]
    meetings_rows = [
        ["1", "풀스택 개발자 채용 회의",
         "신규 풀스택 3명 채용 결정,\n기술면접 2단계 확정",
         "중간", "1"],
        ["2", "AI 서비스 런칭 검토",
         "2분기 내부 베타 런칭,\n보안 감사 선행 필요",
         "높음", "1"],
    ]
    create_styled_table(doc, meetings_headers, meetings_rows,
                        [Cm(1), Cm(3.5), Cm(5), Cm(1.5), Cm(2)],
                        left_align_cols={1, 2})

    add_styled_paragraph(doc, "", space_after=Pt(6))

    # action_items 샘플
    add_heading_styled(doc, "5-3. action_items", level=2)
    ai_headers = ["id", "meeting_id", "content", "assignee", "due_date", "status"]
    ai_rows = [
        ["1", "1", "채용 공고 작성 및 게시", "이영희", "2026-03-01", "pending"],
        ["2", "2", "보안 감사 일정 확정", "김철수", "2026-02-28", "done"],
    ]
    create_styled_table(doc, ai_headers, ai_rows,
                        [Cm(1), Cm(1.5), Cm(4), Cm(2), Cm(2.5), Cm(2)],
                        left_align_cols={2, 3})

    # ================================================================
    # 섹션 7: 운영 고려사항
    # ================================================================
    add_heading_styled(doc, "6. 운영 고려사항", level=1)

    # 쿼리 성능
    add_heading_styled(doc, "6-1. 쿼리 성능", level=2)
    add_bullet_point(doc, "users.email UNIQUE INDEX, chat_logs.session_id INDEX, 모든 FK 컬럼 자동 인덱스",
                     bold_prefix="인덱스: ")
    add_bullet_point(doc, "judgments.regulations_cited, meetings.decisions JSONB 컬럼에 GIN 인덱스 고려",
                     bold_prefix="JSONB: ")
    add_bullet_point(doc, "asyncpg 비동기 드라이버 + async_sessionmaker로 동시성 확보",
                     bold_prefix="비동기 처리: ")
    add_bullet_point(doc, "joinedload / selectinload 옵션으로 N+1 쿼리 문제 방지",
                     bold_prefix="N+1 방지: ")

    add_styled_paragraph(doc, "", space_after=Pt(4))

    # 데이터 보안
    add_heading_styled(doc, "6-2. 데이터 보안", level=2)
    add_bullet_point(doc, "bcrypt 해싱 (72바이트 절단 처리, passlib.CryptContext 사용)",
                     bold_prefix="비밀번호: ")
    add_bullet_point(doc, "Fernet(AES-256) 대칭 암호화로 access_token / refresh_token 저장",
                     bold_prefix="OAuth 토큰: ")
    add_bullet_point(doc, "HS256 알고리즘, 60분 만료, PyJWT 기반",
                     bold_prefix="JWT: ")
    add_bullet_point(doc, "DATABASE_URL, ENCRYPTION_KEY 등 환경변수(.env) 기반 관리",
                     bold_prefix="DB 접속 정보: ")

    add_styled_paragraph(doc, "", space_after=Pt(4))

    # 민감 컬럼 표
    add_styled_paragraph(doc, "민감 컬럼 보호 현황", size=BODY_SIZE, bold=True,
                         color=COLOR_LIGHT_BLUE, space_after=Pt(4))
    sec_headers = ["테이블", "컬럼", "보호 방식"]
    sec_rows = [
        ["users", "hashed_password", "bcrypt (salt 포함, 72바이트 절단)"],
        ["users", "email", "UNIQUE 제약 + 로그 마스킹"],
        ["oauth_tokens", "access_token", "AES-256 (Fernet) 암호화"],
        ["oauth_tokens", "refresh_token", "AES-256 (Fernet) 암호화"],
    ]
    create_styled_table(doc, sec_headers, sec_rows,
                        [Cm(3), Cm(3.5), Cm(8)],
                        left_align_cols={0, 1, 2})

    add_styled_paragraph(doc, "", space_after=Pt(8))

    # 변경 이력
    add_heading_styled(doc, "6-3. 변경 이력 (Alembic 마이그레이션)", level=2)

    add_styled_paragraph(
        doc,
        "Alembic을 통한 스키마 버전 관리 이력입니다. 모든 변경은 revision ID로 추적됩니다.",
        space_after=Pt(8),
    )

    hist_headers = ["날짜", "버전", "변경 내용", "Revision", "비고"]
    hist_rows = [
        ["2026-02-11", "v1.0", "초기 11개 테이블 생성", "77cfec3c68a0", "Initial migration"],
        ["2026-02-11", "v1.1",
         "session_id 추가 (chat_logs),\nassignee_id 추가 (action_items),\nTEXT\u2192JSONB 변환 (decisions, regulations_cited)",
         "ff4b6e2ab2e5", "스키마 보완"],
    ]
    create_styled_table(doc, hist_headers, hist_rows,
                        [Cm(2.5), Cm(1.5), Cm(6), Cm(2.5), Cm(3)],
                        left_align_cols={2, 4})

    # ── 마무리 ──
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_footer.add_run("\u2014 End of Document \u2014")
    set_font(run, size=SMALL_SIZE, color=RGBColor(0x99, 0x99, 0x99), italic=True)

    return doc


# ── 메인 ───────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("데이터베이스 설계문서 생성 시작")
    print("=" * 60)

    print(f"\n출력 파일: {OUTPUT_FILE}")

    print("\n문서 빌드 중...")
    doc = build_document()

    print(f"저장 중: {OUTPUT_FILE}")
    doc.save(str(OUTPUT_FILE))

    file_size = OUTPUT_FILE.stat().st_size / 1024
    print(f"\n완료! 파일 크기: {file_size:.1f} KB")
    print("=" * 60)


if __name__ == "__main__":
    main()
