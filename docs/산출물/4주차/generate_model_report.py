import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

def shade_cell(cell, color="D9E2F3"):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_title(text, size=24):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_heading(text, size=13):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    p.paragraph_format.space_before = Pt(14)
    return p

def add_subheading(text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x2E, 0x4A, 0x7A)
    p.paragraph_format.space_before = Pt(8)
    return p

def add_body(text, size=10):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)
    return p

def add_code(text, size=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
    run._element.get_or_add_rPr().append(shading)
    return p

def add_note(text, size=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = True
    return p

def make_table(headers, rows, header_color="1F3864"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(cell, header_color)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                shade_cell(cell, "EDF2FA")
    doc.add_paragraph()
    return table


# =============================================
# COVER PAGE
# =============================================
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('SK networks  |  Family AI Camp')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_title('데이터 전처리', size=26)
add_title('학습된 인공지능 모델', size=18)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('SKN Family AI Camp 21기 : 최종 프로젝트 3팀')
run.font.size = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
run = p.add_run('WorkFlow Agent (듀듀)')
run.font.size = Pt(14)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('작성: 3조 -- 신지용(PM), 윤경은(AI), 진승언(AI), 안혜빈(Backend), 문지영(Frontend)')
run.font.size = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# =============================================
# 1. 모델 개요
# =============================================
add_heading('1. 모델 개요')

add_body(
    '사용자의 자연어 입력을 8개 intent로 분류하여 적합한 Agent로 라우팅하는 '
    'Intent Classifier 모델입니다.'
)

make_table(
    ['항목', '내용'],
    [
        ['모델명', 'Intent Classifier v2 (Stage 6)'],
        ['Base Model', 'monologg/koelectra-base-v3-discriminator'],
        ['아키텍처', 'ElectraForSequenceClassification'],
        ['파라미터', '112.9M'],
        ['학습 방식', 'Full Fine-tuning + Label Smoothing 0.1'],
        ['학습 데이터', 'Train 2,425건 / Val 285건 / Test 286건'],
        ['모델 파일', 'ai/models/intent_classifier/model.safetensors (431MB)'],
    ]
)

# =============================================
# 2. 모델 구조 (양식 Table 0 대응)
# =============================================
add_heading('2. 모델 구조')

make_table(
    ['레이어', '구성', '상세'],
    [
        ['Embedding',
         'Token Embedding\n+ Position Embedding\n+ Type Embedding',
         'vocab_size: 35,000\nembedding_dim: 768\nmax_position: 512'],
        ['Encoder\n(Transformer x 12)',
         'Multi-Head Self-Attention\n+ Feed-Forward Network\n+ Layer Norm',
         'attention_heads: 12\nhidden_size: 768\nintermediate_size: 3,072\nactivation: gelu\ndropout: 0.1'],
        ['Classification Head',
         'Dropout(0.1)\n+ Dense(768 -> 8)\n+ Softmax',
         'num_labels: 8\nloss: CrossEntropy\n+ Label Smoothing 0.1'],
    ]
)

add_body(
    'ELECTRA의 RTD(Replaced Token Detection) 사전학습 방식을 사용하여, '
    '토큰 간 미세한 차이를 구분하는 능력이 BERT 대비 우수합니다. '
    '입력 텍스트를 WordPiece 토크나이저로 분할한 뒤 12층 Transformer Encoder를 거쳐, '
    '[CLS] 토큰의 768차원 출력을 8개 intent 확률로 변환합니다.'
)

# =============================================
# 3. 학습 설정
# =============================================
add_heading('3. 학습 설정')

make_table(
    ['하이퍼파라미터', '값'],
    [
        ['Epochs', '10'],
        ['Batch Size', '16'],
        ['Learning Rate', '3e-5'],
        ['Optimizer', 'AdamW'],
        ['Loss', 'CrossEntropyLoss + Label Smoothing 0.1'],
        ['Max Length', '64 tokens'],
        ['Seed', '42 (random/numpy/torch/cuda 4중 고정)'],
        ['학습 환경', 'RunPod RTX 4090 (24GB)'],
    ]
)

add_note('* Grid Search 32-point (epochs {3,5,7,10} x lr {1e-5,2e-5,3e-5,5e-5} x batch {16,32}) 결과 선정')

# =============================================
# 4. 학습 결과 (양식 Table 1 대응)
# =============================================
add_heading('4. 학습 결과')

add_subheading('4-1. 전체 성능 메트릭')

make_table(
    ['메트릭', 'Test (286건)', 'Adversarial (450건)'],
    [
        ['Accuracy', '97.90%', '87.56%'],
        ['Precision (Macro)', '97.88%', '87.92%'],
        ['Recall (Macro)', '97.90%', '87.68%'],
        ['F1 Score (Macro)', '97.88%', '87.58%'],
    ]
)

add_subheading('4-2. Intent별 성능 (Adversarial 기준)')

make_table(
    ['Intent', 'Precision', 'Recall', 'F1 Score'],
    [
        ['judgment', '98.2%', '89.8%', '93.8%'],
        ['doc_search', '81.8%', '90.0%', '85.7%'],
        ['doc_generate', '90.2%', '88.5%', '89.3%'],
        ['doc_summary', '89.3%', '94.3%', '91.7%'],
        ['doc_qa', '85.4%', '69.5%', '76.6%'],
        ['schedule_add', '96.4%', '94.6%', '95.5%'],
        ['schedule_view', '78.5%', '91.1%', '84.3%'],
        ['general', '83.6%', '83.6%', '83.6%'],
    ]
)

add_note(
    '* doc_qa F1 76.6%: 라벨 리뷰 결과 오류의 63%가 라벨 문제(Adjusted F1 ~85%). '
    'doc 4개 intent 모두 동일 Document Agent로 라우팅되어 실서비스 영향 제한적.'
)

add_subheading('4-3. 과신뢰 해소 효과 (Label Smoothing)')

make_table(
    ['항목', '적용 전', '적용 후', '변화'],
    [
        ['오분류 중 과신뢰(>90%)', '42건 (66.7%)', '13건 (23.2%)', '-69%'],
        ['오답 confidence 중앙값', '~0.90', '~0.64', '정답/오답 분리 가능'],
    ]
)

add_body(
    'Threshold 0.85 기준으로 정답/오답 분리 가능. '
    '오분류 시 잘못된 Agent 대신 clarify(top-3 후보 재질문) 응답 제공.'
)

doc.add_page_break()

# =============================================
# 5. 모델 파일 및 저장/로드 (양식 Table 2 대응)
# =============================================
add_heading('5. 모델 파일 및 저장/로드')

add_subheading('5-1. 모델 파일 구성')

make_table(
    ['항목', '내용'],
    [
        ['모델 파일', 'model.safetensors (430.8 MB)'],
        ['저장 형식', 'safetensors (HuggingFace 표준)'],
        ['저장 경로', 'ai/models/intent_classifier/'],
        ['설정 파일', 'config.json (아키텍처 + id2label 매핑)'],
        ['토크나이저', 'tokenizer.json + tokenizer_config.json\n(WordPiece, vocab 35,000)'],
        ['라벨 매핑', 'label_map.json (id2label / label2id)'],
        ['메타정보', 'model_info.json\n(base_model, experiment, label_smoothing)'],
    ]
)

add_subheading('5-2. 저장 코드')

add_code(
    '# 학습 완료 후 모델 저장\n'
    'model.save_pretrained("ai/models/intent_classifier")\n'
    'tokenizer.save_pretrained("ai/models/intent_classifier")\n'
    '\n'
    '# 메타정보 저장\n'
    'import json\n'
    'with open("ai/models/intent_classifier/model_info.json", "w") as f:\n'
    '    json.dump({\n'
    '        "base_model": "monologg/koelectra-base-v3-discriminator",\n'
    '        "experiment": "v2_stage6",\n'
    '        "augmented": True,\n'
    '        "label_smoothing": 0.1\n'
    '    }, f, indent=2)'
)

add_subheading('5-3. 로드 코드')

add_code(
    'from transformers import AutoModelForSequenceClassification, AutoTokenizer\n'
    '\n'
    '# 토크나이저 로드 (base model에서)\n'
    'tokenizer = AutoTokenizer.from_pretrained(\n'
    '    "monologg/koelectra-base-v3-discriminator"\n'
    ')\n'
    '\n'
    '# 학습된 모델 로드\n'
    'model = AutoModelForSequenceClassification.from_pretrained(\n'
    '    "ai/models/intent_classifier"\n'
    ')\n'
    'model.eval()  # 추론 모드'
)

add_subheading('5-4. 추론 코드')

add_code(
    'import torch\n'
    '\n'
    'inputs = tokenizer(\n'
    '    text, max_length=64, padding="max_length",\n'
    '    truncation=True, return_tensors="pt"\n'
    ')\n'
    'with torch.no_grad():\n'
    '    logits = model(**inputs).logits\n'
    '\n'
    'probs = torch.softmax(logits, dim=-1)\n'
    'pred_id = probs.argmax().item()\n'
    'intent = id2label[pred_id]         # e.g. "judgment"\n'
    'confidence = probs[0][pred_id].item()  # e.g. 0.9542'
)

# =============================================
# 6. 라벨 매핑 및 서비스 설정
# =============================================
add_heading('6. 라벨 매핑 및 서비스 설정')

add_subheading('6-1. 라벨 매핑')

make_table(
    ['ID', 'Intent', 'Agent', '설명'],
    [
        ['0', 'judgment', 'Judgment Agent', '사내 규정/정책 판단'],
        ['1', 'doc_search', 'Document Agent', '문서 검색/조회'],
        ['2', 'doc_generate', 'Document Agent', '문서 생성'],
        ['3', 'doc_summary', 'Document Agent', '문서 요약'],
        ['4', 'schedule_add', 'Schedule Agent', '일정 등록'],
        ['5', 'schedule_view', 'Schedule Agent', '일정 조회'],
        ['6', 'general', 'General Handler', '일상 대화'],
        ['7', 'doc_qa', 'Document Agent', '문서 기반 Q&A'],
    ]
)

add_subheading('6-2. 서비스 Threshold')

make_table(
    ['설정', '값', '동작'],
    [
        ['INTENT_CONFIDENCE_THRESHOLD', '0.85', 'confidence < 0.85 -> clarify\n(top-3 후보 재질문)'],
        ['INTENT_FALLBACK_THRESHOLD', '0.4', 'confidence < 0.4 -> general 강제'],
    ]
)

add_subheading('6-3. 추론 예시')

make_table(
    ['입력', 'intent', 'confidence', '동작'],
    [
        ['"인턴에게 AWS 권한 줘도 돼?"', 'judgment', '0.9542', 'Judgment Agent 호출'],
        ['"보고서 만들어줘"', 'doc_generate', '0.9821', 'Document Agent 호출'],
        ['"내일 3시에 팀미팅 잡아줘"', 'schedule_add', '0.9734', 'Schedule Agent 호출'],
        ['"ㅂㄱㅅ 써줘"', 'doc_generate', '0.6832', 'clarify (top-3 제시)'],
    ]
)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('최종 수정일: 2026-02-25')
run.bold = True
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('-- End of Document --')
run.italic = True

# Save
out_dir = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(out_dir, '2.데이터 전처리_학습된 인공지능 모델_3팀.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
print(f'Size: {os.path.getsize(out_path)} bytes')
