"""docx 파일 텍스트 추출"""
from docx import Document
import sys

doc = Document(r"c:\SKN21-FINAL-3TEAM\docs\산출물\4주차\2.데이터 전처리_인공지능 학습 결과서_3팀.docx")

# 본문 텍스트
print("=" * 60)
print("본문 텍스트")
print("=" * 60)
for p in doc.paragraphs:
    if p.text.strip():
        print(p.text)

# 표 내용
print("\n" + "=" * 60)
print("표 내용")
print("=" * 60)
for i, table in enumerate(doc.tables):
    print(f"\n--- 표 {i+1} ---")
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        print(" | ".join(cells))
