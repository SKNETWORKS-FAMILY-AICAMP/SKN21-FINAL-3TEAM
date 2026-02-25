import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

# ===== Style setup =====
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

def shade_cell(cell, color="D9E2F3"):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_title(text, size=24, bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_heading(text, size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    p.paragraph_format.space_before = Pt(16)
    return p

def add_subheading(text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x2E, 0x4A, 0x7A)
    p.paragraph_format.space_before = Pt(10)
    return p

def add_body(text, size=10):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)
    return p

def add_note(text, size=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = True
    return p

def make_table(headers, rows, header_color="1F3864"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(cell, header_color)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                shade_cell(cell, "EDF2FA")
    doc.add_paragraph()
    return table


# =============================================
# COVER PAGE
# =============================================
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('SK networks  |  Family AI Camp')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_title('데이터 전처리', size=26)
add_title('인공지능 학습 결과서', size=18)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('SKN Family AI Camp 21기 : 최종 프로젝트 3팀')
run.font.size = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
run = p.add_run('WorkFlow Agent (듀듀)')
run.font.size = Pt(14)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('작성: 3조 -- 신지용(PM), 윤경은(AI), 진승언(AI), 안혜빈(Backend), 문지영(Frontend)')
run.font.size = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()


# =============================================
# 1. 모델 목적 및 비교 / 선정 이유
# =============================================
add_heading('1. 모델 비교 및 선정 이유')

add_body(
    '사용자 입력을 8개 intent(judgment, doc_search, doc_generate, doc_summary, doc_qa, '
    'schedule_add, schedule_view, general)로 분류하여 적절한 Agent로 라우팅하는 Intent Classifier를 학습합니다. '
    '잘못된 분류 = 잘못된 Agent 라우팅이므로, 분류 정확도가 전체 시스템 품질을 결정합니다.'
)

add_subheading('1-1. 후보 모델 비교')

make_table(
    ['모델', '파라미터', '아키텍처', '선정 여부'],
    [
        ['koelectra-base-v3\n(monologg)', '112.9M', 'ELECTRA (RTD)\n12-layer Transformer', '최종 선정'],
        ['klue/bert-base', '110.6M', 'BERT (MLM)\n12-layer Transformer', '탈락'],
        ['monologg/distilkobert', '28.4M', 'DistilBERT (KD)\n6-layer Transformer', '탈락'],
    ]
)

add_subheading('1-2. 선정 근거')

add_body(
    '- Adversarial 강건성 1위: Adv F1 86.04% (bert 85.17%, distilkobert 79.26%)\n'
    '- 추론 속도 균형: 7.9ms (bert 10.4ms 대비 24% 빠름)\n'
    '- Seed 안정성: 0.9874 +/- 0.0033 (3-seed 검증)\n'
    '- ELECTRA의 RTD(Replaced Token Detection) 방식이 토큰 교체 감지에 특화되어,\n'
    '  짧은 한국어 문장의 미세한 의미 차이 구분에 유리\n'
    '- McNemar 검정: koelectra vs bert p>0.05 (통계적 동급) -> 실용적 기준(Adv F1+속도)으로 선정'
)

doc.add_page_break()


# =============================================
# 2. 선정 모델 구조 상세
# =============================================
add_heading('2. 선정 모델 구조')

add_body(
    'koelectra-base-v3-discriminator의 3-layer 구조입니다. '
    '총 112.9M 파라미터, 입력 max_length=64 토큰.'
)

make_table(
    ['레이어', '구성', '설명'],
    [
        ['Embedding', 'Token + Position + Type\nembedding (768-dim)', '입력 텍스트를 768차원 벡터로 변환\nWordPiece 토크나이저 (vocab 35,000)'],
        ['Encoder', 'Transformer x 12층\nMulti-Head Attention (12 heads)\n+ FFN (3072-dim)', '12층 Transformer 블록\n각 층에서 셀프 어텐션 + 피드포워드'],
        ['Classification\nHead', 'Dropout (0.1)\n+ Dense (768 -> 8)\n+ Softmax', '[CLS] 토큰 출력을 8개 intent\n확률 분포로 변환'],
    ]
)


# =============================================
# 3. 학습 설정 및 하이퍼파라미터
# =============================================
add_heading('3. 학습 설정 및 하이퍼파라미터')

add_body(
    'Stage 2(Baseline)에서 고정 HP로 3모델을 비교한 뒤, '
    'Stage 3에서 KoELECTRA 대상 32-point Grid Search를 수행하여 최적 HP를 확정했습니다.'
)

add_subheading('3-1. 최종 학습 설정 (Stage 6 채택)')

make_table(
    ['하이퍼파라미터', '값'],
    [
        ['Base Model', 'monologg/koelectra-base-v3-discriminator'],
        ['Task', 'SequenceClassification (8 labels)'],
        ['Epochs', '10'],
        ['Batch Size', '16'],
        ['Learning Rate', '3e-5'],
        ['Optimizer', 'AdamW'],
        ['Loss Function', 'CrossEntropyLoss + Label Smoothing 0.1'],
        ['Scheduler', 'Linear (warmup_ratio=0.0)'],
        ['Max Length', '64 tokens'],
        ['Weight Decay', '0.01'],
        ['Seed', '42 (random/numpy/torch/cuda 4중 고정)'],
        ['Train Data', '2,425건 (기본 2,327 + 보강 98)'],
        ['Validation Data', '285건'],
    ]
)

add_subheading('3-2. Grid Search 결과 (Stage 3)')

add_body('32-point grid: epochs {3,5,7,10} x lr {1e-5,2e-5,3e-5,5e-5} x batch {16,32}')

make_table(
    ['항목', '탐색 범위', 'Best'],
    [
        ['Epochs', '{3, 5, 7, 10}', '10'],
        ['Learning Rate', '{1e-5, 2e-5, 3e-5, 5e-5}', '3e-5'],
        ['Batch Size', '{16, 32}', '16'],
    ]
)

add_note('* 3-seed 안정성 검증: 0.9874 +/- 0.0033 (seed={42, 123, 2024}). Baseline(0.9825) -> Best(0.9897): +0.72%p.')

doc.add_page_break()


# =============================================
# 4. 학습 결과 및 성능 평가
# =============================================
add_heading('4. 학습 결과 및 성능 평가')

add_subheading('4-1. 3모델 비교 (Stage 4: 각 모델 Best Config 기준)')

make_table(
    ['모델', 'Test Acc', 'Test F1', 'Adv F1', '추론 속도'],
    [
        ['koelectra-base-v3\n(선정)', '0.9720', '0.9726', '0.8604', '7.9ms'],
        ['klue/bert-base', '--', '0.9756', '0.8517', '10.4ms'],
        ['distilkobert', '--', '0.9645', '0.7926', '2.8ms'],
    ]
)

add_note(
    '* koelectra의 Test F1이 bert보다 낮지만, Adversarial(실전) F1에서 koelectra가 우위.\n'
    '* bert/distilkobert는 Grid Search 미수행으로 Baseline HP(ep5, lr2e-5) 기준 평가.'
)

add_subheading('4-2. Stage별 성능 변화 (KoELECTRA)')

make_table(
    ['Stage', '내용', 'Test F1', 'Adv F1'],
    [
        ['Stage 2', 'Baseline (ep5, lr2e-5, bs16)', '-- (Val만 평가)', '-- (Val만 평가)'],
        ['Stage 3', 'Grid Search 32-point', '-- (Val만 평가)', '-- (Val만 평가)'],
        ['Stage 4', '최종 평가 (best config)', '0.9726', '0.8604'],
        ['Stage 5', '타겟 보강 +98건 재학습', '0.9788 (+0.62%p)', '0.8784 (+1.80%p)'],
        ['Stage 6 (최종)', 'Label Smoothing 0.1', '0.9788 (유지)', '0.8758 (-0.26%p)'],
        ['Stage 7 (미채택)', '추가 보강 +25건', '0.9756 (-0.32%p)', '0.8712 (-0.46%p)'],
    ]
)

add_subheading('4-3. 최종 모델 성능 (Stage 6)')

make_table(
    ['평가셋', 'Accuracy', 'Precision', 'Recall', 'F1 Score'],
    [
        ['Validation (285건)', '0.9895', '0.9892', '0.9899', '0.9894'],
        ['Test (286건)', '0.9790', '0.9787', '0.9792', '0.9788'],
        ['Adversarial (450건)', '0.8756', '0.8792', '0.8768', '0.8758'],
    ]
)

add_note(
    '* Adversarial: 초단문/오타/간접표현 등 실사용 환경의 어려운 입력. '
    'doc_qa F1 76.6%가 최저이나, 라벨 리뷰 결과 오류의 63%가 라벨 애매/오류이며 '
    'doc 4개 intent 모두 동일 Agent로 라우팅되므로 실서비스 영향 제한적.'
)

doc.add_page_break()


# =============================================
# 5. 과적합 / 과소적합 대응
# =============================================
add_heading('5. 과적합 / 과소적합 대응')

add_body(
    'Intent 분류기의 학습 과정에서 발생할 수 있는 과적합/과소적합 문제에 대해 '
    '다음과 같은 기법을 적용하였습니다.'
)

make_table(
    ['기법', '적용 내용', '효과'],
    [
        ['Label Smoothing',
         'CrossEntropyLoss에\nsmoothing=0.1 적용 (Stage 6)',
         '과신뢰 오분류 42건 -> 13건 (69% 감소)\n'
         'Threshold 0.85로 정답/오답 분리 가능'],
        ['Dropout',
         'Classification Head에\nDropout(p=0.1) 적용',
         'Transformer 출력에 무작위 뉴런 비활성화\n'
         '과적합 방지 (사전학습 모델 기본 설정)'],
        ['Weight Decay',
         'AdamW optimizer\nweight_decay=0.01',
         'L2 정규화 효과로 가중치 발산 방지'],
        ['Seed 고정',
         'random/numpy/torch/cuda\n4중 고정 (seed=42)',
         '3-seed 검증: 0.9874 +/- 0.0033\n학습 안정성 및 재현성 확보'],
        ['타겟 보강',
         'Adversarial 오분류 분석 후\n98건 취약 패턴 추가 (Stage 5)',
         'Adv F1 +1.80%p, doc_qa +7.9%p\n과소적합 패턴에 대한 데이터 보강'],
    ]
)

add_subheading('5-1. Label Smoothing 효과 상세')

add_body(
    'Stage 6에서 적용한 Label Smoothing 0.1의 핵심 효과는 '
    '"틀릴 때 확신하는" 과신뢰 오분류를 줄인 것입니다.'
)

make_table(
    ['항목', 'Stage 5 (적용 전)', 'Stage 6 (적용 후)', '변화'],
    [
        ['오분류 중 과신뢰(>90%)', '42건 (66.7%)', '13건 (23.2%)', '-69%'],
        ['정답 confidence 중앙값', '0.9968', '0.9366', '부드러운 분포'],
        ['오답 confidence 중앙값', '~0.90', '~0.64', '정답/오답 분리 가능'],
    ]
)

add_body(
    '이를 통해 Confidence Threshold 0.85 기준으로:\n'
    '- confidence >= 0.85: 해당 intent로 Agent 라우팅\n'
    '- 0.4 <= confidence < 0.85: clarify (top-3 후보 제시, 사용자에게 재질문)\n'
    '- confidence < 0.4: general로 강제 라우팅 (분류 불가 판단)'
)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('최종 수정일: 2026-02-25')
run.bold = True
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('-- End of Document --')
run.italic = True

# Save
out_dir = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(out_dir, '2.데이터 전처리_인공지능 학습 결과서_3팀.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
print(f'Size: {os.path.getsize(out_path)} bytes')
