"""
문서 서비스 (팀원 C/D 공동 담당)
- 파일 업로드, 텍스트 추출, 문서 CRUD
- Qdrant 인덱싱 (업로드 시 자동), RAG 검색 (내용)
"""
import os
import re
import uuid
import logging
from datetime import datetime, date, timedelta
from pathlib import Path

from fastapi import UploadFile, HTTPException
from sqlalchemy import select, or_, cast, Date
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.document import Document
from app.config import get_settings

# PDF/DOCX 라이브러리 미리 import
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


def classify_category(title: str, text: str) -> str:
    """제목+본문 키워드 기반 카테고리 분류 (규칙 기반)"""
    t = (title + " " + text[:2000]).lower()

    # 회의록
    if any(k in t for k in ["회의록", "회의 기록", "미팅 노트", "meeting minutes", "회의 결과",
                             "참석자", "안건", "회의 일시", "meeting note"]):
        return "회의록"
    # 계약서
    if any(k in t for k in ["계약서", "계약 조건", "contract", "agreement", "서약서",
                             "비밀유지", "nda", "업무 위탁", "용역 계약"]):
        return "계약서"
    # 제안서
    if any(k in t for k in ["제안서", "proposal", "기획서", "기획안", "사업 제안",
                             "프로젝트 기획", "프로젝트 제안", "제안 배경"]):
        return "제안서"
    # 보고서
    if any(k in t for k in ["보고서", "report", "업무보고", "분석 보고", "결과 보고",
                             "실적 보고", "주간 보고", "월간 보고", "성과 보고"]):
        return "보고서"
    # 정책문서
    if any(k in t for k in ["정책", "규정", "가이드라인", "지침", "policy", "regulation",
                             "내규", "취업규칙", "복무", "보안 정책", "개인정보"]):
        return "정책문서"
    # 인사문서
    if any(k in t for k in ["인사", "채용", "연차", "휴가", "급여", "퇴직", "온보딩",
                             "jd", "job description", "직무기술서", "입사", "경력증명"]):
        return "인사문서"
    return "기타"


async def analyze_document_with_llm(text: str, title: str) -> dict | None:
    """
    문서를 분석하여 summary, tags, category를 추출한다.
    - summary, tags: LLM(sLLM/GPT) 분석
    - category: 규칙 기반 분류 (LLM 불필요)
    실패 시 None 반환 (문서 업로드는 정상 진행).
    """
    category = classify_category(title, text)
    logger.info(f"[DocumentAnalysis] 규칙 기반 카테고리 분류: {category} | title={title}")

    try:
        from ai.agents.document_agent import summarize_document
        print(f"[DocumentAnalysis] summarize_document 호출 | title={title}")

        result = await summarize_document(text)

        tags = result.get("tags", [])
        summary = result.get("summary", "")
        logger.info(f"문서 분석 완료: category={category}, tags={tags}, summary_len={len(summary)}자")

        return {
            "summary": summary,
            "category": category,
            "tags": tags,
        }

    except Exception as e:
        print(f"[DocumentAnalysis] LLM 분석 실패 (카테고리는 규칙 기반 적용): {type(e).__name__}: {e}")
        return {
            "summary": None,
            "category": category,
            "tags": [],
        }


async def save_file(file: UploadFile, upload_dir: str) -> tuple[str, str]:
    """
    UploadFile을 디스크에 저장한다.

    Returns:
        (saved_path, file_type)
    """
    os.makedirs(upload_dir, exist_ok=True)

    original = file.filename or "unknown"
    ext = Path(original).suffix.lower()  # .pdf, .docx, .txt
    file_type = ext.lstrip(".")

    if file_type not in ("pdf", "docx", "txt"):
        raise HTTPException(status_code=400, detail=f"지원하지 않는 파일 형식입니다: {ext}")

    unique_name = f"{uuid.uuid4().hex}{ext}"
    saved_path = os.path.join(upload_dir, unique_name)

    content = await file.read()
    with open(saved_path, "wb") as f:
        f.write(content)

    return saved_path, file_type


def extract_text(file_path: str, file_type: str) -> str:
    """
    PDF / DOCX / TXT 파일에서 텍스트를 추출한다.
    """
    if file_type == "pdf":
        return _extract_pdf(file_path)
    elif file_type == "docx":
        return _extract_docx(file_path)
    elif file_type == "txt":
        return _extract_txt(file_path)
    else:
        raise ValueError(f"지원하지 않는 파일 타입: {file_type}")


def _extract_pdf(file_path: str) -> str:
    if not PDF_AVAILABLE:
        raise ImportError("PyMuPDF (fitz) is not installed. Please install it with: pip install PyMuPDF")

    try:
        text_parts: list[str] = []
        doc = fitz.open(file_path)

        for page in doc:
            text = page.get_text()
            text_parts.append(text)

        doc.close()
        result = "\n".join(text_parts)

        if not result.strip():
            # 텍스트 레이어 없는 스캔 PDF → OCR 폴백
            logger.info(f"PDF 텍스트 레이어 없음, OCR 시도: {file_path}")
            try:
                from ai.document_parser.ocr_parser import OCRParser
                ocr = OCRParser(lang="korean")
                result = ocr.extract_text_from_pdf(file_path)
                if not result.strip():
                    raise ValueError("PDF에서 텍스트를 추출할 수 없습니다 (스캔 이미지 OCR도 실패)")
            except ImportError:
                raise ValueError("PDF 텍스트 레이어가 없고 OCR(PaddleOCR)이 설치되지 않았습니다")

        return result
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"PDF extraction failed: {e}") from e


def _extract_docx(file_path: str) -> str:
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Please install it with: pip install python-docx")

    doc = DocxDocument(file_path)
    parts = []

    # 문단 텍스트 추출 (빈 문단 제외)
    para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    parts.extend(para_texts)

    # 테이블 텍스트 추출 (doc.paragraphs는 테이블 내부 텍스트를 포함하지 않음)
    table_parts = []
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                table_parts.append(" | ".join(row_cells))
    parts.extend(table_parts)

    result = "\n".join(parts)

    # DEBUG: python-docx 파싱 결과 출력
    print(f"\n{'='*60}")
    print(f"[DocxParser] DEBUG 파일: {file_path}")
    print(f"[DocxParser] 전체 단락 수: {len(doc.paragraphs)}, 비어있지 않은 단락: {len(para_texts)}")
    print(f"[DocxParser] 테이블 수: {len(doc.tables)}, 테이블 행 수: {len(table_parts)}")
    print(f"[DocxParser] 추출된 텍스트 총 길이: {len(result)}자")
    print(f"[DocxParser] 추출 내용 미리보기 (앞 500자):\n{result[:500]}")
    print(f"{'='*60}\n")

    return result


def _extract_txt(file_path: str) -> str:
    # UTF-8 시도, 실패하면 여러 인코딩 시도
    encodings = ['utf-8', 'utf-16', 'utf-16-le', 'utf-16-be', 'cp949', 'euc-kr', 'latin-1']
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()
                # 빈 문자열이 아닌 경우에만 반환
                if content.strip():
                    return content
        except (UnicodeDecodeError, UnicodeError):
            continue

    # 모든 인코딩 실패 시 에러
    raise ValueError("Unable to decode text file. File may be corrupted or in an unsupported encoding.")


async def upload_and_parse(
    db: AsyncSession,
    file: UploadFile,
    scope: str,
    user_id: int,
    team_name: str | None = None,
) -> Document:
    """
    파일 업로드 → 텍스트 추출 → DB 저장

    동기적으로 텍스트를 추출하고 status를 바로 completed로 변경한다.
    (Celery 없이 단순 처리)
    """
    settings = get_settings()

    title = Path(file.filename or "untitled").stem
    file_ext = Path(file.filename or "").suffix.lstrip(".").lower() or "txt"

    # 0. 중복 체크: 같은 사용자가 같은 제목+파일타입으로 이미 업로드한 문서가 있으면 기존 문서 반환
    from sqlalchemy import select as sa_select
    existing_result = await db.execute(
        sa_select(Document).where(
            Document.title == title,
            Document.file_type == file_ext,
            Document.uploaded_by == user_id,
            Document.status != "failed",
        ).limit(1)
    )
    existing_doc = existing_result.scalar_one_or_none()
    if existing_doc:
        logger.info(f"중복 문서 감지 — 기존 문서 반환: id={existing_doc.id}, title={title}")
        existing_doc._is_duplicate = True
        return existing_doc

    # 1. 파일 저장
    saved_path, file_type = await save_file(file, settings.UPLOAD_DIR)

    # 2. DB 레코드 생성 (status: processing)
    doc = Document(
        title=title,
        file_path=saved_path,
        file_type=file_type,
        scope=scope,
        team_name=team_name if scope == "team" else None,
        uploaded_by=user_id,
        status="processing",
    )
    db.add(doc)
    await db.flush()  # id 확보

    # 3. 텍스트 추출
    try:
        text = extract_text(saved_path, file_type)
        logger.info(f"텍스트 추출 완료: {len(text)}자, file={saved_path}")

        doc.content = text
        doc.status = "completed"

        # LLM 자동 분석 (요약 + 분류 + 태깅)
        print(f"[DocumentAnalysis] LLM 분석 시작: title={doc.title}, text_len={len(text)}")
        analysis = await analyze_document_with_llm(text, doc.title)
        print(f"[DocumentAnalysis] LLM 분석 결과: {analysis}")
        if analysis:
            doc.summary = analysis.get("summary")
            doc.category = analysis.get("category")
            doc.tags = analysis.get("tags")

        # Qdrant에 인덱싱 (RAG 검색용)
        try:
            from ai.rag.qdrant_pipeline import get_qdrant_pipeline
            pipeline = get_qdrant_pipeline()
            # 분석 결과를 메타데이터에 포함 (RAG 검색 정확도 향상)
            qdrant_meta = {
                "source": "documents",
                "doc_type": doc.category or "general",
                "title": doc.title,
                "scope": doc.scope,
                "team_name": doc.team_name or "",
                "user_id": str(user_id),
                "document_id": doc.id,
                "summary": doc.summary or "",
                "category": doc.category or "",
                "tags": ", ".join(doc.tags) if doc.tags else "",
            }
            # 태그/요약은 메타데이터에만 저장 (BM25 태그 부스트는 hybrid_search에서 처리)
            # content에 prefix를 붙이면 사용자에게 태그 텍스트가 그대로 노출됨
            pipeline.add_documents(
                documents=[text],
                metadatas=[qdrant_meta],
            )
            logger.info(f"문서 Qdrant 인덱싱 완료: document_id={doc.id}, title={doc.title}")
        except Exception as qdrant_err:
            logger.warning(f"Qdrant 인덱싱 실패 (문서는 정상 저장됨): {qdrant_err}")

    except Exception as e:
        logger.error(f"텍스트 추출 실패: {e}", exc_info=True)
        doc.status = "failed"

    return doc


def _parse_date_query(keyword: str) -> tuple[date | None, date | None]:
    """날짜 검색 키워드를 파싱하여 (start_date, end_date) 반환.

    지원 형식:
    - 2026-02-24 / 2026.02.24 → 해당 일
    - 2026-02 / 2026.02 → 해당 월 전체
    - 2026 → 해당 연도 전체
    - 2월 → 현재 연도 해당 월
    """
    keyword = keyword.strip()

    # 2026-02-24 또는 2026.02.24
    m = re.match(r"^(\d{4})[-./](\d{1,2})[-./](\d{1,2})$", keyword)
    if m:
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        try:
            day = date(y, mo, d)
            return day, day
        except ValueError:
            return None, None

    # 2026-02 또는 2026.02
    m = re.match(r"^(\d{4})[-./](\d{1,2})$", keyword)
    if m:
        y, mo = int(m.group(1)), int(m.group(2))
        try:
            start = date(y, mo, 1)
            # 다음 달 1일 - 1일 = 이번 달 마지막 날
            if mo == 12:
                end = date(y, 12, 31)
            else:
                end = date(y, mo + 1, 1)
                end = end - timedelta(days=1)
            return start, end
        except ValueError:
            return None, None

    # 2026
    m = re.match(r"^(\d{4})$", keyword)
    if m:
        y = int(m.group(1))
        return date(y, 1, 1), date(y, 12, 31)

    # N월 (현재 연도)
    m = re.match(r"^(\d{1,2})월$", keyword)
    if m:
        mo = int(m.group(1))
        y = datetime.now().year
        try:
            start = date(y, mo, 1)
            if mo == 12:
                end = date(y, 12, 31)
            else:
                end = date(y, mo + 1, 1)
                end = end - timedelta(days=1)
            return start, end
        except ValueError:
            return None, None

    return None, None


async def list_documents(
    db: AsyncSession,
    user_id: int,
    scope: str | None = None,
    keyword: str | None = None,
    search_type: str = "title",
    user_team: str | None = None,
) -> list[Document]:
    """문서 목록 조회 (scope, keyword, search_type 필터)

    Args:
        search_type: "title" (제목 ILIKE), "title_content" (제목+내용 ILIKE), "date" (날짜 범위)
        user_team: 유저 소속 팀 (team scope 필터링용)
    """
    stmt = select(Document)

    if scope == "team":
        # 팀 문서: 같은 팀의 team scope 문서만
        if user_team:
            stmt = stmt.where(
                (Document.scope == "team") & (Document.team_name == user_team)
            )
        else:
            # 팀 없는 유저는 빈 결과
            stmt = stmt.where(Document.id < 0)
    elif scope:
        stmt = stmt.where(Document.scope == scope)
    else:
        # scope 미지정 시: 회사 문서 + 본인 개인 문서 + 본인 팀 문서
        conditions = [
            Document.scope == "company",
            (Document.scope == "personal") & (Document.uploaded_by == user_id),
        ]
        if user_team:
            conditions.append(
                (Document.scope == "team") & (Document.team_name == user_team)
            )
        stmt = stmt.where(or_(*conditions))

    if keyword:
        if search_type == "title":
            stmt = stmt.where(Document.title.ilike(f"%{keyword}%"))

        elif search_type == "title_content":
            stmt = stmt.where(
                or_(
                    Document.title.ilike(f"%{keyword}%"),
                    Document.content.ilike(f"%{keyword}%"),
                )
            )

        elif search_type == "date":
            start_date, end_date = _parse_date_query(keyword)
            if start_date and end_date:
                stmt = stmt.where(
                    cast(Document.created_at, Date) >= start_date,
                    cast(Document.created_at, Date) <= end_date,
                )
            else:
                # 파싱 실패 시 빈 결과
                stmt = stmt.where(Document.id < 0)

    stmt = stmt.order_by(Document.created_at.desc())
    result = await db.execute(stmt)
    return list(result.scalars().all())


async def get_document(db: AsyncSession, document_id: int) -> Document:
    """문서 상세 조회"""
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if doc is None:
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다")
    return doc


async def delete_document(
    db: AsyncSession,
    document_id: int,
    user_id: int,
) -> dict:
    """문서 삭제 (업로드한 사용자만 가능)"""
    doc = await get_document(db, document_id)

    if doc.uploaded_by != user_id:
        raise HTTPException(status_code=403, detail="본인이 업로드한 문서만 삭제할 수 있습니다")

    # 파일 삭제
    if doc.file_path and os.path.exists(doc.file_path):
        os.remove(doc.file_path)

    # Qdrant에서 벡터 삭제
    try:
        from ai.rag.qdrant_pipeline import get_qdrant_pipeline
        pipeline = get_qdrant_pipeline()
        pipeline.vector_store.delete_by_filter({"document_id": document_id})
        pipeline.searcher.build_bm25_index()
        logger.info(f"Qdrant 문서 삭제 완료: document_id={document_id}")
    except Exception as e:
        logger.warning(f"Qdrant 문서 삭제 실패 (DB 삭제는 진행): {e}")

    await db.delete(doc)
    return {"message": "문서가 삭제되었습니다", "document_id": document_id}


async def analyze_existing_documents(db: AsyncSession) -> dict:
    """기존 문서 중 분석되지 않은 문서를 일괄 LLM 분석한다."""
    stmt = select(Document).where(
        Document.status == "completed",
        Document.summary.is_(None),
        Document.content.isnot(None),
    )
    result = await db.execute(stmt)
    docs = list(result.scalars().all())

    analyzed = 0
    failed = 0
    for doc in docs:
        analysis = await analyze_document_with_llm(doc.content, doc.title)
        if analysis:
            doc.summary = analysis.get("summary")
            doc.category = analysis.get("category")
            doc.tags = analysis.get("tags")
            analyzed += 1
        else:
            failed += 1

    return {"total": len(docs), "analyzed": analyzed, "failed": failed}


async def reindex_all_documents(db: AsyncSession) -> dict:
    """기존 문서를 Qdrant에 재인덱싱 (태그/분류/요약 메타데이터 포함)."""
    stmt = select(Document).where(
        Document.status == "completed",
        Document.content.isnot(None),
    )
    result = await db.execute(stmt)
    docs = list(result.scalars().all())

    indexed = 0
    failed = 0

    try:
        from ai.rag.qdrant_pipeline import get_qdrant_pipeline
        pipeline = get_qdrant_pipeline()
    except Exception as e:
        return {"error": f"Qdrant 파이프라인 로드 실패: {e}"}

    for doc in docs:
        try:
            # 기존 벡터 삭제
            try:
                pipeline.vector_store.delete_by_filter({"document_id": doc.id})
            except Exception:
                pass

            # 태그/요약은 메타데이터에만 저장 (content에 prefix 붙이지 않음)
            indexed_text = doc.content

            qdrant_meta = {
                "source": "documents",
                "doc_type": doc.category or "general",
                "title": doc.title,
                "scope": doc.scope,
                "team_name": doc.team_name or "",
                "user_id": str(doc.uploaded_by),
                "document_id": doc.id,
                "summary": doc.summary or "",
                "category": doc.category or "",
                "tags": ", ".join(doc.tags) if doc.tags else "",
            }

            pipeline.add_documents(
                documents=[indexed_text],
                metadatas=[qdrant_meta],
            )
            indexed += 1
            print(f"[Reindex] 완료: id={doc.id}, title={doc.title}")
        except Exception as e:
            print(f"[Reindex] 실패: id={doc.id}, error={e}")
            failed += 1

    return {"total": len(docs), "indexed": indexed, "failed": failed}


async def generate_and_save(
    db: AsyncSession,
    user_input: str,
    user_id: int,
    template_type: str | None = None,
    template_id: int | None = None,
) -> tuple[Document, dict]:
    """
    Document Agent를 호출하여 문서를 생성하고 DB에 저장한다.

    Returns:
        (Document, agent_response)
    """
    import json as _json

    settings = get_settings()

    # 1. AgentState 구성
    state = {
        "user_input": user_input,
        "user_id": user_id,
        "intent": "doc_generate",
        "stream_mode": False,
        "template_type": template_type,
        "template_id": template_id,
        "context": [],
        "agent_response": {},
        "confidence": 0.0,
    }

    # 2. Document Agent 호출 (lazy import — backend 실행 시 ai/ 의존 최소화)
    from ai.agents.document_agent import document_agent
    state = await document_agent(state)

    agent_response = state.get("agent_response", {})

    # Agent가 에러를 반환한 경우
    if "error" in agent_response:
        raise HTTPException(status_code=500, detail=agent_response["error"])

    # 3. 생성된 데이터를 JSON 파일로 저장
    generated_dir = os.path.join(settings.UPLOAD_DIR, "generated")
    os.makedirs(generated_dir, exist_ok=True)

    file_name = f"{uuid.uuid4().hex}.json"
    file_path = os.path.join(generated_dir, file_name)

    data = agent_response.get("data", agent_response)
    with open(file_path, "w", encoding="utf-8") as f:
        _json.dump(data, f, ensure_ascii=False, indent=2)

    # 4. Document 레코드 생성
    resolved_type = agent_response.get("template_type", template_type or "report")
    title = data.get("title", f"{resolved_type} 문서")

    doc = Document(
        title=title,
        file_path=file_path,
        file_type="json",
        content=agent_response.get("preview", ""),
        scope="personal",
        uploaded_by=user_id,
        status="completed",
    )
    db.add(doc)
    await db.flush()
    await db.refresh(doc)

    # agent_response의 mock ID를 실제 ID로 교체
    agent_response["document_id"] = doc.id
    agent_response["download_url"] = f"/api/v1/documents/{doc.id}/download"

    logger.info(f"문서 생성 완료: document_id={doc.id}, template_type={resolved_type}")
    return doc, agent_response
