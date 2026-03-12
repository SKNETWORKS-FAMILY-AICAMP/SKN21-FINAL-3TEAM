"""테스트 문서 5종 생성 스크립트"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

OUT = Path(__file__).parent


def _style(doc):
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def _table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()
    return t


# ──────────────────────────────────────────────
# 1. 회의록
# ──────────────────────────────────────────────
def create_meeting_minutes():
    doc = Document()
    _style(doc)

    _heading(doc, "2025년 1분기 전략기획 회의록")

    doc.add_paragraph()
    _table(doc, ["항목", "내용"], [
        ["일시", "2025년 3월 10일 (월) 14:00 ~ 16:30"],
        ["장소", "본사 15층 대회의실 (화상 병행: Zoom)"],
        ["참석자", "김태현 대표이사, 박서연 CTO, 이준혁 CPO, 최지은 CFO, 한도윤 마케팅본부장, 정수아 HR본부장"],
        ["작성자", "오민재 전략기획팀 사원"],
        ["배포 범위", "경영진 및 본부장급 이상"],
    ])

    _heading(doc, "1. 2024년 4분기 실적 리뷰", level=2)
    doc.add_paragraph(
        "최지은 CFO가 2024년 4분기 실적을 발표하였다. 매출액은 전년 동기 대비 23.4% 증가한 847억 원을 "
        "기록하였으며, 영업이익률은 14.2%로 전 분기 대비 1.8%p 개선되었다. SaaS 구독 매출 비중이 "
        "전체 매출의 62%로 처음으로 과반을 넘었으며, 이는 B2B 엔터프라이즈 고객 확대에 기인한다."
    )
    doc.add_paragraph(
        "해외 매출은 전체의 18%를 차지하며 전년 대비 41% 성장하였다. 특히 일본 시장에서의 "
        "파트너십 확대(소프트뱅크, NTT데이터)가 주요 성장 동력이었다. 다만, 북미 시장 진출은 "
        "현지 영업 인력 부족으로 당초 목표의 67%에 그쳤다."
    )

    _heading(doc, "2. 2025년 1분기 핵심 전략 논의", level=2)

    _heading(doc, "2-1. AI 제품 로드맵", level=3)
    doc.add_paragraph(
        "박서연 CTO는 AI 기반 업무자동화 플랫폼 'AutoFlow 2.0'의 개발 현황을 보고하였다. "
        "핵심 기능인 자연어 기반 워크플로우 빌더는 3월 말 내부 베타를 시작하며, "
        "LLM 파인튜닝을 통한 산업별 특화 모델(금융, 제조, 물류) 3종을 4월 중 출시 예정이다."
    )
    doc.add_paragraph(
        "이준혁 CPO는 기존 고객 대상 사전 수요조사 결과를 공유하였다. 응답 기업 120곳 중 "
        "78%가 AI 워크플로우 자동화에 관심을 표명하였으며, 특히 문서 처리 자동화(89%)와 "
        "일정 관리 자동화(72%)에 대한 수요가 높았다."
    )

    _heading(doc, "2-2. 글로벌 확장 전략", level=3)
    doc.add_paragraph(
        "한도윤 마케팅본부장은 2025년 해외 매출 비중을 30%까지 확대하는 목표를 제시하였다. "
        "1분기에는 일본 시장 현지법인 설립(4월), 싱가포르 파트너 계약 체결(3월)을 추진하며, "
        "하반기 북미 진출을 위한 사전 마케팅(웨비나, 컨퍼런스)을 2분기부터 시작한다."
    )

    _heading(doc, "2-3. 인력 운영 계획", level=3)
    doc.add_paragraph(
        "정수아 HR본부장은 1분기 채용 계획을 발표하였다. AI/ML 엔지니어 15명, "
        "글로벌 영업 5명, DevOps 3명 등 총 23명 충원 예정이며, "
        "특히 시니어 AI 연구원은 해외 인재 스카우트를 병행한다. "
        "또한 전사 AI 리터러시 교육을 분기 1회 시행하기로 하였다."
    )

    _heading(doc, "3. 의결 사항", level=2)
    items = [
        "AutoFlow 2.0 내부 베타 일정: 2025년 3월 31일 확정",
        "일본 현지법인 설립 예산 15억 원 승인",
        "1분기 신규 채용 23명 승인 (AI/ML 엔지니어 우선)",
        "해외 마케팅 예산 8억 원 배정 (전년 대비 60% 증액)",
        "전사 AI 리터러시 교육 프로그램 4월 시작",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Number")

    _heading(doc, "4. 액션 아이템", level=2)
    _table(doc, ["담당자", "업무", "마감일"], [
        ["박서연 CTO", "AutoFlow 2.0 내부 베타 환경 구축 및 QA 시나리오 작성", "2025-03-25"],
        ["이준혁 CPO", "베타 테스트 참여 고객사 10곳 확보 및 피드백 수집 체계 구축", "2025-03-28"],
        ["한도윤 본부장", "일본 현지법인 설립 법무 검토 착수 및 싱가포르 파트너 MOU 초안", "2025-03-20"],
        ["정수아 본부장", "시니어 AI 연구원 해외 채용 공고 게시 및 헤드헌팅 계약", "2025-03-15"],
        ["최지은 CFO", "1분기 해외 투자 예산 집행 계획 수립 및 이사회 보고 자료 작성", "2025-03-18"],
    ])

    _heading(doc, "5. 차기 회의", level=2)
    doc.add_paragraph("일시: 2025년 4월 14일 (월) 14:00")
    doc.add_paragraph("안건: AutoFlow 2.0 베타 결과 리뷰, 2분기 실행 계획 점검")

    path = OUT / "2025_1분기_전략회의_회의록.docx"
    doc.save(str(path))
    print(f"  [1/5] {path.name}")


# ──────────────────────────────────────────────
# 2. 계약서
# ──────────────────────────────────────────────
def create_contract():
    doc = Document()
    _style(doc)

    title = doc.add_heading("소프트웨어 개발 용역 계약서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph(
        "주식회사 넥스트이노베이션(이하 '갑')과 주식회사 코드랩솔루션(이하 '을')은 "
        "아래와 같이 소프트웨어 개발 용역 계약을 체결한다.",
    )

    articles = [
        ("제1조 (목적)",
         "본 계약은 갑이 을에게 'AI 기반 문서관리 시스템(이하 \"본 시스템\")'의 설계, 개발, "
         "테스트 및 배포에 관한 용역을 위탁하고, 을이 이를 수행함에 있어 필요한 제반 사항을 "
         "정함을 목적으로 한다."),

        ("제2조 (용역 범위)",
         "① 을이 수행할 용역의 범위는 다음 각 호와 같다.\n"
         "  1. 시스템 요구사항 분석 및 설계 문서 작성\n"
         "  2. AI 엔진 개발 (자연어처리, 문서 분류, 요약 기능)\n"
         "  3. 백엔드 API 서버 개발 (FastAPI 기반)\n"
         "  4. 프론트엔드 웹 애플리케이션 개발 (React 기반)\n"
         "  5. 데이터베이스 설계 및 구축 (PostgreSQL)\n"
         "  6. 시스템 통합 테스트 및 성능 최적화\n"
         "  7. 운영 환경 배포 및 안정화 지원 (배포 후 1개월)\n"
         "② 세부 기능 명세는 별첨 '기능요구사항 정의서(SRS)'에 따른다."),

        ("제3조 (계약 기간)",
         "① 본 계약의 수행 기간은 2025년 4월 1일부터 2025년 9월 30일까지 총 6개월로 한다.\n"
         "② 갑과 을의 합의에 의해 계약 기간을 연장할 수 있으며, 이 경우 서면 합의를 원칙으로 한다.\n"
         "③ 불가항력 사유 발생 시 해당 기간만큼 자동 연장되며, 이를 즉시 상대방에게 통지하여야 한다."),

        ("제4조 (계약 금액 및 지급 조건)",
         "① 본 계약의 총 용역 대금은 금 4억 8천만 원(부가가치세 별도)으로 한다.\n"
         "② 대금은 다음과 같이 4회 분할 지급한다.\n"
         "  1. 착수금: 계약 체결 후 7영업일 이내 — 1억 2천만 원 (25%)\n"
         "  2. 1차 중도금: 설계 완료 검수 후 — 1억 2천만 원 (25%)\n"
         "  3. 2차 중도금: 개발 완료 검수 후 — 1억 2천만 원 (25%)\n"
         "  4. 잔금: 최종 검수 완료 후 14영업일 이내 — 1억 2천만 원 (25%)\n"
         "③ 을은 각 지급 시점에 세금계산서를 발행하며, 갑은 수령 후 14영업일 이내에 지급한다."),

        ("제5조 (검수)",
         "① 을은 각 단계별 산출물을 갑에게 제출하고, 갑은 제출일로부터 10영업일 이내에 검수를 완료한다.\n"
         "② 검수 기준은 별첨 '검수 기준서'에 따르며, 주요 결함(Critical/Major)이 없을 것을 기준으로 한다.\n"
         "③ 갑이 검수 기간 내 서면 이의를 제기하지 않을 경우 검수에 합격한 것으로 간주한다.\n"
         "④ 검수 불합격 시 을은 갑의 보완 요청에 따라 10영업일 이내에 수정·보완하여 재제출한다."),

        ("제6조 (지식재산권)",
         "① 본 계약에 의해 개발된 시스템의 소스코드, 설계 문서 등 모든 산출물의 지식재산권은 "
         "최종 대금 지급 완료 시 갑에게 귀속된다.\n"
         "② 을이 기존에 보유한 범용 라이브러리 및 프레임워크의 지식재산권은 을에게 귀속되며, "
         "갑에게 비독점적 사용권을 부여한다.\n"
         "③ 오픈소스 소프트웨어 사용 시 을은 라이선스 목록을 갑에게 사전 제출하여 승인을 받아야 한다."),

        ("제7조 (비밀유지)",
         "① 갑과 을은 본 계약의 이행 과정에서 취득한 상대방의 영업 비밀, 기술 정보, "
         "고객 정보 등 일체의 비밀 정보를 제3자에게 누설하거나 본 계약 목적 외의 용도로 "
         "사용하여서는 아니 된다.\n"
         "② 본 조의 의무는 계약 종료 후 3년간 존속한다.\n"
         "③ 위반 시 상대방에게 발생한 손해를 배상할 책임을 진다."),

        ("제8조 (하자 보수)",
         "① 을은 최종 검수 완료일로부터 12개월간 본 시스템의 하자에 대하여 무상으로 보수할 의무를 진다.\n"
         "② 하자의 범위는 을의 귀책사유로 인한 시스템 오류, 성능 저하, 보안 취약점 등을 포함한다.\n"
         "③ 갑의 요구사항 변경 또는 운영 환경 변경으로 인한 수정 사항은 하자 보수 범위에 포함되지 아니한다."),

        ("제9조 (계약 해제·해지)",
         "① 갑 또는 을이 본 계약의 중요한 의무를 위반하고, 상대방의 서면 최고 후 "
         "30일 이내에 이를 시정하지 아니한 경우 계약을 해제·해지할 수 있다.\n"
         "② 해제·해지 시 이미 완료된 작업에 대한 대금은 정산하여 지급하되, "
         "귀책사유 있는 당사자는 상대방에게 손해배상 책임을 진다."),

        ("제10조 (분쟁 해결)",
         "본 계약과 관련하여 발생하는 분쟁은 서울중앙지방법원을 제1심 관할 법원으로 한다."),
    ]

    for title_text, body in articles:
        _heading(doc, title_text, level=2)
        doc.add_paragraph(body)

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph("본 계약의 성립을 증명하기 위하여 계약서 2부를 작성하고, "
                          "갑과 을이 각각 기명날인한 후 각 1부씩 보관한다.")
    doc.add_paragraph()
    doc.add_paragraph("2025년 3월 12일")
    doc.add_paragraph()

    _table(doc, ["구분", "갑 (위탁자)", "을 (수탁자)"], [
        ["상호", "주식회사 넥스트이노베이션", "주식회사 코드랩솔루션"],
        ["대표이사", "김태현", "이승우"],
        ["사업자등록번호", "123-45-67890", "987-65-43210"],
        ["주소", "서울특별시 강남구 테헤란로 152, 12층", "서울특별시 서초구 서초대로 301, 8층"],
        ["기명날인", "(인)", "(인)"],
    ])

    path = OUT / "소프트웨어_개발_용역_계약서.docx"
    doc.save(str(path))
    print(f"  [2/5] {path.name}")


# ──────────────────────────────────────────────
# 3. 보고서
# ──────────────────────────────────────────────
def create_report():
    doc = Document()
    _style(doc)

    _heading(doc, "2024년 연간 사업실적 보고서")
    p = doc.add_paragraph("주식회사 넥스트이노베이션 | 전략기획실")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph("보고일: 2025년 2월 15일 | 대외비")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    _heading(doc, "I. 경영 실적 요약 (Executive Summary)", level=2)
    doc.add_paragraph(
        "2024년 당사는 매출액 3,128억 원(전년비 +27.3%), 영업이익 439억 원(영업이익률 14.0%), "
        "당기순이익 312억 원을 달성하였다. SaaS 구독 모델 전환이 가속화되면서 반복 매출(ARR) 비중이 "
        "58%로 확대되었으며, 고객 이탈률(Churn Rate)은 3.2%로 업계 평균(8~12%) 대비 현저히 낮은 수준을 "
        "유지하였다. 해외 매출은 563억 원으로 전체의 18%를 차지하며 전년 대비 41% 성장하였다."
    )

    _heading(doc, "II. 부문별 실적", level=2)

    _heading(doc, "1. SaaS 사업부", level=3)
    _table(doc, ["지표", "2023년", "2024년", "증감률"], [
        ["ARR (연간반복매출)", "1,240억", "1,814억", "+46.3%"],
        ["월간 활성 사용자(MAU)", "28만 명", "45만 명", "+60.7%"],
        ["엔터프라이즈 고객 수", "312사", "487사", "+56.1%"],
        ["평균 계약 단가(ACV)", "3,200만 원", "3,720만 원", "+16.3%"],
        ["고객 이탈률(Churn)", "3.8%", "3.2%", "-0.6%p"],
        ["NPS (순추천지수)", "62", "71", "+9"],
    ])
    doc.add_paragraph(
        "주요 성과: AI 문서 자동 분류 기능 출시(6월) 이후 엔터프라이즈 신규 계약이 전년 동기 대비 "
        "83% 증가하였다. 금융권 고객사(KB금융, 신한, 하나) 확보가 ACV 상승의 핵심 요인이었으며, "
        "대기업 전사 도입 사례(삼성SDS, LG CNS) 2건이 레퍼런스로 확보되었다."
    )

    _heading(doc, "2. 해외 사업부", level=3)
    _table(doc, ["시장", "매출", "전년비", "주요 고객"], [
        ["일본", "312억", "+52%", "소프트뱅크, NTT데이터, 라쿠텐"],
        ["동남아", "142억", "+38%", "DBS은행, Grab, Sea Group"],
        ["북미", "109억", "+21%", "Accenture(파일럿), Deloitte(POC)"],
    ])

    _heading(doc, "3. R&D", level=3)
    doc.add_paragraph(
        "2024년 R&D 투자액은 412억 원으로 매출의 13.2%를 집행하였다. 특허 출원 17건(등록 9건), "
        "국제 학회 논문 게재 4편의 성과를 거두었다. 핵심 기술 성과로는 자체 개발 경량 LLM "
        "'NextLM-7B'의 한국어 벤치마크 1위 달성(KoBEST, KLUE), 문서 OCR 정확도 98.7% 달성, "
        "실시간 협업 엔진 지연 시간 50ms 이하 달성이 있다."
    )

    _heading(doc, "III. 재무 현황", level=2)
    _table(doc, ["항목", "2023년", "2024년", "증감"], [
        ["매출액", "2,457억", "3,128억", "+27.3%"],
        ["매출원가", "1,228억", "1,501억", "+22.2%"],
        ["매출총이익", "1,229억", "1,627억", "+32.4%"],
        ["판관비", "891억", "1,188억", "+33.3%"],
        ["영업이익", "338억", "439억", "+29.9%"],
        ["영업이익률", "13.8%", "14.0%", "+0.2%p"],
        ["당기순이익", "241억", "312억", "+29.5%"],
        ["현금및현금성자산", "892억", "1,247억", "+39.8%"],
    ])

    _heading(doc, "IV. 2025년 전망 및 전략", level=2)
    doc.add_paragraph(
        "2025년 매출 목표는 4,200억 원(+34.3%)이며, 영업이익률 15% 이상을 목표로 한다. "
        "핵심 전략은 다음과 같다."
    )
    strategies = [
        "AI 제품 강화: AutoFlow 2.0 출시를 통한 AI 워크플로우 자동화 시장 선점",
        "글로벌 확장: 일본 현지법인 설립, 싱가포르 거점 확보, 북미 본격 진출",
        "엔터프라이즈 심화: 대기업 전사 도입 패키지 출시, 산업별 특화 솔루션 3종 개발",
        "인재 확보: AI/ML 연구인력 30명 충원, 글로벌 인재 채용 확대",
        "M&A: 문서 보안 및 전자서명 분야 스타트업 인수 검토",
    ]
    for s in strategies:
        doc.add_paragraph(s, style="List Bullet")

    _heading(doc, "V. 리스크 요인", level=2)
    _table(doc, ["리스크", "영향도", "대응 방안"], [
        ["글로벌 AI 규제 강화 (EU AI Act 등)", "높음", "컴플라이언스 TF 구성, 법무팀 2명 충원"],
        ["환율 변동 (엔화 약세 지속)", "중간", "선물환 헤지 비율 50%→70% 상향"],
        ["핵심 인력 이탈", "높음", "RSU 보상 확대, 기술 리더십 프로그램 도입"],
        ["경쟁사 AI 제품 출시 가속화", "중간", "차별화된 한국어 특화 + 산업 도메인 전문성 강화"],
    ])

    path = OUT / "2024년_연간_사업실적_보고서.docx"
    doc.save(str(path))
    print(f"  [3/5] {path.name}")


# ──────────────────────────────────────────────
# 4. 인사문서 (채용공고 + 인사발령)
# ──────────────────────────────────────────────
def create_hr_doc():
    doc = Document()
    _style(doc)

    _heading(doc, "2025년 상반기 AI/ML 엔지니어 채용 공고")
    p = doc.add_paragraph("주식회사 넥스트이노베이션 인사팀 | 2025년 3월 1일")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    _heading(doc, "1. 모집 부문", level=2)
    _table(doc, ["직무", "인원", "경력", "고용형태", "근무지"], [
        ["시니어 AI 연구원", "3명", "경력 5년 이상", "정규직", "서울 강남 본사"],
        ["ML 엔지니어", "5명", "경력 3년 이상", "정규직", "서울 강남 본사"],
        ["NLP 엔지니어", "4명", "경력 2년 이상", "정규직", "서울 강남 본사"],
        ["MLOps 엔지니어", "3명", "경력 3년 이상", "정규직", "서울 강남 본사"],
    ])

    _heading(doc, "2. 직무 상세", level=2)

    _heading(doc, "시니어 AI 연구원", level=3)
    doc.add_paragraph("[ 담당 업무 ]")
    tasks = [
        "자체 경량 LLM(NextLM) 아키텍처 설계 및 학습 파이프라인 개발",
        "문서 이해(Document Understanding) 모델 연구 및 고도화",
        "한국어 특화 사전학습 데이터 구축 및 벤치마크 평가 체계 설계",
        "국제 학회(ACL, EMNLP, NeurIPS 등) 논문 투고 및 기술 리더십 발휘",
        "주니어 연구원 멘토링 및 연구 방향 설정",
    ]
    for t in tasks:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_paragraph("[ 자격 요건 ]")
    reqs = [
        "Computer Science, AI, NLP 관련 석사 이상 (박사 우대)",
        "LLM/Foundation Model 학습 경험 (GPT, LLaMA, Gemma 등)",
        "PyTorch 기반 대규모 모델 학습 경험 (Multi-GPU, DeepSpeed/FSDP)",
        "주요 국제 학회 논문 게재 실적 1편 이상",
        "Python, C++ 능숙",
    ]
    for r in reqs:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_paragraph("[ 우대 사항 ]")
    pref = [
        "RAG(Retrieval-Augmented Generation) 시스템 구축 경험",
        "한국어 NLP 데이터셋 구축 또는 평가 경험",
        "RLHF/DPO 등 인간 피드백 기반 학습 경험",
        "오픈소스 AI 프로젝트 기여 경험",
    ]
    for p_item in pref:
        doc.add_paragraph(p_item, style="List Bullet")

    _heading(doc, "3. 보상 및 복리후생", level=2)
    benefits = [
        "연봉: 직무·경력에 따라 개별 협의 (시니어 연구원 기준 1.2억~2억 원 수준)",
        "성과급: 연간 기본급의 0~30% (개인 성과 + 회사 실적 연동)",
        "RSU(양도제한조건부주식): 시니어급 이상 부여",
        "GPU 크레딧: 개인 연구용 A100 GPU 월 500시간 제공",
        "학회 참석 지원: 해외 학회 연 2회 전액 지원 (등록비, 항공, 숙박)",
        "교육비 지원: 연간 500만 원 자기계발비",
        "유연근무: 코어타임(11:00~16:00) 외 자율 출퇴근",
        "원격근무: 주 2일 재택근무 가능",
        "건강검진: 연 1회 종합검진 (배우자 포함)",
        "식대: 점심 식대 월 30만 원 지원",
    ]
    for b in benefits:
        doc.add_paragraph(b, style="List Bullet")

    _heading(doc, "4. 전형 절차", level=2)
    doc.add_paragraph("서류 전형 → 코딩 테스트 → 1차 기술 면접 (90분) → 2차 컬쳐핏 면접 (60분) → 처우 협의 → 최종 합격")
    doc.add_paragraph("전형 기간: 서류 접수 후 약 3~4주 소요")
    doc.add_paragraph("지원 마감: 2025년 4월 15일(화) 23:59 (수시 채용으로 조기 마감될 수 있음)")

    _heading(doc, "5. 지원 방법", level=2)
    doc.add_paragraph("이메일: recruit@nextinnovation.co.kr")
    doc.add_paragraph("제출 서류: 이력서, 자기소개서, 포트폴리오(선택), 논문 목록(해당 시)")

    doc.add_page_break()

    # ─── 인사발령 ───
    _heading(doc, "인사발령 통보서")
    p = doc.add_paragraph("문서번호: HR-2025-0312-001 | 시행일: 2025년 3월 12일")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph(
        "아래와 같이 2025년 상반기 정기 인사발령을 시행하오니 업무에 참고하시기 바랍니다."
    )

    _heading(doc, "승진 발령", level=2)
    _table(doc, ["성명", "현 직급", "발령 직급", "소속", "비고"], [
        ["김수현", "책임연구원", "수석연구원", "AI연구소", "NextLM 프로젝트 리드"],
        ["이하은", "선임개발자", "책임개발자", "플랫폼개발팀", "SaaS 아키텍처 재설계 공로"],
        ["박준서", "대리", "과장", "글로벌영업팀", "일본 시장 매출 목표 150% 달성"],
        ["최예진", "사원", "대리", "마케팅팀", "브랜드 리뉴얼 프로젝트 우수 성과"],
    ])

    _heading(doc, "전보 발령", level=2)
    _table(doc, ["성명", "현 소속", "발령 소속", "직급", "사유"], [
        ["한승우", "백엔드개발팀", "AI플랫폼팀", "책임개발자", "AI 제품 개발 역량 강화"],
        ["정다은", "국내영업팀", "글로벌영업팀", "과장", "일본법인 설립 지원"],
        ["윤서준", "QA팀", "DevOps팀", "선임", "CI/CD 파이프라인 자동화 프로젝트"],
    ])

    doc.add_paragraph()
    doc.add_paragraph("주식회사 넥스트이노베이션")
    doc.add_paragraph("대표이사 김태현")

    path = OUT / "2025_상반기_인사문서.docx"
    doc.save(str(path))
    print(f"  [4/5] {path.name}")


# ──────────────────────────────────────────────
# 5. 제안서
# ──────────────────────────────────────────────
def create_proposal():
    doc = Document()
    _style(doc)

    title = doc.add_heading("AI 기반 스마트 업무자동화 플랫폼 도입 제안서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph("제안사: 주식회사 넥스트이노베이션")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("제안일: 2025년 3월 12일")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("대외비")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    _heading(doc, "I. 제안 개요", level=2)
    doc.add_paragraph(
        "본 제안서는 한국전력공사(이하 '고객사')의 사내 문서 관리 및 업무 프로세스 혁신을 위한 "
        "AI 기반 스마트 업무자동화 플랫폼 'AutoFlow Enterprise' 도입을 제안합니다. "
        "당사는 300여 개 기업의 업무 자동화 경험과 자체 개발 AI 엔진을 바탕으로, "
        "고객사의 디지털 전환 목표 달성을 지원하겠습니다."
    )

    _heading(doc, "1. 프로젝트 배경", level=3)
    doc.add_paragraph(
        "고객사는 연간 약 150만 건의 내부 문서를 생성·관리하고 있으며, "
        "직원 설문 결과 업무 시간의 약 32%가 문서 작성·검색·결재에 소요되는 것으로 파악되었습니다. "
        "또한 부서별 문서 양식 비표준화, 규정 준수 여부 수동 확인, 의사결정 지원 정보 분산 등의 "
        "문제가 업무 효율성을 저하시키고 있습니다."
    )

    _heading(doc, "2. 제안 목표", level=3)
    goals = [
        "문서 처리 시간 50% 단축 (현 평균 45분 → 20분 이하)",
        "규정 준수율 95% 이상 달성 (현 78% → 자동 검증 시스템)",
        "문서 검색 소요 시간 80% 감소 (현 평균 12분 → 2분 이하)",
        "연간 업무 자동화 절감 효과 32억 원 (인건비 환산 기준)",
    ]
    for g in goals:
        doc.add_paragraph(g, style="List Bullet")

    _heading(doc, "II. 솔루션 소개", level=2)

    _heading(doc, "1. 핵심 기능", level=3)
    _table(doc, ["기능", "설명", "기대 효과"], [
        ["AI 문서 자동 분류", "업로드된 문서를 자동으로 카테고리 분류 및 태깅\n(정확도 96.3%)", "문서 정리 시간 70% 절감"],
        ["지능형 문서 검색", "자연어 질의 기반 시맨틱 검색 + 하이브리드 검색\n(BM25 + Vector)", "검색 소요 시간 80% 감소"],
        ["자동 문서 요약", "긴 보고서를 핵심 내용 3~5줄로 자동 요약\n+ 키워드 추출", "문서 검토 시간 60% 절감"],
        ["규정 준수 자동 검증", "작성 문서의 사내 규정·법규 위반 사항 실시간 감지\n+ 조항 출처 제시", "규정 위반 사전 방지 90%"],
        ["회의록 자동 생성", "회의 녹음/텍스트 기반 자동 회의록 작성\n+ 액션 아이템 추출", "회의록 작성 시간 80% 절감"],
        ["스마트 일정 관리", "문서 내 일정 정보 자동 추출\n+ 캘린더 연동 + 리마인더", "일정 누락 방지"],
    ])

    _heading(doc, "2. 기술 아키텍처", level=3)
    doc.add_paragraph(
        "AutoFlow Enterprise는 다음과 같은 기술 스택으로 구성됩니다."
    )
    techs = [
        "AI Engine: 자체 개발 NextLM-7B (한국어 특화) + GPT-4 연동 (하이브리드)",
        "RAG Pipeline: Qdrant Vector DB + BM25 + Cross-Encoder Reranker",
        "Backend: FastAPI + PostgreSQL + Redis (세션/캐시)",
        "Frontend: React + TypeScript + Tailwind CSS",
        "Infrastructure: AWS EKS (Kubernetes) + CloudFront CDN",
        "Security: AES-256 암호화, OAuth 2.0 + SAML SSO, 감사 로그",
    ]
    for t in techs:
        doc.add_paragraph(t, style="List Bullet")

    _heading(doc, "III. 프로젝트 수행 계획", level=2)

    _heading(doc, "1. 추진 일정", level=3)
    _table(doc, ["단계", "기간", "주요 활동", "산출물"], [
        ["1단계: 분석·설계", "4주\n(4월)", "현행 업무 프로세스 분석\n요구사항 정의\n시스템 설계", "SRS, 설계서\nPrototype"],
        ["2단계: 개발", "12주\n(5~7월)", "AI 모델 커스터마이징\n시스템 개발\n데이터 마이그레이션", "개발 산출물\n테스트 보고서"],
        ["3단계: 테스트", "4주\n(8월)", "통합 테스트\n사용자 수용 테스트(UAT)\n성능 튜닝", "테스트 결과서\n성능 보고서"],
        ["4단계: 배포·안정화", "4주\n(9월)", "운영 환경 배포\n사용자 교육\n안정화 모니터링", "운영 매뉴얼\n교육 자료"],
    ])

    _heading(doc, "2. 투입 인력", level=3)
    _table(doc, ["역할", "인원", "등급", "주요 역량"], [
        ["PM(프로젝트 매니저)", "1명", "특급", "대규모 SI 프로젝트 관리 15년"],
        ["AI 아키텍트", "1명", "특급", "LLM/NLP 전문, 논문 게재 12편"],
        ["백엔드 개발", "3명", "고급~특급", "FastAPI, 대용량 시스템 설계"],
        ["프론트엔드 개발", "2명", "고급", "React, UX/UI 전문"],
        ["데이터 엔지니어", "2명", "고급", "ETL, Vector DB, 데이터 파이프라인"],
        ["QA 엔지니어", "1명", "중급", "자동화 테스트, 성능 테스트"],
    ])

    _heading(doc, "IV. 제안 금액", level=2)
    _table(doc, ["항목", "금액 (부가세 별도)", "비고"], [
        ["시스템 개발비", "6억 8천만 원", "커스터마이징 포함"],
        ["AI 모델 라이선스", "연 1억 2천만 원", "NextLM Enterprise + GPT-4"],
        ["클라우드 인프라", "연 4천8백만 원", "AWS 예상 비용"],
        ["유지보수 (연간)", "개발비의 15%", "1억 2백만 원/년"],
        ["교육 및 컨설팅", "3천만 원", "관리자+사용자 교육 3회"],
        ["합계 (초년도)", "9억 7천8백만 원", "-"],
    ])

    _heading(doc, "V. 기대 효과", level=2)
    _table(doc, ["항목", "현재", "도입 후", "개선 효과"], [
        ["문서 처리 시간", "45분/건", "20분/건", "55% 절감"],
        ["문서 검색 시간", "12분/건", "2분/건", "83% 절감"],
        ["규정 준수율", "78%", "96%", "+18%p"],
        ["회의록 작성 시간", "2시간/건", "15분/건", "87% 절감"],
        ["연간 절감 효과", "-", "32억 원", "ROI 327% (3년)"],
    ])

    _heading(doc, "VI. 당사 역량", level=2)
    doc.add_paragraph(
        "당사는 2018년 설립 이래 AI 기반 업무 자동화 분야에서 7년간의 경험을 보유하고 있으며, "
        "487개 기업에 솔루션을 제공하고 있습니다. 주요 레퍼런스로는 삼성SDS 전사 문서관리 시스템, "
        "KB금융그룹 AI 규정 검증 시스템, LG CNS 스마트 워크플로우 등이 있습니다."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "본 제안에 대해 궁금하신 사항이 있으시면 아래 연락처로 문의해 주시기 바랍니다."
    )
    doc.add_paragraph("담당자: 한도윤 상무 (글로벌사업본부)")
    doc.add_paragraph("연락처: doyun.han@nextinnovation.co.kr / 02-1234-5678")

    path = OUT / "AI_업무자동화_플랫폼_도입_제안서.docx"
    doc.save(str(path))
    print(f"  [5/5] {path.name}")


if __name__ == "__main__":
    print("테스트 문서 생성 중...")
    create_meeting_minutes()
    create_contract()
    create_report()
    create_hr_doc()
    create_proposal()
    print("완료!")
